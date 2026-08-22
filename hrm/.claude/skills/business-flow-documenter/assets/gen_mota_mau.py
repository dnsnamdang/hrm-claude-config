# -*- coding: utf-8 -*-
"""Sinh tai lieu mo ta nghiep vu man "Yeu cau kiem tra sua chua - bao hanh"."""
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Mô tả nghiệp vụ - Yêu cầu kiểm tra sửa chữa - bảo hành.docx")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)


def h(text, level):
    doc.add_paragraph(text, style="Heading %d" % level)


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = 1
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.cell(i, j).text = val
    return t


# ───────────────────────── BÌA ─────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("MÔ TẢ NGHIỆP VỤ")
r.bold = True
r.font.size = Pt(20)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Yêu cầu kiểm tra sửa chữa – bảo hành")
r.bold = True
r.font.size = Pt(16)

ip = doc.add_paragraph()
ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
ip.add_run("Phân hệ Bán hàng → Bán dịch vụ · Cập nhật ngày 20/08/2026")

doc.add_page_break()

# ───────────────────────── 1 ─────────────────────────
h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
para("Tài liệu mô tả đầy đủ nghiệp vụ của chức năng “Yêu cầu kiểm tra sửa chữa – bảo hành”: chức "
     "năng này dùng để làm gì, ai tham gia, phiếu đi qua những bước nào, mỗi bước hệ thống gửi "
     "thông báo cho ai, và những quy tắc bắt buộc phải tuân thủ.")
bullet("Nhân viên kinh doanh, nhân viên phòng tiếp nhận: hiểu mình phải làm gì ở từng bước.")
bullet("Cán bộ quản lý, người nghiệm thu: đối chiếu hệ thống có chạy đúng nghiệp vụ không.")
bullet("Bộ phận kiểm thử: dùng kèm file test case của cùng thư mục tài liệu.")

# ───────────────────────── 2 ─────────────────────────
h("2. CHỨC NĂNG NÀY DÙNG ĐỂ LÀM GÌ", 1)
h("2.1. Mục đích", 2)
para("Khi khách hàng báo thiết bị hỏng, cần kiểm tra, cần sửa chữa hoặc cần bảo hành, nhân viên "
     "kinh doanh lập một “Phiếu yêu cầu kiểm tra sửa chữa – bảo hành” để ghi nhận chính thức yêu "
     "cầu đó trên hệ thống, thay cho việc gọi điện hay nhắn tin rời rạc. Phiếu ghi rõ: khách hàng "
     "nào, ai là người liên hệ, sửa ở địa chỉ nào, những thiết bị nào cần xử lý và hiện trạng từng "
     "thiết bị. Sau đó phiếu được gửi cho phòng tiếp nhận để bố trí người xử lý.")
h("2.2. Vị trí trong luồng dịch vụ", 2)
para("Đây là chứng từ MỞ ĐẦU của toàn bộ luồng dịch vụ. Từ phiếu này, phòng tiếp nhận lập tiếp "
     "“Phiếu xử lý yêu cầu”, rồi mới đến các chứng từ phía sau (cung cấp thông tin, báo giá, hợp "
     "đồng, nghiệm thu, thanh toán). Không có phiếu yêu cầu thì không có căn cứ để mở các chứng từ "
     "phía sau.")
h("2.3. Giá trị mang lại", 2)
bullet("Yêu cầu của khách không bị rơi: mọi yêu cầu đều có số phiếu, có người chịu trách nhiệm.")
bullet("Phòng tiếp nhận biết ngay có việc mới nhờ thông báo tự động, không phải chờ ai gọi.")
bullet("Lãnh đạo tra cứu được: khách nào yêu cầu gì, ai tiếp nhận, đang ở bước nào, bao lâu rồi.")
bullet("Có sẵn dữ liệu thiết bị của khách để chọn, không phải gõ tay lại tên/model/serial.")

# ───────────────────────── 3 ─────────────────────────
h("3. NHỮNG AI THAM GIA", 1)
table([
    ["Vai trò", "Làm gì trong nghiệp vụ này"],
    ["Người lập phiếu\n(thường là nhân viên kinh doanh)",
     "Tiếp nhận yêu cầu của khách, lập phiếu, chọn thiết bị cần xử lý, mô tả hiện trạng, chọn "
     "phòng tiếp nhận rồi gửi đi. Sửa và xóa được phiếu của chính mình khi còn ở trạng thái "
     "“Đang tạo”."],
    ["Phòng tiếp nhận xử lý\n(phòng kỹ thuật / bảo hành…)",
     "Nhận thông báo, xem phiếu rồi chọn một trong ba việc: lập Phiếu xử lý yêu cầu để bắt tay "
     "vào làm, chuyển sang phòng khác nếu không thuộc phạm vi của mình, hoặc từ chối và trả lại "
     "người lập kèm lý do."],
    ["Cán bộ quản lý",
     "Theo dõi toàn bộ phiếu trong phạm vi quyền của mình (tổng công ty / công ty / phòng ban), "
     "in và xuất dữ liệu để báo cáo."],
    ["Khách hàng", "Không thao tác trên hệ thống. Thông tin khách được lấy từ danh mục khách hàng."],
])

# ───────────────────────── 4 ─────────────────────────
h("4. VÒNG ĐỜI CỦA PHIẾU", 1)
para("Phiếu có 9 trạng thái. Ba trạng thái đầu do chính màn hình này tạo ra; các trạng thái sau do "
     "những chứng từ phía sau cập nhật ngược về.")
table([
    ["Trạng thái", "Ý nghĩa", "Ai làm phiếu chuyển sang trạng thái này"],
    ["Đang tạo", "Phiếu nháp, chưa gửi đi. Chỉ người lập nhìn thấy.",
     "Người lập bấm “Lưu nháp”, hoặc phòng tiếp nhận bấm “Từ chối”."],
    ["Chờ xử lý", "Đã gửi cho phòng tiếp nhận, đang chờ phòng đó xử lý.",
     "Người lập bấm “Lưu và gửi duyệt”."],
    ["Đang xử lý", "Phòng tiếp nhận đã lập Phiếu xử lý yêu cầu.", "Phòng tiếp nhận."],
    ["Đang CCTT", "Đang cung cấp thông tin cho khách.", "Chứng từ phía sau."],
    ["Đã CCTT báo giá", "Đã cung cấp thông tin phục vụ báo giá.", "Chứng từ phía sau."],
    ["Đã báo giá", "Đã gửi báo giá cho khách.", "Chứng từ phía sau."],
    ["Đã lập hợp đồng", "Đã ký hợp đồng dịch vụ với khách.", "Chứng từ phía sau."],
    ["Đã xử lý", "Yêu cầu đã hoàn tất.", "Chứng từ phía sau."],
    ["Đã tư vấn điện thoại", "Xử lý xong bằng tư vấn qua điện thoại, không cần đi hiện trường.",
     "Chứng từ phía sau."],
])
para("Màu nhãn trạng thái do hệ thống quy định thống nhất: xám (nháp), cam (đang chờ), xanh dương "
     "(đang thực hiện), xanh nhạt (đang theo dõi), xanh lá (hoàn thành).")

# ───────────────────────── 5 ─────────────────────────
h("5. LUỒNG HOẠT ĐỘNG CHI TIẾT", 1)

h("5.1. Bước 1 — Lập phiếu", 2)
para("Người lập vào Bán hàng → Bán dịch vụ → Yêu cầu sửa chữa - bảo hành, bấm “Tạo mới”, rồi:")
bullet("Bấm vào ô “Khách hàng” để mở cửa sổ tìm và chọn khách. Chọn xong hệ thống tự nạp: danh "
       "sách người liên hệ, danh sách địa chỉ giao nhận, loại hình tổ chức và toàn bộ trang thiết "
       "bị hiện có của khách.")
bullet("Chọn người liên hệ (số điện thoại tự điền theo người được chọn), chọn địa chỉ sửa chữa, "
       "nhập ghi chú và chọn phòng tiếp nhận xử lý.")
bullet("Ở bảng “Danh mục trang thiết bị hiện có của khách hàng”, bấm nút thêm ở từng thiết bị cần "
       "xử lý. Thiết bị được đưa lên bảng “Danh sách thiết bị cần kiểm tra sửa chữa – bảo hành”.")
bullet("Với mỗi thiết bị: chọn hoặc gõ serial, nhập mô tả yêu cầu (hỏng gì, cần làm gì) và đính "
       "kèm tài liệu nếu có (ảnh hiện trạng, biên bản…).")
para("Người lập có hai lựa chọn: “Lưu nháp” để lưu tạm và làm tiếp sau, hoặc “Lưu và gửi duyệt” "
     "để gửi ngay cho phòng tiếp nhận. Hệ thống sinh số phiếu ngay từ lần lưu đầu tiên, kể cả lưu "
     "nháp, và số này không đổi về sau.")

h("5.2. Bước 2 — Gửi cho phòng tiếp nhận", 2)
para("Khi bấm “Lưu và gửi duyệt”, hệ thống hỏi xác nhận một lần nữa. Sau khi xác nhận:")
bullet("Phiếu chuyển sang trạng thái “Chờ xử lý”.")
bullet("Hệ thống ghi lại thời điểm gửi yêu cầu.")
bullet("Hệ thống gửi thông báo cho toàn bộ nhân viên của phòng tiếp nhận (xem mục 6).")
bullet("Phiếu khóa lại: người lập không còn sửa hay xóa được nữa.")

h("5.3. Bước 3 — Phòng tiếp nhận xử lý", 2)
para("Người của phòng tiếp nhận mở phiếu và chọn một trong ba hướng:")
table([
    ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
    ["Tạo phiếu xử lý yêu cầu",
     "Bấm “Tạo phiếu xử lý yêu cầu”, hệ thống mở màn lập phiếu xử lý với thông tin lấy sẵn từ "
     "phiếu yêu cầu.",
     "Phiếu chuyển sang “Đang xử lý”. Từ lúc này không chuyển phòng và không từ chối được nữa."],
    ["Chuyển phòng tiếp nhận",
     "Bấm “Chuyển phòng tiếp nhận”, chọn phòng mới rồi xác nhận. Không được chọn lại đúng phòng "
     "đang tiếp nhận.",
     "Phiếu vẫn ở “Chờ xử lý” nhưng đổi phòng tiếp nhận; thời điểm gửi yêu cầu được cập nhật lại "
     "và toàn bộ nhân viên phòng mới nhận được thông báo."],
    ["Từ chối",
     "Bấm “Từ chối”, bắt buộc nhập lý do rồi xác nhận.",
     "Phiếu quay về “Đang tạo”, lý do được lưu lại trên phiếu và người lập nhận được thông báo "
     "kèm lý do. Người lập sửa rồi gửi lại được, số phiếu giữ nguyên."],
])

h("5.4. Bước 4 — Đi tiếp trong luồng dịch vụ", 2)
para("Sau khi có Phiếu xử lý yêu cầu, các bước tiếp theo (cung cấp thông tin, báo giá, hợp đồng, "
     "nghiệm thu…) diễn ra ở những màn hình khác. Trạng thái của phiếu yêu cầu tự cập nhật theo "
     "tiến độ đó để người lập và cán bộ quản lý theo dõi được mà không cần hỏi lại phòng kỹ thuật.")

# ───────────────────────── 6 ─────────────────────────
h("6. THÔNG BÁO — AI NHẬN, KHI NÀO, NỘI DUNG GÌ", 1)
para("Hệ thống gửi thông báo qua chuông thông báo trên thanh công cụ. Có đúng hai sự kiện phát "
     "sinh thông báo:")
table([
    ["Sự kiện", "Ai nhận", "Nội dung thông báo", "Bấm vào thì đi đâu"],
    ["Phiếu được gửi đi (người lập bấm “Lưu và gửi duyệt”)",
     "TOÀN BỘ nhân viên đang thuộc phòng tiếp nhận được chọn — không chỉ trưởng phòng.",
     "[YCSCBH] Chờ duyệt: <số phiếu>. Khách hàng: <tên khách hàng>.",
     "Mở thẳng màn chi tiết của phiếu đó."],
    ["Chuyển phòng tiếp nhận sang phòng khác",
     "TOÀN BỘ nhân viên của phòng tiếp nhận MỚI. Phòng cũ không nhận thêm thông báo nào.",
     "[YCSCBH] Chờ duyệt: <số phiếu>. Khách hàng: <tên khách hàng>.",
     "Mở thẳng màn chi tiết của phiếu đó."],
    ["Phòng tiếp nhận từ chối phiếu",
     "NGƯỜI LẬP phiếu.",
     "[YCSCBH] Từ chối: <số phiếu>. Lý do: <lý do từ chối>.",
     "Mở thẳng màn chi tiết của phiếu đó."],
])
para("Quy ước chung về thông báo:")
bullet("Số phiếu luôn được in đậm để người nhận nhận ra ngay đối tượng.")
bullet("Nội dung giới hạn độ dài; phần bị cắt trước là ghi chú, số phiếu luôn được giữ.")
bullet("Trường hợp người lập cũng đang thuộc chính phòng tiếp nhận thì người đó nhận thông báo "
       "“Chờ duyệt” như mọi thành viên khác của phòng — đây là hành vi đúng.")
bullet("Lỗi gửi thông báo không làm hỏng nghiệp vụ: phiếu vẫn được lưu và vẫn đổi trạng thái bình "
       "thường.")
bullet("Các bước phía sau của luồng dịch vụ (báo giá, hợp đồng…) có thông báo riêng, thuộc phạm "
       "vi tài liệu của những màn hình đó.")

# ───────────────────────── 7 ─────────────────────────
h("7. PHÂN QUYỀN — AI THẤY GÌ, AI LÀM ĐƯỢC GÌ", 1)
h("7.1. Quyền nhìn thấy dữ liệu", 2)
table([
    ["Tên quyền", "Nhìn thấy những phiếu nào"],
    ["Xem yêu cầu đi kiểm tra sửa chữa - bảo hành theo tổng công ty",
     "Phiếu của mọi công ty trong tập đoàn."],
    ["Xem yêu cầu đi kiểm tra sửa chữa - bảo hành theo công ty",
     "Phiếu của công ty mình, cộng thêm phiếu do chính mình lập ở nơi khác."],
    ["Xem yêu cầu đi kiểm tra sửa chữa - bảo hành theo phòng ban",
     "Phiếu của các phòng mình được giao quản lý, cộng thêm phiếu do chính mình lập."],
    ["Xử lý yêu cầu sửa chữa",
     "Ngoài phạm vi trên, còn thấy mọi phiếu gửi về phòng mình đang công tác."],
    ["(Không có quyền nào ở trên)",
     "Vẫn vào được màn hình nhưng chỉ thấy phiếu do chính mình lập."],
])
para("Hai quy tắc luôn đúng, không phụ thuộc quyền:")
bullet("Phiếu “Đang tạo” của người khác không hiện với bất kỳ ai — kể cả người có quyền cao nhất.")
bullet("Ai cũng tự lập được phiếu, và tự sửa/xóa được phiếu nháp của chính mình.")

h("7.2. Quyền thao tác", 2)
table([
    ["Thao tác", "Điều kiện được phép"],
    ["Lập phiếu mới", "Mọi người dùng vào được màn hình."],
    ["Sửa phiếu", "Phiếu ở “Đang tạo” VÀ do chính mình lập."],
    ["Xóa phiếu", "Phiếu ở “Đang tạo” VÀ do chính mình lập. Xóa phiếu là xóa cả các dòng thiết bị."],
    ["Tạo phiếu xử lý yêu cầu", "Có quyền “Xử lý yêu cầu sửa chữa”, phiếu đang “Chờ xử lý”, đúng "
     "phòng tiếp nhận và phiếu chưa có phiếu xử lý nào."],
    ["Chuyển phòng tiếp nhận", "Cùng điều kiện với Tạo phiếu xử lý yêu cầu."],
    ["Từ chối", "Cùng điều kiện với Tạo phiếu xử lý yêu cầu."],
    ["In phiếu / In danh sách / Xuất Excel", "Mọi người dùng thấy được phiếu."],
])
para("Người mang vai trò Quản trị cấp cao thao tác được với phiếu của mọi phòng tiếp nhận, không "
     "bị giới hạn theo phòng đang công tác.")
para("Nút nào không đủ điều kiện thì hệ thống ẨN HẲN, không hiển thị dạng mờ; danh sách nút ở màn "
     "chi tiết luôn bằng đúng danh sách nút ở dòng tương ứng ngoài màn danh sách.")

# ───────────────────────── 8 ─────────────────────────
h("8. QUY TẮC NGHIỆP VỤ BẮT BUỘC", 1)
h("8.1. Bắt buộc nhập", 2)
table([
    ["Khi bấm", "Bắt buộc phải có"],
    ["Lưu nháp", "Chỉ cần chọn Khách hàng."],
    ["Lưu và gửi duyệt",
     "Khách hàng · Người liên hệ · Địa chỉ sửa chữa · Ghi chú · Phòng tiếp nhận xử lý · ít nhất "
     "một thiết bị · mô tả yêu cầu của từng thiết bị · serial của từng thiết bị."],
])
para("Cách chia này để người lập ghi nhận nhanh yêu cầu của khách trước, bổ sung thông tin còn "
     "thiếu sau, nhưng khi đã gửi cho phòng khác thì hồ sơ phải đủ.")

h("8.2. Thiết bị và serial", 2)
bullet("Danh sách thiết bị của khách được gom từ ba nguồn: hàng công ty đã bán cho khách, thiết bị "
       "cũ của khách và thiết bị do nhà cung cấp khác cung cấp.")
bullet("Serial có hai cách nhập: chọn từ danh sách serial hệ thống đang quản lý, hoặc gõ tay serial "
       "mới nếu thiết bị chưa được khai.")
bullet("Cùng một serial của cùng một loại thiết bị không được đưa vào phiếu hai lần.")
bullet("Khi lưu, danh sách thiết bị trên màn hình thay thế toàn bộ danh sách cũ của phiếu.")

h("8.3. Khóa chỉnh sửa", 2)
bullet("Phiếu đã rời trạng thái “Đang tạo” thì không sửa, không xóa được nữa — kể cả khi cố mở "
       "thẳng màn Sửa bằng đường dẫn.")
bullet("Muốn sửa lại phiếu đã gửi, phòng tiếp nhận phải Từ chối để trả phiếu về “Đang tạo”.")

h("8.4. Danh mục bị khóa", 2)
para("Danh sách chọn khách hàng, phòng ban… chỉ liệt kê những bản ghi còn hoạt động. Nhưng nếu một "
     "phiếu cũ đang dùng bản ghi nay đã bị khóa thì giá trị đó vẫn hiển thị đầy đủ khi mở lại "
     "phiếu, không bị mất và không tự đổi sang giá trị khác.")

h("8.5. Số phiếu", 2)
para("Số phiếu sinh tự động theo dạng <mã công ty>.YCSCBH.<năm>.<số thứ tự>, cấp ngay ở lần lưu "
     "đầu tiên và không thay đổi trong suốt vòng đời phiếu, kể cả khi bị từ chối rồi gửi lại.")

# ───────────────────────── 9 ─────────────────────────
h("9. TRA CỨU, IN VÀ XUẤT DỮ LIỆU", 1)
bullet("Tìm nhanh: gõ một ô duy nhất, tìm được theo số phiếu, tên khách hàng hoặc tên người tạo.")
bullet("Bộ lọc nâng cao: trạng thái, khách hàng, người yêu cầu, tên thiết bị, tỉnh/thành của khách, "
       "khoảng ngày tạo, công ty và phòng ban. Chọn xong hệ thống tự lọc, không phải bấm thêm nút.")
bullet("Tùy chỉnh cột: bật/tắt cột hiển thị và hệ thống ghi nhớ lựa chọn cho lần vào sau.")
bullet("In phiếu: in đúng mẫu biểu đang dùng, có đủ khối ký của Người yêu cầu, Trưởng phòng yêu "
       "cầu, Phòng nhận yêu cầu và Ban giám đốc.")
bullet("In danh sách: in theo đúng bộ lọc đang áp dụng, khổ ngang.")
bullet("Xuất Excel: chọn được trường và thứ tự cột; xuất theo đúng bộ lọc đang áp dụng, có báo "
       "tiến độ khi dữ liệu lớn.")

# ───────────────────────── 10 ─────────────────────────
h("10. LIÊN THÔNG VỚI PHẦN MỀM ERP", 1)
para("Màn hình này dùng chung dữ liệu với màn tương ứng bên phần mềm ERP: lập phiếu ở cổng nào thì "
     "cổng còn lại cũng thấy ngay, số phiếu chạy chung một dãy, danh mục khách hàng và trang thiết "
     "bị của khách cũng là một nguồn duy nhất. Nhờ vậy giai đoạn chuyển đổi hai bên vẫn làm việc "
     "song song được, không phải nhập liệu hai lần.")
para("Nghiệp vụ được giữ giống hệt bản ERP. Có ba điểm cố ý làm khác, đều theo hướng chặt hơn "
     "hoặc thuận tiện hơn:")
table([
    ["Điểm khác", "Bên ERP", "Bên hệ thống mới", "Lý do"],
    ["Lưu nháp", "Bắt nhập đủ mọi thông tin ngay cả khi chỉ lưu tạm.",
     "Lưu nháp chỉ cần chọn khách hàng.",
     "Ghi nhận nhanh yêu cầu của khách khi đang nghe điện thoại, bổ sung sau."],
    ["Tìm theo tên thiết bị", "Kết quả trả về bị sai lệch.",
     "Trả về đúng những phiếu có thiết bị khớp từ khóa.", "Sửa lỗi tìm kiếm của bản cũ."],
    ["Thông báo khi từ chối", "Không báo cho ai, người lập phải tự vào xem mới biết.",
     "Báo ngay cho người lập kèm lý do.", "Rút ngắn thời gian phiếu bị treo."],
])

# ───────────────────────── 11 ─────────────────────────
h("11. GIỚI HẠN HIỆN TẠI", 1)
bullet("Màn “Phiếu xử lý yêu cầu” chưa có trên hệ thống mới. Bấm “Tạo phiếu xử lý yêu cầu” hiện "
       "báo hướng dẫn xử lý tạm bên phần mềm ERP; khi màn đó hoàn thành sẽ nối thẳng.")
bullet("In danh sách khi không đặt bộ lọc sẽ in toàn bộ dữ liệu nên rất dài — nên lọc trước khi in.")

doc.save(OUT)
print("Da tao:", OUT)
d = Document(OUT)
print("Heading 1:", sum(1 for p in d.paragraphs if p.style.name == "Heading 1"),
      "| Bang:", len(d.tables), "| Doan:", len(d.paragraphs))
