# -*- coding: utf-8 -*-
"""
Sinh tai lieu HDSD (.docx) cho man "Danh muc dich vu sua chua va chi phi khac".

Dung file mau HDSD_Bomlist.docx lam KHUNG:
  - Giu Bia (logo) + MUC LUC (TOC field) + DANH MUC HINH ANH (TOF field) + toan bo styles
  - Strip body tu heading "TONG QUAN" toi het (giu sectPr)
  - Rebuild noi dung moi bang chinh cac style cua template (khong ep direct formatting)
  - Caption clone tu proto co SEQ field -> tu danh so Hinh 1, Hinh 2...
"""
import copy
import os
import shutil
import sys

# Console Windows mac dinh cp1252 -> print() chuoi tieng Viet se nem UnicodeEncodeError.
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

# Khung tai lieu: uu tien file DA DONG GOI TRONG REPO (ai clone ve cung co).
# `D:\CompanyProject\Document\...` chi ton tai tren may 1 nguoi -> khong dung lam nguon chinh.
TEMPLATE_CANDIDATES = [
    r"d:\CompanyProject\hrm\hrm-claude-config\hrm\.claude\skills\hdsd-documenter\assets\HDSD_MAU.docx",
    r"D:\CompanyProject\Document\HDSD_Bomlist.docx",
]
TEMPLATE = next((t for t in TEMPLATE_CANDIDATES if os.path.exists(t)), None)
assert TEMPLATE, "Khong tim thay file khung HDSD nao"
FEATURE_DIR = r"d:\CompanyProject\hrm\hrm-claude-config\hrm\.plans\gop-db\customer-care-cost-catalog"
SHOTS = os.path.join(FEATURE_DIR, "hdsd_costs_shots")
OUTPUT = os.path.join(FEATURE_DIR, "HDSD_Danh muc dich vu sua chua va chi phi khac.docx")

IMG = lambda n: os.path.join(SHOTS, n)

shutil.copyfile(TEMPLATE, OUTPUT)
doc = Document(OUTPUT)
body = doc.element.body

# ---------------------------------------------------------------- bia
# ⚠️ Word cat dong nay thanh NHIEU run -> replace tren tung run.text KHONG khop chuoi
# ("Luồng nghiệp vụ", ": Tổng hợp", " Bomlist)"...). Phai don het text ve run dau.
BIA_MOI = "(Màn hình: Danh mục dịch vụ sửa chữa và chi phí khác)"
# Dong tieu de bia = dong dang "(...)" nam ngay sau "TÂN PHÁT ETEK" trong ~20 doan dau.
# Bat theo VI TRI, khong bat theo noi dung, de dung duoc voi moi file khung
# ("(Danh mục Khách hàng)" cua HDSD_MAU hay "(Luồng nghiệp vụ: Tổng hợp Bomlist)" cua Bomlist).
for i, p in enumerate(doc.paragraphs[:20]):
    txt = p.text.strip()
    if txt.startswith("(") and txt.endswith(")") and p.runs:
        for run in p.runs[1:]:
            run.text = ""
        p.runs[0].text = BIA_MOI
        print("Da doi dong bia:", txt, "->", BIA_MOI)
        break
else:
    raise AssertionError("Khong tim thay dong tieu de tren bia (dang '(...)')")

# ------------------------------------------------- luu proto Caption
caption_proto = None
for p in doc.paragraphs:
    if p.style.name == "Caption" and "Hình" in p.text:
        caption_proto = copy.deepcopy(p._p)
        break
assert caption_proto is not None, "Khong tim thay paragraph Caption mau"

# ------------------------------- strip body tu Heading 1 THU BA tro di
# Heading 1 #1 = "MỤC LỤC", #2 = "DANH MỤC HÌNH ẢNH" -> giu lai.
# #3 tro di la than bai cua tai lieu mau -> bo het.
# Bat theo VI TRI thay vi theo text ("TỔNG QUAN" / "TỔNG QUAN PHẦN MỀM" moi mau mot khac).
# ⚠️ So khop text bang Paragraph(child, doc).text, KHONG dung child.itertext() —
#    itertext() gom ca text trong field/bookmark nen tra ve "TỔNG QUANTỔNG QUANTỔNG QUAN".
from docx.text.paragraph import Paragraph

sectPr = body.find(qn('w:sectPr'))
children = list(body.iterchildren())


def heading_level(child):
    if child.tag != qn('w:p'):
        return None
    pPr = child.find(qn('w:pPr'))
    if pPr is None:
        return None
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return None
    val = (pStyle.get(qn('w:val')) or "").replace(" ", "")
    return val if val.startswith("Heading") else None


h1_positions = [i for i, c in enumerate(children) if heading_level(c) == "Heading1"]
assert len(h1_positions) >= 3, "File khung khong du 3 Heading 1 (MUC LUC / DANH MUC HINH ANH / than bai)"
start = h1_positions[2]
# Giu lai tieu de than bai cua file khung -> cuoi script assert chung KHONG con trong output.
# Day la cach chac chan nhat de bat loi "muc luc van hien noi dung tai lieu mau".
TEMPLATE_BODY_HEADINGS = [
    Paragraph(children[i], doc).text.strip() for i in h1_positions[2:]
]
print("Strip body tu:", TEMPLATE_BODY_HEADINGS[0][:40])

for child in children[start:]:
    if child is sectPr:
        continue
    body.remove(child)


# ----------------------------------------------------------- helpers
def _append(par):
    """Chen paragraph moi ngay truoc sectPr (sectPr luon la child cuoi)."""
    sectPr.addprevious(par._p)
    return par


def h(text, level):
    p = doc.add_paragraph(text, style="Heading %d" % level)
    return _append(p)


def h1(t): return h(t, 1)
def h2(t): return h(t, 2)
def h3(t): return h(t, 3)


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return _append(p)


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return _append(p)


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = 1  # CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.cell(ri, ci).text = val
    sectPr.addprevious(t._tbl)
    # 1 dong trong sau bang cho thoang
    _append(doc.add_paragraph())
    return t


def image(filename, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(IMG(filename), width=Inches(6.0))
    _append(p)
    cap = copy.deepcopy(caption_proto)
    runs = cap.findall(qn('w:r'))
    # run cuoi cung giu phan ": <mo ta>"
    for r in runs:
        for tnode in r.findall(qn('w:t')):
            if tnode.text and tnode.text.startswith(":"):
                tnode.text = ": " + caption_text
    sectPr.addprevious(cap)


# ================================================================
# TONG QUAN
# ================================================================
h1("TỔNG QUAN")

h2("1. Mục tiêu")
para("Màn hình “Danh mục dịch vụ sửa chữa và chi phí khác” dùng để khai báo và quản lý danh sách các "
     "dịch vụ sửa chữa cùng các khoản chi phí khác phát sinh trong hoạt động dịch vụ. Mỗi dòng danh mục "
     "khai báo 3 thông số dùng cho các nghiệp vụ phía sau: % Tính giá vốn, % VAT và ĐM giảm giá "
     "(định mức giảm giá). Danh mục này là nguồn dữ liệu cho Báo giá hãng và Hợp đồng hãng, vì vậy dòng "
     "đã phát sinh chứng từ sẽ không được xóa mà chỉ được chuyển sang trạng thái Khóa.")
para("Lưu ý quan trọng: danh mục này dùng CHUNG dữ liệu với màn tương ứng bên phần mềm ERP "
     "(Kế toán → Danh mục chi phí, phần “dịch vụ sửa chữa và chi phí khác”). Sửa ở cổng nào thì cổng "
     "còn lại cũng thấy ngay, không cần đồng bộ thủ công.")

h2("2. Đường dẫn truy cập")
para("Trên menu trái của phân hệ CSKH, vào nhóm “Danh mục - Dịch vụ” rồi chọn:")
bullet("“Danh mục dịch vụ sửa chữa và chi phí khác” (đường dẫn: /customer-care/costs).")
image("hdsd_costs_09_menu.png", "Đường dẫn vào màn hình từ nhóm menu “Danh mục - Dịch vụ”")
para("Mục menu này chỉ hiển thị khi tài khoản có ít nhất một trong hai quyền “Quản lý dịch vụ sửa chữa "
     "và chi phí khác” hoặc “Xem dịch vụ sửa chữa và chi phí khác”. Nếu không có quyền nào, mục menu bị "
     "ẩn; trường hợp gõ thẳng đường dẫn thì hệ thống chặn và không trả về dữ liệu.")

h2("3. Vai trò tham gia")
table([
    ["Vai trò", "Thao tác chính"],
    ["Người quản lý danh mục dịch vụ",
     "Tạo mới, Sửa, Xóa, Khóa / Mở khóa dịch vụ và chi phí; khai báo ĐM giảm giá cho công ty của mình."],
    ["Người dùng chỉ xem",
     "Xem danh sách, xem chi tiết, tìm kiếm và Xuất Excel."],
    ["Bộ phận Báo giá / Hợp đồng hãng",
     "Là bên sử dụng danh mục này ở khâu lập Báo giá hãng và Hợp đồng hãng."],
])

h2("4. Các trạng thái")
table([
    ["Trạng thái", "Ý nghĩa"],
    ["Hoạt động", "Dòng danh mục đang dùng được; sửa / khóa / xóa được (nếu không phải dòng hệ thống)."],
    ["Khóa", "Ngừng sử dụng. Không sửa và không xóa được; muốn dùng lại phải Mở khóa trước."],
])
para("Lưu ý: trạng thái ở màn này lưu theo quy ước 1 = Hoạt động và 0 = Khóa, khác một số danh mục khác "
     "trong hệ thống. Người dùng cuối chỉ cần nhìn theo nhãn hiển thị trên màn hình.")

h2("5. Phân loại dịch vụ / chi phí")
table([
    ["Phân loại", "Ý nghĩa"],
    ["Có tính doanh thu",
     "Dịch vụ được tính vào doanh thu. Tương ứng ô “Dịch vụ có tính doanh thu” được tích trên form."],
    ["Chi phí khác",
     "Khoản chi phí không tính doanh thu. Tương ứng ô “Dịch vụ có tính doanh thu” KHÔNG được tích."],
])

h2("6. Luồng sử dụng")
para("Đây là màn danh mục, không có luồng trình duyệt nhiều cấp. Trình tự sử dụng thông thường:")
bullet("Bước 1 — Người quản lý danh mục tạo mới dịch vụ / chi phí, khai báo % Tính giá vốn, % VAT và "
       "ĐM giảm giá.")
bullet("Bước 2 — Bộ phận nghiệp vụ chọn dịch vụ từ danh mục này khi lập Báo giá hãng / Hợp đồng hãng.")
bullet("Bước 3 (khi cần điều chỉnh) — Người quản lý sửa lại thông số; các chứng từ đã lập trước đó "
       "không bị tính lại.")
bullet("Bước 4 (khi ngừng dùng) — Bấm Xóa. Nếu dịch vụ chưa phát sinh chứng từ, hệ thống xóa hẳn; nếu "
       "đã phát sinh, hệ thống tự chuyển sang trạng thái Khóa và giữ nguyên dữ liệu.")

# ================================================================
# PHAN 1
# ================================================================
h1("PHẦN 1: DANH SÁCH & TÌM KIẾM")

h2("1.1 Danh sách")
image("hdsd_costs_01_danhsach.png", "Màn hình danh sách Danh mục dịch vụ sửa chữa và chi phí khác")
para("Màn hình gồm 2 khối: panel bộ lọc ở trên (mặc định thu gọn) và bảng danh sách ở dưới. "
     "Bảng gồm 9 cột:")
table([
    ["Cột", "Nội dung hiển thị"],
    ["STT", "Số thứ tự, đánh liên tục theo trang (trang 2 bắt đầu từ 11)."],
    ["Tên dịch vụ / chi phí", "Tên khai báo, in đậm. Sắp xếp được."],
    ["Phân loại", "Nhãn màu “Có tính doanh thu” (xanh) hoặc “Chi phí khác” (đỏ)."],
    ["ĐM giảm giá",
     "Định mức giảm giá theo % của CÔNG TY ĐANG CHỌN. Chưa khai báo thì hiển thị “—”. Sắp xếp được."],
    ["% Tính giá vốn", "Tỷ lệ tính giá vốn, in đậm. Sắp xếp được."],
    ["% VAT", "Thuế suất VAT. Sắp xếp được."],
    ["Trạng thái",
     "Nhãn “Hoạt động” / “Khóa”, kèm nút ổ khóa để Khóa hoặc Mở khóa (chỉ hiện khi có quyền quản lý)."],
    ["Cập nhật",
     "Thời điểm sửa gần nhất và người sửa (dạng MÃ NV - Họ tên). Bản ghi chưa từng sửa thì hiển thị "
     "người tạo. Sắp xếp được."],
    ["Hành động", "Các nút Xem, Sửa, Xóa."],
])
para("Cột ĐM giảm giá là điểm cần lưu ý: giá trị này khai báo riêng cho từng công ty. Cùng một dịch vụ, "
     "người dùng của công ty A và công ty B có thể nhìn thấy hai mức khác nhau — đó là hành vi đúng, "
     "không phải lỗi dữ liệu.")
para("Chân bảng hiển thị dòng đếm “Hiển thị a–b / N dịch vụ / chi phí”. Số dòng mỗi trang mặc định là "
     "10, chọn được 5 / 10 / 20 / 50. Khi mới vào màn, danh sách sắp xếp theo ngày tạo giảm dần "
     "(bản ghi mới nhất ở đầu).")

h2("1.2 Tìm kiếm và bộ lọc")
image("hdsd_costs_02_boloc.png", "Bộ lọc nâng cao của Danh mục dịch vụ sửa chữa và chi phí khác")
para("Ô tìm nhanh nằm ngay dưới tiêu đề panel, tìm theo tên dịch vụ / chi phí. Lưu ý ô này KHÔNG tự lọc "
     "khi đang gõ — phải bấm nút “Tìm kiếm” (hoặc nhấn Enter) thì danh sách mới được lọc lại.")
para("Bấm “Tìm kiếm nâng cao” để mở thêm 4 tiêu chí:")
bullet("Tên dịch vụ / chi phí — tìm gần đúng theo tên.")
bullet("Phân loại — “Dịch vụ có tính doanh thu” hoặc “Chi phí khác”.")
bullet("Trạng thái — “Hoạt động” hoặc “Khóa”.")
bullet("Người cập nhật — chọn nhân viên. Tiêu chí này bắt cả những bản ghi chưa từng được sửa mà do "
       "chính nhân viên đó tạo ra.")
para("Bốn ô lọc nâng cao có hiệu lực NGAY khi chọn, không cần bấm “Tìm kiếm”. Các tiêu chí kết hợp với "
     "nhau theo quan hệ VÀ. Bấm “Làm mới” để xóa toàn bộ tiêu chí và nạp lại danh sách đầy đủ.")
para("Bộ lọc đang áp dụng được ghi nhớ trong 10 phút: rời màn rồi quay lại trong khoảng thời gian này "
     "thì các tiêu chí cũ vẫn còn; quá 10 phút bộ lọc tự trở về mặc định.")
para("Nút trên thanh công cụ của bảng: “Tạo mới” (chỉ hiện khi có quyền quản lý) và “Xuất Excel”. "
     "Trên từng dòng: Xem (luôn hiện); Sửa và Xóa (chỉ hiện khi có quyền quản lý); nút ổ khóa nằm trong "
     "cột Trạng thái.")

# ================================================================
# PHAN 2 - PHAN QUYEN
# ================================================================
h1("PHẦN 2: PHÂN QUYỀN & HƯỚNG DẪN THEO QUYỀN")
para("Màn hình gắn với 2 quyền. Tài khoản có thể có một trong hai, hoặc cả hai.")
table([
    ["Tên quyền", "Cho phép làm gì", "Nút / khu vực tương ứng", "Ghi chú"],
    ["Quản lý dịch vụ sửa chữa và chi phí khác",
     "Xem danh sách, xem chi tiết, Tạo mới, Sửa, Xóa, Khóa, Mở khóa.",
     "Nút “Tạo mới”; nút Sửa và Xóa trên từng dòng; nút ổ khóa trong cột Trạng thái.",
     "Nút vẫn có thể bị mờ theo trạng thái dòng hoặc do là dòng hệ thống."],
    ["Xem dịch vụ sửa chữa và chi phí khác",
     "Chỉ xem danh sách, xem chi tiết, tìm kiếm và Xuất Excel.",
     "Nút “Xem” trên từng dòng; nút “Xuất Excel”.",
     "Không thấy nút Tạo mới / Sửa / Xóa / ổ khóa."],
])
para("Màn hình KHÔNG phân quyền theo cấp công ty / phòng ban / bộ phận: mọi tài khoản có quyền đều nhìn "
     "thấy toàn bộ danh mục. Công ty đang chọn chỉ ảnh hưởng tới giá trị cột “ĐM giảm giá”.")

h2("2.1 Người dùng có quyền “Xem dịch vụ sửa chữa và chi phí khác”")
bullet("Thấy: mục menu, panel bộ lọc, toàn bộ 9 cột của bảng, nút “Xuất Excel”, nút “Xem” trên từng dòng.")
bullet("Không thấy: nút “Tạo mới”, nút Sửa, nút Xóa và nút ổ khóa trong cột Trạng thái.")
bullet("Làm được: tìm kiếm, lọc, sắp xếp, chuyển trang, mở modal xem chi tiết (mọi ô ở chế độ chỉ đọc), "
       "xuất file Excel theo bộ lọc đang áp dụng.")
para("Nếu không có quyền này, mục menu sẽ không hiển thị; trường hợp truy cập trực tiếp bằng đường dẫn, "
     "hệ thống báo lỗi không có quyền và bảng không có dữ liệu.")

h2("2.2 Người dùng có quyền “Quản lý dịch vụ sửa chữa và chi phí khác”")
bullet("Thấy thêm: nút “Tạo mới” trên thanh công cụ, nút Sửa và Xóa ở cột Hành động, nút ổ khóa ở cột "
       "Trạng thái.")
bullet("Làm được: toàn bộ thao tác Tạo mới / Sửa / Xóa / Khóa / Mở khóa mô tả ở các phần sau.")
bullet("Mỗi thao tác thay đổi dữ liệu (Tạo mới, Sửa, Xóa, Khóa, Mở khóa) đều yêu cầu quyền này. "
       "Nếu quyền bị thu hồi trong lúc form đang mở, thao tác Lưu sẽ bị hệ thống từ chối và dữ liệu "
       "không bị ghi.")

# ================================================================
# PHAN 3 - TAO MOI
# ================================================================
h1("PHẦN 3: TẠO MỚI DỊCH VỤ / CHI PHÍ")
para("Yêu cầu quyền “Quản lý dịch vụ sửa chữa và chi phí khác”.")
para("Tại thanh công cụ của bảng, bấm nút “Tạo mới”. Hệ thống mở cửa sổ “Thêm dịch vụ / chi phí” ngay "
     "trên trang, không chuyển sang màn khác.")
image("hdsd_costs_03_taomoi.png", "Cửa sổ Thêm dịch vụ / chi phí")

h2("3.1 Các trường nhập liệu")
table([
    ["Trường", "Kiểu nhập", "Bắt buộc", "Giá trị điền sẵn khi tạo mới", "Ghi chú"],
    ["Tên dịch vụ / chi phí", "Ô nhập chữ", "Có", "Để trống",
     "Tối đa 255 ký tự. Không được trùng tên với dịch vụ đã có (kể cả dịch vụ đang bị Khóa)."],
    ["% Tính giá vốn", "Ô nhập số", "Có", "Để trống",
     "Phải ≥ 0. KHÔNG bị giới hạn trần 100 — nhập 321 vẫn hợp lệ."],
    ["% VAT", "Ô nhập số", "Có", "Để trống", "Từ 0 đến 100. Nhập lớn hơn 100 sẽ bị báo lỗi."],
    ["ĐM giảm giá (%)", "Ô nhập số", "Không", "Để trống",
     "Từ 0 đến 100. Bỏ trống hoặc nhập 0 nghĩa là không có định mức giảm giá. "
     "Giá trị này được lưu riêng cho công ty đang chọn."],
    ["Dịch vụ có tính doanh thu", "Ô tích chọn", "Không", "ĐÃ TÍCH SẴN",
     "Tích → phân loại “Có tính doanh thu”. Bỏ tích → phân loại “Chi phí khác”."],
])
para("Cửa sổ Tạo mới KHÔNG có ô “Trạng thái”. Dịch vụ mới tạo luôn ở trạng thái Hoạt động; muốn đổi "
     "trạng thái thì mở lại bằng chức năng Sửa.")
para("Cách nhập số thập phân: dùng dấu phẩy hoặc dấu chấm đều được, hệ thống hiểu dấu phẩy là DẤU THẬP "
     "PHÂN. Nhập “12,5” sẽ được lưu là 12,5 phần trăm — không phải 125.")

h2("3.2 Lưu dữ liệu")
bullet("Nút “Lưu”: lưu bản ghi, đóng cửa sổ, hiện thông báo “Thêm mới thành công” và nạp lại danh sách. "
       "Bản ghi mới nằm ở đầu danh sách.", bold_prefix="")
bullet("Nút “Lưu & Tiếp tục”: lưu bản ghi nhưng GIỮ cửa sổ mở và xóa trắng toàn bộ ô nhập, để khai báo "
       "tiếp bản ghi kế theo. Nút này chỉ có ở cửa sổ Tạo mới.")
bullet("Nút “Đóng” hoặc dấu × ở góc: thoát mà không lưu; dữ liệu đang nhập dở bị bỏ.")

h2("3.3 Thông báo lỗi khi nhập thiếu")
para("Nếu bấm “Lưu” khi chưa nhập đủ, hệ thống hiện thông báo đỏ “Bạn chưa nhập đầy đủ thông tin” và "
     "báo lỗi ngay dưới từng ô sai. Cửa sổ không bị đóng, dữ liệu đã nhập vẫn còn.")
image("hdsd_costs_04_validate.png", "Thông báo lỗi khi bỏ trống các trường bắt buộc")
table([
    ["Tình huống nhập sai", "Thông báo hiển thị"],
    ["Bỏ trống Tên / % Tính giá vốn / % VAT", "Bắt buộc phải nhập"],
    ["Tên trùng với dịch vụ đã có", "Đã tồn tại trên hệ thống"],
    ["Tên dài quá 255 ký tự", "Tối đa 255 ký tự"],
    ["Nhập chữ vào ô số", "Phải là số"],
    ["% Tính giá vốn nhỏ hơn 0", "Không được nhỏ hơn 0"],
    ["% VAT hoặc ĐM giảm giá lớn hơn 100", "Tối đa 100"],
])

# ================================================================
# PHAN 4 - SUA / XEM
# ================================================================
h1("PHẦN 4: SỬA VÀ XEM CHI TIẾT")

h2("4.1 Sửa")
para("Yêu cầu quyền “Quản lý dịch vụ sửa chữa và chi phí khác”. Tại cột Hành động, bấm nút hình bút chì "
     "ở dòng cần sửa. Cửa sổ “Sửa dịch vụ / chi phí” mở ra với dữ liệu đã điền sẵn.")
image("hdsd_costs_05_sua.png", "Cửa sổ Sửa dịch vụ / chi phí")
para("Cửa sổ Sửa khác cửa sổ Tạo mới ở ba điểm:")
bullet("Trên tiêu đề có thêm dòng thông tin “Cập nhật: …” và “Bởi: …” cho biết lần sửa gần nhất do ai "
       "thực hiện.")
bullet("Có thêm ô “Trạng thái” để chuyển giữa “Hoạt động” và “Khóa”.")
bullet("Không có nút “Lưu & Tiếp tục”, chỉ có “Lưu” và “Đóng”.")
para("Sau khi bấm “Lưu”, hệ thống hiện thông báo “Cập nhật thành công”, đóng cửa sổ và nạp lại danh "
     "sách. Cột “Cập nhật” của dòng vừa sửa đổi sang thời điểm hiện tại và tên người đang đăng nhập.")
para("Về ô “ĐM giảm giá”: nếu xóa trắng ô này hoặc nhập 0 rồi Lưu, hệ thống sẽ gỡ định mức giảm giá của "
     "CÔNG TY ĐANG CHỌN (cột trên danh sách chuyển thành “—”). Định mức của các công ty khác không bị "
     "ảnh hưởng.")
para("Nút Sửa bị làm mờ trong 2 trường hợp: dòng đang ở trạng thái Khóa (chú thích khi rê chuột: "
     "“Chi phí đang bị khóa, hãy mở khóa trước khi sửa”) hoặc dòng là chi phí hệ thống — cụ thể là "
     "“Chi phí đi lại” và “Chi phí vận chuyển” (chú thích: “Chi phí này không được phép sửa”).")

h2("4.2 Xem chi tiết")
para("Không cần quyền quản lý — chỉ cần quyền xem. Bấm nút hình con mắt ở cột Hành động.")
image("hdsd_costs_08_xem.png", "Cửa sổ Xem dịch vụ / chi phí (chế độ chỉ đọc)")
para("Cửa sổ “Xem dịch vụ / chi phí” hiển thị đầy đủ thông tin nhưng mọi ô đều bị khóa, không nhập được "
     "và không có nút Lưu — chỉ có nút “Đóng”. Chức năng này dùng được với cả dòng đang Hoạt động lẫn "
     "dòng đang Khóa.")

# ================================================================
# PHAN 5 - KHOA / MO KHOA
# ================================================================
h1("PHẦN 5: KHÓA VÀ MỞ KHÓA")
para("Yêu cầu quyền “Quản lý dịch vụ sửa chữa và chi phí khác”.")
para("Nút Khóa / Mở khóa nằm NGAY TRONG cột “Trạng thái”, bên phải nhãn trạng thái (không nằm ở cột "
     "Hành động). Dòng đang Hoạt động hiển thị biểu tượng ổ khóa đóng; dòng đang Khóa hiển thị biểu "
     "tượng ổ khóa mở.")
image("hdsd_costs_07_xacnhankhoa.png", "Hộp thoại xác nhận khóa dịch vụ")
bullet("Khóa: bấm biểu tượng ổ khóa ở dòng đang Hoạt động → hộp thoại “Xác nhận khóa” hiện ra với nội "
       "dung “Bạn có chắc muốn khóa '<tên dịch vụ>'?”. Bấm “Khóa” để xác nhận, hoặc “Hủy” để bỏ qua. "
       "Sau khi khóa, hệ thống báo “Khóa thành công”, nhãn trạng thái chuyển thành “Khóa”, đồng thời "
       "nút Sửa và nút Xóa của dòng đó bị làm mờ.")
bullet("Mở khóa: bấm biểu tượng ổ khóa ở dòng đang Khóa → hộp thoại “Xác nhận mở khóa”. Bấm “Mở khóa” "
       "để xác nhận. Hệ thống báo “Mở khóa thành công”, nhãn trạng thái trở lại “Hoạt động” và nút "
       "Sửa / Xóa sáng lại.")
para("Thao tác Khóa / Mở khóa CHỈ đổi trạng thái, không đụng tới các thông số khác. Đặc biệt, định mức "
     "giảm giá đã khai báo vẫn được giữ nguyên, nên khi mở khóa lại thì dữ liệu cũ còn đầy đủ.")
para("Hai chi phí hệ thống “Chi phí đi lại” và “Chi phí vận chuyển” không khóa và không mở khóa được — "
     "nút ổ khóa của hai dòng này luôn bị làm mờ.")

# ================================================================
# PHAN 6 - XOA
# ================================================================
h1("PHẦN 6: XÓA (HOẶC TỰ CHUYỂN SANG KHÓA)")
para("Yêu cầu quyền “Quản lý dịch vụ sửa chữa và chi phí khác”. Bấm nút hình thùng rác ở cột Hành động.")
para("Trước khi mở hộp thoại, hệ thống tự kiểm tra xem dịch vụ đã được sử dụng ở chứng từ nào chưa. "
     "Kết quả kiểm tra quyết định nội dung hộp thoại và thao tác thực sự xảy ra:")
image("hdsd_costs_06_xacnhanxoa.png", "Hộp thoại xác nhận xóa dịch vụ chưa phát sinh chứng từ")
table([
    ["Tình huống", "Hộp thoại hiển thị", "Kết quả sau khi xác nhận"],
    ["Dịch vụ CHƯA được dùng ở Báo giá hãng và Hợp đồng hãng",
     "Tiêu đề “Xác nhận xóa”; nội dung “Bạn có chắc muốn xóa '<tên>'?”; nút xác nhận ghi “Xóa”.",
     "Xóa hẳn khỏi danh mục, đồng thời xóa cả định mức giảm giá của mọi công ty. "
     "Thông báo “Xóa thành công”."],
    ["Dịch vụ ĐÃ được dùng ở Báo giá hãng và/hoặc Hợp đồng hãng",
     "Tiêu đề đổi thành “Xác nhận khóa”; nội dung nêu rõ dịch vụ đang được dùng ở đâu và cho biết hệ "
     "thống sẽ chuyển sang trạng thái Khóa thay vì xóa; nút xác nhận ghi “Khóa”.",
     "Bản ghi VẪN CÒN trong danh mục nhưng chuyển sang trạng thái Khóa; định mức giảm giá được giữ "
     "nguyên. Thông báo “Khóa thành công”."],
])
para("Đây là cơ chế bảo vệ dữ liệu: chứng từ đã lập trước đó vẫn tham chiếu được tới dịch vụ, không bị "
     "mất tên và mất thông số.")
para("Nút Xóa bị làm mờ khi dòng đang ở trạng thái Khóa (chú thích “Chi phí đã bị khóa”) hoặc khi dòng "
     "là chi phí hệ thống “Chi phí đi lại” / “Chi phí vận chuyển” (chú thích “Chi phí này không được "
     "phép xóa”).")

# ================================================================
# PHAN 7 - XUAT EXCEL
# ================================================================
h1("PHẦN 7: XUẤT EXCEL")
para("Nút “Xuất Excel” nằm cạnh nút “Tạo mới”, dùng được với cả tài khoản chỉ có quyền xem.")
para("Bấm nút, hệ thống tải về file tên danh_muc_dich_vu_sua_chua_va_chi_phi_khac.xlsx và hiện thông "
     "báo “Xuất Excel thành công”. File xuất theo ĐÚNG bộ lọc và thứ tự sắp xếp đang áp dụng trên màn "
     "hình, nhưng lấy TẤT CẢ dòng khớp bộ lọc chứ không chỉ trang đang xem.")
para("File Excel gồm 8 cột: STT, Tên dịch vụ / chi phí, Phân loại, ĐM giảm giá (%), % Tính giá vốn, "
     "% VAT, Trạng thái, Người tạo. Lưu ý hai điểm khác với màn hình:")
bullet("Các cột phần trăm trong file luôn được ghi với 2 chữ số thập phân (ví dụ 80,00) trong khi trên "
       "màn hình hiển thị rút gọn (80%).")
bullet("Cột cuối là “Người tạo” chứ không phải người cập nhật gần nhất như cột “Cập nhật” trên màn hình.")

# ================================================================
# PHAN CHI TIET
# ================================================================
h1("PHẦN CHI TIẾT: THAO TÁC TỪNG BƯỚC")
para("Phần này tóm tắt lại trình tự thao tác để người dùng làm theo.")

h2("A. Khai báo một dịch vụ mới")
bullet("Bước A1: Vào phân hệ CSKH → nhóm menu “Danh mục - Dịch vụ” → “Danh mục dịch vụ sửa chữa và chi "
       "phí khác”.")
bullet("Bước A2: Bấm nút “Tạo mới” ở góc phải trên bảng.")
bullet("Bước A3: Nhập “Tên dịch vụ / chi phí” (bắt buộc, không trùng tên đã có).")
bullet("Bước A4: Nhập “% Tính giá vốn” (bắt buộc, ≥ 0) và “% VAT” (bắt buộc, 0 – 100).")
bullet("Bước A5: Nhập “ĐM giảm giá (%)” nếu công ty có áp dụng định mức giảm giá; bỏ trống nếu không có.")
bullet("Bước A6: Giữ tích ô “Dịch vụ có tính doanh thu” nếu đây là dịch vụ tính doanh thu; bỏ tích nếu "
       "là khoản chi phí khác.")
bullet("Bước A7: Bấm “Lưu” để kết thúc, hoặc “Lưu & Tiếp tục” nếu còn khai báo tiếp dòng khác.")

h2("B. Điều chỉnh một dịch vụ đã có")
bullet("Bước B1: Tìm dòng cần sửa bằng ô tìm nhanh hoặc bộ lọc nâng cao.")
bullet("Bước B2: Bấm nút hình bút chì ở cột Hành động. Nếu nút bị mờ, kiểm tra chú thích khi rê chuột: "
       "dòng đang Khóa thì phải Mở khóa trước; dòng là chi phí hệ thống thì không sửa được.")
bullet("Bước B3: Chỉnh các thông số cần thiết, có thể đổi cả ô “Trạng thái”.")
bullet("Bước B4: Bấm “Lưu”. Kiểm tra lại giá trị trên lưới và cột “Cập nhật”.")

h2("C. Ngừng sử dụng một dịch vụ")
bullet("Bước C1: Nếu chỉ muốn tạm ngừng, bấm biểu tượng ổ khóa trong cột “Trạng thái” rồi xác nhận "
       "“Khóa”. Dữ liệu được giữ nguyên và mở lại được bất cứ lúc nào.")
bullet("Bước C2: Nếu muốn bỏ hẳn, bấm nút thùng rác ở cột Hành động rồi đọc kỹ hộp thoại: hộp thoại ghi "
       "“Xác nhận xóa” nghĩa là dịch vụ sẽ bị xóa hẳn; ghi “Xác nhận khóa” nghĩa là dịch vụ đã phát "
       "sinh chứng từ và chỉ bị chuyển sang trạng thái Khóa.")
bullet("Bước C3: Xác nhận, sau đó kiểm tra lại danh sách.")

h2("D. Kết xuất danh mục ra Excel")
bullet("Bước D1: Đặt bộ lọc và thứ tự sắp xếp mong muốn trên màn hình.")
bullet("Bước D2: Bấm “Xuất Excel” và mở file tải về để kiểm tra.")

# ---------------------------------------------------------------- luu
doc.save(OUTPUT)

# ------------------------------------------------- bat updateFields
from docx.oxml import OxmlElement
doc2 = Document(OUTPUT)
settings = doc2.settings.element
if settings.find(qn('w:updateFields')) is None:
    el = OxmlElement('w:updateFields')
    el.set(qn('w:val'), 'true')
    settings.append(el)
doc2.save(OUTPUT)

# ------------------------------------------------- purge media mo coi
# python-docx khong xoa file anh khi xoa paragraph -> file phinh. Quet r:embed/r:link
# con dung, map qua .rels, bo media + Relationship khong con ai tham chieu.
import re
import zipfile

with zipfile.ZipFile(OUTPUT) as z:
    entries = {n: z.read(n) for n in z.namelist()}

used_rids = set()
for name, data in entries.items():
    if name.endswith(".xml") and (name.startswith("word/document") or
                                  name.startswith("word/header") or
                                  name.startswith("word/footer")):
        text = data.decode("utf-8", "ignore")
        used_rids.update(re.findall(r'r:(?:embed|link)="([^"]+)"', text))

rels_name = "word/_rels/document.xml.rels"
rels_xml = entries[rels_name].decode("utf-8")
keep_targets, drop_rels = set(), set()
for m in re.finditer(r'<Relationship[^>]*/>', rels_xml):
    tag = m.group(0)
    rid = re.search(r'Id="([^"]+)"', tag).group(1)
    target = re.search(r'Target="([^"]+)"', tag).group(1)
    if "media/" not in target:
        continue
    if rid in used_rids:
        keep_targets.add("word/" + target.replace("../", ""))
    else:
        drop_rels.add(tag)

# media duoc header/footer dung qua rels rieng -> giu lai
for name in entries:
    if name.endswith(".rels") and name != rels_name:
        for m in re.finditer(r'Target="([^"]*media/[^"]+)"', entries[name].decode("utf-8")):
            keep_targets.add("word/" + m.group(1).replace("../", ""))

for tag in drop_rels:
    rels_xml = rels_xml.replace(tag, "")
entries[rels_name] = rels_xml.encode("utf-8")

removed = [n for n in list(entries)
           if n.startswith("word/media/") and n not in keep_targets]
for n in removed:
    del entries[n]

with zipfile.ZipFile(OUTPUT, "w", zipfile.ZIP_DEFLATED) as z:
    for name, data in entries.items():
        z.writestr(name, data)
print("Da xoa %d anh mo coi" % len(removed))

# ------------------------------------- cap nhat MUC LUC / DANH MUC HINH ANH
# ⚠️ BAT BUOC. `updateFields=true` chi la LOI MOI Word cap nhat khi mo file — ban than
# document.xml van giu nguyen text cu cua template (muc luc se hien "PHAN 2: TAO BOM LIST"...).
# Phai goi Word cap nhat that roi luu lai thi file ban giao moi dung.
import subprocess

PS = r"""
$p = "{path}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($p, $false, $false)
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
foreach ($tof in $doc.TablesOfFigures) {{ $tof.Update() }}
$doc.Repaginate()
$doc.BuiltInDocumentProperties("Title") = "HDSD - Danh mục dịch vụ sửa chữa và chi phí khác"
$doc.Save()
Write-Output ("Pages=" + $doc.ComputeStatistics(2))
$doc.Close(0)
$word.Quit()
""".format(path=OUTPUT)

res = subprocess.run(["powershell", "-NonInteractive", "-Command", PS],
                     capture_output=True, text=True)
print("Cap nhat field bang Word:", res.stdout.strip() or res.stderr.strip())

# ------------------------------------------------------------ verify
chk = Document(OUTPUT)
raw = zipfile.ZipFile(OUTPUT).read("word/document.xml").decode("utf-8")

# ⚠️ Chot quan trong nhat: muc luc / danh muc hinh anh / bia KHONG duoc con chu nao cua file khung.
sot = [h for h in TEMPLATE_BODY_HEADINGS if h and h in raw]
assert not sot, (
    "Van con tieu de cua file khung trong output: %s\n"
    "-> Muc luc / danh muc hinh anh chua duoc Word cap nhat that." % sot
)
h1_count = sum(1 for p in chk.paragraphs if p.style.name == "Heading 1")
cap_count = sum(1 for p in chk.paragraphs if p.style.name == "Caption")
print("File:", OUTPUT)
print("Heading 1:", h1_count, "| Bang:", len(chk.tables),
      "| Anh:", len(chk.inline_shapes), "| Caption:", cap_count)
