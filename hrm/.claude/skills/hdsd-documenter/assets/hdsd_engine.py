# -*- coding: utf-8 -*-
"""Engine dung tai lieu HDSD (.docx) tu file khung `HDSD_MAU.docx`.

Tach ra tu `gen_hdsd_mau.py` de nhieu man dung chung: generator cua tung man chi con
phan NOI DUNG, moi thu khac (copy khung, doi dong bia, strip body, helper style,
purge media mo coi, cap nhat muc luc bang Word, assert) do engine lo.

Cach dung trong `.plans/<feature>/gen_hdsd.py`:

    import os, sys
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                    "..", "..", "..", ".claude", "skills",
                                    "hdsd-documenter", "assets"))
    from hdsd_engine import HdsdBuilder

    b = HdsdBuilder(output=..., shots_dir=..., cover_title="(Màn hình: ...)",
                    doc_title="HDSD - ...")
    b.h1("TỔNG QUAN")
    b.h2("1. Mục tiêu")
    b.para("...")
    b.table([["Cột A", "Cột B"], ["1", "2"]])
    b.image("01-danh-sach.png", "Màn hình danh sách")
    b.finish()

Luat vang: KHONG ap direct formatting — de style cua template quyet dinh tat ca.
"""
import copy
import os
import re
import shutil
import subprocess
import sys
import zipfile

try:  # console Windows mac dinh cp1252
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph

ASSETS = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(ASSETS, "HDSD_MAU.docx")
assert os.path.exists(TEMPLATE), "Khong tim thay file khung %s" % TEMPLATE


class HdsdBuilder(object):
    """Dung 1 file HDSD tu khung mau. Goi finish() de luu + cap nhat muc luc + verify."""

    def __init__(self, output, shots_dir, cover_title, doc_title):
        self.output = output
        self.shots_dir = shots_dir
        self.doc_title = doc_title

        shutil.copyfile(TEMPLATE, output)
        self.doc = Document(output)
        body = self.doc.element.body

        # --- dong tieu de tren bia -------------------------------------
        # Word cat dong nay thanh NHIEU run -> replace tren tung run.text KHONG khop.
        # Bat theo VI TRI (dong dang "(...)" trong ~20 doan dau), khong bat theo noi dung.
        for p in self.doc.paragraphs[:20]:
            txt = p.text.strip()
            if txt.startswith("(") and txt.endswith(")") and p.runs:
                for run in p.runs[1:]:
                    run.text = ""
                p.runs[0].text = cover_title
                print("Da doi dong bia:", txt, "->", cover_title)
                break
        else:
            raise AssertionError("Khong tim thay dong tieu de tren bia (dang '(...)')")

        # --- proto Caption (co SEQ field -> Word tu danh so) ------------
        self.caption_proto = None
        for p in self.doc.paragraphs:
            if p.style.name == "Caption" and "Hình" in p.text:
                self.caption_proto = copy.deepcopy(p._p)
                break
        assert self.caption_proto is not None, "Khong tim thay paragraph Caption mau"

        # --- strip body tu Heading 1 THU BA tro di ----------------------
        # #1 = MUC LUC, #2 = DANH MUC HINH ANH -> giu; #3 tro di la than bai mau.
        self.sectPr = body.find(qn('w:sectPr'))
        children = list(body.iterchildren())

        def heading_level(child):
            if child.tag != qn('w:p'):
                return None
            pPr = child.find(qn('w:pPr'))
            if pPr is None:
                return None
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                return None
            val = (pStyle.get(qn('w:val')) or "").replace(" ", "")
            return val if val.startswith("Heading") else None

        h1_positions = [i for i, c in enumerate(children) if heading_level(c) == "Heading1"]
        assert len(h1_positions) >= 3, "File khung khong du 3 Heading 1"
        # ⚠️ Doc text bang Paragraph(child, doc).text, KHONG dung child.itertext()
        #    (itertext gom ca text trong field -> "TỔNG QUANTỔNG QUANTỔNG QUAN").
        self.template_headings = [
            Paragraph(children[i], self.doc).text.strip() for i in h1_positions[2:]
        ]
        for child in children[h1_positions[2]:]:
            if child is not self.sectPr:
                body.remove(child)

    # ------------------------------------------------------ helpers
    def _append(self, par):
        """Chen paragraph moi ngay truoc sectPr (sectPr luon la child cuoi)."""
        self.sectPr.addprevious(par._p)
        return par

    def h(self, text, level):
        return self._append(self.doc.add_paragraph(text, style="Heading %d" % level))

    def h1(self, t):
        return self.h(t, 1)

    def h2(self, t):
        return self.h(t, 2)

    def h3(self, t):
        return self.h(t, 3)

    def para(self, text, bold_prefix=None):
        p = self.doc.add_paragraph()
        if bold_prefix:
            p.add_run(bold_prefix).bold = True
        p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return self._append(p)

    def bullet(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            p.add_run(bold_prefix).bold = True
        p.add_run(text)
        return self._append(p)

    def table(self, rows):
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        t.alignment = 1  # CENTER
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                t.cell(ri, ci).text = val
        self.sectPr.addprevious(t._tbl)
        self._append(self.doc.add_paragraph())  # 1 dong trong cho thoang
        return t

    def image(self, filename, caption_text):
        path = os.path.join(self.shots_dir, filename)
        assert os.path.exists(path), "Thieu anh: %s" % path
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(6.0))
        self._append(p)
        cap = copy.deepcopy(self.caption_proto)
        for r in cap.findall(qn('w:r')):
            for tnode in r.findall(qn('w:t')):
                if tnode.text and tnode.text.startswith(":"):
                    tnode.text = ": " + caption_text
        self.sectPr.addprevious(cap)

    # ------------------------------------------------------- finish
    def finish(self):
        self.doc.save(self.output)
        self._set_update_fields()
        removed = self._purge_orphan_media()
        self._update_fields_by_word()
        return self._verify(removed)

    def _set_update_fields(self):
        doc2 = Document(self.output)
        settings = doc2.settings.element
        if settings.find(qn('w:updateFields')) is None:
            el = OxmlElement('w:updateFields')
            el.set(qn('w:val'), 'true')
            settings.append(el)
        doc2.save(self.output)

    def _purge_orphan_media(self):
        """python-docx khong xoa file anh khi xoa paragraph -> file phinh."""
        with zipfile.ZipFile(self.output) as z:
            entries = {n: z.read(n) for n in z.namelist()}

        used_rids = set()
        for name, data in entries.items():
            if name.endswith(".xml") and (name.startswith("word/document") or
                                          name.startswith("word/header") or
                                          name.startswith("word/footer")):
                text = data.decode("utf-8", "ignore")
                used_rids.update(re.findall(r'r:(?:embed|link)="([^"]+)"', text))

        rels_name = "word/_rels/document.xml.rels"
        rels_xml = entries[rels_name].decode("utf-8")
        keep_targets, drop_rels = set(), set()
        for m in re.finditer(r'<Relationship[^>]*/>', rels_xml):
            tag = m.group(0)
            rid = re.search(r'Id="([^"]+)"', tag).group(1)
            target = re.search(r'Target="([^"]+)"', tag).group(1)
            if "media/" not in target:
                continue
            if rid in used_rids:
                keep_targets.add("word/" + target.replace("../", ""))
            else:
                drop_rels.add(tag)

        # media duoc header/footer dung qua rels rieng -> giu lai
        for name in entries:
            if name.endswith(".rels") and name != rels_name:
                for m in re.finditer(r'Target="([^"]*media/[^"]+)"',
                                     entries[name].decode("utf-8")):
                    keep_targets.add("word/" + m.group(1).replace("../", ""))

        for tag in drop_rels:
            rels_xml = rels_xml.replace(tag, "")
        entries[rels_name] = rels_xml.encode("utf-8")

        removed = [n for n in list(entries)
                   if n.startswith("word/media/") and n not in keep_targets]
        for n in removed:
            del entries[n]

        with zipfile.ZipFile(self.output, "w", zipfile.ZIP_DEFLATED) as z:
            for name, data in entries.items():
                z.writestr(name, data)
        print("Da xoa %d anh mo coi" % len(removed))
        return removed

    def _update_fields_by_word(self):
        """⚠️ BAT BUOC. `updateFields=true` chi la LOI MOI Word cap nhat khi mo file —
        document.xml van giu nguyen text muc luc cu cua template."""
        ps = r"""
$p = "{path}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($p, $false, $false)
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
foreach ($tof in $doc.TablesOfFigures) {{ $tof.Update() }}
$doc.Repaginate()
$doc.BuiltInDocumentProperties("Title") = "{title}"
$doc.Save()
Write-Output ("Pages=" + $doc.ComputeStatistics(2))
$doc.Close(0)
$word.Quit()
""".format(path=self.output, title=self.doc_title)
        res = subprocess.run(["powershell", "-NonInteractive", "-Command", ps],
                             capture_output=True, text=True)
        print("Cap nhat field bang Word:", res.stdout.strip() or res.stderr.strip())

    def _verify(self, removed):
        chk = Document(self.output)
        raw = zipfile.ZipFile(self.output).read("word/document.xml").decode("utf-8")

        # Chot quan trong nhat: muc luc / danh muc hinh anh / bia KHONG con chu nao cua khung.
        sot = [h for h in self.template_headings if h and h in raw]
        assert not sot, (
            "Van con tieu de cua file khung trong output: %s\n"
            "-> Muc luc / danh muc hinh anh chua duoc Word cap nhat that." % sot
        )
        h1_count = sum(1 for p in chk.paragraphs if p.style.name == "Heading 1")
        cap_count = sum(1 for p in chk.paragraphs if p.style.name == "Caption")
        print("File:", self.output)
        print("Heading 1:", h1_count, "| Bang:", len(chk.tables),
              "| Anh:", len(chk.inline_shapes), "| Caption:", cap_count)
        return {"h1": h1_count, "tables": len(chk.tables),
                "images": len(chk.inline_shapes), "captions": cap_count,
                "removed_media": len(removed)}
