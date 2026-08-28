# -*- coding: utf-8 -*-
"""Thu vien dung chung de sinh SRS theo FORM CHUAN cua team
(file mau: .claude/skills/srs-documenter/assets/SRS_MAU.docx).

Cach dung: xem cac file .plans/gop-db/<feature>/gen_srs.py

    from srs_docx_lib import SrsDoc
    d = SrsDoc(out=..., menu=..., route=..., full_url=...)
    d.h1('1. Gioi thieu'); d.p('...'); d.bullets([...])
    d.save()

Quy uoc bat buoc (dung theo skill .claude/skills/srs-documenter/SKILL.md):
  - Trang dau: 2 dong tieu de CAN GIUA, KHONG phai Heading -> title_block()
  - Bang "Gioi thieu" 2 cot x 7-8 dong -> intro_table() (bo dong cuoi neu dacbiet=None)
  - Bang "Mo ta chi tiet giao dien" 8 cot -> ui_table();
    chuc nang CHI DOC bo cot "Bat buoc" -> ui_table(rows, required=False)
  - Bang "Danh sach event va xu ly event" 4 cot -> event_table()
  - Muc "Layout man hinh" = URL day du + ANH CHUP THAT -> layout()
  - Bieu do Use Case phai la ANH PNG that -> overview_figure() / uc_figure()

FORM 2026-08-28 (ban mau moi cua QA: "SRS - Danh muc quoc gia", link trong SKILL.md):
  - Muc Layout ghi DUONG DAN MENU (`layout(menu=...)`), KHONG con dong "URL day du"
  - Dau moi muc "Gioi thieu" co 1 doan tro sang tai lieu quy tac dung chung -> rule_ref()
  - "Phan 4. Quy tac nghiep vu" la BANG 5 cot -> rule_table()
  - So do UML tong quan co PHAN CAP «include»/«extend» -> overview_figure2()

FORM 2026-08-17 (ban mau cu SRS_MAU.docx = SRS Danh muc khach hang):
  - Chi con 4 chuong: "Phan 1. Gioi thieu" / "Phan 2. Phan quyen" /
    "Phan 3. Dac ta chi tiet theo tung chuc nang" / "Phan 4. Quy tac nghiep vu"
  - DA BO: bang thong tin trang bia, muc "Pham vi", chuong "Tong quan",
    muc "Quy tac truy cap bat buoc", chuong "Danh muc chuc nang (Function list)",
    muc "Tieu chi nghiem thu", dong "Chuc nang lien quan: FR-xx"
  - Muc Layout CHI con dong "URL day du", bo dong "Menu:" va "Route (FE):"
"""
import os
import tempfile
import sys

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)
import srs_uml_render as uml  # noqa: E402

# Tai lieu quy tac dung chung — MOI chuc nang tro sang day thay vi chep lai quy tac.
COMMON_DOC = ('https://docs.google.com/document/d/'
              '1AiqvNZAg9K4qef45vdbo9oC6o9RS_hbDHCi-ZrKgD7s/edit?tab=t.0')
COMMON_TITLE = 'SRS_Các quy tắc chung_VN_1.0'

# Anchor lay NGUYEN tu ban mau "SRS - Danh muc quoc gia" — KHONG tu bia them anchor moi.
ANCHOR = {
    'list':    '',                            # Man Danh sach / Phan trang / Cau hinh cot
    'search':  '#heading=h.uqrjgqo79fuq',     # Kich ban tim kiem, Bo loc, Dropdown
    'create':  '#heading=h.pr587lqkad5b',     # Man Them moi, Validate du lieu
    'notice':  '#heading=h.nij9n0nvzijj',     # Thong bao / Quy tac Xoa
    'history': '#heading=h.thqm3w6a7nzr',     # Quy tac ghi lich su, Khoa / Mo khoa
    'excel':   '#heading=h.x6yi3popnswy',     # Quy tac Excel va Cau hinh cot
    'detail':  '#heading=h.e51mm7p7jit7',     # Man Xem chi tiet va Phan quyen
    'delete':  '#heading=h.bmqrpa8bs35d',     # Quy tac Xoa (dung trong cot Xu ly event)
}


def add_hyperlink(paragraph, url, text):
    """Chen 1 hyperlink that (xanh, gach chan) vao cuoi doan."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1155CC')
    rpr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(underline)
    run.append(rpr)
    node = OxmlElement('w:t')
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)
    return link


ACTOR_P1 = 'Người quản lý danh mục (P1)'
ACTOR_P2 = 'Người xem danh mục (P2)'
ACTOR_BOTH = 'Người dùng có quyền P1 hoặc P2'


class SrsDoc(object):
    """Mot tai lieu SRS theo form chuan."""

    def __init__(self, out, menu, route, full_url, img_dir=None, img_prefix=''):
        self.out = out
        self.menu = menu
        self.route = route
        self.full_url = full_url
        # Anh UML chi la file TRUNG GIAN — da nhung han vao .docx nen khong can giu.
        # Mac dinh ghi vao thu muc tam cua he dieu hanh de khong rai rac vao repo
        # (truoc day mac dinh la <thu muc lib>/img, gio lib nam trong skill assets).
        # Muon giu lai anh de xem thi truyen img_dir='...' khi khoi tao.
        self.img_dir = img_dir or os.path.join(tempfile.gettempdir(), 'srs_uml_img')
        self.img_prefix = img_prefix
        os.makedirs(self.img_dir, exist_ok=True)
        self._fig = 0

        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

        st = doc.styles['Normal']
        st.font.name = 'Calibri'
        st.font.size = Pt(11)
        for name, size in [('Heading 1', 20), ('Heading 2', 16), ('Heading 3', 14)]:
            hs = doc.styles[name]
            hs.font.size = Pt(size)
            hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        self.doc = doc

    # -------------------------------------------------------- trang dau
    def title_block(self, man_hinh):
        """2 dong tieu de dau tai lieu — CAN GIUA, 24pt, KHONG dung Heading.

        Ban mau khong con dong 'Phan he: ...' va khong con bang thong tin
        (Ma man hinh / Phien ban / Ngay lap / Nguoi lap / ...).
        """
        for text in ('SOFTWARE REQUIREMENTS SPECIFICATION (SRS)',
                     'Màn hình: %s' % man_hinh):
            par = self.doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = par.add_run(text)
            r.bold = True
            r.font.size = Pt(24)

    # ------------------------------------------------------------- text
    def h1(self, t):
        self.doc.add_heading(t, level=1)

    def h2(self, t):
        self.doc.add_heading(t, level=2)

    def h3(self, t):
        self.doc.add_heading(t, level=3)

    def p(self, t=''):
        return self.doc.add_paragraph(t)

    def bullets(self, items):
        for it in items:
            self.doc.add_paragraph(it, style='List Bullet')

    # ------------------------------------------------------------ tables
    def table(self, headers, rows, widths=None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hh in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = str(hh)
            for para in c.paragraphs:
                for r in para.runs:
                    r.bold = True
                    r.font.size = Pt(10)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if v is None else str(v)
                for para in cells[i].paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(10)
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Inches(w)
        self.doc.add_paragraph()
        return t

    def info_table(self, rows):
        """DEPRECATED — form moi (2026-08-17) khong con bang thong tin trang bia.

        Giu lai de cac gen_srs.py cu chay duoc; KHONG dung cho tai lieu moi.
        """
        return self.table(['Thông tin', 'Nội dung'], rows, widths=[1.8, 4.2])

    def intro_table(self, ten, mota, tacnhan, dieukien, chinh, phu, dacbiet=None):
        """Bang 'Gioi thieu' cua tung chuc nang.

        dacbiet=None  -> BO HAN dong 'Yeu cau dac biet' (bang con 7 dong).
        dacbiet=''    -> giu dong nhung de trong (ban mau van co vai cho nhu vay).
        """
        rows = [
            ('Tên chức năng', ten),
            ('Mô tả', mota),
            ('Tác nhân', tacnhan),
            ('Điều kiện ban đầu', dieukien),
            ('Dòng sự kiện chính', chinh),
            ('Dòng sự kiện phụ', phu),
        ]
        if dacbiet is not None:
            rows.append(('Yêu cầu đặc biệt', dacbiet))
        return self.table(['Mục', 'Nội dung'], rows, widths=[1.5, 4.5])

    def ui_table(self, rows, required=True, scope=True):
        """Bang 'Mo ta chi tiet giao dien' (STT tu danh).

        required=True, scope=True  -> 8 cot (mac dinh, dung cho man co nhap lieu)
        required=False             -> 7 cot, bo 'Bat buoc' (chuc nang CHI DOC:
                                      xem danh sach, xem chi tiet, lich su...)
        required=False, scope=False-> 6 cot, bo ca 'Pham vi' (hop xac nhan...)

        So o trong moi dong PHAI khop so cot da chon.
        """
        headers = ['STT', 'Tên đối tượng', 'Loại', 'Trạng thái']
        widths = [0.4, 1.2, 0.8, 0.75]
        if scope:
            headers.append('Phạm vi')
            widths.append(0.85)
        if required:
            headers.append('Bắt buộc')
            widths.append(0.6)
        headers += ['Giá trị ban đầu', 'Mô tả']
        widths += [0.85, 2.2 + (0 if scope else 0.85) + (0 if required else 0.6)]
        return self.table(
            headers,
            [(i + 1,) + tuple(r) for i, r in enumerate(rows)],
            widths=widths)

    def event_table(self, rows):
        """Bang 'Danh sach event va xu ly event' 4 cot (STT tu danh)."""
        return self.table(
            ['STT', 'Event', 'Loại event', 'Xử lý event'],
            [(i + 1,) + tuple(r) for i, r in enumerate(rows)],
            widths=[0.4, 1.6, 0.9, 3.6])

    # ------------------------------------------------- form RUT GON (2026-08-12)
    def field_table(self, rows):
        """Bang 'Mo ta chi tiet' 7 cot cua FORM RUT GON (STT tu danh).

        Cot: Field name | Placeholder | Field type | Content | Required |
             Gioi han ky tu | Ghi chu
        """
        return self.table(
            ['STT', 'Field name', 'Placeholder', 'Field type', 'Content', 'Required',
             'Giới hạn ký tự', 'Ghi chú'],
            [(i + 1,) + tuple(r) for i, r in enumerate(rows)],
            widths=[0.35, 1.15, 1.0, 0.75, 1.15, 0.6, 0.7, 1.9])

    def notice_table(self, rows):
        """Bang 'Danh sach event va thong bao' 4 cot cua FORM RUT GON (STT tu danh)."""
        return self.table(
            ['STT', 'Event', 'Loại event', 'Xử lý và thông báo'],
            [(i + 1,) + tuple(r) for i, r in enumerate(rows)],
            widths=[0.4, 1.5, 0.8, 3.8])

    def toc(self, note='Nhấn chuột phải vào mục lục → Update Field → Update entire table '
                       'để cập nhật số trang.'):
        """Chen truong Muc luc cua Word (tu sinh theo Heading 1-3)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        par = self.doc.add_paragraph()
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
        run = OxmlElement('w:r')
        txt = OxmlElement('w:t')
        txt.text = note
        run.append(txt)
        fld.append(run)
        par._p.append(fld)
        self.doc.add_paragraph()

    # ----------------------------------------------------------- figures
    def figure(self, png_path, caption, width_in=6.0):
        self._fig += 1
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(png_path, width=Inches(width_in))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run('Hình %d: %s' % (self._fig, caption))
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    def _png(self, name):
        return os.path.join(self.img_dir, '%s%s.png' % (self.img_prefix, name))

    def overview_figure2(self, actors, mains, subs, caption):
        """So do UML tong quan CO PHAN CAP — ban dung tu 2026-08-28.

        Chi use case la MAN HINH that su moi noi thang toi actor. Thao tac lam ngay tren
        man do (tim kiem/loc, tuy chinh cot, xoa, in, lich su, popup chon du lieu) phai noi
        vao use case cha bang «include» / «extend» — ve tat ca ngang hang roi noi thang toi
        actor la SAI nghiep vu (user tra tai lieu ve vi loi nay ngay 2026-08-28).

        actors : [(ten_actor, [chi_so_main, ...]), ...]
        mains  : [(ma, ten, nhom), ...]
        subs   : [(ma, ten, nhom, 'include'|'extend', [chi_so_main, ...], ghi_chu|None), ...]
        """
        png = self._png('overview')
        uml.draw_overview2(png, actors, mains, subs)
        self.figure(png, caption, width_in=6.3)

    def overview_figure(self, title, actors, usecases, caption):
        """FORM CU — so do tong quan PHANG, moi use case noi thang toi actor.

        Tai lieu moi dung `overview_figure2()`; giu ham nay cho cac gen_srs.py cu.
        """
        png = self._png('overview')
        uml.draw_overview(png, title, actors, usecases)
        self.figure(png, caption, width_in=6.3)

    def uc_figure(self, code, name, group, relations=(), actor=ACTOR_P1, caption=None):
        """5.2.x.1 Bieu do use case cua 1 chuc nang — anh PNG that."""
        png = self._png('uc_%s' % code.lower().replace('-', ''))
        uml.draw_usecase(png, actor, code, name, group, relations)
        self.figure(png, caption or ('Biểu đồ Use Case — %s %s' % (code, name)), width_in=6.2)

    # ------------------------------------------------------------ layout
    def layout(self, menu=None, modal=None, note='', shot=None, shot_caption=None,
               **_ignored):
        """Muc 'Layout man hinh' — DUONG DAN MENU + ANH CHUP THAT.

        Form 2026-08-28: ghi duong dan MENU, KHONG con dong "URL day du" (form 2026-08-17),
        cang khong con "Route (FE)" (form cu hon nua):

            Duong dan man hinh:
            Menu: Phan he Tai chinh => Khoi tao phieu ... => De nghi thu tien => Them moi

        menu         : duong dan menu day du cua DUNG chuc nang do; bo trong -> lay self.menu
        modal        : ten modal -> them cau "Modal ... duoc mo ngay tren man hinh danh sach"
        shot         : duong dan file .png chup that cua DUNG chuc nang do (6.2 inch)
        shot_caption : chu thich duoi anh
        _ignored     : nuot `route=` / `url=` cua form cu de cac gen_srs.py cu khong vo
        """
        self.p('Đường dẫn màn hình:')
        self.p('Menu: %s' % (menu or self.menu))
        if modal:
            self.p('Modal %s được mở ngay trên màn hình danh sách theo đường dẫn ở trên.' % modal)
        if note:
            self.p(note)
        if shot:
            if not os.path.exists(shot):
                raise IOError('Thieu anh chup cho muc Layout: %s' % shot)
            self.figure(shot, shot_caption or 'Màn hình thực tế', width_in=6.2)

    # ------------------------------------ doan "Quy tac chung: ..." (form 2026-08-28)
    def rule_ref(self, tail, anchor='list', head='Quy tắc chung',
                 lead='Áp dụng SRS Các quy tắc chung '):
        """Doan tro sang tai lieu quy tac dung chung, dat NGAY DAU muc "Gioi thieu".

        Muc dich: tai lieu tung man CHI ghi phan rieng cua man do, khong chep lai quy tac
        dung chung (phan trang, validate, thong bao, ghi lich su...).

            d.rule_ref('- Màn Danh sách, Sắp xếp dữ liệu bảng, Phân trang và Cấu hình cột. '
                       'Chỉ bổ sung các quy tắc riêng của <màn>.', anchor='list')

        tail   : phan viet tiep sau hyperlink
        anchor : key trong ANCHOR — chi dung key co san, KHONG tu bia anchor moi
        head   : nhan dau doan ('Quy tắc chung' o tung chuc nang / 'Quy tắc áp dụng' o Phan 4)
        lead   : phan viet truoc hyperlink (Phan 4 dung cach dan khac)
        """
        par = self.p('%s: %s' % (head, lead))
        add_hyperlink(par, COMMON_DOC + ANCHOR.get(anchor, ''), COMMON_TITLE)
        par.add_run('%s%s' % ('' if tail[:1] in '.,;:' else ' ', tail))
        for run in par.runs:
            run.font.size = Pt(10.5)
        return par

    # ------------------------------- Phan 4: bang quy tac nghiep vu (form 2026-08-28)
    def rule_table(self, rows):
        """Bang "Quy tac nghiep vu" 5 cot (STT tu danh).

        Thay cho dang "BR-0N — <ten>" + gach dau dong cua form cu.
        rows: [(ma, ten, mo_ta, pham_vi_ap_dung), ...]
              mo_ta / pham_vi co the la list -> tu noi bang xuong dong.
        """
        body = []
        for i, (ma, ten, mota, pham_vi) in enumerate(rows):
            if isinstance(mota, (list, tuple)):
                mota = chr(10).join(mota)
            if isinstance(pham_vi, (list, tuple)):
                pham_vi = chr(10).join(pham_vi)
            body.append((i + 1, ma, ten, mota, pham_vi))
        return self.table(['STT', 'Mã quy tắc', 'Tên quy tắc', 'Mô tả', 'Phạm vi áp dụng'],
                          body, widths=[0.35, 0.7, 1.25, 2.6, 1.1])

    # -------------------------------------------------------------- save
    def save(self, verbose=True, update_fields=True):
        os.makedirs(os.path.dirname(self.out), exist_ok=True)
        self.doc.save(self.out)
        if update_fields:
            self._update_fields_by_word()
        if verbose:
            self.selfcheck()
        return self.out

    def _update_fields_by_word(self):
        """Ep Word cap nhat MUC LUC that su.

        ⚠️ Chi chen truong TOC (`toc()`) thi Word MOI CHI hien dong nhac
        "Nhan chuot phai -> Update Field" — nguoi doc phai tu lam. Phai goi Word
        cap nhat roi luu lai thi file ban giao moi co muc luc + so trang san.
        Cung cach `hdsd_engine._update_fields_by_word()` da dung.

        May khong co Word (hoac loi COM) -> chi canh bao, KHONG lam gay generator.
        """
        import subprocess

        ps = r"""
$p = "{path}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($p, $false, $false)
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
foreach ($tof in $doc.TablesOfFigures) {{ $tof.Update() }}
$doc.Repaginate()
$doc.Save()
Write-Output ("Pages=" + $doc.ComputeStatistics(2))
$doc.Close(0)
$word.Quit()
""".format(path=self.out)
        try:
            res = subprocess.run(["powershell", "-NonInteractive", "-Command", ps],
                                 capture_output=True, text=True, timeout=180)
            out = (res.stdout or '').strip() or (res.stderr or '').strip()
        except Exception as exc:  # noqa: BLE001
            out = 'LOI: %s' % exc
        if not out.startswith('Pages='):
            print('!!! Muc luc CHUA duoc Word cap nhat:', out[:200])
        self._word_pages = out

    def selfcheck(self):
        """Buoc 4 cua skill: tu kiem tra truoc khi bao xong."""
        from docx import Document as _D
        d = _D(self.out)
        imgs = sum(1 for r in d.part.rels.values() if 'image' in r.reltype)
        bad = [x.text for x in d.paragraphs
               if any(ch in x.text for ch in ('\u250c', '\u25cb', '\u2502', '\u2514'))]
        pages = getattr(self, '_word_pages', '')
        out = ('OK %s | bang=%d | doan=%d | anh=%d | so-do-ky-tu=%d | %s'
               % (os.path.basename(self.out), len(d.tables), len(d.paragraphs), imgs,
                  len(bad), pages or 'muc-luc=chua-cap-nhat'))
        try:
            print(out)
        except UnicodeEncodeError:
            print(out.encode('ascii', 'replace').decode())
        assert imgs > 0, 'Thieu anh bieu do use case'
        assert not bad, 'Con so do ve bang ky tu: %s' % bad[:3]
