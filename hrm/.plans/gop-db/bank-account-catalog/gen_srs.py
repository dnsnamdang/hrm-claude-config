# -*- coding: utf-8 -*-
"""Sinh SRS man 'Danh muc tai khoan ngan hang' (/finance/account-banks) — FORM MOI 4 chuong.

Bam ban mau da duoc user chot: .plans/gop-db/customer-docs/SRS - Danh muc khach hang.docx
(= .claude/skills/srs-documenter/assets/SRS_MAU.docx). Cac diem bam theo ban mau, KHONG theo
mo ta trong SKILL.md o cho lech nhau:
  - Muc Layout VAN GIU dong "Menu: ..." (ban mau con dong nay; chi bo "Route (FE): ...")
  - Nhan muc con 2.x.y IN DAM, tieu de BR-0N IN DAM
  - 2 dong tieu de trang dau 24pt can giua, KHONG in dam
  - Truong Muc luc de rong roi cho Word cap nhat that (xem cuoi file)

Nguon doi chieu (doc truc tiep tu code nhanh gop_db, 2026-08-17):
  BE  Modules/Finance/Routes/api.php (prefix /account-banks — 7 route deu gan
      checkPermission:'Quan ly danh muc tai khoan ngan hang'; route sua gan them recordNotLocked)
      Modules/Finance/Http/Controllers/V1/CompanyAccountController.php
      Modules/Finance/Services/CompanyAccountService.php
      Modules/Finance/Entities/CompanyAccount/CompanyAccount.php
      Modules/Finance/Http/Requests/CompanyAccount/CompanyAccountRequest.php
      Modules/Finance/Transformers/CompanyAccountResource/*.php
      app/Services/CatalogHistoryService.php (bang `company_accounts`)
  FE  hrm-client/pages/finance/account-banks/{index.vue,AccountBankModal.vue}
      hrm-client/components/subsystem-menu/finance.js
  Anh chup that: tknh_shots/ (cong dev hrm-crm.eteksofts.com, 1440x900) — CHI DE LOCAL.

Chay:  python .plans/gop-db/bank-account-catalog/gen_srs.py
"""
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

sys.path.insert(0, os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "..", ".claude", "skills", "srs-documenter", "assets"))
from srs_docx_lib import SrsDoc  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "tknh_shots")
OUT = os.path.join(HERE, "SRS - Danh mục tài khoản ngân hàng.docx")


def shot(name):
    return os.path.join(SHOTS, name)


ACTOR = 'Người quản lý danh mục tài khoản ngân hàng'
MENU = 'Phân hệ Tài chính → Danh mục → Danh mục tài khoản ngân hàng'
FULL_URL = 'https://<host-hrm>/finance/account-banks'

d = SrsDoc(
    out=OUT,
    menu=MENU,
    route='/finance/account-banks',
    full_url=FULL_URL,
    img_prefix='tknh_')


def sub(text):
    """Nhan muc con 2.x.y — ban mau in DAM (van la doan thuong, khong phai Heading)."""
    par = d.p()
    par.add_run(text).bold = True
    return par


def lay(shot=None, shot_caption=None, modal=None):
    """Muc 'Layout man hinh': dong dan nhap + 2 gach dau dong (Menu / URL) + anh chup that."""
    d.p('Đường dẫn màn hình:')
    d.bullets(['Menu: %s' % MENU, 'URL đầy đủ: %s' % FULL_URL])
    if modal:
        d.p('Modal %s được mở ngay trên màn hình danh sách theo đường dẫn ở trên.' % modal)
    if shot:
        d.figure(shot, shot_caption or 'Màn hình thực tế', width_in=6.2)


# ================================================================ TRANG ĐẦU
d.title_block('Danh mục tài khoản ngân hàng')
# Ban mau: 2 dong tieu de can giua 24pt nhung KHONG in dam -> bo bold cua lib.
for _p in d.doc.paragraphs[:2]:
    for _r in _p.runs:
        _r.bold = None

d.h2('Mục lục')
# Truong TOC de rong, Word se dien noi dung o buoc cuoi file.
d.toc(note='')

# ============================================================ PHẦN 1. GIỚI THIỆU
d.h1('Phần 1. Giới thiệu')

d.h2('1 Mục đích')
d.p('Tài liệu này đặc tả yêu cầu phần mềm cho màn hình Danh mục tài khoản ngân hàng, nhằm:')
d.bullets([
    'Là căn cứ nghiệm thu chức năng, phân quyền và ràng buộc dữ liệu của màn hình.',
    'Làm rõ phạm vi dữ liệu theo công ty: người dùng chỉ thấy và chỉ thao tác được trên tài khoản '
    'ngân hàng của công ty mình.',
    'Làm rõ ràng buộc ba danh mục liên quan — Ngân hàng, Chi nhánh, Loại tiền tệ — trong đó chi '
    'nhánh phải thuộc đúng ngân hàng đã chọn và danh mục đang khóa thì không được chọn mới.',
    'Làm rõ khác biệt giữa cửa sổ Xem và cửa sổ Sửa khi ngân hàng hoặc loại tiền tệ của bản ghi '
    'đã bị khóa — điểm dễ bị hiểu nhầm là mất dữ liệu.',
])

d.h2('2 Thuật ngữ và viết tắt')
d.table(['Thuật ngữ', 'Mô tả'], [
    ('Tài khoản ngân hàng của công ty',
     'Một bản ghi của danh mục: số tài khoản, chủ tài khoản, ngân hàng, chi nhánh, loại tiền tệ.'),
    ('Chủ tài khoản',
     'Tên chủ tài khoản in trên sổ / thẻ ngân hàng. Hệ thống tự chuyển thành CHỮ IN HOA khi lưu.'),
    ('Ngân hàng', 'Bản ghi lấy từ màn Danh mục ngân hàng của phân hệ Danh mục chung.'),
    ('Chi nhánh',
     'Chi nhánh của ngân hàng đã chọn; danh sách chi nhánh lọc theo đúng ngân hàng đó.'),
    ('Loại tiền tệ', 'Bản ghi lấy từ màn Danh mục tiền tệ của phân hệ Tài chính.'),
    ('Trạng thái Hoạt động', 'Tài khoản đang dùng bình thường, sửa được.'),
    ('Trạng thái Khóa',
     'Tài khoản ngừng sử dụng nhưng KHÔNG bị xoá; khi đang khóa thì không sửa được.'),
    ('Công ty của người dùng',
     'Công ty ghi trong hồ sơ nhân sự của người đăng nhập — quyết định phạm vi dữ liệu của màn.'),
    ('Lịch sử thay đổi',
     'Nhật ký thao tác của một bản ghi: ai tạo, sửa, khóa, mở khóa và giá trị trước - sau.'),
    ('SRS', 'Software Requirements Specification.'),
], widths=[1.8, 4.2])

# ============================================================ PHẦN 2. PHÂN QUYỀN
d.h1('Phần 2. Phân quyền')

d.h2('1 Danh sách quyền')

d.p('Nhóm quyền thao tác:')
d.table(['Ký hiệu', 'Tên quyền', 'Tác dụng trên màn hình'], [
    ('Q1', 'Quản lý danh mục tài khoản ngân hàng',
     'Mở màn hình, xem danh sách, tạo mới, sửa, khóa / mở khóa và xem chi tiết. Thiếu quyền này '
     'thì mục menu bị ẩn và mọi thao tác đều bị từ chối.'),
], widths=[0.8, 2.0, 3.2])
d.p('Màn hình dùng đúng MỘT quyền, thuộc nhóm quyền Danh mục tài chính. Quyền này gắn cho cả chức '
    'năng xem danh sách, không riêng các thao tác ghi dữ liệu.')

d.p('Phạm vi dữ liệu (không phải quyền, luôn áp dụng cho mọi người dùng):')
d.bullets([
    'Người dùng chỉ thấy tài khoản ngân hàng thuộc công ty ghi trong hồ sơ nhân sự của mình.',
    'Bản ghi của công ty khác: không hiển thị, tìm kiếm không ra, mở trực tiếp thì hệ thống báo '
    'không tìm thấy.',
    'Người dùng chưa được gắn công ty: danh sách trống hoàn toàn và không tạo mới được.',
    'Màn hình không phân quyền theo phòng ban hay bộ phận.',
])

d.h2('2 Ma trận phân quyền')
d.table(['Chức năng', 'Q1', 'Không có quyền nào'], [
    ('FR-01 Truy cập & xem danh sách', '✅ (chỉ tài khoản của công ty mình)', '❌'),
    ('FR-02 Tìm kiếm và lọc', '✅', '❌'),
    ('FR-03 Tuỳ chỉnh cột hiển thị', '✅', '❌'),
    ('FR-04 Tạo mới tài khoản ngân hàng', '✅', '❌'),
    ('FR-05 Sửa tài khoản ngân hàng', '✅ (trừ tài khoản đang Khóa)', '❌'),
    ('FR-06 Xem chi tiết tài khoản', '✅', '❌'),
    ('FR-07 Khóa / Mở khóa', '✅', '❌'),
    ('FR-08 Xem lịch sử thay đổi', '✅', '❌'),
    ('Xóa tài khoản ngân hàng', 'Màn hình không có chức năng này', '—'),
    ('Xuất Excel / Nhập Excel / In', 'Màn hình không có chức năng này', '—'),
], widths=[2.5, 2.2, 1.3])

# ================================================= PHẦN 3. ĐẶC TẢ CHI TIẾT
d.h1('Phần 3. Đặc tả chi tiết theo từng chức năng')

d.h2('1 Sơ đồ UML tổng quan')
d.overview_figure(
    'HỆ THỐNG HRM — DANH MỤC TÀI KHOẢN NGÂN HÀNG',
    [(ACTOR, list(range(8)))],
    [('FR-01', 'Truy cập & xem danh sách', 'view', '«include» Lọc theo công ty người dùng'),
     ('FR-02', 'Tìm kiếm & Lọc', 'view', None),
     ('FR-03', 'Tuỳ chỉnh cột', 'view', None),
     ('FR-04', 'Tạo mới tài khoản', 'crud', '«include» Kiểm tra trùng số tài khoản'),
     ('FR-05', 'Sửa tài khoản', 'crud', '«include» Chặn khi bản ghi đang Khóa'),
     ('FR-06', 'Xem chi tiết', 'view', None),
     ('FR-07', 'Khóa / Mở khóa', 'action', None),
     ('FR-08', 'Xem lịch sử thay đổi', 'view', None)],
    'Sơ đồ Use Case tổng quan màn Danh mục tài khoản ngân hàng')

d.h2('2 Đặc tả chi tiết từng chức năng')

# ---------------------------------------------------------------- 2.1
d.h3('2.1 Truy cập và xem danh sách tài khoản ngân hàng')

sub('2.1.1 Giới thiệu')
d.intro_table(
    ten='Truy cập và xem danh sách tài khoản ngân hàng',
    mota='Hiển thị danh sách tài khoản ngân hàng thuộc công ty của người đăng nhập, kèm trạng '
         'thái, thông tin người tạo / người cập nhật và bộ nút thao tác của từng dòng.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đã đăng nhập và có quyền Quản lý danh mục tài khoản ngân hàng.',
    chinh='1. Người dùng vào menu Tài chính → Danh mục → Danh mục tài khoản ngân hàng.\n'
          '2. Hệ thống kiểm tra quyền và xác định công ty của người đăng nhập.\n'
          '3. Hệ thống nạp cấu hình cột hiển thị của tài khoản và khôi phục bộ lọc đã lưu nếu còn '
          'trong 10 phút.\n'
          '4. Hệ thống trả về trang đầu tiên của danh sách trong phạm vi công ty và tổng số bản '
          'ghi.\n'
          '5. Bảng hiển thị dữ liệu, ô “Hiển thị a–b / N” hiển thị đúng khoảng và tổng.',
    phu='• Không có quyền → mục menu bị ẩn; gõ thẳng đường dẫn thì hệ thống từ chối truy cập.\n'
        '• Người dùng chưa được gắn công ty → danh sách trống hoàn toàn.\n'
        '• Không có bản ghi khớp bộ lọc → bảng hiện “Không có dữ liệu phù hợp bộ lọc.”.\n'
        '• Phiên đăng nhập hết hạn → điều hướng về màn đăng nhập.',
    dacbiet='Danh sách mặc định xếp theo Ngày tạo mới nhất trước. Bảng có 12 cột khả dụng, hiển '
            'thị sẵn 8 cột; 4 cột Loại tiền tệ, Chi nhánh, Người cập nhật, Ngày cập nhật mặc định '
            'ẩn, bật ở cửa sổ Tuỳ chỉnh cột.')

sub('2.1.2 Layout màn hình')
lay(shot=shot('01-danh-sach.png'),
    shot_caption='Màn Danh mục tài khoản ngân hàng lúc mới truy cập')
d.figure(shot('02-cot-hanh-dong.png'),
         'Bảng danh sách khi thu gọn menu bên trái — thấy đủ cột Hành động', width_in=6.2)

sub('2.1.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Bảng danh sách tài khoản', 'Table/Grid', 'Read-only', '–', '–',
     'Hiển thị danh sách theo phân trang, cột Hành động chốt ở cuối bảng.'),
    ('Cột STT', 'Label', 'Read-only', '–', 'Số thứ tự theo trang',
     'Luôn hiển thị, không tắt được ở cửa sổ Tuỳ chỉnh cột.'),
    ('Cột Số tài khoản', 'Button', 'Enable', '0–255 ký tự', 'Theo dữ liệu',
     'Cột định danh, bấm vào mở cửa sổ Xem chi tiết. Luôn hiển thị, không tắt được. Cho phép sắp '
     'xếp.'),
    ('Cột Chủ tài khoản', 'Text', 'Read-only', '0–255 ký tự', 'Theo dữ liệu',
     'Luôn hiển thị dạng CHỮ IN HOA vì hệ thống chuyển khi lưu. Cho phép sắp xếp.'),
    ('Cột Ngân hàng', 'Text', 'Read-only', '0–255 ký tự', 'Theo dữ liệu',
     'Tên ngân hàng lưu sẵn trên bản ghi nên vẫn hiện đúng khi ngân hàng đã bị khóa.'),
    ('Cột Loại tiền tệ', 'Text', 'Read-only', '–', 'Ẩn',
     'Hiển thị mã tiền tệ; bản ghi cũ chưa có tiền tệ hiện dấu gạch ngang. Mặc định tắt.'),
    ('Cột Chi nhánh', 'Text', 'Read-only', '0–255 ký tự', 'Ẩn', 'Mặc định tắt.'),
    ('Cột Người cập nhật', 'Text', 'Read-only', '–', 'Ẩn', 'Người sửa gần nhất. Mặc định tắt.'),
    ('Cột Ngày cập nhật', 'Text', 'Read-only', 'dd/mm/yyyy HH:mm', 'Ẩn',
     'Mặc định tắt. Cho phép sắp xếp.'),
    ('Cột Người tạo', 'Text', 'Read-only', '–', 'Theo dữ liệu', 'Họ tên người tạo bản ghi.'),
    ('Cột Ngày tạo', 'Text', 'Read-only', 'dd/mm/yyyy HH:mm', 'Theo dữ liệu',
     'Cho phép sắp xếp; là tiêu chí sắp xếp mặc định của danh sách.'),
    ('Cột Trạng thái', 'Badge', 'Read-only', 'Hoạt động / Khóa', 'Theo dữ liệu',
     'Nhãn xanh cho Hoạt động, nhãn đỏ cho Khóa.'),
    ('Cột Hành động', 'Icon Button', 'Enable / Ẩn', '–', 'Hiển thị',
     'Gồm Sửa, Khóa / Mở khóa và Lịch sử. Nút không dùng được thì ẩn hẳn chứ không làm mờ.'),
    ('Nút Tạo mới', 'Button', 'Enable / Ẩn', '–', 'Hiển thị khi có quyền',
     'Nằm ở góc phải trên bảng.'),
    ('Nút Cấu hình cột hiển thị', 'Icon Button', 'Enable', '–', 'Hiển thị',
     'Mở cửa sổ Tuỳ chỉnh cột.'),
    ('Phân trang', 'Pagination', 'Enable', '–', 'Trang 1',
     'Có nút về đầu / lùi / số trang / tiến / về cuối.'),
    ('Ô Số dòng/trang', 'Dropdown', 'Enable', '5 / 10 / 20 / 50 / 100', '10',
     'Đổi số dòng thì danh sách quay về trang 1.'),
    ('Ô “Hiển thị a–b / N”', 'Label', 'Read-only', '–', 'Theo kết quả',
     'N là tổng số tài khoản khớp bộ lọc trong công ty của người đăng nhập.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', 'Ẩn',
     'Nội dung “Không có dữ liệu phù hợp bộ lọc.”.'),
    ('Vòng quay chờ', 'Loading', 'Hiển thị', '–', 'Ẩn', 'Hiện trong lúc nạp danh sách.'),
], required=False)

sub('2.1.4 Danh sách event và xử lý event')
d.event_table([
    ('Mở màn hình', 'System',
     'Before:\n– Kiểm tra quyền Quản lý danh mục tài khoản ngân hàng; không có quyền → từ chối '
     'truy cập.\n'
     'During:\n– Xác định công ty của người đăng nhập; chưa gắn công ty → trả danh sách rỗng.\n'
     '– Nạp cấu hình cột và khôi phục bộ lọc đã lưu nếu chưa quá 10 phút.\n'
     'After:\n– Trả trang 1 với 10 dòng/trang và tổng số bản ghi; hiển thị bảng.'),
    ('Đổi trang', 'Click', 'Tải lại danh sách theo trang mới, giữ nguyên bộ lọc và thứ tự sắp xếp.'),
    ('Đổi số dòng/trang', 'Change', 'Đặt lại về trang 1 và tải lại danh sách.'),
    ('Bấm tiêu đề cột có mũi tên', 'Click',
     'Đổi cột và chiều sắp xếp rồi tải lại danh sách. Cột được phép sắp xếp: Số tài khoản, Chủ tài '
     'khoản, Ngày tạo, Ngày cập nhật; cột khác bấm vào không có tác dụng.'),
    ('Bấm số tài khoản', 'Click', 'Mở cửa sổ Xem tài khoản ngân hàng ở chế độ chỉ đọc.'),
])

# ---------------------------------------------------------------- 2.2
d.h3('2.2 Tìm kiếm và lọc danh sách')

sub('2.2.1 Giới thiệu')
d.intro_table(
    ten='Tìm kiếm và lọc danh sách tài khoản ngân hàng',
    mota='Thu hẹp danh sách theo từ khoá (số tài khoản, chủ tài khoản, ngân hàng), theo tên chi '
         'nhánh và theo trạng thái.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đang ở màn hình danh sách tài khoản ngân hàng.',
    chinh='1. Người dùng nhập từ khoá vào ô tìm nhanh và/hoặc nhập, chọn các ô lọc.\n'
          '2. Người dùng bấm Tìm kiếm hoặc nhấn Enter trong ô tìm nhanh.\n'
          '3. Hệ thống đặt lại về trang 1 và tải danh sách khớp toàn bộ điều kiện, trong phạm vi '
          'công ty của người đăng nhập.\n'
          '4. Hệ thống ghi nhớ bộ lọc trong 10 phút để dùng lại khi quay về màn hình.',
    phu='• Thay đổi ô Chi nhánh hoặc Trạng thái → hệ thống tự tải lại ngay, không cần bấm Tìm '
        'kiếm.\n'
        '• Chỉ gõ từ khoá mà chưa bấm Tìm kiếm / Enter → danh sách giữ nguyên.\n'
        '• Bấm Làm mới → xoá toàn bộ điều kiện và tải lại danh sách đầy đủ.\n'
        '• Không có kết quả → hiển thị trạng thái rỗng.',
    dacbiet='Tổng số ô lọc kể cả ô tìm nhanh là 3 nên bộ lọc bày hết trên một hàng ngang; màn hình '
            'không có nút “Tìm kiếm nâng cao” và không có cửa sổ “Cài đặt bộ lọc”.')

sub('2.2.2 Layout màn hình')
lay(shot=shot('03-tim-kiem.png'),
    shot_caption='Kết quả sau khi tìm nhanh theo từ khoá “TRANG TEST”')

sub('2.2.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Khối Bộ lọc danh sách', 'Modal', 'Enable', '–', '–', 'Mở sẵn',
     'Hàng ngang gồm 3 ô lọc và 2 nút.'),
    ('Ô tìm nhanh', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Gợi ý “Tìm theo số tài khoản, chủ tài khoản, ngân hàng”; khớp gần đúng đồng thời ba trường; '
     'có nút x để xoá nhanh từ khoá.'),
    ('Ô Chi nhánh', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Khớp gần đúng theo tên chi nhánh lưu trên bản ghi.'),
    ('Ô Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Không', 'Trống',
     'Để trống là lấy tất cả; khớp chính xác.'),
    ('Nút Tìm kiếm', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Áp dụng bộ lọc và quay về trang 1.'),
    ('Nút Làm mới', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Xoá toàn bộ điều kiện lọc, đặt lại trang 1 và tải lại danh sách.'),
])

sub('2.2.4 Danh sách event và xử lý event')
d.event_table([
    ('Nhập ô tìm nhanh', 'Keypress', 'Ghi nhận từ khoá, chưa tải lại danh sách.'),
    ('Nhấn Enter trong ô tìm nhanh', 'Keypress', 'Đặt lại trang 1 và tải lại danh sách.'),
    ('Bấm Tìm kiếm', 'Click', 'Đặt lại trang 1 và tải lại danh sách theo toàn bộ điều kiện.'),
    ('Đổi ô Chi nhánh hoặc Trạng thái', 'Change', 'Đặt lại trang 1 và tải lại danh sách ngay.'),
    ('Bấm Làm mới', 'Click',
     'After:\n– Đặt toàn bộ ô lọc về giá trị khởi tạo.\n– Tải lại danh sách đầy đủ từ trang 1.'),
    ('Rời màn hình rồi quay lại', 'System',
     'Khôi phục bộ lọc đã lưu nếu chưa quá 10 phút; khoá lọc đã bị bỏ khỏi màn hình thì không khôi '
     'phục để tránh lọc ngầm.'),
])

# ---------------------------------------------------------------- 2.3
d.h3('2.3 Tuỳ chỉnh cột hiển thị')

sub('2.3.1 Biểu đồ Usecase')
d.uc_figure('FR-03', 'Tuỳ chỉnh cột hiển thị', 'view',
            [('include', 'Ghi nhớ cấu hình theo người dùng')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-03 Tuỳ chỉnh cột hiển thị')

sub('2.3.2 Giới thiệu')
d.intro_table(
    ten='Tuỳ chỉnh cột hiển thị',
    mota='Cho phép mỗi tài khoản tự chọn cột nào hiển thị trên bảng và thứ tự các cột.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đang ở màn hình danh sách tài khoản ngân hàng.',
    chinh='1. Người dùng bấm nút Cấu hình cột hiển thị bên cạnh nút Tạo mới.\n'
          '2. Hệ thống mở cửa sổ “Tuỳ chỉnh cột” liệt kê 12 cột với trạng thái bật/tắt hiện tại.\n'
          '3. Người dùng tích hoặc bỏ tích cột, hoặc kéo thả để đổi vị trí.\n'
          '4. Người dùng bấm Lưu; hệ thống lưu cấu hình theo tài khoản và vẽ lại bảng.',
    phu='• Bấm Đóng → đóng cửa sổ, bỏ mọi thay đổi chưa lưu.\n'
        '• Cột STT, Số tài khoản và Hành động bị khoá → không bỏ tích được.',
    dacbiet='Cấu hình lưu riêng theo từng tài khoản và từng màn hình; tải lại trang vẫn giữ nguyên '
            'và không ảnh hưởng tới người dùng khác.')

sub('2.3.3 Layout màn hình')
lay(modal='Tuỳ chỉnh cột', shot=shot('04-cau-hinh-cot.png'),
    shot_caption='Cửa sổ Tuỳ chỉnh cột — cột STT và Số tài khoản bị khoá')

sub('2.3.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Cấu hình cột hiển thị', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Nằm cạnh nút Tạo mới, chỉ có biểu tượng cột.'),
    ('Danh sách cột', 'Table/Grid', 'Enable', '12 dòng', '–', 'Theo cấu hình đã lưu',
     'Mỗi dòng gồm ô tích chọn, tên cột và tay nắm kéo thả.'),
    ('Ô tích chọn cột', 'Icon Button', 'Enable / Disable', '–', 'Không', 'Theo cấu hình đã lưu',
     'Cột khoá (STT, Số tài khoản, Hành động) hiển thị ổ khoá và không đổi được.'),
    ('Tay nắm kéo thả', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Kéo để đổi thứ tự cột trên bảng.'),
    ('Nút Lưu', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Lưu cấu hình theo tài khoản và đóng cửa sổ.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ, bỏ thay đổi chưa lưu.'),
])

sub('2.3.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Cấu hình cột hiển thị', 'Click', 'Mở cửa sổ Tuỳ chỉnh cột với cấu hình hiện tại.'),
    ('Tích hoặc bỏ tích một cột', 'Change',
     'Cập nhật trạng thái tạm trong cửa sổ, chưa áp dụng lên bảng.'),
    ('Kéo thả đổi vị trí', 'Change', 'Cập nhật thứ tự tạm trong cửa sổ.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Bỏ qua thay đổi trên cột bị khoá.\n'
     'After:\n– Lưu cấu hình theo tài khoản, vẽ lại bảng và đóng cửa sổ.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ và bỏ toàn bộ thay đổi chưa lưu.'),
])

# ---------------------------------------------------------------- 2.4
d.h3('2.4 Tạo mới tài khoản ngân hàng')

sub('2.4.1 Biểu đồ Usecase')
d.uc_figure('FR-04', 'Tạo mới tài khoản ngân hàng', 'crud',
            [('include', 'Kiểm tra trùng số tài khoản'),
             ('include', 'Kiểm tra chi nhánh thuộc ngân hàng đã chọn'),
             ('extend', 'Gán công ty của người đăng nhập')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-04 Tạo mới tài khoản ngân hàng')

sub('2.4.2 Giới thiệu')
d.intro_table(
    ten='Tạo mới tài khoản ngân hàng',
    mota='Thêm một tài khoản ngân hàng mới cho công ty của người đăng nhập.',
    tacnhan=ACTOR,
    dieukien='Người dùng có quyền Quản lý danh mục tài khoản ngân hàng và hồ sơ nhân sự đã gắn '
             'công ty.',
    chinh='1. Người dùng bấm nút Tạo mới.\n'
          '2. Hệ thống mở cửa sổ “Tạo tài khoản ngân hàng”, ô Trạng thái điền sẵn Hoạt động.\n'
          '3. Người dùng nhập Số tài khoản, chọn Loại tiền tệ, nhập Chủ tài khoản.\n'
          '4. Người dùng chọn Ngân hàng rồi chọn Chi nhánh của ngân hàng đó.\n'
          '5. Người dùng bấm Lưu.\n'
          '6. Hệ thống kiểm tra dữ liệu, chuyển Chủ tài khoản và tên ngân hàng thành chữ in hoa, '
          'gán công ty của người đăng nhập và ghi bản ghi mới.\n'
          '7. Hệ thống hiển thị “Thêm mới thành công”, đóng cửa sổ và tải lại danh sách.',
    phu='• Bỏ trống Chủ tài khoản → giao diện chặn ngay, hiện lỗi đỏ dưới ô và không gửi dữ liệu '
        'đi.\n'
        '• Bỏ trống Số tài khoản, Loại tiền tệ, Ngân hàng hoặc Chi nhánh → hệ thống trả lỗi, hiện '
        'lỗi đỏ dưới đúng từng ô.\n'
        '• Số tài khoản đã tồn tại (kể cả ở công ty khác) → báo “Số tài khoản đã tồn tại”.\n'
        '• Người dùng chưa gắn công ty → báo “Tài khoản đăng nhập chưa gắn công ty, không thể thao '
        'tác”.\n'
        '• Bấm Đóng → đóng cửa sổ, không lưu.',
    dacbiet='Ô Chi nhánh bị khoá cho tới khi chọn Ngân hàng; đổi Ngân hàng thì Chi nhánh tự xoá '
            'trắng. Ô chọn Ngân hàng và Loại tiền tệ chỉ liệt kê bản ghi đang Hoạt động.')

sub('2.4.3 Layout màn hình')
lay(modal='Tạo tài khoản ngân hàng', shot=shot('05-tao-moi.png'),
    shot_caption='Cửa sổ Tạo tài khoản ngân hàng khi mới mở')
d.figure(shot('06-loi-validate.png'),
         'Bấm Lưu khi chưa nhập gì — giao diện chặn ngay tại ô Chủ tài khoản', width_in=6.2)
d.figure(shot('07-loi-bat-buoc-be.png'),
         'Bốn ô còn lại báo lỗi đỏ sau khi hệ thống kiểm tra dữ liệu', width_in=6.2)

sub('2.4.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '–', '“Tạo tài khoản ngân hàng”',
     'Đổi thành “Sửa tài khoản ngân hàng” / “Xem tài khoản ngân hàng” tuỳ chế độ mở.'),
    ('Số tài khoản', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Duy nhất toàn hệ thống; trống báo “Bắt buộc phải nhập”, trùng báo “Số tài khoản đã tồn tại”.'),
    ('Loại tiền tệ', 'Dropdown', 'Enable', 'Danh sách tiền tệ đang Hoạt động', 'Có', 'Trống',
     'Mỗi lựa chọn hiển thị dạng “Mã — Tên”, ví dụ “VNĐ — VietNamDong”.'),
    ('Chủ tài khoản', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Gợi ý “Hệ thống tự chuyển thành CHỮ IN HOA khi lưu”; giao diện chặn trước khi gửi nếu để '
     'trống.'),
    ('Ngân hàng', 'Dropdown', 'Enable', 'Danh sách ngân hàng đang Hoạt động', 'Có', 'Trống',
     'Ngân hàng đang khóa không xuất hiện trong danh sách chọn.'),
    ('Chi nhánh', 'Dropdown', 'Enable / Disable', 'Chi nhánh của ngân hàng đã chọn', 'Có', 'Trống',
     'Bị khoá khi chưa chọn Ngân hàng; đổi Ngân hàng thì tự xoá trắng.'),
    ('Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Có', 'Hoạt động',
     'Không cho xoá trắng lựa chọn; đổi được ngay khi tạo mới.'),
    ('Nút Lưu', 'Button', 'Enable / Disable', '–', '–', 'Enable',
     'Bị khóa trong lúc đang lưu để không tạo trùng bản ghi.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ, không lưu.'),
    ('Thông báo lỗi inline', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     'Chữ đỏ kèm biểu tượng cảnh báo, hiện ngay dưới ô tương ứng; ô lỗi có viền đỏ.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     '“Thêm mới thành công” hoặc thông báo lỗi trả về từ hệ thống.'),
])

sub('2.4.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Tạo mới', 'Click',
     'Before:\n– Kiểm tra quyền; không có quyền thì nút không hiển thị.\n'
     'After:\n– Đặt lại toàn bộ ô nhập, điền sẵn Trạng thái = Hoạt động và mở cửa sổ.'),
    ('Chọn Ngân hàng', 'Change',
     'Lọc lại danh sách Chi nhánh theo ngân hàng vừa chọn; nếu chi nhánh đang chọn không thuộc '
     'ngân hàng mới thì xoá trắng ô Chi nhánh.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Kiểm tra quyền Quản lý danh mục tài khoản ngân hàng; không có quyền → từ chối và '
     'dừng xử lý.\n'
     '– Kiểm tra ô Chủ tài khoản; trống → hiện lỗi đỏ dưới ô và dừng xử lý.\n'
     '– Khóa nút Lưu, bật hiệu ứng chờ.\n'
     'During:\n– Người dùng chưa gắn công ty → hiển thị “Tài khoản đăng nhập chưa gắn công ty, '
     'không thể thao tác”.\n'
     '– Số tài khoản trống → “Bắt buộc phải nhập”; đã tồn tại → “Số tài khoản đã tồn tại”.\n'
     '– Loại tiền tệ trống → “Bắt buộc phải nhập”; đang khóa → “Loại tiền tệ không tồn tại hoặc đã '
     'bị khóa”.\n'
     '– Ngân hàng trống → “Bắt buộc phải nhập”; đang khóa → “Ngân hàng không tồn tại hoặc đã bị '
     'khóa”.\n'
     '– Chi nhánh trống → “Bắt buộc phải nhập”; không thuộc ngân hàng đã chọn → “Chi nhánh không '
     'thuộc ngân hàng đã chọn”.\n'
     '– Trạng thái không hợp lệ → “Trạng thái không hợp lệ”.\n'
     '– Nếu có lỗi kiểm tra → không thực hiện bước After, cửa sổ vẫn mở, dữ liệu đã nhập giữ '
     'nguyên.\n'
     'After:\n– Cắt khoảng trắng đầu cuối, chuyển Chủ tài khoản và tên ngân hàng thành chữ in '
     'hoa.\n'
     '– Chép tên ngân hàng, tên chi nhánh vào bản ghi và gán công ty của người đăng nhập.\n'
     '– Ghi một dòng lịch sử “Tạo mới”.\n'
     '– Hiển thị “Thêm mới thành công”, đóng cửa sổ và tải lại danh sách.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ và đặt lại dữ liệu tạm của cửa sổ.'),
])

# ---------------------------------------------------------------- 2.5
d.h3('2.5 Sửa tài khoản ngân hàng')

sub('2.5.1 Biểu đồ Usecase')
d.uc_figure('FR-05', 'Sửa tài khoản ngân hàng', 'crud',
            [('include', 'Chặn khi bản ghi đang Khóa'),
             ('include', 'Kiểm tra trùng số tài khoản, loại trừ chính bản ghi'),
             ('extend', 'Bắt chọn lại ngân hàng / tiền tệ đã bị khóa')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-05 Sửa tài khoản ngân hàng')

sub('2.5.2 Giới thiệu')
d.intro_table(
    ten='Sửa tài khoản ngân hàng',
    mota='Cập nhật thông tin của một tài khoản ngân hàng đang ở trạng thái Hoạt động, thuộc công '
         'ty của người đăng nhập.',
    tacnhan=ACTOR,
    dieukien='Tài khoản cần sửa thuộc công ty của người đăng nhập và đang ở trạng thái Hoạt động; '
             'tài khoản đang Khóa không hiển thị nút Sửa.',
    chinh='1. Người dùng bấm nút Sửa ở dòng tương ứng.\n'
          '2. Hệ thống nạp dữ liệu hiện tại và mở cửa sổ “Sửa tài khoản ngân hàng”.\n'
          '3. Người dùng chỉnh sửa thông tin cần thiết.\n'
          '4. Người dùng bấm Lưu.\n'
          '5. Hệ thống kiểm tra dữ liệu, cập nhật bản ghi, làm mới tên ngân hàng / chi nhánh theo '
          'lựa chọn mới và ghi lịch sử thay đổi.\n'
          '6. Hệ thống hiển thị “Cập nhật thành công”, đóng cửa sổ và tải lại danh sách.',
    phu='• Ngân hàng của bản ghi đã bị khóa → ô Ngân hàng bị xoá trắng kèm thông báo “Ngân hàng '
        'của tài khoản này đã bị khóa, vui lòng chọn ngân hàng khác”.\n'
        '• Loại tiền tệ của bản ghi đã bị khóa → ô Loại tiền tệ bị xoá trắng kèm thông báo tương '
        'ứng.\n'
        '• Bản ghi vừa bị người khác khóa → hệ thống từ chối, yêu cầu mở khoá trước khi cập nhật.\n'
        '• Bản ghi thuộc công ty khác → hệ thống báo không tìm thấy tài khoản ngân hàng.\n'
        '• Vi phạm bắt buộc hoặc trùng số tài khoản → hiện lỗi đỏ dưới ô tương ứng, cửa sổ không '
        'đóng.',
    dacbiet='Đổi trạng thái sang Khóa ngay trong cửa sổ Sửa cũng được ghi thành mốc lịch sử “Thay '
            'đổi trạng thái”, giống như bấm nút Khóa ngoài danh sách.')

sub('2.5.3 Layout màn hình')
lay(modal='Sửa tài khoản ngân hàng', shot=shot('09-sua.png'),
    shot_caption='Cửa sổ Sửa tài khoản ngân hàng — ô Ngân hàng trống vì ngân hàng của bản ghi đã '
                 'bị khóa')

sub('2.5.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '–', '“Sửa tài khoản ngân hàng”',
     'Khác chế độ Tạo mới ở tiêu đề.'),
    ('Các ô thông tin', 'Textbox', 'Enable', '0–255 ký tự', 'Theo mục 2.4', 'Theo dữ liệu',
     'Nạp sẵn dữ liệu hiện tại; ràng buộc bắt buộc và không trùng giống Tạo mới.'),
    ('Ô Ngân hàng', 'Dropdown', 'Enable', 'Danh sách ngân hàng đang Hoạt động', 'Có',
     'Theo dữ liệu', 'Bị xoá trắng nếu ngân hàng của bản ghi đã bị khóa.'),
    ('Ô Loại tiền tệ', 'Dropdown', 'Enable', 'Danh sách tiền tệ đang Hoạt động', 'Có',
     'Theo dữ liệu', 'Bị xoá trắng nếu loại tiền tệ của bản ghi đã bị khóa.'),
    ('Nút Lưu', 'Button', 'Enable / Disable', '–', '–', 'Enable',
     'Bị khóa trong lúc đang lưu.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ, không lưu.'),
    ('Thông báo chặn sửa', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     '“Tài khoản ngân hàng đang bị khóa, không thể sửa” khi bản ghi đã bị khóa.'),
])

sub('2.5.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Sửa', 'Click',
     'Before:\n– Bản ghi đang Khóa → hiển thị “Tài khoản ngân hàng đang bị khóa, không thể sửa” và '
     'không mở cửa sổ.\n'
     'After:\n– Mở cửa sổ, nạp chi tiết bản ghi; ngân hàng hoặc tiền tệ đã bị khóa thì xoá trắng ô '
     'tương ứng kèm thông báo.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Kiểm tra quyền; không có quyền → từ chối và dừng xử lý.\n'
     '– Bản ghi không thuộc công ty người dùng → báo không tìm thấy và dừng xử lý.\n'
     '– Bản ghi đang Khóa → từ chối, yêu cầu mở khoá trước khi cập nhật.\n'
     'During:\n– Kiểm tra bắt buộc và trùng số tài khoản, có loại trừ chính bản ghi đang sửa.\n'
     '– Kiểm tra ngân hàng, chi nhánh, loại tiền tệ như mục 2.4.\n'
     '– Nếu có lỗi → hiển thị lỗi đỏ dưới từng ô, không thực hiện bước After.\n'
     'After:\n– Chụp giá trị trước khi ghi rồi cập nhật bản ghi, làm mới tên ngân hàng và tên chi '
     'nhánh.\n'
     '– Ghi lịch sử: đổi trạng thái ghi thành mốc “Thay đổi trạng thái”, các trường còn lại ghi '
     'thành mốc “Thay đổi thông tin”; không đổi gì thì không ghi.\n'
     '– Hiển thị “Cập nhật thành công”, đóng cửa sổ và tải lại danh sách.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ, không ghi dữ liệu.'),
])

# ---------------------------------------------------------------- 2.6
d.h3('2.6 Xem chi tiết tài khoản ngân hàng')

sub('2.6.1 Giới thiệu')
d.intro_table(
    ten='Xem chi tiết tài khoản ngân hàng',
    mota='Xem toàn bộ thông tin của một tài khoản ngân hàng ở chế độ chỉ đọc, kèm khối lịch sử '
         'thay đổi.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đang ở màn hình danh sách và bản ghi thuộc công ty của mình.',
    chinh='1. Người dùng bấm vào số tài khoản ở cột định danh.\n'
          '2. Hệ thống mở cửa sổ “Xem tài khoản ngân hàng” và nạp dữ liệu của bản ghi.\n'
          '3. Người dùng xem thông tin; có thể bấm “Xem lịch sử” để mở khối lịch sử ngay trong cửa '
          'sổ.\n'
          '4. Người dùng bấm Đóng để thoát.',
    phu='• Bản ghi đang Khóa → vẫn xem được, ô Trạng thái hiển thị “Khóa”.\n'
        '• Bản ghi chưa có lịch sử → khối lịch sử hiển thị “Chưa có lịch sử thao tác nào.”.\n'
        '• Bản ghi thuộc công ty khác → hệ thống báo không tìm thấy tài khoản ngân hàng.',
    dacbiet='Khác cửa sổ Sửa: nếu ngân hàng của bản ghi đã bị khóa thì cửa sổ Xem VẪN hiển thị tên '
            'ngân hàng đó, lấy từ tên lưu sẵn trên bản ghi.')

sub('2.6.2 Layout màn hình')
lay(modal='Xem tài khoản ngân hàng', shot=shot('10-xem-lich-su.png'),
    shot_caption='Cửa sổ Xem tài khoản ngân hàng với khối Lịch sử đã mở')

sub('2.6.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '“Xem tài khoản ngân hàng”',
     'Tiêu đề của cửa sổ chỉ đọc.'),
    ('Số tài khoản / Chủ tài khoản', 'Textbox', 'Disable', '–', 'Theo dữ liệu',
     'Chỉ đọc, không sửa được.'),
    ('Loại tiền tệ / Ngân hàng / Chi nhánh', 'Dropdown', 'Disable', '–', 'Theo dữ liệu',
     'Chỉ đọc; ngân hàng đã bị khóa vẫn hiển thị đúng tên.'),
    ('Trạng thái', 'Dropdown', 'Disable', 'Hoạt động / Khóa', 'Theo dữ liệu', 'Chỉ đọc.'),
    ('Khối Lịch sử', 'Table/Grid', 'Enable', '–', 'Thu gọn',
     'Bấm “Xem lịch sử” để mở; có nút Làm mới và nút Thu gọn.'),
    ('Nút Đóng', 'Button', 'Enable', '–', 'Hiển thị', 'Nút duy nhất ở chân cửa sổ.'),
], required=False)

sub('2.6.4 Danh sách event và xử lý event')
d.event_table([
    ('Bấm số tài khoản', 'Click', 'Bật chế độ chỉ đọc, nạp chi tiết bản ghi rồi mở cửa sổ.'),
    ('Bấm Xem lịch sử', 'Click', 'Gọi lấy lịch sử của bản ghi và mở khối lịch sử.'),
    ('Bấm Đóng', 'Click',
     'Đóng cửa sổ và đặt lại chế độ để lần mở sau không bị kẹt ở chế độ chỉ đọc.'),
])

# ---------------------------------------------------------------- 2.7
d.h3('2.7 Khóa và Mở khóa tài khoản ngân hàng')

sub('2.7.1 Biểu đồ Usecase')
d.uc_figure('FR-07', 'Khóa / Mở khóa tài khoản ngân hàng', 'action',
            [('include', 'Xác nhận thao tác'),
             ('include', 'Ghi lịch sử thay đổi trạng thái')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-07 Khóa / Mở khóa tài khoản ngân hàng')

sub('2.7.2 Giới thiệu')
d.intro_table(
    ten='Khóa / Mở khóa tài khoản ngân hàng',
    mota='Ngừng sử dụng một tài khoản ngân hàng mà vẫn giữ nguyên dữ liệu, hoặc đưa tài khoản đã '
         'khóa trở lại sử dụng.',
    tacnhan=ACTOR,
    dieukien='Tài khoản cần thao tác thuộc công ty của người đăng nhập và đang hiển thị trên danh '
             'sách.',
    chinh='1. Người dùng bấm nút Khóa hoặc Mở khóa ở dòng tương ứng.\n'
          '2. Hệ thống mở hộp thoại xác nhận có nêu số tài khoản.\n'
          '3. Người dùng bấm nút xác nhận.\n'
          '4. Hệ thống đổi trạng thái và ghi một dòng lịch sử nhóm “Thay đổi trạng thái”.\n'
          '5. Hệ thống hiển thị “Khóa thành công” hoặc “Mở khóa thành công” và tải lại danh sách.',
    phu='• Bấm Hủy → đóng hộp thoại, không thay đổi gì.\n'
        '• Bản ghi vừa bị xoá hoặc thuộc công ty khác → hệ thống báo “Dữ liệu đã thay đổi, vui '
        'lòng tải lại”.\n'
        '• Không có quyền → nút không hiển thị; gọi thẳng chức năng thì bị từ chối.',
    dacbiet='Tài khoản sau khi khóa sẽ ẩn nút Sửa; muốn sửa phải Mở khóa trước.')

sub('2.7.3 Layout màn hình')
lay(shot=shot('12-xac-nhan-khoa.png'),
    shot_caption='Hộp thoại Xác nhận khóa có nêu rõ số tài khoản')

sub('2.7.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Khóa', 'Icon Button', 'Enable / Ẩn', 'Hiển thị khi Hoạt động',
     'Biểu tượng ổ khoá; chỉ hiện khi có quyền quản lý.'),
    ('Nút Mở khóa', 'Icon Button', 'Enable / Ẩn', 'Hiển thị khi Khóa',
     'Biểu tượng ổ khoá mở; chỉ hiện khi có quyền quản lý.'),
    ('Hộp thoại xác nhận', 'Modal', 'Hiển thị', 'Ẩn',
     'Tiêu đề “Xác nhận khóa” hoặc “Xác nhận mở khóa”, nội dung nêu số tài khoản.'),
    ('Nút xác nhận', 'Button', 'Enable', 'Hiển thị', 'Nhãn “Khóa” hoặc “Mở khóa”.'),
    ('Nút Hủy', 'Button', 'Enable', 'Hiển thị', 'Đóng hộp thoại, không thực hiện gì.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', 'Ẩn',
     '“Khóa thành công” / “Mở khóa thành công” / “Dữ liệu đã thay đổi, vui lòng tải lại”.'),
], required=False, scope=False)

sub('2.7.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm Khóa hoặc Mở khóa ở dòng', 'Click',
     'Ghi nhận dòng đang thao tác và mở hộp thoại xác nhận.'),
    ('Bấm nút xác nhận', 'Click',
     'Before:\n– Kiểm tra quyền Quản lý danh mục tài khoản ngân hàng.\n'
     '– Kiểm tra bản ghi thuộc công ty của người đăng nhập; không thuộc → báo dữ liệu đã thay '
     'đổi.\n'
     'During:\n– Đổi trạng thái sang Khóa hoặc Hoạt động theo trạng thái hiện tại của dòng.\n'
     'After:\n– Ghi một dòng lịch sử nhóm Thay đổi trạng thái kèm giá trị cũ → mới.\n'
     '– Hiển thị thông báo thành công và tải lại danh sách.'),
    ('Bấm Hủy', 'Click', 'Đóng hộp thoại, không gọi xử lý.'),
])

# ---------------------------------------------------------------- 2.8
d.h3('2.8 Xem lịch sử thay đổi')

sub('2.8.1 Giới thiệu')
d.intro_table(
    ten='Xem lịch sử thay đổi của tài khoản ngân hàng',
    mota='Truy vết các thao tác đã thực hiện trên một tài khoản ngân hàng: ai làm, lúc nào, thay '
         'đổi gì.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đang ở màn hình danh sách hoặc đang mở cửa sổ Xem của một tài khoản.',
    chinh='1. Người dùng bấm nút Lịch sử ở dòng tương ứng, hoặc mở khối Lịch sử trong cửa sổ Xem.\n'
          '2. Hệ thống tải các mốc lịch sử của bản ghi, sắp xếp mới nhất trước.\n'
          '3. Người dùng lọc theo loại hoạt động nếu cần.\n'
          '4. Người dùng bấm Đóng để thoát.',
    phu='• Bản ghi chưa có lịch sử → hiển thị “Chưa có lịch sử thao tác nào.”.\n'
        '• Lọc theo một nhóm hoạt động → chỉ hiển thị các mốc thuộc nhóm đó.',
    dacbiet='Các trường được theo dõi: Số tài khoản, Chủ tài khoản, Ngân hàng, Chi nhánh, Loại '
            'tiền tệ và Trạng thái. Loại tiền tệ được lưu theo TÊN nên đổi tên danh mục về sau '
            'không làm sai log cũ.')

sub('2.8.2 Layout màn hình')
lay(modal='Lịch sử thay đổi', shot=shot('11-lich-su.png'),
    shot_caption='Cửa sổ Lịch sử thay đổi mở từ nút Lịch sử của một dòng')

sub('2.8.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '“Lịch sử thay đổi”',
     'Dòng phụ ghi “Tài khoản ngân hàng: <số tài khoản> - <chủ tài khoản>”.'),
    ('Nút Bộ lọc', 'Button', 'Enable', '–', 'Hiển thị',
     'Mở khối chọn loại hoạt động và người thực hiện.'),
    ('Ô Loại hoạt động', 'Dropdown', 'Enable',
     'Tạo mới / Thay đổi thông tin / Thay đổi trạng thái', 'Trống',
     'Ba nhóm cố định dùng chung cho mọi màn danh mục.'),
    ('Dòng mốc lịch sử', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Gồm ngày giờ, nhãn hành động có màu, người thực hiện kèm phòng ban và chi tiết thay đổi.'),
    ('Chi tiết thay đổi', 'Label', 'Read-only', '–', 'Theo dữ liệu',
     'Định dạng “<Tên trường>: <giá trị cũ> → <giá trị mới>” với nhãn trường bằng tiếng Việt.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', 'Ẩn', '“Chưa có lịch sử thao tác nào.”.'),
    ('Nút Đóng', 'Button', 'Enable', '–', 'Hiển thị', 'Đóng cửa sổ lịch sử.'),
], required=False)

sub('2.8.4 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Lịch sử ở dòng', 'Click', 'Mở cửa sổ và tải lịch sử của đúng bản ghi đang chọn.'),
    ('Bấm Xem lịch sử trong cửa sổ Xem', 'Click',
     'Mở khối lịch sử ngay trong cửa sổ, cùng nội dung với cửa sổ Lịch sử.'),
    ('Chọn loại hoạt động', 'Change', 'Lọc danh sách mốc theo nhóm hoạt động đã chọn.'),
    ('Đóng cửa sổ', 'Click', 'Xoá bản ghi đang xem để lần mở sau luôn tải lại dữ liệu mới.'),
])

# ======================================================= PHẦN 4. QUY TẮC NGHIỆP VỤ
d.h1('Phần 4. Quy tắc nghiệp vụ')

sub('BR-01 — Phạm vi dữ liệu theo công ty')
d.bullets([
    'Người dùng chỉ thấy và chỉ thao tác được trên tài khoản ngân hàng thuộc công ty ghi trong hồ '
    'sơ nhân sự của mình.',
    'Bản ghi của công ty khác được xử lý như không tồn tại: không hiển thị, không tìm ra, mở trực '
    'tiếp thì báo không tìm thấy.',
    'Người dùng chưa được gắn công ty thì danh sách trống hoàn toàn và không tạo mới được — hệ '
    'thống báo “Tài khoản đăng nhập chưa gắn công ty, không thể thao tác”.',
    'Bản ghi mới luôn được gán công ty của người tạo, người dùng không tự chọn công ty.',
])

sub('BR-02 — Định danh tài khoản')
d.bullets([
    'Số tài khoản là duy nhất trên TOÀN HỆ THỐNG, không phân biệt công ty; trùng thì báo “Số tài '
    'khoản đã tồn tại”.',
    'Khi sửa, phép kiểm tra trùng loại trừ chính bản ghi đang sửa.',
    'Năm trường bắt buộc khi ghi dữ liệu: Số tài khoản, Chủ tài khoản, Ngân hàng, Chi nhánh và '
    'Loại tiền tệ; trạng thái bắt buộc là Hoạt động hoặc Khóa.',
    'Khoảng trắng đầu cuối của Số tài khoản và Chủ tài khoản được cắt bỏ trước khi lưu.',
])

sub('BR-03 — Chữ in hoa và tên lưu kèm')
d.bullets([
    'Chủ tài khoản và tên ngân hàng luôn được chuyển thành CHỮ IN HOA khi lưu.',
    'Tên ngân hàng và tên chi nhánh được chép sẵn vào bản ghi tại thời điểm lưu; danh mục ngân '
    'hàng đổi tên về sau thì danh sách vẫn hiện tên cũ cho tới khi mở Sửa và lưu lại.',
    'Nhờ tên lưu sẵn, cửa sổ Xem vẫn hiển thị đúng tên ngân hàng kể cả khi ngân hàng đó đã bị khóa.',
])

sub('BR-04 — Ràng buộc với ba danh mục liên quan')
d.bullets([
    'Chi nhánh phải thuộc đúng ngân hàng đã chọn; hệ thống kiểm tra lại ở máy chủ chứ không chỉ '
    'lọc ở giao diện.',
    'Chỉ chọn được ngân hàng và loại tiền tệ đang ở trạng thái Hoạt động.',
    'Mở Sửa một bản ghi có ngân hàng hoặc loại tiền tệ đã bị khóa thì ô đó bị xoá trắng, bắt người '
    'dùng chọn lại — tránh lưu ngược giá trị đã khóa vào bản ghi.',
    'Đổi ngân hàng thì ô Chi nhánh tự xoá trắng.',
])

sub('BR-05 — Khóa không phải Xóa')
d.bullets([
    'Màn hình không có chức năng xóa; tài khoản không dùng nữa thì chuyển sang trạng thái Khóa.',
    'Tài khoản đang Khóa không sửa được: giao diện ẩn nút Sửa, máy chủ chặn lại nếu gọi thẳng chức '
    'năng Sửa và yêu cầu mở khoá trước.',
    'Tài khoản đã khóa vẫn nằm trong danh sách với nhãn Khóa và vẫn xem được chi tiết, lịch sử.',
    'Có hai đường đổi trạng thái cho cùng một kết quả: nút Khóa / Mở khóa ngoài danh sách và ô '
    'Trạng thái trong cửa sổ Sửa.',
])

sub('BR-06 — Lịch sử thay đổi')
d.bullets([
    'Các trường được theo dõi: Số tài khoản, Chủ tài khoản, Ngân hàng, Chi nhánh, Loại tiền tệ, '
    'Trạng thái.',
    'Ghi mốc cho các thao tác: Tạo mới, Thay đổi thông tin, Khóa, Mở khóa.',
    'Không ghi mốc nếu nội dung không thay đổi.',
    'Trạng thái và loại tiền tệ được lưu dưới dạng CHỮ (Hoạt động / Khóa, tên tiền tệ) để bản ghi '
    'lịch sử tự chứa nghĩa, không phụ thuộc danh mục sau này.',
    'Bộ lọc loại hoạt động dùng chung 3 nhóm cố định cho mọi màn danh mục.',
])

sub('BR-07 — Phân quyền')
d.bullets([
    'Toàn bộ chức năng của màn dùng chung một quyền “Quản lý danh mục tài khoản ngân hàng”, kể cả '
    'chức năng xem danh sách.',
    'Thiếu quyền: mục menu bị ẩn, gõ thẳng đường dẫn bị chặn, gọi thẳng bất kỳ chức năng nào cũng '
    'bị từ chối.',
    'Chức năng xem lịch sử không đòi quyền riêng — vào được màn là xem được.',
])

sub('BR-08 — Dữ liệu dùng chung với hệ thống ERP')
d.bullets([
    'Danh mục này dùng chung một nguồn dữ liệu với màn tương ứng của hệ thống ERP đang chạy song '
    'song, không đổi cấu trúc dữ liệu gốc.',
    'Bốn khác biệt có chủ đích so với màn ERP: lọc số tài khoản theo khớp gần đúng thay vì khớp '
    'chính xác; bắt buộc Loại tiền tệ cả khi sửa; kiểm tra chi nhánh thuộc ngân hàng ở máy chủ; '
    'sửa lại thông báo trùng cho đúng là “Số tài khoản đã tồn tại”.',
    'Vì hai màn chạy song song trên cùng dữ liệu nên vẫn có khả năng hai người ở hai hệ thống tạo '
    'trùng số tài khoản cùng lúc — đây là hiện trạng được chấp nhận, giữ nguyên như ERP.',
])

d.save()


# ----------------------------------------------------------------- MỤC LỤC
# `d.toc()` chi CHEN TRUONG TOC — mo file len thay TRONG cho toi khi bam Update Field.
# Cho Word cap nhat that (giong hdsd_engine._update_fields_by_word).
def _update_toc_by_word(path, title):
    import subprocess
    ps = """
$p = "{path}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($p, $false, $false)
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
$doc.Repaginate()
$doc.BuiltInDocumentProperties("Title") = "{title}"
$doc.Save()
Write-Output ("Pages=" + $doc.ComputeStatistics(2))
$doc.Close(0)
$word.Quit()
""".format(path=path, title=title)
    res = subprocess.run(["powershell", "-NonInteractive", "-Command", ps],
                         capture_output=True, text=True)
    print('Cap nhat muc luc bang Word:', res.stdout.strip() or res.stderr.strip())


_update_toc_by_word(OUT, 'SRS - Danh mục tài khoản ngân hàng')

# --------------------------------------- Bước 4 của skill: tự kiểm tra form mới
from docx import Document  # noqa: E402

_chk = Document(OUT)
for _s in ['Tổng quan', 'Mini-Spec', 'Tiêu chí nghiệm thu', 'Ngoài phạm vi',
           'Chức năng liên quan', 'Route (FE)', 'Phân hệ:']:
    assert not any(_s in _p.text for _p in _chk.paragraphs), 'Con muc cua form cu: %s' % _s
print('OK - khong con muc nao cua form cu')
