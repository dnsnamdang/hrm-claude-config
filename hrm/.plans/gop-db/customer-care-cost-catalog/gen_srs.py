# -*- coding: utf-8 -*-
"""Sinh SRS theo FORM CHUAN cua 'SRS - Linh vuc.docx'
cho man 'Danh muc dich vu sua chua va chi phi khac' (phan he CSKH).

Khac ban mau: KHONG chen anh chup man hinh — muc 'Layout man hinh' chi ghi duong dan vao man.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import sys, os as _os

# Module dung chung srs_uml_render nam trong assets cua skill srs-documenter
# -> di theo repo, ai clone ve cung chay duoc.
_sys_dir = _os.path.join(
    _os.path.dirname(_os.path.abspath(__file__)),
    "..", "..", "..", ".claude", "skills", "srs-documenter", "assets")
if _sys_dir not in sys.path:
    sys.path.insert(0, _sys_dir)
import srs_uml_render as uml

OUT = r"d:\CompanyProject\hrm\hrm-claude-config\hrm\.plans\gop-db\customer-care-cost-catalog\SRS - Danh mục dịch vụ sửa chữa và chi phí khác.docx"
IMG = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'img')
os.makedirs(IMG, exist_ok=True)

ROUTE = "/customer-care/costs"
FULL_URL = "https://<host-hrm>/customer-care/costs"
MENU = "Phân hệ Chăm sóc khách hàng → Danh mục - Dịch vụ → Danh mục dịch vụ sửa chữa và chi phí khác"

doc = Document()

# ---------- page + base style (bam theo ban mau: Letter 8.5x11, le 1.25in, chu 11pt) ----------
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = Inches(1.25)
sec.right_margin = Inches(1.25)

st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)

for name, size in [('Heading 1', 20), ('Heading 2', 16), ('Heading 3', 14)]:
    hs = doc.styles[name]
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


# ---------------------------------------------------------------- helpers
def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def h3(t):
    doc.add_heading(t, level=3)


def p(t=''):
    return doc.add_paragraph(t)


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(hh)
        for para in c.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = '' if v is None else str(v)
            for para in cells[i].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def intro_table(ten, mota, tacnhan, dieukien, chinh, phu, dacbiet=''):
    """Bang 'Gioi thieu' 8 dong nhu ban mau."""
    table(['Mục', 'Nội dung'], [
        ('Tên chức năng', ten),
        ('Mô tả', mota),
        ('Tác nhân', tacnhan),
        ('Điều kiện ban đầu', dieukien),
        ('Dòng sự kiện chính', chinh),
        ('Dòng sự kiện phụ', phu),
        ('Yêu cầu đặc biệt', dacbiet),
    ], widths=[1.5, 4.5])


def ui_table(rows):
    """Bang 'Mo ta chi tiet giao dien' 8 cot nhu ban mau."""
    table(['STT', 'Tên đối tượng', 'Loại', 'Trạng thái', 'Phạm vi', 'Bắt buộc',
           'Giá trị ban đầu', 'Mô tả'],
          [(i + 1,) + tuple(r) for i, r in enumerate(rows)],
          widths=[0.4, 1.2, 0.8, 0.75, 0.85, 0.6, 0.85, 2.2])


def event_table(rows):
    """Bang 'Danh sach event va xu ly event' 4 cot nhu ban mau."""
    table(['STT', 'Event', 'Loại event', 'Xử lý event'],
          [(i + 1,) + tuple(r) for i, r in enumerate(rows)],
          widths=[0.4, 1.6, 0.9, 3.6])


FIG = {'n': 0}


def figure(png_path, caption, width_in=6.0):
    """Chen anh can giua + chu thich 'Hinh N: ...' ben duoi."""
    FIG['n'] += 1
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(png_path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run('Hình %d: %s' % (FIG['n'], caption))
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


ACTOR_P1 = 'Người quản lý danh mục (P1)'
ACTOR_BOTH = 'Người dùng có quyền P1 hoặc P2'


def uc_figure(code, name, group, relations, actor=ACTOR_P1, caption=None):
    """Ve + chen bieu do use case cho 1 chuc nang."""
    png = os.path.join(IMG, 'uc_%s.png' % code.lower().replace('-', ''))
    uml.draw_usecase(png, actor, code, name, group, relations)
    figure(png, caption or ('Biểu đồ Use Case — %s %s' % (code, name)), width_in=6.2)


def layout(note=''):
    """Muc 'Layout man hinh' — KHONG chen anh, chi ghi duong dan vao man."""
    p('Đường dẫn màn hình:')
    bullets([
        'Menu: %s' % MENU,
        'Route (FE): %s' % ROUTE,
        'URL đầy đủ: %s' % FULL_URL,
    ])
    if note:
        p(note)


# ================================================================ TRANG BIA
h1('SOFTWARE REQUIREMENTS SPECIFICATION (SRS)')
h2('Màn hình: Danh mục dịch vụ sửa chữa và chi phí khác')
h2('Phân hệ: Chăm sóc khách hàng (CSKH) – nhóm Danh mục - Dịch vụ')

table(['Thông tin', 'Nội dung'], [
    ('Mã màn hình', 'CSKH-DM-COST'),
    ('Đường dẫn', ROUTE),
    ('Phiên bản', '1.0'),
    ('Ngày lập', '07/08/2026'),
    ('Người lập', '@junfoke'),
    ('Trạng thái tài liệu', 'Draft'),
    ('Nguồn đối chiếu', 'Màn ERP admin/accounting/costs?kind_of=2 (bảng costs)'),
], widths=[1.8, 4.2])

# ================================================================ 1. GIOI THIEU
h1('1. Giới thiệu')

h2('1.1 Mục đích')
p('Tài liệu này đặc tả yêu cầu phần mềm (SRS) cho màn hình quản lý danh mục dịch vụ sửa chữa '
  'và chi phí khác, nhằm:')
bullets([
    'Thống nhất yêu cầu giữa BA/PO/Dev/Test',
    'Là căn cứ nghiệm thu chức năng và phân quyền',
    'Đảm bảo ràng buộc nghiệp vụ (không xoá dịch vụ đã phát sinh chứng từ, không sửa/xoá chi phí hệ thống) '
    'được thực thi đúng',
    'Làm rõ đặc thù chiết khấu tách theo từng công ty — cùng một dịch vụ, mỗi công ty một mức chiết khấu',
])

h2('1.2 Phạm vi')
p('Màn hình Danh mục dịch vụ sửa chữa và chi phí khác cung cấp chức năng:')
bullets([
    'Xem danh sách dịch vụ sửa chữa và chi phí khác, tra cứu nhanh và lọc nâng cao',
    'Sắp xếp theo nhiều cột (kể cả cột Chiết khấu), phân trang',
    'Quản lý dịch vụ (tạo/sửa/khoá/xoá) theo phân quyền',
    'Ghi nhận chiết khấu riêng cho công ty người dùng đang đăng nhập',
    'Xuất danh sách ra file Excel',
])
p('Ngoài phạm vi:')
bullets([
    'Nhóm chi phí phải trả / chi phí bán hàng (cùng bảng dữ liệu nhưng khác phân loại) — vẫn quản lý bên ERP',
    'Đồng bộ dữ liệu sang hệ thống CRM ngoài (đã được quyết định loại bỏ)',
    'Nhập dữ liệu từ file Excel (Import), In danh sách, Lịch sử chỉnh sửa — màn không có các chức năng này',
    'Thiết kế backend/API chi tiết và mapping file Excel',
])

h2('1.3 Thuật ngữ và viết tắt')
table(['Thuật ngữ', 'Mô tả'], [
    ('Dịch vụ sửa chữa', 'Hạng mục dịch vụ do đơn vị cung cấp, có tính doanh thu'),
    ('Chi phí khác', 'Hạng mục chi phí phát sinh, không tính doanh thu'),
    ('Tỷ lệ giá vốn', 'Tỷ lệ phần trăm dùng để tính giá vốn của dịch vụ'),
    ('Tính doanh thu', 'Cờ phân biệt: Dịch vụ có tính doanh thu / Chi phí khác'),
    ('Chiết khấu', 'Tỷ lệ phần trăm giảm giá, khai báo RIÊNG theo từng công ty'),
    ('Công ty hiện tại', 'Công ty người dùng đang chọn khi đăng nhập, quyết định chiết khấu được hiển thị và ghi'),
    ('P1', 'Quyền “Quản lý dịch vụ sửa chữa và chi phí khác”'),
    ('P2', 'Quyền “Xem dịch vụ sửa chữa và chi phí khác”'),
    ('Quick Search', 'Tìm kiếm nhanh'),
    ('Advanced Filter', 'Bộ lọc nâng cao'),
    ('Chứng từ phát sinh', 'Báo giá hãng / Hợp đồng hãng có sử dụng dịch vụ'),
    ('SRS', 'Software Requirements Specification'),
], widths=[1.8, 4.2])

# ================================================================ 2. TONG QUAN
h1('2. Tổng quan')

h2('2.1 Bối cảnh nghiệp vụ')
p('Danh mục dịch vụ sửa chữa và chi phí khác là dữ liệu nền tảng của nghiệp vụ dịch vụ, dùng để:')
bullets([
    'Khai báo các hạng mục dịch vụ và chi phí đưa vào báo giá, hợp đồng',
    'Cung cấp tỷ lệ giá vốn và thuế VAT làm đầu vào tính giá',
    'Cung cấp mức chiết khấu áp dụng riêng cho từng công ty trong tập đoàn',
])
p('Do đó cần:')
bullets([
    'Phân quyền rõ ràng giữa người quản lý và người chỉ tra cứu',
    'Ràng buộc nghiêm: không cho xoá khi dịch vụ đã phát sinh ở báo giá hãng hoặc hợp đồng hãng — '
    'trường hợp này chỉ được chuyển sang trạng thái Khóa',
    'Bảo vệ các chi phí hệ thống: hai hạng mục “Chi phí đi lại” và “Chi phí vận chuyển” không cho sửa và không cho xoá',
    'Cô lập dữ liệu chiết khấu giữa các công ty: thao tác của công ty này không được làm ảnh hưởng công ty khác',
])

h2('2.2 Nhóm người dùng')
bullets([
    'Người dùng có quyền P1: được quản lý danh mục (tạo/sửa/xoá/khoá) và khai báo chiết khấu cho công ty mình',
    'Người dùng có quyền P2: chỉ được xem/tra cứu và xuất Excel',
    'Người dùng không có P1/P2: bị chặn truy cập',
])

# ================================================================ 3. PHAN QUYEN
h1('3. Phân quyền và kiểm soát truy cập')

h2('3.1 Danh sách quyền')
table(['Ký hiệu', 'Tên quyền', 'Mã quyền', 'Nhóm quyền'], [
    ('P1', 'Quản lý dịch vụ sửa chữa và chi phí khác', '1123', 'Danh mục dịch vụ bảo dưỡng'),
    ('P2', 'Xem dịch vụ sửa chữa và chi phí khác', '1124', 'Danh mục dịch vụ bảo dưỡng'),
], widths=[0.8, 2.8, 0.9, 1.5])

h2('3.2 Quy tắc truy cập bắt buộc')
bullets([
    'Chỉ user có P1 hoặc P2 mới được truy cập màn hình.',
    'User không có P1/P2: không hiển thị menu điều hướng tới màn hình.',
    'User không có P1/P2: truy cập trực tiếp URL bị chặn (điều hướng về trang 404), gọi API trả về lỗi 403.',
    'User chỉ có P2: mọi thao tác ghi (tạo/sửa/xoá) bị chặn ở cả giao diện lẫn API (403), '
    'không phụ thuộc vào việc giao diện có ẩn nút hay không.',
    'Chiết khấu hiển thị và ghi nhận theo công ty người dùng đang đăng nhập; '
    'người dùng không xác định được công ty thì cột Chiết khấu để trống và không ghi được chiết khấu.',
])

h2('3.3 Ma trận phân quyền')
table(['Chức năng', 'P1', 'P2', 'Không có quyền'], [
    ('Truy cập màn', '✅', '✅', '❌'),
    ('Xem danh sách tất cả dữ liệu', '✅', '✅', '❌'),
    ('Tìm kiếm nhanh / Lọc nâng cao', '✅', '✅', '❌'),
    ('Sắp xếp / Phân trang', '✅', '✅', '❌'),
    ('Tạo mới', '✅', '❌', '❌'),
    ('Sửa', '✅', '❌', '❌'),
    ('Xoá / Khoá', '✅', '❌', '❌'),
    ('Khai báo chiết khấu theo công ty', '✅', '❌', '❌'),
    ('Xuất Excel', '✅', '✅', '❌'),
], widths=[3.0, 0.8, 0.8, 1.4])

# ================================================================ 4. FUNCTION LIST
h1('4. Danh mục chức năng (Function list)')
table(['ID', 'Chức năng', 'Mô tả đặc tả thu nhỏ (Mini-Spec)', 'Quyền'], [
    ('FR-01', 'Truy cập màn hình',
     'Kiểm tra quyền P1/P2. Không có quyền sẽ bị chặn (ẩn menu, chặn URL, API trả 403).', 'P1, P2'),
    ('FR-02', 'Xem danh sách',
     'Hiển thị bảng dữ liệu 9 cột, phân trang, mặc định sắp xếp theo ngày tạo mới nhất trước. '
     'Cột Chiết khấu lấy theo công ty người dùng đang đăng nhập.', 'P1, P2'),
    ('FR-03', 'Tìm kiếm & Lọc',
     'Kết hợp Quick Search (theo tên) và Advanced Filter (Tên, Trạng thái, Tính doanh thu, Người cập nhật). '
     'Hỗ trợ sắp xếp trên 7 cột, kể cả cột Chiết khấu.', 'P1, P2'),
    ('FR-04', 'Tạo mới',
     'Mở Modal, nhập Tên (*), Tỷ lệ giá vốn (*), Thuế VAT (*), Tính doanh thu, Chiết khấu. '
     'Bản ghi mới luôn ở trạng thái Hoạt động.', 'P1'),
    ('FR-05', 'Chỉnh sửa',
     'Cho phép sửa toàn bộ trường và Trạng thái. Hai chi phí hệ thống bị chặn sửa. '
     'Chiết khấu chỉ ghi/xoá cho công ty hiện tại.', 'P1'),
    ('FR-06', 'Xoá / Khoá',
     'Xoá bản ghi sau khi xác nhận. Nếu dịch vụ đã phát sinh ở báo giá/hợp đồng hãng thì hệ thống '
     'tự chuyển sang trạng thái Khoá thay vì xoá.', 'P1'),
    ('FR-07', 'Xuất Excel',
     'Xuất danh sách theo đúng bộ lọc đang áp dụng ra file Excel.', 'P1, P2'),
], widths=[0.7, 1.4, 3.4, 0.8])

# ================================================================ 5. DAC TA CHI TIET
h1('5. Đặc tả chi tiết theo từng chức năng (FUNCTIONAL PACKAGING)')

h2('5.1 Sơ đồ UML tổng quan')
p('Sơ đồ Use Case tổng quan của màn hình, thể hiện quan hệ giữa hai nhóm người dùng và bảy chức năng:')

_ov = os.path.join(IMG, 'overview.png')
uml.draw_overview(
    _ov,
    'HỆ THỐNG HRM — Danh mục dịch vụ sửa chữa và chi phí khác',
    [(ACTOR_P1, [0, 1, 2, 3, 4, 5, 6]),
     ('Người xem danh mục (P2)', [0, 1, 2, 6])],
    [('FR-01', 'Truy cập màn hình', 'view', None),
     ('FR-02', 'Xem danh sách', 'view', None),
     ('FR-03', 'Tìm kiếm & Lọc', 'view', None),
     ('FR-04', 'Tạo mới', 'crud', None),
     ('FR-05', 'Chỉnh sửa', 'crud', None),
     ('FR-06', 'Xoá / Khoá', 'action', '«extend» Khoá khi đã phát sinh chứng từ'),
     ('FR-07', 'Xuất Excel', 'io', None)])
figure(_ov, 'Sơ đồ Use Case tổng quan màn hình Danh mục dịch vụ sửa chữa và chi phí khác', width_in=6.3)

h2('5.2 Đặc tả chi tiết từng chức năng')

# ------------------------------------------------ 5.2.1 TRUY CAP
h2('5.2.1 Truy cập màn hình danh mục dịch vụ sửa chữa và chi phí khác')

h3('5.2.1.1 Biểu đồ Usecase')
uc_figure('FR-01', 'Truy cập màn hình danh mục dịch vụ sửa chữa và chi phí khác', 'view',
          [('include', 'Kiểm tra quyền truy cập'),
           ('include', 'Xác định công ty hiện tại')],
          actor=ACTOR_BOTH)

h3('5.2.1.2 Giới thiệu')
intro_table(
    'Truy cập màn hình danh mục dịch vụ sửa chữa và chi phí khác',
    'Cho phép người dùng truy cập vào màn hình quản lý danh mục dịch vụ sửa chữa và chi phí khác '
    'để tra cứu và quản lý dữ liệu.',
    'Admin; User được phân quyền (P1 hoặc P2)',
    'Người dùng đã đăng nhập thành công vào hệ thống.',
    '1. Người dùng chọn menu Chăm sóc khách hàng → Danh mục - Dịch vụ → Danh mục dịch vụ sửa chữa và chi phí khác.\n'
    '2. Hệ thống xác thực quyền truy cập (P1 hoặc P2).\n'
    '3. Hệ thống xác định công ty hiện tại của người dùng để lấy dữ liệu chiết khấu.\n'
    '4. Hệ thống điều hướng tới màn hình danh sách và tải dữ liệu trang đầu tiên.',
    '• Người dùng không có quyền → Hệ thống ẩn menu; nếu truy cập trực tiếp URL thì điều hướng về trang 404.\n'
    '• Gọi trực tiếp API khi không có quyền → Hệ thống trả về lỗi 403.\n'
    '• Người dùng không xác định được công ty → Màn hình vẫn hiển thị, cột Chiết khấu để trống ở mọi dòng.',
    '')

h3('5.2.1.3 Layout màn hình')
layout('Ghi chú: tài liệu này không đính kèm ảnh chụp màn hình, người đọc truy cập trực tiếp đường dẫn trên để đối chiếu giao diện.')

h3('5.2.1.4 Tiêu chí nghiệm thu')
p('Người dùng có quyền truy cập:')
bullets([
    'Nhìn thấy menu Danh mục dịch vụ sửa chữa và chi phí khác.',
    'Truy cập được màn hình danh sách.',
    'Hiển thị mặc định: danh sách phải hiển thị đúng cấu trúc bảng: STT, Tên, Tỷ lệ giá vốn, '
    'Tính doanh thu, Thuế VAT, Chiết khấu, Trạng thái, Người cập nhật, Thao tác.',
])
p('Người dùng không có quyền:')
bullets([
    'Không nhìn thấy menu.',
    'Truy cập trực tiếp URL bị chặn (điều hướng 404), gọi API trả về 403.',
])

h3('5.2.1.5 Danh sách event và xử lý event')
event_table([
    ('Click menu Danh mục dịch vụ sửa chữa và chi phí khác', 'Click',
     'Kiểm tra quyền (P1 hoặc P2) và điều hướng tới màn hình danh sách.'),
    ('Truy cập URL trực tiếp', 'System',
     'Kiểm tra quyền; nếu không hợp lệ → điều hướng về trang 404 (giao diện) và trả về lỗi 403 (API).'),
    ('Load màn hình', 'System',
     'Xác định công ty hiện tại của người dùng; tải danh sách trang 1 kèm cột Chiết khấu theo công ty đó.'),
])

# ------------------------------------------------ 5.2.2 XEM DANH SACH
h2('5.2.2 Xem danh sách dịch vụ sửa chữa và chi phí khác')

h3('5.2.2.1 Giới thiệu')
intro_table(
    'Xem danh sách dịch vụ sửa chữa và chi phí khác',
    'Hiển thị danh sách các dịch vụ sửa chữa và chi phí khác đã được khai báo, kèm tỷ lệ giá vốn, '
    'thuế VAT và mức chiết khấu áp dụng cho công ty người dùng đang đăng nhập.',
    'Admin; User được phân quyền (P1 hoặc P2)',
    'Người dùng truy cập thành công màn hình danh mục dịch vụ sửa chữa và chi phí khác.',
    '1. Hệ thống lấy danh sách bản ghi thuộc nhóm “Dịch vụ sửa chữa và chi phí khác”.\n'
    '2. Hệ thống lấy mức chiết khấu tương ứng của công ty hiện tại cho từng bản ghi.\n'
    '3. Hệ thống hiển thị danh sách theo phân trang, mặc định sắp xếp theo ngày tạo giảm dần.',
    '• Không có dữ liệu → Hiển thị danh sách trống.\n'
    '• Bản ghi chưa khai báo chiết khấu cho công ty hiện tại → Cột Chiết khấu để trống.\n'
    '• Người dùng không xác định được công ty → Toàn bộ cột Chiết khấu để trống.',
    '• Danh sách chỉ hiển thị bản ghi thuộc nhóm “Dịch vụ sửa chữa và chi phí khác”; '
    'các bản ghi thuộc nhóm “Chi phí phải trả / Chi phí bán hàng” không được hiển thị ở màn này.\n'
    '• Trạng thái của màn này quy ước 1 = Hoạt động, 0 = Khóa.')

h3('5.2.2.2 Layout màn hình')
layout()

h3('5.2.2.3 Mô tả chi tiết giao diện')
ui_table([
    ('Bảng danh sách dịch vụ', 'Table/Grid', 'Enable', '–', '–', '–',
     'Hiển thị danh sách dịch vụ và chi phí theo phân trang'),
    ('STT', 'Label', 'Enable', '–', '–', '–', 'Số thứ tự bản ghi trên trang'),
    ('Tên', 'Text', 'Enable', '–', '–', 'Lấy từ cơ sở dữ liệu', 'Hiển thị tên dịch vụ / chi phí'),
    ('Tỷ lệ giá vốn', 'Number', 'Enable', '≥ 0', '–', '–', 'Tỷ lệ giá vốn (%), hiển thị theo định dạng số thập phân'),
    ('Tính doanh thu', 'Badge', 'Enable', 'Dịch vụ có tính doanh thu / Chi phí khác', '–', '–',
     'Hiển thị nhãn tương ứng, không hiển thị giá trị số'),
    ('Thuế VAT', 'Number', 'Enable', '0 – 100', '–', '–', 'Thuế VAT (%)'),
    ('Chiết khấu', 'Number', 'Enable', '0 – 100', '–', 'Trống nếu chưa khai báo',
     'Mức chiết khấu của CÔNG TY người dùng đang đăng nhập; hai người dùng khác công ty sẽ thấy giá trị khác nhau'),
    ('Trạng thái', 'Badge', 'Enable', 'Hoạt động / Khóa', '–', 'Hoạt động', 'Hiển thị trạng thái dịch vụ'),
    ('Người cập nhật', 'Text', 'Enable', '–', '–', '–', 'Hiển thị dạng “MÃ - Họ tên”'),
    ('Sửa', 'Icon Button', 'Enable / Ẩn', '–', '–', '–',
     'Mở modal chỉnh sửa. Ẩn khi người dùng không có P1 hoặc bản ghi là chi phí hệ thống'),
    ('Xoá', 'Icon Button', 'Enable / Ẩn', '–', '–', '–',
     'Xoá hoặc khoá dịch vụ. Ẩn khi không có P1, khi bản ghi đang ở trạng thái Khóa, '
     'hoặc khi bản ghi là chi phí hệ thống'),
    ('Phân trang', 'Pagination', 'Enable', '–', '–', 'Trang 1', 'Điều hướng giữa các trang danh sách'),
    ('Chọn số dòng/trang', 'Dropdown', 'Enable', 'Danh sách', 'Không', '10', 'Thay đổi số bản ghi hiển thị mỗi trang'),
    ('Trạng thái rỗng', 'Label', 'Enable', '–', '–', 'Ẩn', 'Hiển thị “Không có dữ liệu” khi danh sách trống'),
])

h3('5.2.2.4 Tiêu chí nghiệm thu')
bullets([
    'Danh sách chỉ hiển thị bản ghi thuộc nhóm “Dịch vụ sửa chữa và chi phí khác”, '
    'không lẫn bản ghi của nhóm chi phí phải trả / chi phí bán hàng.',
    'Danh sách hiển thị đủ cả bản ghi Hoạt động lẫn bản ghi Khóa.',
    'Mặc định sắp xếp theo ngày tạo giảm dần (bản ghi mới nhất nằm đầu danh sách).',
    'Cột Chiết khấu hiển thị đúng mức của công ty người dùng đang đăng nhập; '
    'hai người dùng thuộc hai công ty khác nhau nhìn cùng một dịch vụ phải thấy hai giá trị khác nhau.',
    'Cột Tính doanh thu hiển thị nhãn tiếng Việt, không hiển thị giá trị số 0/1.',
    'Cột Tỷ lệ giá vốn hiển thị đúng phần thập phân (ví dụ 12,5 không được hiển thị thành 125).',
    'Phân trang hoạt động đúng, không trùng hoặc thiếu dữ liệu.',
    'Người không có quyền P1 không nhìn thấy nút Sửa / Xoá.',
])

h3('5.2.2.5 Danh sách event và xử lý event')
event_table([
    ('Load màn hình', 'System',
     'Lấy dữ liệu danh sách kèm mức chiết khấu của công ty hiện tại và hiển thị.'),
    ('Chuyển trang', 'Click', 'Load dữ liệu theo trang được chọn, giữ nguyên toàn bộ điều kiện lọc và sắp xếp.'),
    ('Thay đổi số dòng', 'Change', 'Load lại danh sách theo số dòng/trang, đưa về trang 1.'),
    ('Hover dòng dữ liệu', 'Hover', 'Hiển thị các nút thao tác được phép trên dòng đó.'),
])

# ------------------------------------------------ 5.2.3 TIM KIEM & LOC
h2('5.2.3 Tìm kiếm & Lọc dịch vụ sửa chữa và chi phí khác')

h3('5.2.3.1 Giới thiệu')
intro_table(
    'Tìm kiếm & Lọc dịch vụ sửa chữa và chi phí khác',
    'Cho phép người dùng tìm kiếm nhanh và lọc danh sách theo nhiều tiêu chí, đồng thời sắp xếp '
    'kết quả trên các cột được hỗ trợ.',
    'Admin; User được phân quyền (P1 hoặc P2)',
    'Danh sách dịch vụ sửa chữa và chi phí khác đã được hiển thị.',
    '1. Người dùng nhập từ khoá tìm kiếm nhanh hoặc mở khu vực lọc nâng cao.\n'
    '2. Người dùng chọn các điều kiện lọc.\n'
    '3. Người dùng nhấn nút Tìm kiếm.\n'
    '4. Hệ thống áp dụng đồng thời tất cả điều kiện (AND) và trả về kết quả ở trang 1.',
    '• Không nhập điều kiện nào → Hệ thống hiển thị toàn bộ danh sách.\n'
    '• Không có kết quả phù hợp → Hiển thị danh sách trống.\n'
    '• Nhấn Làm mới → Xoá toàn bộ điều kiện lọc và hiển thị lại danh sách mặc định.',
    '• Bộ lọc Trạng thái và Tính doanh thu phải nhận được giá trị 0 (Khóa / Chi phí khác), '
    'không được bỏ qua điều kiện khi giá trị bằng 0.\n'
    '• Bộ lọc Người cập nhật phải bắt được cả bản ghi chưa từng được sửa (khi đó lấy theo người tạo).')

h3('5.2.3.2 Layout màn hình')
layout()

h3('5.2.3.3 Mô tả chi tiết giao diện')
ui_table([
    ('Tiêu đề “Bộ lọc danh mục dịch vụ sửa chữa và chi phí khác”', 'Label', 'Enable', '–', 'Không', '–',
     'Tiêu đề khu vực tìm kiếm'),
    ('Nút Tìm kiếm nâng cao / Ẩn tìm kiếm nâng cao', 'Button', 'Enable', '–', 'Không', 'Tìm kiếm nâng cao',
     'Mở / thu gọn khu vực tìm kiếm nâng cao'),
    ('Ô tìm kiếm nhanh', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Nhập từ khoá tìm theo Tên dịch vụ / chi phí'),
    ('Tên', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống', 'Nhập tên dịch vụ / chi phí để tìm kiếm'),
    ('Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Không', 'Trống', 'Lọc theo trạng thái'),
    ('Tính doanh thu', 'Dropdown', 'Enable', 'Dịch vụ có tính doanh thu / Chi phí khác', 'Không', 'Trống',
     'Lọc theo phân loại tính doanh thu'),
    ('Người cập nhật', 'Dropdown', 'Enable', 'Danh sách', 'Không', 'Trống',
     'Lọc theo người cập nhật gần nhất; bao gồm cả bản ghi chưa từng sửa (tính theo người tạo)'),
    ('Nút Tìm kiếm', 'Button', 'Enable', '–', 'Không', '–', 'Thực hiện tìm kiếm theo các tiêu chí đã chọn'),
    ('Nút Làm mới', 'Button', 'Enable', '–', 'Không', '–', 'Xoá toàn bộ điều kiện lọc và tải lại danh sách mặc định'),
    ('Tiêu đề cột có thể sắp xếp', 'Table Header', 'Enable',
     'Tên, Trạng thái, Tỷ lệ giá vốn, Thuế VAT, Chiết khấu, Ngày tạo, Ngày cập nhật', 'Không', 'Ngày tạo giảm dần',
     'Nhấn để đổi chiều sắp xếp; sắp xếp thực hiện trên toàn bộ dữ liệu, không chỉ trang hiện tại'),
])

h3('5.2.3.4 Tiêu chí nghiệm thu')
bullets([
    'Tìm kiếm đúng theo từng trường và tổ hợp điều kiện (các điều kiện nối với nhau bằng AND).',
    'Khi không nhập điều kiện, hệ thống hiển thị toàn bộ danh sách.',
    'Lọc Trạng thái = Khóa và Tính doanh thu = Chi phí khác (giá trị 0) phải trả về đúng kết quả, '
    'không bị bỏ qua điều kiện.',
    'Lọc Người cập nhật phải bao gồm cả bản ghi chưa từng được sửa nhưng do người đó tạo.',
    'Sắp xếp theo cột Chiết khấu cho kết quả đúng theo mức chiết khấu của công ty người dùng đang đăng nhập.',
    'Chuyển trang sau khi lọc vẫn giữ nguyên toàn bộ điều kiện lọc và chiều sắp xếp.',
    'Sắp xếp bằng cột không được hỗ trợ thì hệ thống bỏ qua và giữ thứ tự mặc định, không phát sinh lỗi.',
    'Không reload lại toàn trang khi tìm kiếm.',
])

h3('5.2.3.5 Danh sách event và xử lý event')
event_table([
    ('Nhập ô tìm kiếm nhanh', 'Change', 'Lưu tạm giá trị tìm kiếm.'),
    ('Nhấn Enter trong ô tìm kiếm', 'Keypress', 'Thực hiện tìm kiếm nhanh theo tên.'),
    ('Click “Tìm kiếm nâng cao”', 'Click', 'Hiển thị / Ẩn khu vực lọc nâng cao.'),
    ('Chọn Trạng thái', 'Change', 'Cập nhật điều kiện lọc (nhận cả giá trị 0 = Khóa).'),
    ('Chọn Tính doanh thu', 'Change', 'Cập nhật điều kiện lọc (nhận cả giá trị 0 = Chi phí khác).'),
    ('Chọn Người cập nhật', 'Change', 'Cập nhật điều kiện lọc.'),
    ('Click nút Tìm kiếm', 'Click',
     'Đưa về trang 1 và thực hiện tìm kiếm theo tất cả điều kiện đang chọn.'),
    ('Click nút Làm mới', 'Click', 'Xoá toàn bộ điều kiện lọc, đưa về trang 1, hiển thị danh sách mặc định.'),
    ('Click tiêu đề cột', 'Click',
     'Before:\n– Kiểm tra cột có nằm trong danh sách cột được phép sắp xếp không.\n'
     'During:\n– Nếu không hợp lệ → bỏ qua, giữ nguyên thứ tự mặc định.\n'
     'After:\n– Nếu hợp lệ → tải lại dữ liệu theo cột và chiều sắp xếp mới, giữ nguyên điều kiện lọc.'),
])

# ------------------------------------------------ 5.2.4 TAO MOI
h2('5.2.4 Tạo mới dịch vụ sửa chữa và chi phí khác')

h3('5.2.4.1 Biểu đồ Usecase')
uc_figure('FR-04', 'Tạo mới dịch vụ sửa chữa và chi phí khác', 'crud',
          [('include', 'Kiểm tra trùng tên'),
           ('include', 'Ghi nhận chiết khấu theo công ty')])

h3('5.2.4.2 Giới thiệu')
intro_table(
    'Tạo mới dịch vụ sửa chữa và chi phí khác',
    'Chức năng cho phép người dùng khai báo mới một dịch vụ sửa chữa hoặc một khoản chi phí khác, '
    'phục vụ lập báo giá và hợp đồng dịch vụ.',
    'Admin; User được phân quyền quản lý danh mục dịch vụ sửa chữa và chi phí khác (P1)',
    'Người dùng đã đăng nhập thành công vào hệ thống và có quyền P1.',
    '1. Người dùng nhấn nút Thêm mới trên màn hình danh sách.\n'
    '2. Hệ thống mở modal khai báo với các trường trống.\n'
    '3. Người dùng nhập Tên, Tỷ lệ giá vốn, Thuế VAT, chọn Tính doanh thu và nhập Chiết khấu (nếu có).\n'
    '4. Người dùng nhấn nút Lưu.\n'
    '5. Hệ thống kiểm tra quyền và tính hợp lệ của dữ liệu.\n'
    '6. Hệ thống lưu bản ghi ở trạng thái Hoạt động và ghi nhận chiết khấu cho công ty hiện tại.\n'
    '7. Hệ thống đóng modal, tải lại danh sách và hiển thị thông báo thành công.',
    '• Lưu → Hệ thống lưu dữ liệu và cập nhật danh sách.\n'
    '• Hủy → Hệ thống đóng modal và không lưu dữ liệu.\n'
    '• Bỏ trống Chiết khấu hoặc nhập 0 → Hệ thống không tạo bản ghi chiết khấu cho công ty hiện tại.',
    '• Ba trường Tỷ lệ giá vốn, Thuế VAT, Chiết khấu đều là phần trăm, nên dấu phẩy được hiểu là '
    'DẤU THẬP PHÂN (nhập “12,5” phải lưu thành 12.5, không phải 125).\n'
    '• Bản ghi tạo mới luôn ở trạng thái Hoạt động; muốn khoá phải dùng chức năng Chỉnh sửa.\n'
    '• Chiết khấu chỉ áp dụng cho công ty người dùng đang đăng nhập, không ảnh hưởng công ty khác.')

h3('5.2.4.3 Layout màn hình')
layout('Modal Thêm mới được mở ngay trên màn hình danh sách theo đường dẫn ở trên.')

h3('5.2.4.4 Mô tả chi tiết giao diện')
ui_table([
    ('Nút Thêm mới', 'Button', 'Enable / Ẩn', '–', 'Không', '–', 'Mở modal khai báo mới; ẩn khi không có quyền P1'),
    ('Tiêu đề modal', 'Label', 'Hiển thị', '–', '–', 'Thêm mới dịch vụ / chi phí', 'Tiêu đề chính của modal'),
    ('Tên', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Ngầm định trống',
     'Nhập tên dịch vụ / chi phí (duy nhất trong nhóm này); tự động cắt khoảng trắng đầu và cuối'),
    ('Tỷ lệ giá vốn (%)', 'Textbox', 'Enable', '≥ 0', 'Có', 'Ngầm định trống',
     'Nhập tỷ lệ giá vốn; chấp nhận dấu phẩy làm dấu thập phân'),
    ('Tính doanh thu', 'Dropdown', 'Enable', 'Dịch vụ có tính doanh thu / Chi phí khác', 'Không', 'Chi phí khác',
     'Chọn phân loại tính doanh thu'),
    ('Thuế VAT (%)', 'Textbox', 'Enable', '0 – 100', 'Có', 'Ngầm định trống',
     'Nhập thuế VAT; chấp nhận dấu phẩy làm dấu thập phân'),
    ('Chiết khấu (%)', 'Textbox', 'Enable', '0 – 100', 'Không', 'Ngầm định trống',
     'Nhập chiết khấu áp dụng cho công ty hiện tại; bỏ trống hoặc nhập 0 thì không lưu chiết khấu'),
    ('Trạng thái', 'Dropdown', 'Disable', 'Hoạt động', 'Không', 'Hoạt động',
     'Bản ghi tạo mới luôn ở trạng thái Hoạt động, không cho chọn khi thêm mới'),
    ('Nút Lưu', 'Button', 'Enable', '–', 'Không', '–', 'Lưu dữ liệu dịch vụ / chi phí'),
    ('Nút Hủy', 'Button', 'Enable', '–', 'Không', '–', 'Đóng modal và không lưu dữ liệu'),
    ('Thông báo lỗi từng trường', 'Label', 'Enable', '–', '–', 'Ẩn',
     'Hiển thị viền đỏ và dòng lỗi ngay dưới ô nhập, chỉ hiện sau lần nhấn Lưu đầu tiên'),
    ('Thông báo kết quả', 'Toast / Alert', '–', '–', '–', 'Ẩn', 'Hiển thị kết quả lưu thành công hoặc thất bại'),
])

h3('5.2.4.5 Tiêu chí nghiệm thu')
bullets([
    'Lưu thành công khi dữ liệu hợp lệ và Tên chưa tồn tại trong nhóm dịch vụ sửa chữa và chi phí khác.',
    'Hệ thống báo lỗi và không cho lưu nếu bỏ trống các trường đánh dấu (*): Tên, Tỷ lệ giá vốn, Thuế VAT.',
    'Nhập “12,5” vào các trường phần trăm phải lưu đúng giá trị 12,5 (không được thành 125).',
    'Bản ghi vừa tạo luôn ở trạng thái Hoạt động và xuất hiện ở đầu danh sách.',
    'Có nhập Chiết khấu > 0 → hệ thống ghi nhận chiết khấu cho công ty hiện tại; '
    'bỏ trống hoặc nhập 0 → không phát sinh bản ghi chiết khấu.',
    'Chiết khấu vừa khai báo không làm thay đổi chiết khấu của các công ty khác.',
    'Người dùng không có quyền P1 không nhìn thấy nút “Thêm mới”; gọi trực tiếp API bị chặn 403.',
])

h3('5.2.4.6 Danh sách event và xử lý event')
event_table([
    ('Nhấn nút Thêm mới', 'Click',
     'Mở modal với toàn bộ trường trống; nếu vừa mở modal Sửa trước đó thì phải xoá sạch dữ liệu cũ.'),
    ('Nhấn nút Lưu', 'Click',
     'Before:\n– Kiểm tra quyền P1.\n'
     '– Nếu không có quyền → hiển thị thông báo “Bạn không có quyền thực hiện chức năng này.” và dừng xử lý.\n'
     'During:\n'
     '– Tên trống → hiển thị “Bắt buộc phải nhập.”\n'
     '– Tên trùng → hiển thị “Đã tồn tại trên hệ thống.”\n'
     '– Tỷ lệ giá vốn trống → hiển thị “Bắt buộc phải nhập.”; nhỏ hơn 0 → “Không được nhỏ hơn 0.”\n'
     '– Thuế VAT trống → hiển thị “Bắt buộc phải nhập.”; lớn hơn 100 → “Tối đa 100.”\n'
     '– Chiết khấu lớn hơn 100 → hiển thị “Tối đa 100.”\n'
     '– Giá trị không phải số → hiển thị “Phải là số.”\n'
     '– Nếu có lỗi validate → không thực hiện bước After.\n'
     'After:\n'
     '– Chuẩn hoá dấu phẩy thành dấu thập phân cho ba trường phần trăm.\n'
     '– Lưu bản ghi ở trạng thái Hoạt động, ghi nhận người tạo và người cập nhật.\n'
     '– Nếu Chiết khấu > 0 → ghi nhận chiết khấu cho công ty hiện tại.\n'
     '– Đóng modal, tải lại danh sách, hiển thị thông báo “Thêm mới thành công.”'),
    ('Nhấn nút Hủy', 'Click', 'Đóng modal, không lưu dữ liệu, danh sách giữ nguyên.'),
    ('Nhập Tên', 'Change / Blur',
     '- Nếu trống → không kiểm tra trùng.\n'
     '- Nếu có giá trị → cắt khoảng trắng đầu/cuối rồi kiểm tra tồn tại.\n'
     '- Nếu trùng → “Đã tồn tại trên hệ thống”.'),
    ('Nhập Tỷ lệ giá vốn / Thuế VAT / Chiết khấu', 'Change / Blur',
     '– Chấp nhận dấu phẩy hoặc dấu chấm làm dấu thập phân.\n'
     '– Validate khoảng giá trị theo từng trường.\n'
     '– Không kiểm tra trùng.'),
    ('Nhấn Lưu hai lần liên tiếp', 'Click',
     'Vô hiệu hoá nút Lưu trong khi đang xử lý để chỉ tạo đúng một bản ghi.'),
])

# ------------------------------------------------ 5.2.5 CHINH SUA
h2('5.2.5 Chỉnh sửa dịch vụ sửa chữa và chi phí khác')

h3('5.2.5.1 Biểu đồ Usecase')
uc_figure('FR-05', 'Chỉnh sửa dịch vụ sửa chữa và chi phí khác', 'crud',
          [('include', 'Kiểm tra chi phí hệ thống'),
           ('include', 'Kiểm tra trùng tên'),
           ('include', 'Cập nhật chiết khấu theo công ty')])

h3('5.2.5.2 Giới thiệu')
intro_table(
    'Chỉnh sửa dịch vụ sửa chữa và chi phí khác',
    'Chức năng cho phép người dùng cập nhật thông tin của một dịch vụ / chi phí đã khai báo, '
    'bao gồm cả trạng thái và mức chiết khấu áp dụng cho công ty mình.',
    'Admin; User được phân quyền quản lý danh mục dịch vụ sửa chữa và chi phí khác (P1)',
    'Người dùng đã đăng nhập, có quyền P1 và bản ghi không thuộc nhóm chi phí hệ thống.',
    '1. Người dùng chọn một bản ghi trong danh sách và nhấn nút Sửa.\n'
    '2. Hệ thống mở modal, nạp sẵn dữ liệu của bản ghi kèm mức chiết khấu của công ty hiện tại.\n'
    '3. Người dùng chỉnh sửa thông tin cần thay đổi.\n'
    '4. Người dùng nhấn nút Lưu.\n'
    '5. Hệ thống kiểm tra quyền và tính hợp lệ của dữ liệu.\n'
    '6. Hệ thống cập nhật bản ghi và xử lý chiết khấu của công ty hiện tại.\n'
    '7. Hệ thống đóng modal, tải lại danh sách và hiển thị thông báo thành công.',
    '• Lưu → Hệ thống lưu dữ liệu và cập nhật danh sách.\n'
    '• Hủy → Hệ thống đóng modal và không lưu dữ liệu.\n'
    '• Xoá trắng ô Chiết khấu rồi Lưu → Hệ thống xoá mức chiết khấu của công ty hiện tại, '
    'giữ nguyên chiết khấu của các công ty khác.\n'
    '• Chuyển Trạng thái sang Hoạt động → Bản ghi được mở khoá trở lại.',
    '• Hai hạng mục “Chi phí đi lại” và “Chi phí vận chuyển” là chi phí hệ thống: nút Sửa bị ẩn.\n'
    '• Đây là chức năng duy nhất cho phép mở khoá một bản ghi đang ở trạng thái Khóa '
    '(màn hình không có nút Mở khoá riêng).')

h3('5.2.5.3 Layout màn hình')
layout('Modal Chỉnh sửa được mở ngay trên màn hình danh sách theo đường dẫn ở trên.')

h3('5.2.5.4 Mô tả chi tiết giao diện')
ui_table([
    ('Nút Sửa', 'Icon Button', 'Enable / Ẩn', '–', '–', '–',
     'Mở modal chỉnh sửa; ẩn khi không có P1 hoặc bản ghi là chi phí hệ thống'),
    ('Tiêu đề modal', 'Label', 'Hiển thị', '–', '–', 'Chỉnh sửa dịch vụ / chi phí', 'Tiêu đề chính của modal'),
    ('Tên', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Lấy từ hệ thống',
     'Sửa tên dịch vụ / chi phí; vẫn phải duy nhất trong nhóm này'),
    ('Tỷ lệ giá vốn (%)', 'Textbox', 'Enable', '≥ 0', 'Có', 'Lấy từ hệ thống', 'Sửa tỷ lệ giá vốn'),
    ('Tính doanh thu', 'Dropdown', 'Enable', 'Dịch vụ có tính doanh thu / Chi phí khác', 'Không', 'Lấy từ hệ thống',
     'Sửa phân loại tính doanh thu'),
    ('Thuế VAT (%)', 'Textbox', 'Enable', '0 – 100', 'Có', 'Lấy từ hệ thống', 'Sửa thuế VAT'),
    ('Chiết khấu (%)', 'Textbox', 'Enable', '0 – 100', 'Không', 'Chiết khấu của công ty hiện tại',
     'Sửa hoặc xoá trắng để bỏ chiết khấu của công ty hiện tại; không ảnh hưởng công ty khác'),
    ('Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Có', 'Lấy từ hệ thống',
     'Đổi trạng thái; đây là nơi duy nhất mở khoá được bản ghi đang Khóa'),
    ('Nút Lưu', 'Button', 'Enable', '–', 'Không', '–', 'Lưu thay đổi'),
    ('Nút Hủy', 'Button', 'Enable', '–', 'Không', '–', 'Đóng modal và không lưu dữ liệu'),
    ('Thông báo lỗi từng trường', 'Label', 'Enable', '–', '–', 'Ẩn',
     'Hiển thị viền đỏ và dòng lỗi ngay dưới ô nhập, chỉ hiện sau lần nhấn Lưu đầu tiên'),
    ('Thông báo kết quả', 'Toast / Alert', '–', '–', '–', 'Ẩn', 'Hiển thị kết quả cập nhật'),
])

h3('5.2.5.5 Tiêu chí nghiệm thu')
bullets([
    'Cập nhật thành công khi dữ liệu hợp lệ.',
    'Modal nạp đúng dữ liệu của bản ghi được chọn, bao gồm mức chiết khấu của công ty hiện tại.',
    'Giữ nguyên tên cũ khi lưu không bị báo trùng; đổi sang tên đã tồn tại thì bị chặn.',
    'Sửa chiết khấu chỉ thay đổi mức của công ty hiện tại; mức của công ty khác giữ nguyên.',
    'Xoá trắng ô Chiết khấu rồi lưu thì mức chiết khấu của công ty hiện tại bị gỡ bỏ, '
    'công ty khác không bị ảnh hưởng.',
    'Đổi Trạng thái sang Hoạt động thì bản ghi được mở khoá và nút Xoá xuất hiện trở lại.',
    'Không cho chỉnh sửa hai chi phí hệ thống “Chi phí đi lại” và “Chi phí vận chuyển”.',
    'Không cho chỉnh sửa khi không có quyền P1 (ẩn nút ở giao diện và chặn 403 ở API).',
])

h3('5.2.5.6 Danh sách event và xử lý event')
event_table([
    ('Nhấn nút Sửa', 'Click',
     'Before:\n– Kiểm tra quyền P1 và kiểm tra bản ghi có phải chi phí hệ thống không.\n'
     'After:\n– Mở modal và nạp dữ liệu bản ghi kèm chiết khấu của công ty hiện tại.'),
    ('Nhấn nút Lưu', 'Click',
     'Before:\n– Kiểm tra quyền P1.\n'
     '– Nếu không có quyền → hiển thị “Bạn không có quyền thực hiện chức năng này.” và dừng xử lý.\n'
     'During:\n'
     '– Áp dụng toàn bộ quy tắc kiểm tra như chức năng Tạo mới (bỏ qua chính bản ghi khi kiểm trùng tên).\n'
     '– Trạng thái ngoài hai giá trị Hoạt động / Khóa → hiển thị “Trạng thái không hợp lệ.”\n'
     'After:\n'
     '– Cập nhật bản ghi và ghi nhận người cập nhật.\n'
     '– Nếu Chiết khấu > 0 → cập nhật mức chiết khấu của công ty hiện tại.\n'
     '– Nếu Chiết khấu trống hoặc bằng 0 → xoá mức chiết khấu của công ty hiện tại.\n'
     '– Đóng modal, tải lại danh sách, hiển thị “Cập nhật thành công.”'),
    ('Nhấn nút Hủy', 'Click', 'Đóng modal, không lưu dữ liệu.'),
    ('Đổi Trạng thái sang Hoạt động', 'Change',
     'Cập nhật giá trị trạng thái; sau khi lưu, bản ghi được mở khoá và hiển thị lại nút Xoá trên danh sách.'),
    ('Nhập Tên', 'Change / Blur',
     '- Nếu trống → không kiểm tra trùng.\n'
     '- Nếu có giá trị → kiểm tra trùng nhưng bỏ qua chính bản ghi đang sửa.\n'
     '- Nếu trùng bản ghi khác → “Đã tồn tại trên hệ thống”.'),
])

# ------------------------------------------------ 5.2.6 XOA / KHOA
h2('5.2.6 Xoá / Khoá dịch vụ sửa chữa và chi phí khác')

h3('5.2.6.1 Biểu đồ Usecase')
uc_figure('FR-06', 'Xoá / Khoá dịch vụ sửa chữa và chi phí khác', 'action',
          [('include', 'Kiểm tra chứng từ phát sinh'),
           ('extend', 'Khoá dịch vụ khi đã phát sinh chứng từ')])

h3('5.2.6.2 Giới thiệu')
intro_table(
    'Xoá / Khoá dịch vụ sửa chữa và chi phí khác',
    'Chức năng cho phép người dùng loại bỏ một dịch vụ / chi phí khỏi danh mục. '
    'Nếu bản ghi đã phát sinh ở báo giá hãng hoặc hợp đồng hãng thì hệ thống chỉ chuyển sang '
    'trạng thái Khoá thay vì xoá khỏi cơ sở dữ liệu.',
    'Admin; User được phân quyền quản lý danh mục dịch vụ sửa chữa và chi phí khác (P1)',
    'Người dùng đã đăng nhập, có quyền P1; bản ghi đang ở trạng thái Hoạt động và không thuộc '
    'nhóm chi phí hệ thống.',
    '1. Người dùng nhấn nút Xoá trên dòng dữ liệu.\n'
    '2. Hệ thống kiểm tra bản ghi đã phát sinh ở báo giá hãng / hợp đồng hãng chưa.\n'
    '3. Hệ thống hiển thị popup xác nhận, nêu rõ kết quả sẽ là Xoá hẳn hay chuyển sang Khoá.\n'
    '4. Người dùng nhấn Xác nhận.\n'
    '5. Nếu chưa phát sinh chứng từ → Hệ thống xoá bản ghi và xoá toàn bộ dữ liệu chiết khấu của mọi công ty.\n'
    '6. Nếu đã phát sinh chứng từ → Hệ thống chuyển bản ghi sang trạng thái Khoá và giữ nguyên dữ liệu chiết khấu.\n'
    '7. Hệ thống tải lại danh sách và hiển thị thông báo kết quả.',
    '• Nhấn Hủy → Đóng popup, không thay đổi dữ liệu.\n'
    '• Bản ghi đang ở trạng thái Khoá → Nút Xoá không hiển thị.\n'
    '• Bản ghi là chi phí hệ thống → Nút Xoá không hiển thị.',
    '• Đây là thao tác “xoá hoặc khoá”: kết quả phụ thuộc dữ liệu đã phát sinh, người dùng không tự chọn được.\n'
    '• Khi xoá hẳn, dữ liệu chiết khấu của TẤT CẢ các công ty gắn với bản ghi đều bị xoá theo, '
    'không chỉ công ty hiện tại.')

h3('5.2.6.3 Layout màn hình')
layout('Popup xác nhận được mở ngay trên màn hình danh sách theo đường dẫn ở trên.')

h3('5.2.6.4 Mô tả chi tiết giao diện')
ui_table([
    ('Nút Xoá', 'Button / Icon', 'Enable / Ẩn', '–', '–', '–',
     'Thực hiện thao tác xoá dịch vụ đã chọn; ẩn khi không có P1, khi bản ghi đang Khoá '
     'hoặc là chi phí hệ thống'),
    ('Popup xác nhận', 'Modal', 'Enable', '–', 'Có', 'Ẩn', 'Hiển thị yêu cầu xác nhận trước khi xoá'),
    ('Nội dung popup', 'Text', 'Read-only', '–', 'Có', '–',
     'Nêu rõ bản ghi sẽ bị xoá hẳn hay chỉ chuyển sang trạng thái Khoá, kèm nơi đang phát sinh chứng từ'),
    ('Nút Xoá (trong popup)', 'Button', 'Enable', '–', '–', '–', 'Xác nhận thực hiện thao tác'),
    ('Nút Hủy (trong popup)', 'Button', 'Enable', '–', '–', '–', 'Đóng popup và không thay đổi dữ liệu'),
    ('Thông báo kết quả', 'Toast / Alert', '–', '–', '–', 'Ẩn',
     'Hiển thị kết quả: xoá thành công hoặc đã chuyển sang trạng thái Khoá'),
])

h3('5.2.6.5 Tiêu chí nghiệm thu')
p('Người dùng có quyền P1:')
bullets([
    'Xoá thành công dịch vụ chưa phát sinh ở báo giá hãng và hợp đồng hãng.',
    'Bản ghi biến mất khỏi danh sách và toàn bộ dữ liệu chiết khấu của mọi công ty gắn với nó cũng bị xoá.',
])
p('Nếu dịch vụ đã phát sinh chứng từ:')
bullets([
    'Hệ thống không xoá mà chuyển bản ghi sang trạng thái Khoá.',
    'Popup xác nhận phải báo trước điều này, nêu rõ nơi đang sử dụng (Báo giá hãng / Hợp đồng hãng).',
    'Bản ghi vẫn còn trong danh sách với trạng thái Khoá và không còn hiển thị nút Xoá.',
])
p('Người dùng không có quyền P1:')
bullets([
    'Không hiển thị nút Xoá.',
    'Gọi trực tiếp API bị chặn (403).',
])
p('Chi phí hệ thống:')
bullets([
    'Hai hạng mục “Chi phí đi lại” và “Chi phí vận chuyển” không hiển thị nút Xoá.',
])

h3('5.2.6.6 Danh sách event và xử lý event')
event_table([
    ('Nhấn nút Xoá', 'Click',
     'Before:\n– Kiểm tra quyền P1.\n'
     '– Kiểm tra bản ghi có đang ở trạng thái Hoạt động và không phải chi phí hệ thống.\n'
     'During:\n– Kiểm tra bản ghi đã phát sinh ở báo giá hãng / hợp đồng hãng chưa.\n'
     'After:\n– Mở popup xác nhận, nêu rõ kết quả sẽ là Xoá hẳn hay chuyển sang Khoá.'),
    ('Nhấn Xác nhận', 'Click',
     'Before:\n– Kiểm tra lại quyền và điều kiện ở phía máy chủ.\n'
     'During:\n'
     '– Chưa phát sinh chứng từ → xoá dữ liệu chiết khấu của mọi công ty rồi xoá bản ghi.\n'
     '– Đã phát sinh chứng từ → cập nhật trạng thái sang Khoá và ghi nhận người cập nhật.\n'
     'After:\n– Tải lại danh sách và hiển thị thông báo tương ứng với kết quả thực tế.'),
    ('Nhấn nút Hủy', 'Click', 'Đóng popup, không thay đổi dữ liệu.'),
    ('Xoá thất bại', 'System',
     'Hiển thị thông báo lỗi và giữ nguyên dữ liệu; danh sách không thay đổi.'),
])

# ------------------------------------------------ 5.2.7 XUAT EXCEL
h2('5.2.7 Xuất Excel danh mục dịch vụ sửa chữa và chi phí khác')

h3('5.2.7.1 Giới thiệu')
intro_table(
    'Xuất Excel danh mục dịch vụ sửa chữa và chi phí khác',
    'Chức năng cho phép người dùng xuất dữ liệu danh mục dịch vụ sửa chữa và chi phí khác ra file Excel '
    'theo đúng bộ lọc đang áp dụng trên màn hình.',
    'Admin; User được phân quyền (P1 hoặc P2)',
    'Người dùng đã đăng nhập, có quyền truy cập màn hình và danh sách đã được hiển thị.',
    '1. Người dùng áp dụng bộ lọc mong muốn trên màn hình danh sách.\n'
    '2. Người dùng nhấn nút Xuất Excel.\n'
    '3. Hệ thống kiểm tra quyền truy cập.\n'
    '4. Hệ thống dựng file Excel theo đúng bộ lọc đang áp dụng và mức chiết khấu của công ty hiện tại.\n'
    '5. Trình duyệt tải file về máy người dùng.',
    '• Không có kết quả phù hợp → File vẫn được tải về, chỉ chứa dòng tiêu đề.\n'
    '• Người dùng không có quyền → Không nhìn thấy nút hoặc thao tác bị chặn.',
    '• Cột Chiết khấu trong file phản ánh mức của công ty người dùng đang đăng nhập, '
    'nên hai người dùng khác công ty sẽ tải về hai file có giá trị chiết khấu khác nhau.\n'
    '• Màn hình KHÔNG có chức năng Nhập dữ liệu từ file Excel (Import).')

h3('5.2.7.2 Layout màn hình')
layout()

h3('5.2.7.3 Mô tả chi tiết giao diện')
ui_table([
    ('Nút Xuất Excel', 'Button', 'Enable / Ẩn', '–', 'Không', 'Enable',
     'Xuất danh sách ra file Excel theo bộ lọc hiện tại; ẩn khi người dùng không có quyền truy cập màn hình'),
    ('Chỉ báo đang xử lý', 'Loading', '–', '–', '–', 'Ẩn', 'Hiển thị trong lúc hệ thống dựng file'),
    ('Thông báo kết quả', 'Toast / Alert', '–', '–', '–', 'Ẩn', 'Hiển thị kết quả xuất file thành công hoặc thất bại'),
])

h3('5.2.7.4 Tiêu chí nghiệm thu')
bullets([
    'File tải về đúng định dạng Excel và mở được bằng phần mềm bảng tính.',
    'Nội dung file khớp đúng với bộ lọc đang áp dụng trên màn hình, không xuất toàn bộ danh sách.',
    'Cột và thứ tự cột trong file khớp với bảng hiển thị trên màn hình.',
    'Cột Chiết khấu trong file đúng theo công ty người dùng đang đăng nhập.',
    'Khi bộ lọc không có kết quả, file vẫn tải về được và chỉ chứa dòng tiêu đề, không phát sinh lỗi.',
    'Người dùng không có quyền truy cập màn hình thì không sử dụng được chức năng này.',
])

h3('5.2.7.5 Danh sách event và xử lý event')
event_table([
    ('Nhấn nút Xuất Excel', 'Click',
     'Before:\n– Kiểm tra quyền truy cập màn hình (P1 hoặc P2).\n'
     '– Gắn đầy đủ thông tin xác thực vào yêu cầu tải file.\n'
     'During:\n– Gửi kèm toàn bộ điều kiện lọc đang áp dụng.\n'
     'After:\n– Trình duyệt tải file Excel về máy; hiển thị thông báo kết quả.'),
    ('Xuất file thất bại', 'System',
     'Hiển thị thông báo lỗi; không tải về file rỗng hoặc file sai định dạng.'),
])

# ================================================================ 6. BUSINESS RULES
h1('6. Quy tắc nghiệp vụ (Business Rules)')

p('BR-01  Phạm vi dữ liệu của màn hình')
bullets([
    'Màn hình chỉ hiển thị và thao tác trên các bản ghi thuộc nhóm “Dịch vụ sửa chữa và chi phí khác”.',
    'Các bản ghi thuộc nhóm “Chi phí phải trả / Chi phí bán hàng” tuy nằm cùng bảng dữ liệu '
    'nhưng không hiển thị ở màn này và vẫn được quản lý bên hệ thống ERP.',
])

p('BR-02  Duy nhất tên dịch vụ / chi phí')
bullets([
    'Tên phải duy nhất trong phạm vi nhóm “Dịch vụ sửa chữa và chi phí khác”.',
    'Cho phép trùng tên với bản ghi thuộc nhóm khác (chi phí phải trả / chi phí bán hàng).',
    'Tên được cắt khoảng trắng ở đầu và cuối trước khi kiểm tra trùng và trước khi lưu.',
])

p('BR-03  Quy tắc nhập giá trị phần trăm')
bullets([
    'Ba trường Tỷ lệ giá vốn, Thuế VAT và Chiết khấu đều là phần trăm.',
    'Dấu phẩy được hiểu là DẤU THẬP PHÂN, không phải phân cách hàng nghìn: “12,5” lưu thành 12,5.',
    'Tỷ lệ giá vốn: bắt buộc, giá trị ≥ 0, không giới hạn trần.',
    'Thuế VAT: bắt buộc, giá trị trong khoảng 0 – 100.',
    'Chiết khấu: không bắt buộc, giá trị trong khoảng 0 – 100.',
])

p('BR-04  Chiết khấu tách theo công ty')
bullets([
    'Chiết khấu không phải thuộc tính chung của dịch vụ mà được khai báo riêng cho từng công ty.',
    'Màn hình luôn hiển thị và ghi nhận chiết khấu theo công ty người dùng đang đăng nhập.',
    'Thao tác khai báo, sửa hoặc xoá chiết khấu chỉ tác động tới công ty hiện tại, '
    'không làm thay đổi chiết khấu của các công ty khác.',
    'Bỏ trống hoặc nhập giá trị 0 được hiểu là “không áp dụng chiết khấu”, hệ thống sẽ gỡ bỏ '
    'bản ghi chiết khấu của công ty hiện tại (không lưu giá trị 0).',
    'Người dùng không xác định được công ty thì cột Chiết khấu để trống và mọi thao tác ghi chiết khấu bị bỏ qua.',
])

p('BR-05  Quy tắc trạng thái')
bullets([
    'Trạng thái của màn hình này quy ước: 1 = Hoạt động, 0 = Khoá.',
    'Bản ghi tạo mới luôn ở trạng thái Hoạt động, không cho chọn trạng thái khi thêm mới.',
    'Việc mở khoá một bản ghi đang Khoá được thực hiện qua chức năng Chỉnh sửa; '
    'màn hình không có nút Mở khoá riêng.',
])

p('BR-06  Ràng buộc xoá — “xoá hoặc khoá”')
bullets([
    'Nếu dịch vụ đã phát sinh ở Báo giá hãng hoặc Hợp đồng hãng: không cho xoá, hệ thống tự chuyển '
    'bản ghi sang trạng thái Khoá.',
    'Nếu chưa phát sinh chứng từ: xoá bản ghi khỏi cơ sở dữ liệu, đồng thời xoá dữ liệu chiết khấu '
    'của TẤT CẢ các công ty gắn với bản ghi đó.',
    'Chỉ được xoá bản ghi đang ở trạng thái Hoạt động; bản ghi đang Khoá không hiển thị nút Xoá.',
    'Điều kiện phải được kiểm tra lại ở phía máy chủ, không chỉ dựa vào việc giao diện ẩn nút.',
])

p('BR-07  Bảo vệ chi phí hệ thống')
bullets([
    'Hai hạng mục có tên “Chi phí đi lại” và “Chi phí vận chuyển” là chi phí hệ thống.',
    'Không cho phép sửa và không cho phép xoá / khoá hai hạng mục này.',
    'Điều kiện nhận diện dựa trên TÊN của bản ghi, do đó việc đổi tên có thể làm mất hoặc thêm '
    'sự bảo vệ này — cần lưu ý khi vận hành.',
])

p('BR-08  Quy tắc lọc và sắp xếp')
bullets([
    'Các điều kiện lọc kết hợp với nhau bằng phép AND.',
    'Bộ lọc Trạng thái và Tính doanh thu phải nhận được giá trị 0, không được bỏ qua điều kiện khi giá trị bằng 0.',
    'Bộ lọc Người cập nhật bao gồm cả bản ghi chưa từng được sửa: khi đó lấy theo người tạo bản ghi.',
    'Chỉ các cột Tên, Trạng thái, Tỷ lệ giá vốn, Thuế VAT, Chiết khấu, Ngày tạo, Ngày cập nhật '
    'được phép sắp xếp; cột khác bị bỏ qua và giữ thứ tự mặc định.',
    'Thứ tự mặc định của danh sách là theo ngày tạo giảm dần.',
])

p('BR-09  Cập nhật thông tin hệ thống')
bullets([
    'Khi tạo mới, hệ thống ghi nhận người tạo và người cập nhật là người đang đăng nhập.',
    'Khi sửa hoặc khi chuyển trạng thái sang Khoá qua thao tác xoá, hệ thống cập nhật lại người cập nhật.',
    'Tên người tạo / người cập nhật hiển thị theo định dạng “MÃ - Họ tên”.',
])

p('BR-10  Phạm vi tích hợp')
bullets([
    'Màn hình không thực hiện đồng bộ dữ liệu sang hệ thống CRM ngoài. '
    'Nếu môi trường vận hành có bật tính năng đồng bộ CRM thì dữ liệu ghi từ màn hình này sẽ không được đẩy sang CRM.',
    'Dữ liệu của màn hình dùng chung với hệ thống ERP; hai cổng cùng đọc ghi trên một nguồn dữ liệu.',
])

p('Chức năng liên quan: FR-04 Tạo mới, FR-05 Chỉnh sửa, FR-06 Xoá / Khoá.')

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print('DONE ->', OUT)
