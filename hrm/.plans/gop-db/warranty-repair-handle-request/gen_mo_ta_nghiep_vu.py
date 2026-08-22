# -*- coding: utf-8 -*-
"""Sinh tai lieu mo ta nghiep vu man "Phieu xu ly yeu cau"."""
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Mô tả nghiệp vụ - Phiếu xử lý yêu cầu.docx")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)


def h(text, level):
    doc.add_paragraph(text, style="Heading %d" % level)


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = 1
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.cell(i, j).text = val
    return t


# ───────────────────────── BÌA ─────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("MÔ TẢ NGHIỆP VỤ")
r.bold = True
r.font.size = Pt(20)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Phiếu xử lý yêu cầu")
r.bold = True
r.font.size = Pt(16)

ip = doc.add_paragraph()
ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
ip.add_run("Phân hệ CSKH → Kiểm tra bảo hành sửa chữa · Cập nhật ngày 21/08/2026")

doc.add_page_break()

# ───────────────────────── 1 ─────────────────────────
h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
para("Tài liệu mô tả đầy đủ nghiệp vụ của chức năng “Phiếu xử lý yêu cầu”: dùng để làm gì, ai tham "
     "gia, phiếu đi qua những bước nào, mỗi bước hệ thống báo cho ai, và những quy tắc bắt buộc "
     "phải tuân thủ.")
bullet("Nhân viên phòng tiếp nhận (kỹ thuật, bảo hành): hiểu mình phải làm gì với từng yêu cầu.")
bullet("Nhân viên kinh doanh — người lập phiếu yêu cầu: biết yêu cầu của mình được xử lý ra sao.")
bullet("Cán bộ quản lý, người nghiệm thu: đối chiếu hệ thống có chạy đúng nghiệp vụ không.")
bullet("Bộ phận kiểm thử: dùng kèm file test case của cùng thư mục tài liệu.")

# ───────────────────────── 2 ─────────────────────────
h("2. CHỨC NĂNG NÀY DÙNG ĐỂ LÀM GÌ", 1)
h("2.1. Mục đích", 2)
para("Khi nhận được yêu cầu kiểm tra sửa chữa – bảo hành của khách, phòng tiếp nhận phải trả lời "
     "hai câu: hỏng vì NGUYÊN NHÂN gì, và XỬ LÝ theo hướng nào. Phiếu xử lý yêu cầu là nơi ghi lại "
     "hai câu trả lời đó cho từng thiết bị, làm căn cứ cho mọi việc phía sau.")
para("Với mỗi thiết bị, người xử lý chọn một trong hai hướng: tư vấn qua điện thoại là khách tự "
     "khắc phục được (việc kết thúc tại đây), hoặc cần cung cấp thông tin để làm báo giá (chuyển "
     "tiếp cho bước sau).")
h("2.2. Vị trí trong luồng dịch vụ", 2)
para("Đây là chứng từ THỨ HAI. Trước nó là Phiếu yêu cầu kiểm tra sửa chữa – bảo hành (nơi khách "
     "báo hỏng). Sau nó là Phiếu cung cấp thông tin làm báo giá, rồi tới báo giá và hợp đồng dịch "
     "vụ. Phiếu xử lý không tự tạo mới được — nó luôn sinh ra từ một phiếu yêu cầu đang chờ xử lý.")
h("2.3. Giá trị mang lại", 2)
bullet("Yêu cầu của khách được phân loại ngay: việc nào tư vấn là xong, việc nào phải báo giá.")
bullet("Nguyên nhân hỏng được ghi theo danh mục chuẩn, không mỗi người mô tả một kiểu.")
bullet("Người lập phiếu yêu cầu biết yêu cầu của mình đã được xử lý, không phải hỏi lại phòng kỹ thuật.")
bullet("Bộ phận báo giá nhận được đủ dữ liệu đầu vào ngay khi phiếu chuyển sang chờ cung cấp thông tin.")

# ───────────────────────── 3 ─────────────────────────
h("3. NHỮNG AI THAM GIA", 1)
table([
    ["Vai trò", "Làm gì trong nghiệp vụ này"],
    ["Người xử lý\n(nhân viên phòng tiếp nhận)",
     "Mở phiếu yêu cầu đang chờ xử lý, bấm “Tạo phiếu xử lý yêu cầu”, chọn nguyên nhân và hành "
     "động cho từng thiết bị rồi gửi đi. Sửa và xóa được phiếu của chính mình khi còn ở trạng thái "
     "“Đang tạo”."],
    ["Người lập phiếu cung cấp thông tin\n(bộ phận báo giá)",
     "Nhận thông báo khi có phiếu chờ, mở phiếu rồi chọn: lập Phiếu cung cấp thông tin để đi tiếp, "
     "hoặc “Không duyệt” và trả phiếu lại kèm lý do."],
    ["Người lập phiếu yêu cầu\n(nhân viên kinh doanh)",
     "Không thao tác trên màn này, nhưng theo dõi được trạng thái yêu cầu của mình và nhận thông "
     "báo khi phiếu xử lý bị trả lại."],
    ["Cán bộ quản lý",
     "Theo dõi toàn bộ phiếu trong phạm vi quyền của mình, in và xuất dữ liệu để báo cáo."],
])

# ───────────────────────── 4 ─────────────────────────
h("4. VÒNG ĐỜI CỦA PHIẾU", 1)
para("Phiếu có 6 trạng thái. Ba trạng thái đầu do chính màn hình này tạo ra; các trạng thái còn lại "
     "do Phiếu cung cấp thông tin cập nhật ngược về.")
table([
    ["Trạng thái", "Ý nghĩa", "Ai làm phiếu chuyển sang trạng thái này"],
    ["Đang tạo", "Phiếu nháp, chưa gửi đi. Chỉ người lập nhìn thấy.",
     "Người xử lý bấm “Lưu nháp”, hoặc bộ phận báo giá bấm “Không duyệt”."],
    ["Chờ CCTT", "Đã gửi đi, chờ lập Phiếu cung cấp thông tin làm báo giá.",
     "Người xử lý bấm “Lưu và gửi” khi có ít nhất một thiết bị cần báo giá."],
    ["Đã tư vấn điện thoại",
     "Mọi thiết bị đều xử lý xong bằng tư vấn qua điện thoại. Luồng dịch vụ KẾT THÚC tại đây.",
     "Người xử lý, khi chọn “Tư vấn điện thoại” cho TẤT CẢ thiết bị."],
    ["Đang CCTT", "Đang lập Phiếu cung cấp thông tin.", "Chứng từ phía sau."],
    ["Đã CCTT", "Đã cung cấp xong thông tin làm báo giá.", "Chứng từ phía sau."],
    ["Chờ CCTT bổ sung", "Cần bổ sung thêm thông tin.", "Chứng từ phía sau."],
])

# ───────────────────────── 5 ─────────────────────────
h("5. LUỒNG HOẠT ĐỘNG CHI TIẾT", 1)

h("5.1. Bước 1 — Mở phiếu xử lý từ phiếu yêu cầu", 2)
para("Người của phòng tiếp nhận mở màn “Yêu cầu kiểm tra sửa chữa – bảo hành”, tìm phiếu đang ở "
     "trạng thái “Chờ xử lý” gửi về phòng mình rồi bấm “Tạo phiếu xử lý yêu cầu”.")
para("Hệ thống chỉ cho lập phiếu khi đủ ba điều kiện: phiếu yêu cầu đang “Chờ xử lý”, đúng phòng "
     "tiếp nhận của người bấm, và phiếu đó chưa có phiếu xử lý nào. Thiếu một điều kiện là bị từ "
     "chối, kể cả khi gõ thẳng đường dẫn.")

h("5.2. Bước 2 — Điền nguyên nhân và hành động", 2)
para("Màn lập phiếu chép sẵn và khóa toàn bộ thông tin của phiếu yêu cầu: số phiếu yêu cầu, người "
     "yêu cầu, phòng yêu cầu, ngày nhận yêu cầu, khách hàng, người liên hệ, số điện thoại, địa chỉ "
     "sửa chữa và danh sách thiết bị. Người xử lý chỉ điền phần việc của mình:")
bullet("**Nguyên nhân** — chọn được NHIỀU công việc / lỗi thiết bị cho một thiết bị. Danh sách chỉ "
       "gồm những lỗi đã khai cho đúng hàng hóa đó; chưa có thì dùng nút “Thêm nhanh” để khai ngay "
       "tại chỗ, lỗi mới sẽ được gắn cho hàng hóa và tự chọn vào dòng.")
bullet("**Hành động** — chọn MỘT trong hai: “Tư vấn điện thoại” (bắt buộc ghi nội dung đã tư vấn) "
       "hoặc “Cung cấp thông tin làm báo giá”.")
bullet("**Tệp đính kèm** — ảnh hiện trạng, biên bản… nếu cần.")
para("Thiết bị nào do người dùng tự gõ tên, chưa gắn hàng hóa trong danh mục, thì bắt buộc phải "
     "chọn “hàng hóa tương đương” — nếu không, các chứng từ phía sau không có mã hàng để chạy tiếp.")

h("5.3. Bước 3 — Lưu và gửi đi", 2)
table([
    ["Nút bấm", "Điều kiện", "Kết quả"],
    ["Lưu nháp", "Chỉ cần có phiếu yêu cầu gốc.",
     "Phiếu ở “Đang tạo”, sinh số phiếu ngay lần lưu đầu, chỉ người lập nhìn thấy."],
    ["Lưu và gửi",
     "Bắt buộc: mọi thiết bị đều có nguyên nhân và hành động; hành động “Tư vấn điện thoại” phải "
     "có nội dung xử lý; thiết bị tự gõ phải có hàng hóa tương đương.",
     "Phiếu sang “Chờ CCTT”, ghi nhận thời điểm gửi, báo cho bộ phận báo giá, và phiếu yêu cầu gốc "
     "chuyển sang “Đã xử lý”."],
])
para("Có một quy tắc đặc biệt cần nhớ: nếu MỌI thiết bị trên phiếu đều chọn “Tư vấn điện thoại” "
     "thì phiếu tự chuyển thành “Đã tư vấn điện thoại” — bất kể người dùng bấm “Lưu nháp” hay "
     "“Lưu và gửi”. Phiếu yêu cầu gốc cũng chuyển theo và luồng dịch vụ kết thúc, không đi "
     "tiếp sang báo giá.")

h("5.4. Bước 4 — Bộ phận báo giá xử lý", 2)
table([
    ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
    ["Tạo phiếu cung cấp thông tin",
     "Bấm “Tạo phiếu cung cấp thông tin”, hệ thống mở màn lập chứng từ tiếp theo với dữ liệu lấy "
     "sẵn từ phiếu xử lý.",
     "Phiếu đi tiếp trong luồng dịch vụ; trạng thái do chứng từ sau cập nhật."],
    ["Không duyệt",
     "Bấm “Không duyệt”, bắt buộc nhập lý do rồi xác nhận.",
     "Phiếu trở lại “Đang tạo”, lý do được lưu và hiển thị trên phiếu; người xử lý nhận thông báo, "
     "sửa lại rồi gửi lần nữa. Số phiếu giữ nguyên."],
])
para("Lưu ý quan trọng: “Không duyệt” KHÔNG trả phiếu yêu cầu gốc về “Chờ xử lý” — phiếu yêu cầu "
     "vẫn là “Đã xử lý”, vì việc tiếp nhận đã diễn ra rồi. Chỉ khi XÓA hẳn phiếu xử lý thì phiếu "
     "yêu cầu mới quay lại “Chờ xử lý” và mở lại nút “Tạo phiếu xử lý yêu cầu”.")

# ───────────────────────── 6 ─────────────────────────
h("6. THÔNG BÁO — AI NHẬN, KHI NÀO, NỘI DUNG GÌ", 1)
para("Hệ thống gửi thông báo qua chuông trên thanh công cụ. Có hai sự kiện phát sinh thông báo:")
table([
    ["Sự kiện", "Ai nhận", "Nội dung thông báo", "Bấm vào thì đi đâu"],
    ["Phiếu được gửi đi (chuyển sang “Chờ CCTT”)",
     "MỌI nhân viên CÓ QUYỀN “Tạo phiếu cung cấp thông tin” và CÙNG CÔNG TY với người gửi — không "
     "giới hạn theo phòng ban.",
     "[PXL] Chờ duyệt: <số phiếu>. Khách hàng: <tên khách hàng>.",
     "Mở thẳng màn chi tiết của phiếu đó."],
    ["Phiếu bị “Không duyệt”",
     "NGƯỜI LẬP phiếu xử lý.",
     "[PXL] Từ chối: <số phiếu>. Lý do: <lý do không duyệt>.",
     "Mở thẳng màn chi tiết của phiếu đó."],
])
para("Quy ước chung về thông báo:")
bullet("Số phiếu luôn in đậm; nội dung có giới hạn độ dài, phần bị cắt trước là ghi chú.")
bullet("Người nhận đi theo QUYỀN, không theo phòng ban — khác với màn Phiếu yêu cầu (màn đó báo cho "
       "toàn bộ phòng tiếp nhận).")
bullet("Gửi lại phiếu sau khi bị không duyệt cũng phát sinh thông báo mới cho bộ phận báo giá.")
bullet("Các sự kiện KHÔNG gửi thông báo: lưu nháp, sửa phiếu nháp, xóa phiếu, và trường hợp mọi "
       "thiết bị đều tư vấn qua điện thoại (việc đã xong, không ai phải làm tiếp).")
bullet("Lỗi gửi thông báo không làm hỏng nghiệp vụ: phiếu vẫn lưu và vẫn đổi trạng thái bình thường.")

# ───────────────────────── 7 ─────────────────────────
h("7. PHÂN QUYỀN — AI THẤY GÌ, AI LÀM ĐƯỢC GÌ", 1)
h("7.1. Quyền nhìn thấy dữ liệu", 2)
table([
    ["Tên quyền", "Nhìn thấy những phiếu nào"],
    ["Xem phiếu xử lý yêu cầu đi kiểm tra sửa chữa - bảo hành theo tổng công ty",
     "Phiếu của mọi công ty."],
    ["Xem phiếu xử lý yêu cầu đi kiểm tra sửa chữa - bảo hành theo công ty",
     "Phiếu của công ty mình."],
    ["Xem phiếu xử lý yêu cầu đi kiểm tra sửa chữa - bảo hành theo phòng ban",
     "Phiếu của các phòng mình được giao quản lý VÀ phòng mình đang công tác, cộng phiếu do chính "
     "mình lập."],
    ["Tạo phiếu cung cấp thông tin",
     "Ngoài quyền xem, đây là quyền để lập chứng từ tiếp theo và để “Không duyệt”. Người có quyền "
     "này cũng là người nhận thông báo khi có phiếu chờ."],
    ["(Không có quyền nào ở trên)", "Chỉ thấy phiếu do chính mình lập."],
])
para("Hai quy tắc luôn đúng, không phụ thuộc quyền:")
bullet("Phiếu “Đang tạo” của người khác KHÔNG ai nhìn thấy — không hiện ở danh sách, mở bằng đường "
       "dẫn trực tiếp cũng bị từ chối, in cũng bị chặn. Kể cả tài khoản quản trị cấp cao.")
bullet("Ai lập phiếu thì người đó tự sửa và xóa được phiếu nháp của mình.")

h("7.2. Quyền thao tác", 2)
table([
    ["Thao tác", "Điều kiện được phép"],
    ["Lập phiếu xử lý", "Phiếu yêu cầu đang “Chờ xử lý”, đúng phòng tiếp nhận, chưa có phiếu xử lý nào."],
    ["Sửa phiếu", "Phiếu ở “Đang tạo” VÀ do chính mình lập."],
    ["Xóa phiếu", "Phiếu ở “Đang tạo” VÀ do chính mình lập. Xóa xong phiếu yêu cầu gốc quay lại “Chờ xử lý”."],
    ["Tạo phiếu cung cấp thông tin", "Phiếu ở “Chờ CCTT” VÀ có quyền “Tạo phiếu cung cấp thông tin”."],
    ["Không duyệt", "Cùng điều kiện với Tạo phiếu cung cấp thông tin."],
    ["Thêm nhanh nguyên nhân", "Có thêm quyền “Quản lý danh mục công việc - lỗi thiết bị”."],
    ["Xem Lịch sử", "Không cần quyền riêng — ai vào được màn hình thì xem được."],
    ["In phiếu / In danh sách / Xuất Excel", "Mọi người dùng thấy được phiếu."],
])
para("Nút nào không đủ điều kiện thì hệ thống ẨN HẲN, không hiển thị dạng mờ; danh sách nút ở màn "
     "chi tiết luôn khớp với danh sách nút ở dòng tương ứng ngoài màn danh sách.")

# ───────────────────────── 8 ─────────────────────────
h("8. QUY TẮC NGHIỆP VỤ BẮT BUỘC", 1)
h("8.1. Bắt buộc nhập", 2)
table([
    ["Khi bấm", "Bắt buộc phải có"],
    ["Lưu nháp", "Chỉ cần phiếu yêu cầu gốc."],
    ["Lưu và gửi",
     "Mọi thiết bị phải có ít nhất một Nguyên nhân và một Hành động · hành động “Tư vấn điện "
     "thoại” phải có Nội dung xử lý · thiết bị tự gõ phải chọn Hàng hóa tương đương · phiếu phải "
     "có ít nhất một thiết bị."],
])

h("8.2. Nguyên nhân và hành động", 2)
bullet("Một thiết bị chọn được NHIỀU nguyên nhân nhưng chỉ MỘT hành động.")
bullet("Danh sách nguyên nhân của mỗi dòng chỉ gồm lỗi đã khai cho chính hàng hóa đó, không phải "
       "toàn bộ danh mục — tránh chọn nhầm lỗi của thiết bị khác.")
bullet("Không được đưa cùng một thiết bị (cùng serial) vào phiếu hai lần với nguyên nhân trùng nhau.")
bullet("Lỗi đã bị khóa trong danh mục vẫn hiển thị ở phiếu đang dùng nó (có dấu ổ khóa), không bị "
       "mất khi mở lại phiếu.")

h("8.3. Khóa chỉnh sửa", 2)
bullet("Phiếu đã rời trạng thái “Đang tạo” thì không sửa, không xóa được — kể cả khi gõ thẳng "
       "đường dẫn màn Sửa.")
bullet("Muốn sửa lại phiếu đã gửi, bộ phận báo giá phải “Không duyệt” để trả phiếu về “Đang tạo”.")

h("8.4. Số phiếu", 2)
para("Số phiếu sinh tự động theo dạng <mã công ty>.PXL.<năm><số thứ tự>, cấp ngay ở lần lưu đầu "
     "tiên và không thay đổi trong suốt vòng đời phiếu, kể cả khi bị không duyệt rồi gửi lại.")

h("8.5. Lịch sử thay đổi", 2)
para("Hệ thống ghi lại các mốc: Tạo mới, Thay đổi thông tin, Không duyệt (kèm lý do) và Xóa. Xem "
     "được ở hai nơi với cùng một bố cục: mục “Lịch sử” trong menu hành động của mỗi dòng ngoài "
     "danh sách, và khối “Lịch sử” trong thân màn chi tiết. Danh sách sắp xếp mới → cũ, lọc được "
     "theo loại hoạt động, người thực hiện và khoảng ngày.")
para("Giới hạn hiện tại: lịch sử theo dõi thông tin chung và trạng thái của phiếu, chưa theo dõi "
     "thay đổi ở từng dòng thiết bị.")

# ───────────────────────── 9 ─────────────────────────
h("9. TRA CỨU, IN VÀ XUẤT DỮ LIỆU", 1)
bullet("Tìm nhanh: một ô duy nhất, tìm được theo số phiếu xử lý, số phiếu yêu cầu, tên khách hàng "
       "và người xử lý.")
bullet("Bộ lọc nâng cao: trạng thái, số phiếu yêu cầu, khách hàng, tên thiết bị, model, khoảng "
       "ngày tạo, công ty và phòng ban. Chọn xong hệ thống tự lọc.")
bullet("Tùy chỉnh cột: bảng mặc định 7 cột gọn; các cột Số phiếu yêu cầu, Người yêu cầu, Ngày nhận "
       "yêu cầu, Tên thiết bị, Địa chỉ sửa chữa, Ngày xử lý… bật thêm khi cần và được ghi nhớ.")
bullet("In phiếu: in đúng mẫu biểu đang dùng, bảng chi tiết có đủ cột Nguyên nhân và Hành động.")
bullet("In danh sách: in theo đúng bộ lọc đang áp dụng, khổ ngang.")
bullet("Xuất Excel: chọn được trường và thứ tự cột, xuất theo đúng bộ lọc, có báo tiến độ.")

# ───────────────────────── 10 ─────────────────────────
h("10. LIÊN THÔNG VỚI PHẦN MỀM ERP", 1)
para("Màn hình dùng CHUNG dữ liệu với màn tương ứng bên phần mềm ERP: lập phiếu ở cổng nào thì "
     "cổng còn lại cũng thấy ngay, số phiếu chạy chung một dãy. Nghiệp vụ được giữ giống bản ERP, "
     "trừ bốn điểm cố ý làm khác:")
table([
    ["Điểm khác", "Bên ERP", "Bên hệ thống mới", "Lý do"],
    ["Lưu nháp", "Nút “Lưu” cũng bắt nhập đủ nguyên nhân và hành động.",
     "Lưu nháp chỉ cần phiếu yêu cầu gốc.",
     "Người xử lý ghi nhận dở dang rồi bổ sung sau, không phải điền một mạch."],
    ["Xem phiếu nháp của người khác",
     "Quản trị cấp cao mở link trực tiếp vẫn đọc được, dù danh sách đã ẩn phiếu đó.",
     "Không ai xem được, kể cả quản trị; in cũng bị chặn.",
     "Phiếu nháp là việc riêng của người lập; hai chỗ của ERP đang nói ngược nhau."],
    ["Thông báo khi bị không duyệt", "Không báo cho ai, người lập phải tự vào xem.",
     "Báo ngay cho người lập kèm lý do.", "Rút ngắn thời gian phiếu bị treo."],
    ["Vị trí nút “Không duyệt”", "Chỉ có ở màn xem chi tiết.",
     "Có ở cả màn danh sách lẫn màn chi tiết.", "Hai màn phải khớp nhau về số hành động."],
])

# ───────────────────────── 11 ─────────────────────────
h("11. GIỚI HẠN HIỆN TẠI", 1)
bullet("Màn “Phiếu cung cấp thông tin làm báo giá” chưa có trên hệ thống mới. Bấm “Tạo phiếu cung "
       "cấp thông tin” hiện báo hướng dẫn xử lý tạm bên phần mềm ERP; khi màn đó hoàn thành sẽ nối "
       "thẳng.")
bullet("Popup “Thêm nhanh” nguyên nhân chỉ khai phần thông tin chính. Muốn khai thêm thiết bị áp "
       "dụng và dịch vụ sửa chữa đi kèm thì vào màn Danh mục → Công việc, lỗi thiết bị.")
bullet("Lịch sử chưa theo dõi thay đổi ở từng dòng thiết bị (xem mục 8.5).")

doc.save(OUT)
print("Da tao:", OUT)
d = Document(OUT)
print("Heading 1:", sum(1 for p in d.paragraphs if p.style.name == "Heading 1"),
      "| Bang:", len(d.tables), "| Doan:", len(d.paragraphs))
