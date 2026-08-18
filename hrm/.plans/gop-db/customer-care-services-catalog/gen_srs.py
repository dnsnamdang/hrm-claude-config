# -*- coding: utf-8 -*-
"""Sinh SRS man 'Danh muc goi bao duong' (/customer-care/services) — FORM MOI 4 chuong.

Bam ban mau da duoc user chot: .plans/gop-db/customer-docs/SRS - Danh muc khach hang.docx
(= .claude/skills/srs-documenter/assets/SRS_MAU.docx). Cac diem bam theo ban mau, KHONG theo
mo ta trong SKILL.md o cho lech nhau:
  - Muc Layout VAN GIU dong "Menu: ..." (ban mau con dong nay; chi bo "Route (FE): ...")
  - Nhan muc con 2.x.y IN DAM, tieu de BR-0N IN DAM
  - 2 dong tieu de trang dau 24pt can giua, KHONG in dam
  - Truong Muc luc de rong roi cho Word cap nhat that (xem cuoi file)

Nguon doi chieu (doc truc tiep tu code nhanh gop_db, 2026-08-17):
  BE  Modules/CustomerCare/Routes/api.php (prefix /services — 3 quyen Them/Sua/Xoa;
      danh sach/xem/in/xuat Excel KHONG gate; route sua+xoa gan `serviceNotLocked`)
      Modules/CustomerCare/Http/Controllers/V1/ServiceController.php
      Modules/CustomerCare/Services/ServiceService.php
      Modules/CustomerCare/Entities/Service/{Service,ServiceLevel,ServiceMaintain,
      ServiceMaintainLevel}.php
      Modules/CustomerCare/Http/Requests/Service/ServiceRequest.php
      Modules/CustomerCare/Transformers/ServiceResource/ServiceListResource.php
      app/ExcelExport/ServiceExport.php · app/Services/CatalogHistoryService.php (bang `services`)
  FE  hrm-client/pages/customer-care/services/{index.vue,create.vue,_id/index.vue,_id/edit.vue,
      _id/print.vue,components/ServiceFormComponent.vue}
  Anh chup that: gbd_shots/ (cong dev hrm-crm.eteksofts.com, 1440x900) — CHI DE LOCAL.

Chay:  python .plans/gop-db/customer-care-services-catalog/gen_srs.py
"""
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

sys.path.insert(0, os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "..", ".claude", "skills", "srs-documenter", "assets"))
from srs_docx_lib import SrsDoc  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "gbd_shots")
OUT = os.path.join(HERE, "SRS - Danh mục gói bảo dưỡng.docx")


def shot(name):
    return os.path.join(SHOTS, name)


ACTOR = 'Người quản lý danh mục gói bảo dưỡng'
MENU = 'Phân hệ Chăm sóc khách hàng → Danh mục → Danh mục gói bảo dưỡng'
FULL_URL = 'https://<host-hrm>/customer-care/services'

d = SrsDoc(
    out=OUT,
    menu=MENU,
    route='/customer-care/services',
    full_url=FULL_URL,
    img_prefix='gbd_')


def sub(text):
    """Nhan muc con 2.x.y — ban mau in DAM (van la doan thuong, khong phai Heading)."""
    par = d.p()
    par.add_run(text).bold = True
    return par


def lay(shot=None, shot_caption=None, modal=None, url=None, url_note=None):
    """Muc 'Layout man hinh': dong dan nhap + 2 gach dau dong (Menu / URL) + anh chup that."""
    d.p('Đường dẫn màn hình:')
    items = ['Menu: %s' % MENU, 'URL đầy đủ: %s' % (url or FULL_URL)]
    if url_note:
        items.append(url_note)
    d.bullets(items)
    if modal:
        d.p('Cửa sổ %s được mở ngay trên màn hình danh sách theo đường dẫn ở trên.' % modal)
    if shot:
        d.figure(shot, shot_caption or 'Màn hình thực tế', width_in=6.2)


# ================================================================ TRANG ĐẦU
d.title_block('Danh mục gói bảo dưỡng')
# Ban mau: 2 dong tieu de can giua 24pt nhung KHONG in dam -> bo bold cua lib.
for _p in d.doc.paragraphs[:2]:
    for _r in _p.runs:
        _r.bold = None

d.h2('Mục lục')
# Truong TOC de rong, Word se dien noi dung o buoc cuoi file.
d.toc(note='')

# ============================================================ PHẦN 1. GIỚI THIỆU
d.h1('Phần 1. Giới thiệu')

d.h2('1 Mục đích')
d.p('Tài liệu này đặc tả yêu cầu phần mềm cho màn hình Danh mục gói bảo dưỡng, nhằm:')
d.bullets([
    'Là căn cứ nghiệm thu chức năng, phân quyền và ràng buộc dữ liệu của màn hình.',
    'Làm rõ cấu trúc dữ liệu nhiều tầng của một gói bảo dưỡng: thông tin chung, ma trận nội dung '
    'kiểm tra theo từng cấp bảo dưỡng, hệ số giá bán theo từng công ty, hàng hoá áp dụng và file '
    'đính kèm.',
    'Làm rõ công thức tính Giá vốn, Giá công thức, Giá bán cơ sở và giá bán theo từng công ty — '
    'phần dễ hiểu nhầm nhất của màn hình.',
    'Làm rõ ranh giới giữa Xóa và Khóa: gói đã được sử dụng thì không bị xoá mà chuyển sang trạng '
    'thái Khóa.',
    'Làm rõ hiện trạng phân quyền: xem danh sách, xem chi tiết, in phiếu và xuất Excel không đòi '
    'quyền; chỉ ba thao tác ghi dữ liệu mới cần quyền.',
])

d.h2('2 Thuật ngữ và viết tắt')
d.table(['Thuật ngữ', 'Mô tả'], [
    ('Gói bảo dưỡng',
     'Một bản ghi của danh mục, mô tả trọn gói công việc bảo dưỡng: nội dung kiểm tra, cấp bảo '
     'dưỡng, giá bán, hàng hoá áp dụng và tài liệu kèm theo.'),
    ('Cấp bảo dưỡng',
     'Mức độ bảo dưỡng lấy từ danh mục Cấp dịch vụ; mỗi cấp là MỘT CỘT trong bảng nội dung kiểm '
     'tra và có bộ thông số giá riêng.'),
    ('Nội dung kiểm tra bảo dưỡng',
     'Một hạng mục phải làm khi bảo dưỡng; mỗi hạng mục là MỘT DÒNG trong bảng, kèm đơn vị tính và '
     'số lượng.'),
    ('Ghi chú kiểm tra',
     'Ký hiệu công việc phải thực hiện tại ô giao giữa một nội dung kiểm tra và một cấp bảo dưỡng, '
     'ví dụ DK (định kỳ), VS (vệ sinh), CC (căn chỉnh).'),
    ('Định mức công', 'Số công quy đổi cần cho một cấp bảo dưỡng; dùng để tính giá vốn.'),
    ('Hệ số công nghệ', 'Hệ số nhân thêm vào giá vốn theo độ phức tạp công nghệ của cấp đó.'),
    ('Giá vốn',
     'Đơn giá công của công ty quản lý × Định mức công × Hệ số công nghệ, làm tròn xuống.'),
    ('Giá công thức', 'Giá vốn × Hệ số giá bán của gói bảo dưỡng.'),
    ('Giá bán cơ sở',
     'Giá bán chuẩn của một cấp; mặc định bằng Giá công thức nhưng người dùng sửa tay được.'),
    ('Hệ số giá bán theo công ty',
     'Hệ số riêng của từng công ty; giá bán của công ty đó = Giá bán cơ sở × hệ số này.'),
    ('Công ty quản lý gói bảo dưỡng',
     'Công ty mà hệ thống lấy đơn giá công để tính giá vốn của gói.'),
    ('Trạng thái Hoạt động', 'Gói đang dùng bình thường, sửa và xoá được.'),
    ('Trạng thái Khóa',
     'Gói ngừng sử dụng nhưng KHÔNG bị xoá; khi đang khóa thì không sửa được.'),
    ('Gói đã được sử dụng',
     'Gói đã gắn hàng hoá hoặc đã được chọn trong báo giá dịch vụ; gói này không xoá được.'),
    ('Lịch sử thay đổi',
     'Nhật ký thao tác của một bản ghi: ai tạo, sửa, khóa, mở khóa và giá trị trước - sau.'),
    ('SRS', 'Software Requirements Specification.'),
], widths=[1.8, 4.2])

# ============================================================ PHẦN 2. PHÂN QUYỀN
d.h1('Phần 2. Phân quyền')

d.h2('1 Danh sách quyền')

d.p('Nhóm quyền thao tác:')
d.table(['Ký hiệu', 'Tên quyền', 'Tác dụng trên màn hình'], [
    ('Q1', 'Thêm danh mục gói bảo dưỡng',
     'Hiện nút Tạo mới và nút Nhân bản; cho phép ghi gói mới. Thiếu quyền này thì hai nút bị ẩn và '
     'thao tác tạo bị từ chối.'),
    ('Q2', 'Sửa danh mục gói bảo dưỡng',
     'Hiện nút Sửa và nút Mở khóa; cho phép cập nhật gói và mở khoá gói đang Khóa.'),
    ('Q3', 'Xóa danh mục gói bảo dưỡng',
     'Hiện nút Xóa ở những gói chưa được sử dụng; cho phép xoá gói.'),
], widths=[0.8, 2.0, 3.2])
d.p('Ba quyền trên thuộc nhóm quyền Danh mục dịch vụ bảo dưỡng và độc lập với nhau: một tài khoản '
    'có thể chỉ có một trong ba.')

d.p('Chức năng KHÔNG đòi quyền (hiện trạng giữ nguyên theo phần mềm cũ):')
d.bullets([
    'Xem danh sách, xem chi tiết, in phiếu Danh mục kiểm tra bảo dưỡng định kỳ và xuất Excel: mọi '
    'tài khoản đã đăng nhập đều thực hiện được.',
    'Xem lịch sử thay đổi: không có quyền riêng, vào được màn hình là xem được.',
    'Nếu nghiệp vụ muốn siết quyền xem hoặc quyền xuất dữ liệu thì phải bổ sung quyền mới, đây là '
    'điểm cần quyết định chứ không phải lỗi cấu hình môi trường.',
])

d.p('Phạm vi dữ liệu (không phải quyền, luôn áp dụng cho mọi người dùng):')
d.bullets([
    'Danh sách hiển thị TẤT CẢ gói bảo dưỡng của hệ thống, không cắt theo công ty của người đăng '
    'nhập.',
    'Màn hình không phân quyền theo công ty, phòng ban hay bộ phận.',
    'Công ty của người đăng nhập chỉ dùng để điền sẵn ô Công ty quản lý gói bảo dưỡng khi tạo mới.',
])

d.h2('2 Ma trận phân quyền')
d.table(['Chức năng', 'Q1', 'Q2', 'Q3', 'Không có quyền nào'], [
    ('FR-01 Truy cập & xem danh sách', '✅', '✅', '✅', '✅'),
    ('FR-02 Tìm kiếm và lọc', '✅', '✅', '✅', '✅'),
    ('FR-03 Tuỳ chỉnh cột hiển thị', '✅', '✅', '✅', '✅'),
    ('FR-04 Tạo mới gói bảo dưỡng', '✅', '❌', '❌', '❌'),
    ('FR-05 Sửa gói bảo dưỡng', '❌', '✅ (trừ gói đang Khóa)', '❌', '❌'),
    ('FR-06 Xem chi tiết gói', '✅', '✅', '✅', '✅'),
    ('FR-07 Nhân bản gói', '✅', '❌', '❌', '❌'),
    ('FR-08 Xóa gói / Mở khóa gói', 'Mở khóa: ❌', 'Mở khóa: ✅', 'Xóa: ✅', '❌'),
    ('FR-09 Xuất Excel', '✅', '✅', '✅', '✅'),
    ('FR-10 In phiếu kiểm tra bảo dưỡng', '✅', '✅', '✅', '✅'),
    ('FR-11 Xem lịch sử thay đổi', '✅', '✅', '✅', '✅'),
    ('Nhập từ Excel', 'Màn hình không có chức năng này', '—', '—', '—'),
], widths=[2.0, 1.0, 1.4, 0.9, 1.1])

# ================================================= PHẦN 3. ĐẶC TẢ CHI TIẾT
d.h1('Phần 3. Đặc tả chi tiết theo từng chức năng')

d.h2('1 Sơ đồ UML tổng quan')
d.overview_figure(
    'HỆ THỐNG HRM — DANH MỤC GÓI BẢO DƯỠNG',
    [(ACTOR, list(range(11)))],
    [('FR-01', 'Truy cập & xem danh sách', 'view', None),
     ('FR-02', 'Tìm kiếm & Lọc', 'view', None),
     ('FR-03', 'Tuỳ chỉnh cột', 'view', None),
     ('FR-04', 'Tạo mới gói bảo dưỡng', 'crud', '«include» Kiểm tra trùng tên và mã'),
     ('FR-05', 'Sửa gói bảo dưỡng', 'crud', '«include» Chặn khi gói đang Khóa'),
     ('FR-06', 'Xem chi tiết gói', 'view', None),
     ('FR-07', 'Nhân bản gói', 'crud', None),
     ('FR-08', 'Xóa / Mở khóa gói', 'action', '«include» Kiểm tra gói đã được sử dụng'),
     ('FR-09', 'Xuất Excel', 'export', None),
     ('FR-10', 'In phiếu kiểm tra bảo dưỡng', 'export', None),
     ('FR-11', 'Xem lịch sử thay đổi', 'view', None)],
    'Sơ đồ Use Case tổng quan màn Danh mục gói bảo dưỡng')

d.h2('2 Đặc tả chi tiết từng chức năng')

# ---------------------------------------------------------------- 2.1
d.h3('2.1 Truy cập và xem danh sách gói bảo dưỡng')

sub('2.1.1 Giới thiệu')
d.intro_table(
    ten='Truy cập và xem danh sách gói bảo dưỡng',
    mota='Hiển thị danh sách gói bảo dưỡng của toàn hệ thống kèm trạng thái, thông tin người tạo / '
         'người sửa và bộ nút thao tác của từng dòng.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đã đăng nhập vào hệ thống.',
    chinh='1. Người dùng vào menu Chăm sóc khách hàng → Danh mục → Danh mục gói bảo dưỡng.\n'
          '2. Hệ thống nạp cấu hình cột hiển thị của tài khoản và khôi phục bộ lọc đã lưu nếu còn '
          'trong 10 phút.\n'
          '3. Hệ thống trả về trang đầu tiên của danh sách và tổng số bản ghi.\n'
          '4. Bảng hiển thị dữ liệu, ô “Hiển thị a–b / N” hiển thị đúng khoảng và tổng.\n'
          '5. Với mỗi dòng, hệ thống tính sẵn cờ cho phép xoá và giá bán theo từng cấp để hiển thị '
          'ở biểu tượng chữ i cạnh tên gói.',
    phu='• Chưa đăng nhập → điều hướng về màn đăng nhập.\n'
        '• Không có bản ghi khớp bộ lọc → bảng hiện “Không có dữ liệu phù hợp bộ lọc.”.\n'
        '• Gói chưa khai cấp bảo dưỡng → không hiển thị biểu tượng chữ i.\n'
        '• Phiên đăng nhập hết hạn → điều hướng về màn đăng nhập.',
    dacbiet='Màn hình không đòi quyền để xem: tài khoản không có quyền nào vẫn vào được và vẫn thấy '
            'đủ danh sách, chỉ thiếu các nút ghi dữ liệu. Bảng có 10 cột khả dụng, hiển thị sẵn 7 '
            'cột; 3 cột Công ty quản lý gói bảo dưỡng, Người sửa, Ngày sửa mặc định ẩn.')

sub('2.1.2 Layout màn hình')
lay(shot=shot('01-danh-sach.png'),
    shot_caption='Màn Danh mục gói bảo dưỡng lúc mới truy cập')
d.figure(shot('02-menu-hanh-dong.png'),
         'Menu ba chấm của một dòng — Nhân bản, In, Lịch sử', width_in=6.2)

sub('2.1.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Bảng danh sách gói bảo dưỡng', 'Table/Grid', 'Read-only', '–', '–',
     'Hiển thị danh sách theo phân trang, cột Hành động chốt ở cuối bảng.'),
    ('Cột STT', 'Label', 'Read-only', '–', 'Số thứ tự theo trang',
     'Luôn hiển thị, không tắt được ở cửa sổ Tuỳ chỉnh cột.'),
    ('Cột Mã', 'Button', 'Enable', '0–255 ký tự', 'Theo dữ liệu',
     'Cột định danh, bấm vào mở màn Chi tiết gói bảo dưỡng. Luôn hiển thị, không tắt được. Cho '
     'phép sắp xếp.'),
    ('Cột Tên gói bảo dưỡng', 'Text', 'Read-only', '0–255 ký tự', 'Theo dữ liệu',
     'Kèm biểu tượng chữ i khi gói đã khai cấp bảo dưỡng; rê chuột hiện bảng giá bán theo từng '
     'cấp. Cho phép sắp xếp.'),
    ('Cột Công ty quản lý gói bảo dưỡng', 'Text', 'Read-only', '–', 'Ẩn',
     'Công ty dùng để lấy đơn giá công. Mặc định tắt.'),
    ('Cột Người tạo', 'Text', 'Read-only', '–', 'Theo dữ liệu', 'Họ tên người tạo bản ghi.'),
    ('Cột Ngày tạo', 'Text', 'Read-only', 'dd/mm/yyyy HH:mm', 'Theo dữ liệu',
     'Cho phép sắp xếp; là tiêu chí sắp xếp mặc định của danh sách.'),
    ('Cột Người sửa', 'Text', 'Read-only', '–', 'Ẩn', 'Người sửa gần nhất. Mặc định tắt.'),
    ('Cột Ngày sửa', 'Text', 'Read-only', 'dd/mm/yyyy HH:mm', 'Ẩn',
     'Mặc định tắt. Cho phép sắp xếp.'),
    ('Cột Trạng thái', 'Badge', 'Read-only', 'Hoạt động / Khóa', 'Theo dữ liệu',
     'Nhãn xanh cho Hoạt động, nhãn đỏ cho Khóa.'),
    ('Cột Hành động', 'Icon Button', 'Enable / Ẩn', '–', 'Hiển thị',
     'Gồm Sửa, Xóa, Mở khóa và menu ba chấm (Nhân bản, In, Lịch sử). Nút không dùng được thì ẩn '
     'hẳn chứ không làm mờ.'),
    ('Nút Tạo mới', 'Button', 'Enable / Ẩn', '–', 'Hiển thị khi có quyền Thêm',
     'Nằm ở góc phải trên bảng.'),
    ('Nút Xuất Excel', 'Button', 'Enable', '–', 'Hiển thị',
     'Mở cửa sổ Chọn trường xuất file; không đòi quyền.'),
    ('Nút Cấu hình cột hiển thị', 'Icon Button', 'Enable', '–', 'Hiển thị',
     'Mở cửa sổ Tuỳ chỉnh cột.'),
    ('Phân trang', 'Pagination', 'Enable', '–', 'Trang 1',
     'Có nút về đầu / lùi / số trang / tiến / về cuối.'),
    ('Ô Số dòng/trang', 'Dropdown', 'Enable', '5 / 10 / 20 / 50 / 100', '10',
     'Đổi số dòng thì danh sách quay về trang 1.'),
    ('Ô “Hiển thị a–b / N”', 'Label', 'Read-only', '–', 'Theo kết quả',
     'N là tổng số gói khớp bộ lọc.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', 'Ẩn',
     'Nội dung “Không có dữ liệu phù hợp bộ lọc.”.'),
    ('Vòng quay chờ', 'Loading', 'Hiển thị', '–', 'Ẩn', 'Hiện trong lúc nạp danh sách.'),
], required=False)

sub('2.1.4 Danh sách event và xử lý event')
d.event_table([
    ('Mở màn hình', 'System',
     'Before:\n– Kiểm tra phiên đăng nhập; chưa đăng nhập → chuyển về màn đăng nhập.\n'
     'During:\n– Nạp cấu hình cột và khôi phục bộ lọc đã lưu nếu chưa quá 10 phút.\n'
     '– Tính giá bán theo từng cấp và cờ cho phép xoá của từng dòng.\n'
     'After:\n– Trả trang 1 với 10 dòng/trang và tổng số bản ghi; hiển thị bảng.'),
    ('Đổi trang', 'Click', 'Tải lại danh sách theo trang mới, giữ nguyên bộ lọc và thứ tự sắp xếp.'),
    ('Đổi số dòng/trang', 'Change', 'Đặt lại về trang 1 và tải lại danh sách.'),
    ('Bấm tiêu đề cột có mũi tên', 'Click',
     'Đổi cột và chiều sắp xếp rồi tải lại danh sách. Cột được phép sắp xếp: Mã, Tên gói bảo '
     'dưỡng, Trạng thái, Ngày tạo, Ngày sửa; cột khác bấm vào không có tác dụng.'),
    ('Bấm mã gói', 'Click', 'Mở màn Chi tiết gói bảo dưỡng ở chế độ chỉ đọc.'),
    ('Rê chuột vào biểu tượng chữ i', 'Hover',
     'Hiện bảng nhỏ liệt kê từng cấp bảo dưỡng kèm giá bán tương ứng, số tiền có dấu phân cách '
     'hàng nghìn.'),
])

# ---------------------------------------------------------------- 2.2
d.h3('2.2 Tìm kiếm và lọc danh sách')

sub('2.2.1 Giới thiệu')
d.intro_table(
    ten='Tìm kiếm và lọc danh sách gói bảo dưỡng',
    mota='Thu hẹp danh sách theo từ khoá (tên hoặc mã gói), theo trạng thái và theo người tạo.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đang ở màn hình danh sách gói bảo dưỡng.',
    chinh='1. Người dùng nhập từ khoá vào ô tìm nhanh và/hoặc chọn các ô lọc.\n'
          '2. Hệ thống tự lọc sau khi người dùng ngừng gõ khoảng nửa giây; người dùng cũng có thể '
          'bấm Tìm kiếm hoặc nhấn Enter để lọc ngay.\n'
          '3. Hệ thống đặt lại về trang 1 và tải danh sách khớp toàn bộ điều kiện.\n'
          '4. Kết quả tìm theo từ khoá được xếp theo độ khớp: trùng khít trước, khớp đầu chuỗi '
          'tiếp theo, khớp giữa chuỗi cuối cùng.\n'
          '5. Hệ thống ghi nhớ bộ lọc trong 10 phút để dùng lại khi quay về màn hình.',
    phu='• Thay đổi ô Trạng thái hoặc Người tạo → hệ thống tải lại ngay, không cần bấm Tìm kiếm.\n'
        '• Gõ liên tiếp nhiều ký tự → chỉ gọi lấy dữ liệu MỘT lần sau khi ngừng gõ.\n'
        '• Bấm Làm mới → xoá toàn bộ điều kiện và tải lại danh sách đầy đủ.\n'
        '• Không có kết quả → hiển thị trạng thái rỗng.\n'
        '• Người dùng tự bấm sắp xếp theo cột → thứ tự theo độ khớp bị bỏ qua.',
    dacbiet='Từ khoá chỉ 1 ký tự thì hệ thống không xếp theo độ khớp mà trả về thứ tự mặc định. Ô '
            'tìm nhanh quét đồng thời Tên và Mã; một bản ghi khớp cả hai vẫn chỉ hiện một dòng.')

sub('2.2.2 Layout màn hình')
lay(shot=shot('03-tim-kiem.png'),
    shot_caption='Kết quả sau khi tìm nhanh theo từ khoá “GBDT”')

sub('2.2.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Khối Bộ lọc danh sách', 'Modal', 'Enable', '–', '–', 'Mở sẵn',
     'Hàng ngang gồm 3 ô lọc và 2 nút.'),
    ('Ô tìm nhanh', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Gợi ý “Tìm theo tên hoặc mã gói bảo dưỡng...”; khớp gần đúng đồng thời tên và mã; tự lọc sau '
     'khi ngừng gõ khoảng nửa giây.'),
    ('Ô Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Không', 'Trống',
     'Để trống là lấy tất cả; khớp chính xác.'),
    ('Ô Người tạo', 'Dropdown', 'Enable', 'Danh sách nhân viên đã tạo gói', 'Không', 'Trống',
     'Chọn một nhân viên để chỉ xem gói do người đó tạo.'),
    ('Nút Tìm kiếm', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Áp dụng bộ lọc ngay và quay về trang 1, không phải chờ hết thời gian trễ.'),
    ('Nút Làm mới', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Xoá toàn bộ điều kiện lọc, đặt lại trang 1 và tải lại danh sách.'),
])

sub('2.2.4 Danh sách event và xử lý event')
d.event_table([
    ('Nhập ô tìm nhanh', 'Keypress',
     'Chờ khoảng nửa giây sau ký tự cuối rồi mới lọc; gõ liên tiếp chỉ gọi dữ liệu một lần.'),
    ('Nhấn Enter trong ô tìm nhanh', 'Keypress', 'Đặt lại trang 1 và tải lại danh sách ngay.'),
    ('Bấm Tìm kiếm', 'Click', 'Đặt lại trang 1 và tải lại danh sách theo toàn bộ điều kiện.'),
    ('Đổi ô Trạng thái hoặc Người tạo', 'Change', 'Đặt lại trang 1 và tải lại danh sách ngay.'),
    ('Bấm Làm mới', 'Click',
     'After:\n– Đặt toàn bộ ô lọc về giá trị khởi tạo.\n– Tải lại danh sách đầy đủ từ trang 1.'),
    ('Rời màn hình rồi quay lại', 'System',
     'Khôi phục bộ lọc đã lưu nếu chưa quá 10 phút; ô lọc đã bị bỏ khỏi màn hình thì không khôi '
     'phục để tránh lọc ngầm.'),
])

# ---------------------------------------------------------------- 2.3
d.h3('2.3 Tuỳ chỉnh cột hiển thị')

sub('2.3.1 Biểu đồ Usecase')
d.uc_figure('FR-03', 'Tuỳ chỉnh cột hiển thị', 'view',
            [('include', 'Ghi nhớ cấu hình theo người dùng')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-03 Tuỳ chỉnh cột hiển thị')

sub('2.3.2 Giới thiệu')
d.intro_table(
    ten='Tuỳ chỉnh cột hiển thị',
    mota='Cho phép mỗi tài khoản tự chọn cột nào hiển thị trên bảng và thứ tự các cột.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đang ở màn hình danh sách gói bảo dưỡng.',
    chinh='1. Người dùng bấm nút Cấu hình cột hiển thị bên cạnh nút Xuất Excel.\n'
          '2. Hệ thống mở cửa sổ “Tuỳ chỉnh cột” liệt kê 10 cột với trạng thái bật/tắt hiện tại.\n'
          '3. Người dùng tích hoặc bỏ tích cột, hoặc kéo thả để đổi vị trí.\n'
          '4. Người dùng bấm Lưu; hệ thống lưu cấu hình theo tài khoản và vẽ lại bảng.',
    phu='• Bấm Đóng → đóng cửa sổ, bỏ mọi thay đổi chưa lưu.\n'
        '• Cột STT, Mã và Hành động bị khoá → không bỏ tích được.',
    dacbiet='Cấu hình lưu riêng theo từng tài khoản và từng màn hình; tải lại trang vẫn giữ nguyên '
            'và không ảnh hưởng tới người dùng khác.')

sub('2.3.3 Layout màn hình')
lay(modal='Tuỳ chỉnh cột', shot=shot('04-cau-hinh-cot.png'),
    shot_caption='Cửa sổ Tuỳ chỉnh cột — cột STT, Mã và Hành động bị khoá')

sub('2.3.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Cấu hình cột hiển thị', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Nằm cạnh nút Xuất Excel, chỉ có biểu tượng cột.'),
    ('Danh sách cột', 'Table/Grid', 'Enable', '10 dòng', '–', 'Theo cấu hình đã lưu',
     'Mỗi dòng gồm ô tích chọn, tên cột và tay nắm kéo thả.'),
    ('Ô tích chọn cột', 'Icon Button', 'Enable / Disable', '–', 'Không', 'Theo cấu hình đã lưu',
     'Cột khoá (STT, Mã, Hành động) hiển thị ổ khoá và không đổi được.'),
    ('Tay nắm kéo thả', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Kéo để đổi thứ tự cột trên bảng.'),
    ('Nút Lưu', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Lưu cấu hình theo tài khoản và đóng cửa sổ.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ, bỏ thay đổi chưa lưu.'),
])

sub('2.3.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Cấu hình cột hiển thị', 'Click', 'Mở cửa sổ Tuỳ chỉnh cột với cấu hình hiện tại.'),
    ('Tích hoặc bỏ tích một cột', 'Change',
     'Cập nhật trạng thái tạm trong cửa sổ, chưa áp dụng lên bảng.'),
    ('Kéo thả đổi vị trí', 'Change', 'Cập nhật thứ tự tạm trong cửa sổ.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Bỏ qua thay đổi trên cột bị khoá.\n'
     'After:\n– Lưu cấu hình theo tài khoản, vẽ lại bảng và đóng cửa sổ.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ và bỏ toàn bộ thay đổi chưa lưu.'),
])

# ---------------------------------------------------------------- 2.4
d.h3('2.4 Tạo mới gói bảo dưỡng')

sub('2.4.1 Biểu đồ Usecase')
d.uc_figure('FR-04', 'Tạo mới gói bảo dưỡng', 'crud',
            [('include', 'Kiểm tra trùng tên và mã'),
             ('include', 'Kiểm tra bắt buộc đính kèm file PDF'),
             ('extend', 'Tính giá vốn và giá bán theo cấp')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-04 Tạo mới gói bảo dưỡng')

sub('2.4.2 Giới thiệu')
d.intro_table(
    ten='Tạo mới gói bảo dưỡng',
    mota='Khai báo một gói bảo dưỡng mới gồm thông tin chung, ma trận nội dung kiểm tra theo cấp, '
         'hệ số giá bán theo công ty, hàng hoá áp dụng và file đính kèm.',
    tacnhan=ACTOR,
    dieukien='Người dùng có quyền Thêm danh mục gói bảo dưỡng.',
    chinh='1. Người dùng bấm nút Tạo mới; hệ thống mở TRANG RIÊNG “Thêm gói bảo dưỡng”.\n'
          '2. Người dùng nhập Tên gói, Mã gói và các thông tin chung; ô Công ty quản lý gói bảo '
          'dưỡng đã điền sẵn công ty của người đăng nhập.\n'
          '3. Người dùng bấm “+ Thêm danh mục kiểm tra bảo dưỡng” để thêm dòng nội dung kiểm tra, '
          'nhập Nội dung, chọn Đơn vị tính và Số lượng.\n'
          '4. Người dùng bấm dấu + ở tiêu đề bảng để thêm cột cấp bảo dưỡng và chọn cấp cho cột đó.\n'
          '5. Người dùng chọn ghi chú kiểm tra tại các ô giao nhau, nhập Định mức công và Hệ số '
          'công nghệ cho từng cấp; hệ thống tính ngay Giá vốn, Giá công thức, Giá bán cơ sở và giá '
          'bán theo từng công ty.\n'
          '6. Người dùng chọn hàng hoá áp dụng và đính kèm ít nhất một file PDF.\n'
          '7. Người dùng bấm Lưu; hệ thống kiểm tra dữ liệu, chuyển Mã gói thành chữ in hoa và ghi '
          'bản ghi mới ở trạng thái Hoạt động.\n'
          '8. Hệ thống hiển thị “Tạo gói bảo dưỡng thành công” và quay về danh sách.',
    phu='• Bỏ trống ô bắt buộc → hiện lỗi đỏ ngay dưới đúng ô, không rời trang.\n'
        '• Trùng Tên hoặc Mã (kể cả chỉ khác chữ hoa chữ thường) → báo “Đã tồn tại”.\n'
        '• Chọn trùng cấp bảo dưỡng ở hai cột → cảnh báo “Cấp bảo dưỡng này đã được chọn ở cột '
        'khác” và không nhận cấp trùng.\n'
        '• Chưa đính kèm file PDF nào → báo bắt buộc đính kèm ít nhất 1 file PDF.\n'
        '• Đính kèm file không phải PDF → báo chỉ nhận file PDF.\n'
        '• Bấm Quay lại khi đang nhập dở → hỏi lại “Thông tin chưa lưu”.',
    dacbiet='Màn Thêm mới KHÔNG có ô Trạng thái: gói mới luôn ở trạng thái Hoạt động. Các ô hệ số '
            'và phần trăm nhận DẤU PHẨY làm dấu thập phân, ví dụ 12,5 nghĩa là mười hai phẩy năm.')

sub('2.4.3 Layout màn hình')
lay(shot=shot('06-tao-moi-thong-tin-chung.png'),
    shot_caption='Trang Thêm gói bảo dưỡng — khối Thông tin chung và bảng nội dung kiểm tra',
    url='https://<host-hrm>/customer-care/services/create',
    url_note='Trang Thêm mới là một trang riêng, không phải cửa sổ bật lên.')
d.figure(shot('07-tao-moi-hang-hoa-file.png'),
         'Khối Giá vốn theo công ty, Áp dụng cho hàng hóa và File đính kèm (PDF)', width_in=6.2)
d.figure(shot('14-chon-hang-hoa.png'),
         'Cửa sổ Chọn hàng hóa áp dụng', width_in=6.2)
d.figure(shot('08-loi-validate.png'),
         'Bấm Lưu khi chưa nhập gì — lỗi đỏ hiện ngay tại ô bắt buộc', width_in=6.2)

sub('2.4.4 Mô tả chi tiết giao diện — khối Thông tin chung')
d.ui_table([
    ('Tiêu đề trang', 'Label', 'Hiển thị', '–', '–', '“Thêm gói bảo dưỡng”',
     'Đổi thành “Sửa gói bảo dưỡng” / “Chi tiết gói bảo dưỡng” / “Sao chép gói bảo dưỡng” tuỳ chế '
     'độ mở.'),
    ('Tên gói bảo dưỡng', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Duy nhất toàn hệ thống; trống báo “Bắt buộc phải nhập”, trùng báo “Đã tồn tại”. Cắt khoảng '
     'trắng đầu cuối trước khi lưu.'),
    ('Mã gói bảo dưỡng', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Duy nhất toàn hệ thống; hệ thống chuyển thành CHỮ IN HOA trước khi kiểm tra trùng và lưu.'),
    ('Công ty quản lý gói bảo dưỡng', 'Dropdown', 'Enable', 'Danh sách công ty', 'Có',
     'Công ty của người đăng nhập',
     'Quyết định đơn giá công dùng để tính giá vốn của mọi cấp trong gói.'),
    ('Định mức đàm phán giá (%)', 'Textbox', 'Enable', 'Số, 0–99', 'Không', 'Trống',
     'Vượt trần báo “Tối đa 99”; nhận dấu phẩy làm dấu thập phân.'),
    ('VAT (%)', 'Textbox', 'Enable', 'Số, 0–100', 'Không', 'Trống',
     'Vượt trần báo “Tối đa 100”, số âm báo không được nhỏ hơn 0.'),
    ('Hệ số giá bán gói bảo dưỡng', 'Textbox', 'Enable', 'Số, 1–100', 'Không', 'Trống',
     'Nhỏ hơn 1 báo “Không được nhỏ hơn 1”; dùng để tính Giá công thức.'),
    ('Ghi chú', 'Textarea', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Nội dung này được in ở cuối phiếu Danh mục kiểm tra bảo dưỡng định kỳ.'),
    ('Trạng thái', 'Dropdown', 'Ẩn', 'Hoạt động / Khóa', 'Không', 'Hoạt động',
     'Chỉ xuất hiện ở màn Sửa; màn Thêm mới không có ô này.'),
])

sub('2.4.5 Mô tả chi tiết giao diện — bảng Danh mục kiểm tra bảo dưỡng định kỳ')
d.ui_table([
    ('Nút “+ Thêm danh mục kiểm tra bảo dưỡng”', 'Button', 'Enable', '–', 'Không', 'Hiển thị',
     'Thêm một DÒNG nội dung kiểm tra vào bảng.'),
    ('Nút thêm cột cấp', 'Icon Button', 'Enable', '–', 'Không', 'Hiển thị',
     'Dấu + ở góc phải tiêu đề bảng; thêm một CỘT cấp bảo dưỡng.'),
    ('Ô Nội dung kiểm tra', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Mỗi dòng một hạng mục kiểm tra.'),
    ('Ô Đơn vị tính', 'Dropdown', 'Enable', 'Danh mục đơn vị tính đang dùng', 'Có', 'Trống',
     'Đơn vị đang ngừng sử dụng không xuất hiện trong danh sách chọn.'),
    ('Ô Số lượng', 'Textbox', 'Enable', 'Số', 'Có', 'Trống', 'Nhập chữ thì báo “Phải là số”.'),
    ('Ô chọn Cấp bảo dưỡng của cột', 'Dropdown', 'Enable', 'Danh mục Cấp dịch vụ', 'Có', 'Trống',
     'Mỗi cấp chỉ được chọn MỘT lần trên toàn bảng; chọn trùng báo “Cấp bảo dưỡng này đã được chọn '
     'ở cột khác”.'),
    ('Ô giao nhau — Ghi chú kiểm tra', 'Dropdown', 'Enable', 'Danh mục ghi chú kiểm tra', 'Có',
     'Trống', 'Chọn được nhiều ghi chú; ô hiển thị ký hiệu đã chọn, ví dụ DK, VS.'),
    ('Dòng Định mức công', 'Textbox', 'Enable', 'Số', 'Có', 'Trống',
     'Bỏ trống báo “Bắt buộc phải nhập”; là thừa số của công thức tính giá vốn.'),
    ('Dòng Hệ số công nghệ', 'Textbox', 'Enable', 'Số', 'Không', 'Trống',
     'Nhân thêm vào giá vốn của cấp đó.'),
    ('Dòng Giá vốn', 'Label', 'Read-only', 'Số tiền', 'Không', 'Tự tính',
     'Đơn giá công của công ty quản lý × Định mức công × Hệ số công nghệ, làm tròn xuống.'),
    ('Dòng Giá công thức', 'Label', 'Read-only', 'Số tiền', 'Không', 'Tự tính',
     'Giá vốn × Hệ số giá bán gói bảo dưỡng.'),
    ('Dòng Giá bán cơ sở', 'Textbox', 'Enable', 'Số tiền', 'Không', 'Bằng Giá công thức',
     'Sửa tay được; là căn cứ tính giá bán của từng công ty.'),
    ('Dòng Gợi ý hàng hoá', 'Dropdown', 'Enable', 'Danh mục hàng hoá', 'Không', 'Trống',
     'Hàng hoá gợi ý dùng kèm cho cấp bảo dưỡng đó.'),
    ('Nút xoá dòng', 'Icon Button', 'Enable', '–', 'Không', 'Hiển thị',
     'Xoá một nội dung kiểm tra; STT các dòng còn lại được đánh lại liên tục.'),
    ('Nút xoá cột', 'Icon Button', 'Enable', '–', 'Không', 'Hiển thị',
     'Xoá cột cấp cùng toàn bộ thông số của cấp đó.'),
])

sub('2.4.6 Mô tả chi tiết giao diện — Giá vốn theo công ty, Hàng hoá và File đính kèm')
d.ui_table([
    ('Bảng Giá vốn theo công ty', 'Table/Grid', 'Enable', '–', 'Không', 'Liệt kê sẵn các công ty',
     'Mỗi công ty một dòng, hiển thị đơn giá công và ô nhập hệ số giá bán riêng.'),
    ('Ô Hệ số giá bán của công ty', 'Textbox', 'Enable', 'Số, tối đa 99.999.999,99', 'Không', '1',
     'Để trống thì hệ thống hiểu là 1, không phải 0; giá bán công ty = Giá bán cơ sở × hệ số này.'),
    ('Nút “Chọn hàng hóa”', 'Button', 'Enable', '–', 'Không', 'Hiển thị',
     'Mở cửa sổ Chọn hàng hóa áp dụng, có bộ lọc và ô tìm kiếm, tích chọn nhiều hàng hoá.'),
    ('Nút “Chọn nhóm hàng”', 'Button', 'Enable', '–', 'Không', 'Hiển thị',
     'Chọn một nhóm hàng để thêm toàn bộ hàng hoá của nhóm vào bảng.'),
    ('Bảng Áp dụng cho hàng hóa', 'Table/Grid', 'Enable', '–', 'Không', 'Trống',
     'Hiển thị hình ảnh, tên hàng, mã hàng và nút bỏ khỏi danh sách.'),
    ('Khối File đính kèm (PDF)', 'File', 'Enable', 'Chỉ nhận file PDF', 'Có', 'Trống',
     'Bắt buộc ít nhất 1 file; thiếu file thì khối viền đỏ và không lưu được. Thêm được nhiều '
     'file.'),
    ('Nút Lưu', 'Button', 'Enable / Disable', '–', '–', 'Enable',
     'Bị khóa trong lúc đang lưu để không tạo trùng bản ghi.'),
    ('Nút Quay lại', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Về danh sách; hỏi lại nếu còn thay đổi chưa lưu.'),
    ('Thông báo lỗi inline', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     'Chữ đỏ hiện ngay dưới ô tương ứng; ô lỗi có viền đỏ.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     '“Tạo gói bảo dưỡng thành công” hoặc thông báo lỗi trả về từ hệ thống.'),
])

sub('2.4.7 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Tạo mới', 'Click',
     'Before:\n– Kiểm tra quyền Thêm danh mục gói bảo dưỡng; không có quyền thì nút không hiển '
     'thị.\n'
     'After:\n– Mở trang Thêm gói bảo dưỡng, điền sẵn Công ty quản lý theo người đăng nhập và liệt '
     'kê sẵn các công ty ở khối Giá vốn theo công ty.'),
    ('Chọn cấp bảo dưỡng cho một cột', 'Change',
     'Kiểm tra cấp đã dùng ở cột khác chưa; đã dùng thì cảnh báo và không nhận.'),
    ('Nhập Định mức công / Hệ số công nghệ / Hệ số giá bán', 'Change',
     'Tính lại ngay Giá vốn, Giá công thức, Giá bán cơ sở của cấp đó và giá bán theo từng công ty.'),
    ('Sửa tay Giá bán cơ sở', 'Change', 'Tính lại giá bán của từng công ty theo hệ số của công ty.'),
    ('Chọn hàng hoá hoặc nhóm hàng', 'Click',
     'Thêm hàng hoá vào bảng Áp dụng cho hàng hóa, bỏ qua hàng hoá đã có trong bảng.'),
    ('Thêm file đính kèm', 'Change',
     'Chỉ nhận file PDF; file khác định dạng bị từ chối kèm thông báo.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Kiểm tra các ô bắt buộc ở cả bốn khối; thiếu thì hiện lỗi đỏ và dừng.\n'
     'During:\n– Chuyển Mã gói thành chữ in hoa, cắt khoảng trắng đầu cuối của Tên và Mã.\n'
     '– Kiểm tra trùng Tên và trùng Mã trên toàn hệ thống.\n'
     '– Ghi gói, ma trận nội dung kiểm tra, thông số theo cấp, hệ số theo công ty, hàng hoá áp '
     'dụng và file đính kèm trong cùng một lần ghi.\n'
     'After:\n– Báo “Tạo gói bảo dưỡng thành công”, quay về danh sách và ghi mốc Tạo mới vào lịch '
     'sử.'),
    ('Bấm Quay lại khi còn thay đổi chưa lưu', 'Click',
     'Hỏi lại “Thông tin chưa lưu”; chọn ở lại thì giữ nguyên dữ liệu đang nhập.'),
])

# ---------------------------------------------------------------- 2.5
d.h3('2.5 Sửa gói bảo dưỡng')

sub('2.5.1 Biểu đồ Usecase')
d.uc_figure('FR-05', 'Sửa gói bảo dưỡng', 'crud',
            [('include', 'Chặn khi gói đang Khóa'),
             ('include', 'Chặn bỏ cấp đã phát sinh báo giá')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-05 Sửa gói bảo dưỡng')

sub('2.5.2 Giới thiệu')
d.intro_table(
    ten='Sửa gói bảo dưỡng',
    mota='Cập nhật thông tin của một gói bảo dưỡng đang ở trạng thái Hoạt động, kể cả đổi trạng '
         'thái sang Khóa.',
    tacnhan=ACTOR,
    dieukien='Người dùng có quyền Sửa danh mục gói bảo dưỡng và gói đang ở trạng thái Hoạt động.',
    chinh='1. Người dùng bấm nút Sửa ở dòng tương ứng, hoặc bấm Sửa ở chân trang chi tiết.\n'
          '2. Hệ thống mở trang “Sửa gói bảo dưỡng” và nạp sẵn toàn bộ dữ liệu của gói.\n'
          '3. Người dùng chỉnh sửa các khối cần thay đổi.\n'
          '4. Người dùng bấm Lưu; hệ thống kiểm tra dữ liệu và ghi thay đổi.\n'
          '5. Hệ thống báo “Cập nhật gói bảo dưỡng thành công”, quay về danh sách và cập nhật cột '
          'Người sửa, Ngày sửa.',
    phu='• Gói đang Khóa → nút Sửa bị ẩn; gõ thẳng đường dẫn sửa hoặc gọi thẳng chức năng Sửa đều '
        'bị hệ thống chặn kèm yêu cầu mở khoá trước.\n'
        '• Đổi Tên hoặc Mã trùng gói khác → báo “Đã tồn tại”; giữ nguyên tên và mã của chính nó thì '
        'không báo trùng.\n'
        '• Bỏ một cột cấp đã phát sinh báo giá dịch vụ → hệ thống chặn khi lưu, dữ liệu gói giữ '
        'nguyên.\n'
        '• Đổi ô Trạng thái sang Khóa rồi Lưu → gói chuyển sang Khóa và mất nút Sửa ngoài danh '
        'sách.\n'
        '• Bấm Quay lại khi còn thay đổi chưa lưu → hỏi lại “Thông tin chưa lưu”.',
    dacbiet='Khác màn Thêm mới, màn Sửa CÓ ô Trạng thái với hai lựa chọn Hoạt động / Khóa. Lịch sử '
            'thay đổi chỉ ghi các trường của khối Thông tin chung; chỉnh sửa ma trận nội dung kiểm '
            'tra không sinh mốc lịch sử.')

sub('2.5.3 Layout màn hình')
lay(shot=shot('13-sua.png'),
    shot_caption='Trang Sửa gói bảo dưỡng — dữ liệu hiện tại đã được nạp sẵn',
    url='https://<host-hrm>/customer-care/services/<mã bản ghi>/edit')

sub('2.5.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề trang', 'Label', 'Hiển thị', '–', '–', '“Sửa gói bảo dưỡng”', 'Kèm mã gói đang sửa.'),
    ('Các ô của khối Thông tin chung', 'Textbox', 'Enable', 'Như màn Thêm mới', 'Như màn Thêm mới',
     'Giá trị hiện tại', 'Ràng buộc nhập liệu giống hệt màn Thêm mới.'),
    ('Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Có', 'Giá trị hiện tại',
     'Chỉ có ở màn Sửa; đổi sang Khóa là một trong hai cách khoá gói.'),
    ('Bảng Danh mục kiểm tra bảo dưỡng định kỳ', 'Table/Grid', 'Enable', '–', 'Có',
     'Ma trận hiện tại',
     'Thêm / bớt dòng và cột được; bỏ cột cấp đã phát sinh báo giá thì bị chặn khi lưu.'),
    ('Bảng Giá vốn theo công ty', 'Table/Grid', 'Enable', '–', 'Không', 'Hệ số hiện tại',
     'Sửa hệ số của từng công ty; giá bán tính lại ngay.'),
    ('Bảng Áp dụng cho hàng hóa', 'Table/Grid', 'Enable', '–', 'Không', 'Danh sách hiện tại',
     'Thêm hoặc bỏ hàng hoá; gói đã gắn hàng hoá thì không xoá được nữa.'),
    ('Khối File đính kèm (PDF)', 'File', 'Enable', 'Chỉ nhận file PDF', 'Có', 'File hiện tại',
     'Bỏ hết file cũ mà không thêm file mới thì không lưu được.'),
    ('Nút Lưu', 'Button', 'Enable / Disable', '–', '–', 'Enable', 'Bị khóa trong lúc đang lưu.'),
    ('Nút Quay lại', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Về danh sách; hỏi lại nếu còn thay đổi chưa lưu.'),
])

sub('2.5.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Sửa', 'Click',
     'Before:\n– Kiểm tra quyền Sửa và trạng thái gói; gói đang Khóa thì nút không hiển thị.\n'
     'After:\n– Mở trang Sửa và nạp toàn bộ dữ liệu hiện tại của gói.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Kiểm tra quyền Sửa và kiểm tra gói có đang bị khoá không; đang khoá thì từ chối '
     'kèm yêu cầu mở khoá.\n'
     'During:\n– Kiểm tra trùng Tên và Mã, bỏ qua chính bản ghi đang sửa.\n'
     '– Kiểm tra các cấp bị bỏ đã phát sinh báo giá dịch vụ chưa; đã phát sinh thì huỷ toàn bộ '
     'thay đổi và báo lỗi.\n'
     'After:\n– Ghi thay đổi, báo “Cập nhật gói bảo dưỡng thành công” và ghi mốc lịch sử nếu có '
     'trường theo dõi thay đổi.'),
    ('Lưu mà không sửa ô nào', 'Click',
     'Vẫn báo cập nhật thành công nhưng KHÔNG phát sinh mốc lịch sử mới.'),
    ('Lưu thất bại giữa chừng', 'System',
     'Huỷ toàn bộ thay đổi của lần lưu đó; ma trận nội dung kiểm tra của gói phải nguyên vẹn như '
     'trước khi sửa.'),
])

# ---------------------------------------------------------------- 2.6
d.h3('2.6 Xem chi tiết gói bảo dưỡng')

sub('2.6.1 Giới thiệu')
d.intro_table(
    ten='Xem chi tiết gói bảo dưỡng',
    mota='Xem toàn bộ dữ liệu của một gói ở chế độ chỉ đọc, kèm các nút chuyển sang Sửa, In và '
         'Nhân bản.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đã đăng nhập; không cần quyền riêng.',
    chinh='1. Người dùng bấm vào mã gói ở danh sách.\n'
          '2. Hệ thống mở trang “Chi tiết gói bảo dưỡng: <mã gói>”.\n'
          '3. Hệ thống hiển thị đủ năm khối dữ liệu ở chế độ chỉ đọc.\n'
          '4. Người dùng bấm Quay lại để về danh sách.',
    phu='• Gói đang Khóa → chân trang không có nút Sửa, vẫn còn In, Nhân bản và Quay lại.\n'
        '• Tài khoản không có quyền Sửa → chân trang không có nút Sửa.\n'
        '• Tài khoản không có quyền Thêm → chân trang không có nút Nhân bản.\n'
        '• Bấm vào tên file đính kèm → file PDF mở ở tab mới.',
    dacbiet='Trang chi tiết có đường dẫn riêng nên mở được ở tab mới bằng chuột phải; mọi ô đều '
            'khoá nhập, kể cả các ô số của bảng ma trận.')

sub('2.6.2 Layout màn hình')
lay(shot=shot('09-chi-tiet.png'),
    shot_caption='Trang Chi tiết gói bảo dưỡng ở chế độ chỉ đọc',
    url='https://<host-hrm>/customer-care/services/<mã bản ghi>')

sub('2.6.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề trang', 'Label', 'Read-only', '–', '“Chi tiết gói bảo dưỡng: <mã gói>”',
     'Hiển thị mã của gói đang xem.'),
    ('Khối Thông tin chung', 'Label', 'Read-only', '–', 'Theo dữ liệu',
     'Đầy đủ tên, mã, công ty quản lý, định mức đàm phán giá, VAT, hệ số giá bán, ghi chú, trạng '
     'thái.'),
    ('Bảng Danh mục kiểm tra bảo dưỡng định kỳ', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Kèm các dòng Định mức công, Hệ số công nghệ, Giá vốn, Giá công thức, Giá bán cơ sở.'),
    ('Bảng Giá bán theo công ty', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Hệ số và giá bán tương ứng của từng công ty.'),
    ('Bảng Áp dụng cho hàng hóa', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Hình ảnh, tên hàng và mã hàng.'),
    ('Khối File đính kèm', 'File', 'Read-only', '–', 'Theo dữ liệu',
     'Bấm vào tên file để mở PDF ở tab mới.'),
    ('Nút Sửa', 'Button', 'Enable / Ẩn', '–', 'Hiển thị khi có quyền và gói đang Hoạt động',
     'Chuyển sang trang Sửa của đúng gói đang xem.'),
    ('Nút In', 'Button', 'Enable', '–', 'Hiển thị', 'Mở màn In phiếu ở tab mới.'),
    ('Nút Nhân bản', 'Button', 'Enable / Ẩn', '–', 'Hiển thị khi có quyền Thêm',
     'Mở trang Sao chép gói bảo dưỡng.'),
    ('Nút Quay lại', 'Button', 'Enable', '–', 'Hiển thị', 'Về màn danh sách.'),
], required=False)

sub('2.6.4 Danh sách event và xử lý event')
d.event_table([
    ('Bấm mã gói ở danh sách', 'Click', 'Mở trang chi tiết của gói đó ở chế độ chỉ đọc.'),
    ('Bấm nút Sửa ở chân trang', 'Click', 'Chuyển sang trang Sửa của đúng gói đang xem.'),
    ('Bấm tên file đính kèm', 'Click', 'Mở file PDF ở tab mới.'),
    ('Bấm Quay lại', 'Click', 'Về màn danh sách, giữ nguyên bộ lọc đang có.'),
])

# ---------------------------------------------------------------- 2.7
d.h3('2.7 Nhân bản gói bảo dưỡng')

sub('2.7.1 Biểu đồ Usecase')
d.uc_figure('FR-07', 'Nhân bản gói bảo dưỡng', 'crud',
            [('include', 'Sao chép toàn bộ dữ liệu gói nguồn'),
             ('include', 'Kiểm tra trùng tên và mã')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-07 Nhân bản gói bảo dưỡng')

sub('2.7.2 Giới thiệu')
d.intro_table(
    ten='Nhân bản gói bảo dưỡng',
    mota='Tạo một gói mới bằng cách sao chép toàn bộ dữ liệu của gói đang có, rồi sửa lại tên và mã '
         'trước khi lưu.',
    tacnhan=ACTOR,
    dieukien='Người dùng có quyền Thêm danh mục gói bảo dưỡng.',
    chinh='1. Người dùng mở menu ba chấm của một dòng và bấm Nhân bản, hoặc bấm Nhân bản ở chân '
          'trang chi tiết.\n'
          '2. Hệ thống mở trang “Sao chép gói bảo dưỡng” và điền sẵn toàn bộ dữ liệu của gói '
          'nguồn, kể cả ma trận nội dung kiểm tra, hệ số theo công ty, hàng hoá áp dụng và file '
          'đính kèm.\n'
          '3. Người dùng sửa Tên gói và Mã gói cho khác gói nguồn.\n'
          '4. Người dùng bấm Lưu; hệ thống ghi một gói mới ở trạng thái Hoạt động.',
    phu='• Không sửa Tên và Mã rồi bấm Lưu → cả hai ô báo “Đã tồn tại”, không tạo bản ghi.\n'
        '• Nhân bản từ gói đang Khóa vẫn thực hiện được; gói mới ở trạng thái Hoạt động, không kế '
        'thừa trạng thái Khóa.\n'
        '• Gói nguồn không bị thay đổi trong mọi trường hợp.',
    dacbiet='Vì bản nhân bản mang theo cả danh sách hàng hoá áp dụng nên gói mới cũng bị tính là '
            '“đã được sử dụng” và sẽ không có nút Xóa — cần lưu ý khi dọn dữ liệu kiểm thử.')

sub('2.7.3 Layout màn hình')
lay(url='https://<host-hrm>/customer-care/services/create?cloneId=<mã bản ghi nguồn>',
    url_note='Trang Sao chép dùng chung khuôn với trang Thêm mới, chỉ khác tiêu đề và dữ liệu '
             'điền sẵn.')

sub('2.7.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Nhân bản', 'Button', 'Enable / Ẩn', '–', 'Hiển thị khi có quyền Thêm',
     'Nằm trong menu ba chấm ngoài danh sách và ở chân trang chi tiết.'),
    ('Tiêu đề trang', 'Label', 'Read-only', '–', '“Sao chép gói bảo dưỡng”',
     'Phân biệt với trang Thêm gói bảo dưỡng.'),
    ('Toàn bộ ô nhập', 'Textbox', 'Enable', 'Như màn Thêm mới', 'Dữ liệu gói nguồn',
     'Sửa được tất cả trước khi lưu.'),
    ('Ô Tên và Mã gói', 'Textbox', 'Enable', '0–255 ký tự', 'Giữ nguyên của gói nguồn',
     'BẮT BUỘC sửa vì hai trường này là duy nhất toàn hệ thống.'),
], required=False)

sub('2.7.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm Nhân bản', 'Click',
     'Before:\n– Kiểm tra quyền Thêm; không có quyền thì nút không hiển thị.\n'
     'After:\n– Mở trang Sao chép và điền sẵn toàn bộ dữ liệu của gói nguồn.'),
    ('Bấm Lưu', 'Click',
     'Kiểm tra trùng Tên và Mã như khi tạo mới; hợp lệ thì ghi gói mới ở trạng thái Hoạt động và '
     'ghi mốc Tạo mới vào lịch sử của gói mới.'),
])

# ---------------------------------------------------------------- 2.8
d.h3('2.8 Xóa gói bảo dưỡng và Mở khóa gói')

sub('2.8.1 Biểu đồ Usecase')
d.uc_figure('FR-08', 'Xóa gói / Mở khóa gói', 'action',
            [('include', 'Kiểm tra gói đã được sử dụng'),
             ('extend', 'Chuyển sang trạng thái Khóa thay vì xoá')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-08 Xóa gói bảo dưỡng và Mở khóa gói')

sub('2.8.2 Giới thiệu')
d.intro_table(
    ten='Xóa gói bảo dưỡng và Mở khóa gói',
    mota='Xoá hẳn một gói chưa được sử dụng; với gói đã được sử dụng thì hệ thống chuyển sang trạng '
         'thái Khóa. Gói đang Khóa có thể được mở khoá để dùng lại.',
    tacnhan=ACTOR,
    dieukien='Xóa cần quyền Xóa danh mục gói bảo dưỡng; Mở khóa cần quyền Sửa danh mục gói bảo '
             'dưỡng.',
    chinh='1. Người dùng bấm nút Xóa ở dòng tương ứng.\n'
          '2. Hệ thống mở hộp thoại “Xác nhận xóa” nêu rõ tên gói.\n'
          '3. Người dùng bấm Xóa để xác nhận.\n'
          '4. Hệ thống kiểm tra gói đã được sử dụng chưa; chưa dùng thì xoá gói cùng toàn bộ nội '
          'dung kiểm tra và cấp của gói.\n'
          '5. Hệ thống báo “Xóa gói bảo dưỡng thành công”, dòng biến mất và tổng số bản ghi giảm.',
    phu='• Gói đã gắn hàng hoá hoặc đã dùng ở báo giá dịch vụ → nút Xóa bị ẩn ngoài danh sách.\n'
        '• Gọi thẳng chức năng Xóa với gói đã được sử dụng → hệ thống KHÔNG xoá mà chuyển sang '
        'trạng thái Khóa và báo “Gói bảo dưỡng đang được sử dụng nên đã được chuyển sang trạng '
        'thái Khóa”.\n'
        '• Bấm Hủy ở hộp xác nhận → đóng hộp thoại, không thay đổi gì.\n'
        '• Gọi Mở khóa cho gói đang Hoạt động → báo “Gói bảo dưỡng đang hoạt động, không cần mở '
        'khóa.”.\n'
        '• Xoá dòng cuối cùng của trang cuối → danh sách tự lùi về trang trước.',
    dacbiet='Có hai đường đưa gói về trạng thái Khóa cho cùng một kết quả: đổi ô Trạng thái trong '
            'màn Sửa, hoặc thao tác Xóa trên gói đã được sử dụng. Bản ghi đã khóa vẫn nằm trong '
            'danh sách với nhãn Khóa và vẫn xem, in, nhân bản được.')

sub('2.8.3 Layout màn hình')
lay(modal='Xác nhận xóa', shot=shot('11-xac-nhan-xoa.png'),
    shot_caption='Hộp thoại Xác nhận xóa gói bảo dưỡng')

sub('2.8.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Xóa', 'Icon Button', 'Enable / Ẩn', 'Hiển thị khi có quyền Xóa, gói đang Hoạt động và '
     'chưa được sử dụng', 'Biểu tượng thùng rác màu đỏ ở cột Hành động.'),
    ('Hộp thoại Xác nhận xóa', 'Modal', 'Hiển thị', 'Ẩn',
     'Tiêu đề “Xác nhận xóa”; nội dung “Bạn có chắc chắn muốn xóa gói bảo dưỡng ‘<tên gói>’? Hành '
     'động này không thể hoàn tác.”.'),
    ('Nút Xóa trong hộp thoại', 'Button', 'Enable', 'Hiển thị', 'Nền đỏ, thực hiện xoá.'),
    ('Nút Hủy trong hộp thoại', 'Button', 'Enable', 'Hiển thị', 'Đóng hộp thoại, không xoá.'),
    ('Nút Mở khóa', 'Icon Button', 'Enable / Ẩn',
     'Hiển thị khi có quyền Sửa và gói đang Khóa', 'Biểu tượng ổ khoá mở ở cột Hành động.'),
    ('Hộp thoại Xác nhận mở khóa', 'Modal', 'Hiển thị', 'Ẩn',
     'Tiêu đề “Xác nhận mở khóa”, nội dung nêu tên gói; hai nút Mở khóa và Hủy.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', 'Ẩn',
     '“Xóa gói bảo dưỡng thành công”, “Mở khóa gói bảo dưỡng thành công” hoặc thông báo gói đã '
     'chuyển sang trạng thái Khóa.'),
], required=False, scope=False)

sub('2.8.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Xóa', 'Click', 'Mở hộp thoại Xác nhận xóa kèm tên gói.'),
    ('Bấm Xóa trong hộp thoại', 'Click',
     'Before:\n– Kiểm tra quyền Xóa và kiểm tra gói có đang bị khoá không.\n'
     'During:\n– Kiểm tra gói đã gắn hàng hoá hoặc đã dùng ở báo giá dịch vụ chưa.\n'
     '– Chưa dùng → xoá gói cùng nội dung kiểm tra và các cấp của gói.\n'
     '– Đã dùng → chuyển gói sang trạng thái Khóa, không xoá dữ liệu.\n'
     'After:\n– Báo kết quả tương ứng, tải lại danh sách và ghi mốc lịch sử nếu gói chuyển sang '
     'Khóa.'),
    ('Bấm Hủy trong hộp thoại', 'Click', 'Đóng hộp thoại, không thay đổi dữ liệu.'),
    ('Bấm nút Mở khóa', 'Click', 'Mở hộp thoại Xác nhận mở khóa kèm tên gói.'),
    ('Bấm Mở khóa trong hộp thoại', 'Click',
     'Before:\n– Kiểm tra quyền Sửa và trạng thái hiện tại của gói.\n'
     'During:\n– Gói đang Hoạt động → báo không cần mở khóa và dừng.\n'
     'After:\n– Chuyển gói về Hoạt động, hiện lại nút Sửa và ghi mốc Mở khóa vào lịch sử.'),
])

# ---------------------------------------------------------------- 2.9
d.h3('2.9 Xuất Excel danh sách gói bảo dưỡng')

sub('2.9.1 Biểu đồ Usecase')
d.uc_figure('FR-09', 'Xuất Excel', 'export',
            [('include', 'Chọn trường và thứ tự cột xuất')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-09 Xuất Excel danh sách gói bảo dưỡng')

sub('2.9.2 Giới thiệu')
d.intro_table(
    ten='Xuất Excel danh sách gói bảo dưỡng',
    mota='Tải về file Excel danh sách gói bảo dưỡng với các trường do người dùng chọn.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đã đăng nhập; không cần quyền riêng.',
    chinh='1. Người dùng bấm nút Xuất Excel.\n'
          '2. Hệ thống mở cửa sổ “Chọn trường xuất file” với 6 trường, mặc định chọn sẵn cả 6.\n'
          '3. Người dùng bỏ bớt hoặc chọn lại các trường theo thứ tự mong muốn.\n'
          '4. Người dùng bấm Xuất file; hệ thống tạo file và trình duyệt tải về.',
    phu='• Bỏ chọn hết trường → nút Xuất file không dùng được.\n'
        '• Bấm Chọn tất cả → quay lại đủ 6 trường.\n'
        '• Danh sách đang lọc rỗng → file vẫn ra đủ dữ liệu toàn danh mục.',
    dacbiet='⚠ File xuất ra KHÔNG áp bộ lọc đang có trên màn hình: dù đang lọc còn vài dòng, file '
            'vẫn chứa toàn bộ gói bảo dưỡng của danh mục. Thứ tự cột trong file chạy theo THỨ TỰ '
            'NGƯỜI DÙNG CHỌN trong cửa sổ, không theo thứ tự cột trên bảng.')

sub('2.9.3 Layout màn hình')
lay(modal='Chọn trường xuất file', shot=shot('05-chon-truong-xuat.png'),
    shot_caption='Cửa sổ Chọn trường xuất file — mặc định chọn sẵn cả 6 trường')

sub('2.9.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Xuất Excel', 'Button', 'Enable', '–', 'Hiển thị',
     'Nền xanh lá, nằm cạnh nút Tạo mới; không đòi quyền.'),
    ('Cửa sổ Chọn trường xuất file', 'Modal', 'Hiển thị', '–', 'Ẩn',
     'Liệt kê 6 trường: Mã, Tên gói bảo dưỡng, Công ty quản lý, Trạng thái, Người tạo, Ngày tạo.'),
    ('Ô tích chọn trường', 'Icon Button', 'Enable', '–', 'Chọn sẵn cả 6',
     'Thứ tự chọn quyết định thứ tự cột trong file.'),
    ('Dòng “Đang chọn n/6 trường”', 'Label', 'Read-only', '–', '“Đang chọn 6/6 trường”',
     'Cập nhật theo số trường đang chọn.'),
    ('Nút Chọn tất cả', 'Button', 'Enable', '–', 'Hiển thị', 'Chọn lại đủ 6 trường.'),
    ('Nút Bỏ chọn hết', 'Button', 'Enable', '–', 'Hiển thị', 'Bỏ chọn toàn bộ trường.'),
    ('Nút Xuất file', 'Button', 'Enable / Disable', '–', 'Enable',
     'Bị khoá khi không còn trường nào được chọn.'),
    ('Nút Đóng', 'Button', 'Enable', '–', 'Hiển thị', 'Đóng cửa sổ, không xuất file.'),
], required=False)

sub('2.9.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm Xuất Excel', 'Click', 'Mở cửa sổ Chọn trường xuất file với 6 trường chọn sẵn.'),
    ('Tích / bỏ tích một trường', 'Change',
     'Cập nhật dòng “Đang chọn n/6 trường” và ghi nhận thứ tự chọn.'),
    ('Bấm Xuất file', 'Click',
     'During:\n– Lấy TOÀN BỘ danh mục gói bảo dưỡng, không áp bộ lọc đang có trên màn hình.\n'
     '– Sắp cột theo đúng thứ tự người dùng đã chọn.\n'
     'After:\n– Trình duyệt tải về file Danh_sach_goi_bao_duong.xlsx; mã gói toàn số giữ nguyên số '
     '0 ở đầu.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ, không xuất file.'),
])

# ---------------------------------------------------------------- 2.10
d.h3('2.10 In phiếu Danh mục kiểm tra bảo dưỡng định kỳ')

sub('2.10.1 Biểu đồ Usecase')
d.uc_figure('FR-10', 'In phiếu kiểm tra bảo dưỡng', 'export',
            [('include', 'Lấy mẫu in Danh mục kiểm tra bảo dưỡng')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-10 In phiếu Danh mục kiểm tra bảo dưỡng định kỳ')

sub('2.10.2 Giới thiệu')
d.intro_table(
    ten='In phiếu Danh mục kiểm tra bảo dưỡng định kỳ',
    mota='Mở bản in ma trận nội dung kiểm tra của một gói để in ra giấy dùng ngoài hiện trường.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đã đăng nhập; không cần quyền riêng.',
    chinh='1. Người dùng mở menu ba chấm của một dòng và bấm In, hoặc bấm In ở chân trang chi tiết.\n'
          '2. Hệ thống mở TAB MỚI với màn in của gói đó.\n'
          '3. Hệ thống lấy mẫu in Danh mục kiểm tra bảo dưỡng và đổ dữ liệu gói vào mẫu.\n'
          '4. Người dùng bấm nút In để mở hộp thoại in của trình duyệt.',
    phu='• Gói chưa khai nội dung kiểm tra → phiếu vẫn mở được, bảng nội dung để trống.\n'
        '• Mẫu in bị thiếu trong hệ thống → báo không tìm thấy mẫu in, không để trang trắng.\n'
        '• Bản xem trước khi in không có menu bên trái và không có nút In.',
    dacbiet='Màn in mở ở tab mới nên danh sách ở tab cũ giữ nguyên bộ lọc và trang đang xem.')

sub('2.10.3 Layout màn hình')
lay(shot=shot('10-in-phieu.png'),
    shot_caption='Màn In phiếu Danh mục kiểm tra bảo dưỡng định kỳ',
    url='https://<host-hrm>/customer-care/services/<mã bản ghi>/print',
    url_note='Màn in luôn được mở ở tab mới.')

sub('2.10.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút In', 'Button', 'Enable', '–', 'Hiển thị',
     'Có ở menu ba chấm ngoài danh sách và ở chân trang chi tiết; không đòi quyền.'),
    ('Đầu phiếu', 'Label', 'Read-only', '–', 'Theo dữ liệu',
     'Logo và thông tin công ty.'),
    ('Tiêu đề phiếu', 'Label', 'Read-only', '–',
     '“DANH MỤC KIỂM TRA BẢO DƯỠNG ĐỊNH KỲ”',
     'Kèm dòng “TÊN DỊCH VỤ: <tên gói in hoa>”.'),
    ('Bảng nội dung kiểm tra', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Gồm STT, Nội dung kiểm tra bảo dưỡng, SL, các cột cấp bảo dưỡng, cột Kiểm tra (Có/Không) và '
     'Ghi chú.'),
    ('Phần Ghi chú cuối phiếu', 'Label', 'Read-only', '–', 'Theo dữ liệu',
     'Lấy từ ô Ghi chú của khối Thông tin chung.'),
    ('Bảng giải thích ký hiệu', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Diễn giải các ký hiệu ghi chú kiểm tra như KTBM, DK, CC, VS.'),
    ('Nút In trên màn in', 'Button', 'Enable', '–', 'Hiển thị',
     'Nằm ở góc trái phía trên; mở hộp thoại in của trình duyệt và tự ẩn khi in.'),
], required=False)

sub('2.10.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút In ngoài danh sách hoặc ở chi tiết', 'Click',
     'Mở tab mới tới màn in của đúng gói được chọn.'),
    ('Mở màn in', 'System',
     'Before:\n– Lấy mẫu in Danh mục kiểm tra bảo dưỡng; không có mẫu thì báo lỗi đọc hiểu được.\n'
     'After:\n– Đổ dữ liệu gói vào mẫu và hiển thị bản xem trước.'),
    ('Bấm nút In trên màn in', 'Click',
     'Mở hộp thoại in của trình duyệt; bản in ẩn menu bên trái và ẩn chính nút In.'),
])

# ---------------------------------------------------------------- 2.11
d.h3('2.11 Xem lịch sử thay đổi')

sub('2.11.1 Biểu đồ Usecase')
d.uc_figure('FR-11', 'Xem lịch sử thay đổi', 'view',
            [('include', 'Lọc theo loại hoạt động')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-11 Xem lịch sử thay đổi')

sub('2.11.2 Giới thiệu')
d.intro_table(
    ten='Xem lịch sử thay đổi của gói bảo dưỡng',
    mota='Xem nhật ký thao tác của một gói: ai tạo, ai sửa, ai khóa hoặc mở khóa và giá trị trước '
         '- sau của từng trường.',
    tacnhan='%s; Người dùng đã đăng nhập' % ACTOR,
    dieukien='Người dùng đã đăng nhập; không cần quyền riêng.',
    chinh='1. Người dùng mở menu ba chấm của một dòng và bấm Lịch sử.\n'
          '2. Hệ thống mở cửa sổ “Lịch sử thay đổi” kèm dòng phụ “Gói bảo dưỡng: <mã> - <tên>”.\n'
          '3. Hệ thống hiển thị các mốc theo thứ tự mới nhất trước.\n'
          '4. Người dùng bấm Bộ lọc để lọc theo loại hoạt động nếu cần.',
    phu='• Gói chưa phát sinh thao tác nào → hiển thị “Chưa có lịch sử thao tác nào.”.\n'
        '• Sửa mà không thay đổi ô nào → không phát sinh mốc mới.\n'
        '• Xoá gói đang được sử dụng → chỉ có mốc Khóa, không có mốc Xóa vì bản ghi không bị xoá.',
    dacbiet='Lịch sử chỉ theo dõi các trường của khối Thông tin chung: Mã, Tên gói bảo dưỡng, Định '
            'mức đàm phán giá, VAT, Ghi chú, Trạng thái. Thay đổi trong ma trận nội dung kiểm tra, '
            'hệ số theo công ty, hàng hoá áp dụng hay file đính kèm KHÔNG sinh mốc lịch sử.')

sub('2.11.3 Layout màn hình')
lay(modal='Lịch sử thay đổi', shot=shot('12-lich-su.png'),
    shot_caption='Cửa sổ Lịch sử thay đổi của một gói bảo dưỡng')

sub('2.11.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Lịch sử', 'Icon Button', 'Enable', '–', 'Hiển thị',
     'Nằm trong menu ba chấm của từng dòng; không đòi quyền riêng.'),
    ('Tiêu đề cửa sổ', 'Label', 'Read-only', '–', '“Lịch sử thay đổi”',
     'Dòng phụ ghi mã và tên gói đang xem.'),
    ('Danh sách mốc lịch sử', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Mỗi mốc gồm ngày giờ, nhãn hành động, người thực hiện kèm phòng ban và chi tiết thay đổi.'),
    ('Dòng chi tiết thay đổi', 'Label', 'Read-only', '–', 'Theo dữ liệu',
     'Dạng “Tên trường: giá trị cũ → giá trị mới”, nhãn trường bằng tiếng Việt.'),
    ('Nút Bộ lọc', 'Button', 'Enable', '–', 'Hiển thị',
     'Lọc theo 3 nhóm hoạt động cố định: Tạo mới, Thay đổi thông tin, Thay đổi trạng thái.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', 'Ẩn', 'Nội dung “Chưa có lịch sử thao tác nào.”.'),
    ('Nút Đóng', 'Button', 'Enable', '–', 'Hiển thị', 'Đóng cửa sổ.'),
], required=False)

sub('2.11.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Lịch sử', 'Click', 'Mở cửa sổ Lịch sử thay đổi và nạp các mốc của bản ghi đó.'),
    ('Bấm Bộ lọc và chọn loại hoạt động', 'Change',
     'Lọc lại danh sách mốc theo nhóm hoạt động đã chọn.'),
    ('Ghi mốc khi tạo / sửa / khóa / mở khóa', 'System',
     'Chỉ ghi khi nội dung theo dõi thực sự thay đổi; trạng thái được lưu dưới dạng chữ (Hoạt '
     'động / Khóa) để bản ghi lịch sử tự chứa nghĩa.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ, danh sách phía sau giữ nguyên.'),
])

# ============================================================ PHẦN 4. QUY TẮC
d.h1('Phần 4. Quy tắc nghiệp vụ')

sub('BR-01 — Định danh gói bảo dưỡng')
d.bullets([
    'Tên gói bảo dưỡng và Mã gói bảo dưỡng đều DUY NHẤT trên toàn hệ thống.',
    'Mã gói luôn được chuyển thành CHỮ IN HOA trước khi kiểm tra trùng và trước khi lưu, nên '
    '“bdt001” và “BDT001” là một.',
    'Tên gói giữ nguyên như người dùng nhập, chỉ cắt khoảng trắng ở đầu và cuối.',
    'Khi sửa, hệ thống bỏ qua chính bản ghi đang sửa nên giữ nguyên tên và mã cũ không bị báo '
    'trùng.',
])

sub('BR-02 — Cấu trúc ma trận nội dung kiểm tra')
d.bullets([
    'Mỗi DÒNG là một nội dung kiểm tra kèm đơn vị tính và số lượng; mỗi CỘT là một cấp bảo dưỡng.',
    'Ô giao nhau giữa một dòng và một cột chứa các ghi chú kiểm tra phải thực hiện; ô để trống thì '
    'không lưu được.',
    'Một cấp bảo dưỡng chỉ được chọn MỘT lần trên toàn bảng; chọn trùng bị chặn kèm cảnh báo “Cấp '
    'bảo dưỡng này đã được chọn ở cột khác”.',
    'Không được bỏ một cột cấp đã phát sinh báo giá dịch vụ; hệ thống chặn khi lưu và giữ nguyên '
    'toàn bộ dữ liệu cũ của gói.',
])

sub('BR-03 — Công thức tính giá')
d.bullets([
    'Giá vốn của một cấp = Đơn giá công của công ty quản lý × Định mức công × Hệ số công nghệ, làm '
    'tròn xuống.',
    'Giá công thức = Giá vốn × Hệ số giá bán của gói bảo dưỡng.',
    'Giá bán cơ sở mặc định bằng Giá công thức nhưng người dùng sửa tay được.',
    'Giá bán của từng công ty = Giá bán cơ sở × Hệ số giá bán của công ty đó; hệ số để trống được '
    'hiểu là 1, không phải 0.',
    'Giá hiển thị ở biểu tượng chữ i ngoài danh sách = Đơn giá công của công ty quản lý × Định mức '
    'công của cấp × Hệ số giá bán gói, làm tròn xuống.',
    'Mọi ô hệ số và phần trăm nhận DẤU PHẨY làm dấu thập phân: nhập 12,5 nghĩa là mười hai phẩy '
    'năm.',
])

sub('BR-04 — Ràng buộc nhập liệu')
d.bullets([
    'Định mức đàm phán giá tối đa 99; VAT trong khoảng 0–100; Hệ số giá bán gói trong khoảng '
    '1–100; hệ số theo công ty tối đa 99.999.999,99.',
    'Tên gói, Mã gói và Ghi chú tối đa 255 ký tự; vượt quá thì báo lỗi chứ không âm thầm cắt chuỗi.',
    'Bắt buộc đính kèm ít nhất MỘT file PDF thì mới lưu được gói; file khác định dạng bị từ chối.',
    'Ô Công ty quản lý gói bảo dưỡng bắt buộc có giá trị vì quyết định đơn giá công dùng để tính '
    'giá vốn.',
])

sub('BR-05 — Xóa hay Khóa')
d.bullets([
    'Gói CHƯA được sử dụng: xoá hẳn cùng toàn bộ nội dung kiểm tra và các cấp của gói.',
    'Gói ĐÃ được sử dụng (đã gắn hàng hoá hoặc đã dùng ở báo giá dịch vụ): không xoá mà chuyển '
    'sang trạng thái Khóa kèm thông báo “Gói bảo dưỡng đang được sử dụng nên đã được chuyển sang '
    'trạng thái Khóa”.',
    'Giao diện chủ động ẩn nút Xóa với gói đã được sử dụng, nhưng chốt chặn thật nằm ở máy chủ.',
    'Gói đang Khóa không sửa được: nút Sửa bị ẩn và máy chủ từ chối nếu gọi thẳng chức năng Sửa.',
    'Gói đang Khóa vẫn nằm trong danh sách với nhãn Khóa và vẫn xem, in, nhân bản, xem lịch sử '
    'được.',
])

sub('BR-06 — Lịch sử thay đổi')
d.bullets([
    'Các trường được theo dõi: Mã, Tên gói bảo dưỡng, Định mức đàm phán giá, VAT, Ghi chú, Trạng '
    'thái.',
    'Ghi mốc cho các thao tác: Tạo mới, Thay đổi thông tin, Khóa, Mở khóa.',
    'Không ghi mốc nếu nội dung theo dõi không thay đổi.',
    'Thay đổi ma trận nội dung kiểm tra, hệ số theo công ty, hàng hoá áp dụng và file đính kèm '
    'KHÔNG sinh mốc lịch sử.',
    'Trạng thái được lưu dưới dạng CHỮ (Hoạt động / Khóa) để bản ghi lịch sử tự chứa nghĩa.',
    'Bộ lọc loại hoạt động dùng chung 3 nhóm cố định cho mọi màn danh mục.',
])

sub('BR-07 — Phân quyền')
d.bullets([
    'Ba quyền độc lập: Thêm danh mục gói bảo dưỡng, Sửa danh mục gói bảo dưỡng, Xóa danh mục gói '
    'bảo dưỡng.',
    'Nút Nhân bản đi theo quyền Thêm; nút Mở khóa đi theo quyền Sửa.',
    'Xem danh sách, xem chi tiết, in phiếu, xuất Excel và xem lịch sử KHÔNG đòi quyền — mọi tài '
    'khoản đã đăng nhập đều thực hiện được.',
    'Nút không dùng được thì ẩn hẳn chứ không làm mờ; mọi thao tác ghi dữ liệu đều được kiểm tra '
    'lại ở máy chủ.',
])

sub('BR-08 — Dữ liệu dùng chung với hệ thống ERP')
d.bullets([
    'Danh mục này dùng chung một nguồn dữ liệu với màn tương ứng của hệ thống ERP đang chạy song '
    'song, không đổi cấu trúc dữ liệu gốc.',
    'Ba khác biệt có chủ đích so với màn ERP: bắt buộc đính kèm ít nhất một file PDF; bổ sung '
    'chức năng xem lịch sử thay đổi; bổ sung tuỳ chỉnh cột hiển thị theo từng người dùng.',
    'Hành vi “xoá gói đang dùng thì chuyển sang Khóa” và “xuất Excel không áp bộ lọc màn hình” '
    'được giữ nguyên như ERP; muốn đổi thì phải là yêu cầu nghiệp vụ mới.',
])

d.save()


# ----------------------------------------------------------------- MỤC LỤC
# `d.toc()` chi CHEN TRUONG TOC — mo file len thay TRONG cho toi khi bam Update Field.
# Cho Word cap nhat that (giong hdsd_engine._update_fields_by_word).
def _update_toc_by_word(path, title):
    import subprocess
    ps = """
$p = "{path}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($p, $false, $false)
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
$doc.Repaginate()
$doc.BuiltInDocumentProperties("Title") = "{title}"
$doc.Save()
Write-Output ("Pages=" + $doc.ComputeStatistics(2))
$doc.Close(0)
$word.Quit()
""".format(path=path, title=title)
    res = subprocess.run(["powershell", "-NonInteractive", "-Command", ps],
                         capture_output=True, text=True)
    print('Cap nhat muc luc bang Word:', res.stdout.strip() or res.stderr.strip())


_update_toc_by_word(OUT, 'SRS - Danh mục gói bảo dưỡng')

# --------------------------------------- Bước 4 của skill: tự kiểm tra form mới
from docx import Document  # noqa: E402

_chk = Document(OUT)
for _s in ['Tổng quan', 'Mini-Spec', 'Tiêu chí nghiệm thu', 'Ngoài phạm vi',
           'Chức năng liên quan', 'Route (FE)', 'Phân hệ:']:
    assert not any(_s in _p.text for _p in _chk.paragraphs), 'Con muc cua form cu: %s' % _s
print('OK - khong con muc nao cua form cu')
