# -*- coding: utf-8 -*-
"""Sinh 2 tai lieu mo ta nghiep vu: Phieu cung cap thong tin lam bao gia + Phieu bao hanh."""
import os
import re
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
NGAY = "28/08/2026"

# Bo kiem tra thuat ngu ky thuat (copy tu tc_engine.check_terms)
BANNED = [
    r"`[a-z_]{3,}`", r"\bBE\b", r"\bFE\b", r"\bHTTP\b",
    r"trả (400|403|404|422|423)",
    # Ma loi ky thuat. (?<![\d.]) / (?![\d.]) de KHONG bat nham so tien: "400.000" hop le,
    # con "loi 400" thi van bi bat.
    r"(?<![\d.])(400|403|404|422|423)(?![\d.])",
    r"permission id", r"\bAPI /", r"/api/v1", r"localStorage",
    r"number_format", r"meta\.", r"sort_by", r"per_page",
    r"role_has_permissions", r"current_company_role",
]

_texts = []


def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    return doc


class Writer(object):
    def __init__(self, doc):
        self.doc = doc

    def h(self, text, level):
        _texts.append(text)
        self.doc.add_paragraph(text, style="Heading %d" % level)

    def para(self, text, bold_prefix=None):
        _texts.append((bold_prefix or "") + text)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
        p.add_run(text)

    def bullet(self, text, bold_prefix=None):
        _texts.append((bold_prefix or "") + text)
        p = self.doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
        p.add_run(text)

    def table(self, rows):
        for row in rows:
            _texts.extend(row)
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        t.alignment = 1
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                t.cell(i, j).text = val
        return t

    def cover(self, title, subtitle):
        tp = self.doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = tp.add_run("MÔ TẢ NGHIỆP VỤ")
        r.bold = True
        r.font.size = Pt(20)

        sp = self.doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sp.add_run(title)
        r.bold = True
        r.font.size = Pt(16)

        ip = self.doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run(subtitle + " · Cập nhật ngày " + NGAY)

        self.doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# TÀI LIỆU 1 — PHIẾU CUNG CẤP THÔNG TIN LÀM BÁO GIÁ
# ══════════════════════════════════════════════════════════════════════════
def doc_cctt():
    doc = new_doc()
    w = Writer(doc)
    w.cover("Phiếu cung cấp thông tin làm báo giá",
            "Phân hệ CSKH → Kiểm tra bảo hành sửa chữa")

    w.h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
    w.para("Tài liệu mô tả đầy đủ nghiệp vụ của chức năng “Phiếu cung cấp thông tin làm báo giá”: "
           "dùng để làm gì, ai tham gia, phiếu đi qua những bước nào, mỗi bước hệ thống báo cho ai, "
           "và những quy tắc bắt buộc phải tuân thủ.")
    w.bullet("Cán bộ kỹ thuật lập phiếu và cán bộ kinh doanh nhận phiếu để làm báo giá.")
    w.bullet("Trưởng phòng, giám đốc cần nắm luồng để phân quyền và theo dõi tiến độ.")
    w.bullet("Người nghiệm thu và kiểm thử, để biết đâu là hành vi đúng.")
    w.bullet("Người viết tài liệu hướng dẫn sử dụng và người đào tạo người dùng mới.")

    w.h("2. CHỨC NĂNG NÀY DÙNG ĐỂ LÀM GÌ", 1)
    w.h("2.1. Mục đích", 2)
    w.para("Sau khi kỹ thuật đã khảo sát thiết bị hỏng của khách hàng, cần có một chứng từ liệt kê "
           "đầy đủ: thiết bị nào được bảo hành miễn phí, thiết bị nào phải sửa chữa có phí, cần "
           "những dịch vụ và vật tư gì, kèm các khoản chi phí phát sinh. Chứng từ đó chính là "
           "Phiếu cung cấp thông tin làm báo giá.")
    w.para("Trước khi có chứng từ này, thông tin khảo sát được truyền miệng hoặc ghi tay, dẫn tới "
           "hai hậu quả thường gặp: báo giá gửi khách bị thiếu hạng mục, và không truy được ai đã "
           "cung cấp con số nào.")
    w.h("2.2. Vị trí trong luồng lớn", 2)
    w.para("Đây là chứng từ thứ ba của dây chuyền dịch vụ sửa chữa – bảo hành:")
    w.table([
        ["Thứ tự", "Chứng từ", "Ai lập"],
        ["1", "Yêu cầu kiểm tra sửa chữa – bảo hành", "Người tiếp nhận yêu cầu của khách"],
        ["2", "Phiếu xử lý yêu cầu", "Cán bộ tiếp nhận, quyết định hướng xử lý từng thiết bị"],
        ["3", "Phiếu cung cấp thông tin làm báo giá", "Cán bộ kỹ thuật đã khảo sát"],
        ["4", "Báo giá dịch vụ", "Người đã lập Phiếu yêu cầu ở bước 1"],
    ])
    w.para("Phiếu này không tự sinh ra: chỉ lập được từ một Phiếu xử lý yêu cầu đang ở trạng thái "
           "chờ cung cấp thông tin. Ngược lại, khi phiếu được gửi đi, nó vừa đẩy trạng thái hai "
           "chứng từ phía trước, vừa mở đường cho bước làm báo giá.")
    w.h("2.3. Giá trị mang lại", 2)
    w.bullet("Mọi con số gửi cho khách đều có nguồn gốc, tra ngược được về đúng người khảo sát.")
    w.bullet("Tách bạch phần khách được bảo hành miễn phí và phần khách phải trả tiền.")
    w.bullet("Người làm báo giá nhận được thông báo ngay, không phải hỏi lại kỹ thuật.")
    w.bullet("Tự động sinh Phiếu bảo hành cho phần thiết bị còn bảo hành, không phải nhập lại.")

    w.h("3. NHỮNG AI THAM GIA", 1)
    w.table([
        ["Vai trò", "Làm gì trong luồng này"],
        ["Cán bộ kỹ thuật (người lập phiếu)",
         "Lập phiếu từ Phiếu xử lý yêu cầu, nhập dịch vụ, vật tư, chi phí; lưu nháp và gửi đi."],
        ["Người lập Phiếu yêu cầu (thường là kinh doanh)",
         "Nhận phiếu, kiểm tra thông tin. Đồng ý thì làm Báo giá dịch vụ; không đồng ý thì Từ chối "
         "tiếp nhận kèm lý do để kỹ thuật sửa lại."],
        ["Trưởng phòng, giám đốc", "Theo dõi qua danh sách theo phạm vi quyền được cấp; không thao "
                                  "tác trực tiếp trên phiếu."],
        ["Khách hàng", "Không thao tác trên hệ thống. Đây là người sẽ nhận báo giá ở bước sau, nên "
                       "mọi số liệu trên phiếu phải chính xác."],
    ])

    w.h("4. VÒNG ĐỜI CỦA PHIẾU", 1)
    w.table([
        ["Trạng thái", "Ý nghĩa", "Ai làm nó chuyển sang trạng thái này"],
        ["Đang tạo", "Phiếu nháp, còn sửa và xoá được", "Người lập phiếu bấm Lưu"],
        ["Chờ làm báo giá", "Đã gửi cho người làm báo giá, khoá không sửa nữa",
         "Người lập phiếu bấm Lưu và gửi"],
        ["Không duyệt", "Bị trả lại kèm lý do, người lập sửa rồi gửi lại được",
         "Người lập Phiếu yêu cầu bấm Từ chối tiếp nhận"],
        ["Đang báo giá", "Báo giá đã được lập ở dạng nháp", "Chứng từ Báo giá dịch vụ cập nhật ngược về"],
        ["Báo giá đã duyệt", "Báo giá đã duyệt xong", "Chứng từ Báo giá dịch vụ cập nhật ngược về"],
        ["Đã lập hợp đồng", "Đã ký hợp đồng dịch vụ từ báo giá", "Chứng từ Hợp đồng dịch vụ cập nhật ngược về"],
        ["Đã hoàn thành / Kết thúc", "Các bước cuối của dây chuyền", "Các chứng từ phía sau cập nhật ngược về"],
    ])
    w.para("Ba trạng thái đầu do chính màn này tạo ra. Các trạng thái còn lại do chứng từ phía sau "
           "đẩy ngược về, người dùng của màn này không tự đặt được.", bold_prefix="Lưu ý: ")
    w.para("Ngoài trạng thái phiếu còn một cột riêng là Trạng thái bảo hành, chỉ có giá trị khi "
           "phiếu có thiết bị ở khối bảo hành: “Chờ tạo phiếu bảo hành” và “Đã tạo phiếu bảo hành”.")

    w.h("5. LUỒNG HOẠT ĐỘNG CHI TIẾT", 1)
    w.h("5.1. Bước 1 — Mở phiếu từ Phiếu xử lý yêu cầu", 2)
    w.para("Người lập phiếu vào Phiếu xử lý yêu cầu đang chờ cung cấp thông tin và chọn thao tác "
           "Tạo phiếu cung cấp thông tin. Hệ thống chép sẵn sang phiếu mới: số phiếu xử lý, người "
           "yêu cầu, phòng yêu cầu, ngày nhận yêu cầu, khách hàng, người liên hệ, số điện thoại và "
           "địa chỉ sửa chữa. Toàn bộ khối này chỉ để đọc, chỉ ô Ghi chú cho nhập.")
    w.para("Chỉ những thiết bị được đánh dấu hướng xử lý “Cung cấp thông tin làm báo giá” ở bước "
           "trước mới được chép sang. Thiết bị đã xử lý bằng tư vấn điện thoại thì dừng ở bước "
           "trước, không đưa vào phiếu này.")
    w.h("5.2. Bước 2 — Phân loại thiết bị: bảo hành hay sửa chữa", 2)
    w.para("Mỗi dòng thiết bị có ô Loại công việc với hai lựa chọn. Đổi lựa chọn này sẽ chuyển "
           "nguyên dòng thiết bị cùng toàn bộ dịch vụ và vật tư của nó giữa hai khối:")
    w.table([
        ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
        ["Chọn Bảo hành", "Dòng thiết bị chuyển sang khối A – Bảo hành thiết bị",
         "Hệ thống tự đặt Chi phí được bảo hành bằng đúng Thành tiền của công, dịch vụ và vật tư; "
         "cột Khách hàng phải trả về 0"],
        ["Chọn Sửa chữa", "Dòng thiết bị chuyển sang khối B – Dịch vụ kiểm tra, sửa chữa",
         "Hệ thống đặt lại toàn bộ cột Cho bảo hành và Miễn phí ở khối Chi phí khác về 0"],
    ])
    w.h("5.3. Bước 3 — Nhập dịch vụ, vật tư và thiết bị cần bảo dưỡng", 2)
    w.bullet("Với từng thiết bị, người lập thêm dịch vụ từ danh mục dịch vụ sửa chữa và thêm vật "
             "tư từ danh mục hàng hoá; mỗi dòng nhập số lượng, đơn giá bán và chiết khấu.")
    w.bullet("Khối Danh mục thiết bị cần bảo dưỡng cho phép thêm thiết bị khách đang có, kèm các "
             "gói dịch vụ bảo dưỡng và vật tư của từng gói.")
    w.bullet("Khối Chi phí khác gồm năm khoản dựng sẵn, trong đó có chi phí vận chuyển tách riêng.")
    w.bullet("Khối Tổng hợp báo giá tự cộng lại ngay khi gõ, không phải tải lại trang.")
    w.h("5.4. Bước 4 — Lưu nháp hoặc gửi đi", 2)
    w.table([
        ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
        ["Bấm Lưu", "Lưu phiếu ở trạng thái Đang tạo",
         "Phiếu xử lý và Phiếu yêu cầu gốc chuyển sang trạng thái đang cung cấp thông tin. Người "
         "lập vẫn sửa và xoá được phiếu."],
        ["Bấm Lưu và gửi", "Phiếu chuyển sang Chờ làm báo giá",
         "Phiếu xử lý chuyển sang đã cung cấp thông tin kèm dấu người và ngày xử lý; Phiếu yêu cầu "
         "gốc chuyển sang đã có phiếu cung cấp thông tin. Phiếu khoá, không sửa được nữa."],
    ])
    w.para("Nếu phiếu có thiết bị ở khối bảo hành, ngay khi gửi đi hệ thống tự sinh một Phiếu bảo "
           "hành riêng và đánh dấu phiếu này là đã tạo phiếu bảo hành. Người dùng không phải nhập "
           "lại gì. Chi tiết xem tài liệu Mô tả nghiệp vụ – Phiếu bảo hành.")
    w.h("5.5. Bước 5 — Người làm báo giá tiếp nhận hoặc từ chối", 2)
    w.table([
        ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
        ["Tạo báo giá dịch vụ", "Người lập Phiếu yêu cầu mở phiếu và chọn thao tác này",
         "Chuyển sang màn lập Báo giá dịch vụ với dữ liệu lấy sẵn"],
        ["Từ chối tiếp nhận", "Bắt buộc nhập lý do vào ô Ghi chú duyệt",
         "Phiếu chuyển sang Không duyệt; hai chứng từ phía trước quay lại trạng thái đang cung cấp "
         "thông tin; người lập phiếu nhận thông báo kèm lý do và sửa lại rồi gửi lại được"],
    ])

    w.h("6. THÔNG BÁO — AI NHẬN, KHI NÀO, NỘI DUNG GÌ", 1)
    w.table([
        ["Sự kiện", "Ai nhận", "Nội dung", "Bấm vào thì đi đâu"],
        ["Phiếu được gửi đi (chuyển sang Chờ làm báo giá)",
         "Đúng MỘT người: người đã lập Phiếu yêu cầu ở bước 1. Người lập phiếu này không nhận.",
         "[PCCTT] Chờ duyệt: <số phiếu>. Khách hàng: <tên khách hàng>.",
         "Màn chi tiết đúng phiếu vừa gửi"],
        ["Bị từ chối tiếp nhận",
         "Đúng MỘT người: người đã lập phiếu cung cấp thông tin.",
         "[PCCTT] Từ chối: <số phiếu>. Lý do: <lý do người từ chối nhập>.",
         "Màn chi tiết đúng phiếu bị từ chối"],
    ])
    w.h("Quy ước chung về thông báo", 2)
    w.bullet("Thông báo hiện ở chuông thông báo trên thanh công cụ và luôn kèm đường dẫn tới đúng "
             "phiếu, không phải về danh sách chung.")
    w.bullet("Tên đối tượng được in đậm và cắt bớt nếu quá dài; toàn bộ nội dung giới hạn 120 ký "
             "tự, phần ghi chú bị cắt trước.")
    w.bullet("Nếu người lập phiếu cũng chính là người lập Phiếu yêu cầu thì họ sẽ tự nhận thông "
             "báo của chính mình — đúng thiết kế, vì đây là lời nhắc sang việc tiếp theo.")
    w.bullet("Gửi thông báo thất bại KHÔNG làm hỏng nghiệp vụ: phiếu vẫn được lưu và vẫn chuyển "
             "trạng thái, lỗi chỉ được ghi lại để quản trị xem sau.")
    w.para("Các sự kiện KHÔNG phát sinh thông báo, tránh báo thiếu khi kiểm thử: lưu nháp; sửa "
           "phiếu nháp; xoá phiếu; gửi đi phiếu chỉ có phần bảo hành mà không có dịch vụ sửa chữa "
           "và không có thiết bị bảo dưỡng.", bold_prefix="Lưu ý: ")

    w.h("7. PHÂN QUYỀN", 1)
    w.h("7.1. Nhìn thấy dữ liệu nào", 2)
    w.table([
        ["Quyền được cấp", "Nhìn thấy gì"],
        ["Xem phiếu cung cấp thông tin theo tổng công ty", "Toàn bộ phiếu của mọi công ty"],
        ["Xem phiếu cung cấp thông tin theo công ty", "Phiếu của công ty mình, cộng phiếu do chính mình lập"],
        ["Xem phiếu cung cấp thông tin theo phòng ban",
         "Phiếu của các phòng ban mình quản lý, cộng phiếu do chính mình lập"],
        ["Không có quyền nào ở trên", "Chỉ thấy phiếu do chính mình lập"],
    ])
    w.para("Phiếu đang ở trạng thái Đang tạo của người khác thì KHÔNG AI nhìn thấy, kể cả người có "
           "quyền cao nhất. Mở thẳng đường dẫn cũng bị chặn.", bold_prefix="Ngoại lệ quan trọng: ")
    w.h("7.2. Được phép làm gì", 2)
    w.table([
        ["Thao tác", "Điều kiện được phép"],
        ["Lập phiếu", "Có quyền Tạo phiếu cung cấp thông tin, và Phiếu xử lý yêu cầu đang chờ cung "
                      "cấp thông tin"],
        ["Sửa, Xoá", "Là người lập phiếu, và phiếu đang ở Đang tạo hoặc Không duyệt"],
        ["Tạo báo giá dịch vụ", "Là người lập Phiếu yêu cầu, phiếu đang Chờ làm báo giá, phiếu có "
                                "phần cần báo giá, và mỗi thiết bị chỉ mang một lỗi"],
        ["Từ chối tiếp nhận", "Cùng điều kiện với Tạo báo giá dịch vụ, và bắt buộc nhập lý do"],
        ["In phiếu, xem lịch sử", "Ai xem được phiếu thì làm được"],
    ])
    w.para("Nút nào không dùng được thì bị ẩn hẳn khỏi màn hình, không hiện dạng chữ mờ. Danh sách "
           "nút ở màn chi tiết luôn khớp với màn danh sách của cùng phiếu đó.")

    w.h("8. QUY TẮC NGHIỆP VỤ BẮT BUỘC", 1)
    w.h("8.1. Bắt buộc nhập — tách theo từng nút bấm", 2)
    w.table([
        ["Bấm nút", "Bắt buộc có gì"],
        ["Lưu (nháp)", "Chỉ cần phiếu xử lý gốc và khách hàng. Các phần còn lại để trống vẫn lưu được."],
        ["Lưu và gửi", "Phải có điều khoản báo giá, người liên hệ, địa chỉ sửa chữa; mỗi dòng thiết "
                       "bị, dịch vụ, vật tư phải có số lượng và đơn giá; các khoản chi phí phải có "
                       "giá trị."],
        ["Từ chối tiếp nhận", "Bắt buộc nhập lý do."],
    ])
    w.h("8.2. Ràng buộc dữ liệu", 2)
    w.bullet("Chi phí được bảo hành không nhận số âm; nhập số âm thì tự đưa về 0.")
    w.bullet("Hai dòng cùng một thiết bị ở khối bảo dưỡng không được trùng số sê-ri.")
    w.bullet("Thiết bị chỉ được thêm tối đa bằng số lượng khách hàng đang có.")
    w.bullet("Danh mục dịch vụ đã bị khoá vẫn hiện đúng tên trên phiếu cũ đang dùng nó, kèm dấu ổ "
             "khoá, để phiếu cũ không bị mất dữ liệu.")
    w.h("8.3. Khoá chỉnh sửa", 2)
    w.para("Phiếu đã gửi đi thì không sửa, không xoá được nữa — chặn ngay ở máy chủ chứ không chỉ "
           "ẩn nút. Muốn sửa thì phải được người làm báo giá từ chối tiếp nhận để phiếu quay lại.")
    w.h("8.4. Sinh số phiếu", 2)
    w.para("Số phiếu do hệ thống tự sinh theo khuôn: mã công ty, chữ viết tắt của loại phiếu, năm "
           "lập và số thứ tự sáu chữ số. Người dùng không sửa được số phiếu.")

    w.h("8.5. Cách tính tiền trên phiếu", 2)
    w.para("Toàn bộ số tiền trên phiếu đều được tính ra từ những ô người lập nhập, không có ô tiền "
           "nào vừa nhập tay vừa tự tính. Phiếu chia tiền làm hai dòng chảy tách bạch: phần khách "
           "được BẢO HÀNH (công ty chịu) và phần khách phải TRẢ TIỀN. Hai dòng chảy này đi riêng "
           "suốt phiếu, chỉ gặp nhau ở bảng tổng hợp cuối cùng.")

    w.h("8.5.1. Khối A — Bảo hành thiết bị", 2)
    w.para("Mỗi thiết bị có ba nhóm tiền: công sửa chữa, danh sách dịch vụ và danh sách vật tư. "
           "Khác với khối sửa chữa, ở đây phần được bảo hành là MỘT SỐ TIỀN người lập tự gõ vào, "
           "không phải tỷ lệ phần trăm.")
    w.table([
        ["Cột", "Cách ra số"],
        ["Thành tiền", "Đơn giá bán × Số lượng"],
        ["Chi phí được bảo hành", "Người lập tự nhập số tiền công ty chịu cho dòng đó"],
        ["Khách hàng phải trả", "Thành tiền − Chi phí được bảo hành"],
    ])
    w.para("Thiết bị được bảo hành toàn bộ thì người lập nhập Chi phí được bảo hành đúng bằng "
           "Thành tiền, khi đó Khách hàng phải trả bằng 0 — đây là trường hợp phổ biến nhất.")

    w.h("8.5.2. Khối B-I — Dịch vụ kiểm tra, sửa chữa", 2)
    w.para("Đây là phần khách trả tiền, nên có chiết khấu và thuế. Chiết khấu ở khối này là TỶ LỆ "
           "phần trăm, không phải số tiền.")
    w.table([
        ["Cột", "Cách ra số"],
        ["Thành tiền", "Đơn giá bán × Số lượng"],
        ["Chiết khấu", "Thành tiền × Tỷ lệ chiết khấu của dòng"],
        ["Thành tiền sau chiết khấu", "Thành tiền − Chiết khấu"],
        ["Tiền thuế", "Thành tiền sau chiết khấu × Thuế suất của dòng"],
        ["Thành tiền sau thuế", "Thành tiền sau chiết khấu + Tiền thuế"],
    ])
    w.bullet("Dòng công: số lượng công để trống thì hệ thống hiểu là 1, nếu không dòng công sẽ ra "
             "0 đồng.")
    w.bullet("Thuế suất lấy theo TỪNG DÒNG (mỗi dịch vụ, mỗi vật tư một mức riêng theo danh mục), "
             "không phải một mức chung cho cả phiếu.")

    w.h("8.5.3. Khối B-II — Thiết bị cần bảo dưỡng", 2)
    w.para("Tính theo GÓI bảo dưỡng: mỗi gói có đơn giá và số lượng gói, bên trong gói có thể kèm "
           "vật tư. Công thức từng dòng giống khối B-I.")
    w.bullet("Gói có số lượng bằng 0 thì bị bỏ qua HOÀN TOÀN, kể cả vật tư nằm trong gói đó — "
             "không cộng một đồng nào vào tổng.")

    w.h("8.5.4. Khối C — Chi phí khác", 2)
    w.para("Năm khoản chi phí cố định, chia hai bảng: bốn khoản liên quan và riêng chi phí vận "
           "chuyển. Người lập chỉ nhập ba ô đầu, ba cột còn lại hệ thống tự tính — đây là chỗ hay "
           "bị hiểu nhầm nhất của cả màn hình.")
    w.table([
        ["Cột", "Nhập hay tự tính", "Cách ra số"],
        ["Giá trị", "Nhập", "Tổng số tiền của khoản chi phí đó"],
        ["Cho bảo hành", "Nhập", "Phần trong Giá trị được tính vào diện bảo hành. Không được lớn "
                                 "hơn Giá trị"],
        ["Miễn phí", "Nhập", "Phần trong Cho bảo hành mà công ty miễn hẳn cho khách. Không được "
                             "lớn hơn Cho bảo hành"],
        ["Trả phí", "Tự tính", "Cho bảo hành − Miễn phí"],
        ["Chi phí cho SC - BD", "Tự tính", "Giá trị − Cho bảo hành"],
        ["Khách hàng phải trả", "Tự tính", "Trả phí + Chi phí cho SC - BD"],
    ])
    w.bullet("Nhóm ba cột Cho bảo hành / Miễn phí / Trả phí CHỈ hiện khi phiếu có thiết bị thuộc "
             "diện bảo hành. Phiếu thuần sửa chữa thì bảng rút gọn còn Giá trị, Chi phí cho "
             "SC - BD, Khách hàng phải trả và Ghi chú.")
    w.bullet("Thuế suất của dòng chi phí để trống thì hệ thống hiểu là 8%.")
    w.bullet("Chi phí khác không có chiết khấu.")

    w.h("8.5.5. Khối D — Tổng hợp báo giá", 2)
    w.para("Gồm ba bảng. Hai bảng đầu là tổng của hai dòng chảy tiền, bảng thứ ba cộng lại thành "
           "số khách phải thanh toán.")
    w.table([
        ["Bảng", "Gồm những dòng nào"],
        ["I — Bảo hành", "Tổng chi phí bảo hành (công + dịch vụ + vật tư), Các khoản chi phí liên "
                         "quan và Chi phí vận chuyển — lấy phần thuộc diện bảo hành"],
        ["II — Sửa chữa, Bảo dưỡng", "Tổng chi phí sửa chữa (công + dịch vụ + vật tư), Tổng chi "
                                     "phí bảo dưỡng (gói + vật tư trong gói), Các khoản chi phí "
                                     "liên quan và Chi phí vận chuyển — lấy phần khách phải trả"],
        ["III — Tổng hợp báo giá", "Hai dòng Bảo hành và Sửa chữa bảo dưỡng, rồi Thành tiền, Thuế "
                                   "và Tổng thanh toán"],
    ])
    w.para("Hai dòng chi phí trong bảng tổng hợp KHÔNG phải nhập lại — chúng lấy thẳng từ khối "
           "\u201cC - Chi phí khác\u201d, nhưng mỗi bảng lấy một cột khác nhau. Đây là chỗ hay bị "
           "hỏi nhất khi đối chiếu số:")
    w.table([
        ["Dòng ở khối D", "Lấy từ đâu", "Lấy cột nào"],
        ["I — Các khoản chi phí liên quan", "Bảng C-I (bốn khoản: đi lại, lưu trú, làm thêm giờ, "
                                            "chi phí khác)",
         "Thành tiền lấy cột Cho bảo hành · Chi phí được bảo hành lấy cột Miễn phí · Khách hàng "
         "phải trả lấy cột Trả phí"],
        ["I — Chi phí vận chuyển", "Bảng C-II (riêng khoản vận chuyển)", "Ba cột như trên"],
        ["II — Các khoản chi phí liên quan", "Bảng C-I", "Lấy cột Chi phí cho SC - BD; cột Giảm "
                                                         "giá luôn bằng 0 vì chi phí không có "
                                                         "chiết khấu"],
        ["II — Chi phí vận chuyển", "Bảng C-II", "Lấy cột Chi phí cho SC - BD"],
    ])
    w.para("Nói gọn: một khoản chi phí được CHẺ ĐÔI theo đúng hai ô người lập đã nhập — phần ghi "
           "vào Cho bảo hành chạy lên bảng I, phần còn lại (Giá trị trừ Cho bảo hành) chạy xuống "
           "bảng II. Cộng hai bảng lại luôn bằng đúng cột Giá trị ban đầu, không đội lên cũng "
           "không hụt đi.")
    w.para("Ví dụ một khoản 1.000.000 đồng, người lập ghi 400.000 vào Cho bảo hành và 100.000 vào "
           "Miễn phí: bảng I nhận 400.000 (được miễn 100.000, khách trả 300.000), bảng II nhận "
           "600.000. Khách hàng phải trả tổng cộng 900.000 đồng.")

    w.para("Ba dòng cuối cùng của phiếu ra số như sau:")
    w.table([
        ["Dòng", "Cách ra số"],
        ["Thành tiền", "Tổng phần bảo hành sau khi trừ bảo hành, cộng tổng phần sửa chữa - bảo "
                       "dưỡng sau chiết khấu"],
        ["Thuế", "Cộng tiền thuế của tất cả các dòng. Phần sửa chữa - bảo dưỡng tính theo thuế "
                 "suất từng dòng; phần bảo hành tính theo thuế suất chung của phiếu"],
        ["Tổng thanh toán", "Thành tiền + Thuế. Đây là số đọc thành chữ ở dòng Bằng chữ và là số "
                            "khách hàng phải trả"],
    ])
    w.bullet("Phiếu chỉ có thiết bị bảo hành, được bảo hành toàn bộ thì Tổng thanh toán bằng 0 — "
             "đúng nghiệp vụ, không phải lỗi thiếu dữ liệu.")
    w.bullet("Bản in dùng đúng công thức của màn hình, nên số trên giấy luôn khớp số trên màn.")

    w.h("9. CÁC LỐI VÀO MÀN HÌNH", 1)
    w.para("Cùng một màn hình nhưng vào bằng những đường dẫn khác nhau thì danh sách hiện ra khác "
           "nhau — đây là cách hệ thống cũ đang chạy và hệ thống mới giữ nguyên. Đường dẫn đầy đủ "
           "là địa chỉ máy chủ ghép với phần dưới đây.")
    w.para("Màn hình: /customer-care/wr-information-requests")
    w.table([
        ["Vào bằng", "Danh sách hiện ra", "Dùng khi nào"],
        ["(không kèm gì)", "Chỉ phiếu do chính tôi lập", "Xem lại việc của mình"],
        ["?type=all", "Toàn bộ phiếu trong phạm vi quyền của tôi", "Lối vào chính từ menu"],
        ["?type=waiting_create_quotation", "Phiếu đang chờ CHÍNH TÔI làm báo giá", "Danh sách việc của người làm báo giá"],
    ])
    w.para("Ba điều cần biết:")
    w.bullet("Đường dẫn quyết định PHẠM VI xem, không phải quyền: người không có quyền xem theo "
             "cấp mà mở đường dẫn xem tất cả thì vẫn chỉ thấy phiếu của chính mình.")
    w.bullet("Bấm nút Làm mới chỉ xoá điều kiện lọc, không đưa người dùng sang phạm vi khác.")
    w.bullet("Đường dẫn kèm giá trị lạ thì hệ thống bỏ qua và giữ phạm vi mặc định.")

    w.h("10. TRA CỨU, IN VÀ XUẤT DỮ LIỆU", 1)
    w.bullet("Ô tìm nhanh tìm theo: số phiếu cung cấp thông tin, số phiếu xử lý, tên khách hàng, "
             "người tạo.")
    w.bullet("Bộ lọc nâng cao gồm: trạng thái, số phiếu xử lý, khách hàng, người yêu cầu, người "
             "tạo, trạng thái bảo hành, tên hoặc mã thiết bị, khoảng ngày tạo, và khối công ty – "
             "phòng ban theo phạm vi quyền.")
    w.bullet("In một phiếu và in cả danh sách theo đúng điều kiện đang lọc; bản in có phần đầu thư "
             "của đúng công ty ghi trên phiếu.")
    w.bullet("Xuất Excel danh sách, tự chọn những cột cần xuất.")
    w.bullet("Mỗi phiếu có mục Lịch sử ghi lại ai sửa gì, giá trị cũ và giá trị mới, lúc nào.")

    w.h("11. LIÊN THÔNG VỚI HỆ THỐNG KHÁC", 1)
    w.para("Chức năng này dùng chung dữ liệu với hệ thống cũ: cùng danh mục khách hàng, danh mục "
           "hàng hoá, danh mục dịch vụ và cùng kho dữ liệu phiếu. Phiếu lập ở hệ thống mới tra cứu "
           "được ở hệ thống cũ và ngược lại.")
    w.table([
        ["Điểm khác", "Hệ thống cũ", "Hệ thống mới", "Lý do"],
        ["Cột Giá vốn", "Hiện cho mọi người mở được phiếu",
         "Ẩn hẳn nếu không có quyền xem giá vốn hàng hoá",
         "Giá vốn là dữ liệu nhạy cảm; người không có quyền vẫn lưu được phiếu mà không làm mất "
         "giá vốn cũ"],
        ["Phiếu nháp của người khác", "Người quản trị cao nhất vẫn mở được bằng đường dẫn trực tiếp",
         "Không ai mở được, kể cả đường in", "Phiếu nháp là bản làm dở, chưa phải dữ liệu chính thức"],
        ["Thiết bị mang nhiều lỗi",
         "Tách thành nhiều dòng nhưng các dòng tách ra bị mất phần dịch vụ",
         "Dòng tách ra giữ đúng dịch vụ của lỗi tương ứng",
         "Mất dịch vụ làm báo giá gửi khách bị thiếu tiền"],
        ["Bắt buộc nhập khi lưu nháp", "Bắt buộc đầy đủ mọi trường ở cả hai nút",
         "Lưu nháp chỉ cần phiếu xử lý gốc và khách hàng",
         "Đồng nhất với hai chứng từ trước, cho phép ghi lại việc đang làm dở"],
        ["Bốn cột tiền tổng của phiếu", "Luôn để 0", "Ghi số thật",
         "Để báo cáo đọc được ngay mà không phải tính lại"],
    ])

    w.h("12. GIỚI HẠN HIỆN TẠI", 1)
    w.bullet("Thao tác Tạo báo giá dịch vụ chưa chuyển sang được màn Báo giá dịch vụ vì màn đó "
             "chưa xây xong; hiện hệ thống báo rõ để người dùng làm tạm trên hệ thống cũ. Bỏ giới "
             "hạn này ngay khi màn Báo giá dịch vụ hoàn thành.")
    w.bullet("Khối bảo hành thiết bị và khối sản phẩm mở rộng chưa được kiểm thử trên dữ liệu thật "
             "có phát sinh bảo hành, do phiếu dùng để kiểm thử không có phần này.")
    w.bullet("Ba quyền xem của màn hiện mới được gán cho vai trò quản trị cao nhất; các vai trò "
             "nghiệp vụ cần được cấp bổ sung thì mới thấy đủ dữ liệu và bộ lọc theo công ty – "
             "phòng ban.")

    out = os.path.join(HERE, "Mô tả nghiệp vụ - Phiếu cung cấp thông tin làm báo giá.docx")
    doc.save(out)
    return out


# ══════════════════════════════════════════════════════════════════════════
# TÀI LIỆU 2 — PHIẾU BẢO HÀNH
# ══════════════════════════════════════════════════════════════════════════
def doc_bao_hanh():
    doc = new_doc()
    w = Writer(doc)
    w.cover("Phiếu bảo hành", "Phân hệ CSKH → Kiểm tra bảo hành sửa chữa")

    w.h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
    w.para("Tài liệu mô tả nghiệp vụ của chức năng “Phiếu bảo hành”: phiếu này ở đâu ra, ai xem "
           "được, dùng để làm gì và vì sao nó không có nút thêm hay sửa như các màn khác.")
    w.bullet("Cán bộ kỹ thuật và cán bộ chăm sóc khách hàng cần tra cứu phần việc bảo hành.")
    w.bullet("Trưởng phòng, giám đốc theo dõi khối lượng bảo hành đang phải thực hiện.")
    w.bullet("Người nghiệm thu và kiểm thử.")

    w.h("2. CHỨC NĂNG NÀY DÙNG ĐỂ LÀM GÌ", 1)
    w.h("2.1. Mục đích", 2)
    w.para("Khi kỹ thuật khảo sát xong và xác định một phần thiết bị của khách vẫn còn trong thời "
           "hạn bảo hành, phần việc đó công ty làm miễn phí. Nó không đi vào báo giá gửi khách, "
           "nhưng vẫn phải giao cho kỹ thuật làm và vẫn phải theo dõi chi phí. Phiếu bảo hành là "
           "chứng từ giữ phần việc miễn phí đó.")
    w.h("2.2. Vị trí trong luồng lớn", 2)
    w.para("Ở bước Phiếu cung cấp thông tin làm báo giá, dây chuyền tách làm hai nhánh theo từng "
           "dòng thiết bị:")
    w.table([
        ["Nhánh", "Dành cho", "Chứng từ sinh ra", "Khách hàng trả tiền?"],
        ["Sửa chữa", "Thiết bị đã hết bảo hành", "Báo giá dịch vụ, rồi Hợp đồng dịch vụ", "Có"],
        ["Bảo hành", "Thiết bị còn trong hạn bảo hành", "Phiếu bảo hành", "Không"],
    ])
    w.para("Hai nhánh nhập lại ở bước giao việc: dù là hợp đồng hay phiếu bảo hành, kỹ thuật đều "
           "nhận việc và nhập kết quả theo cùng một cách.")
    w.h("2.3. Giá trị mang lại", 2)
    w.bullet("Phần việc bảo hành miễn phí không bị bỏ sót, vì được lập thành chứng từ riêng.")
    w.bullet("Không phải nhập lại lần thứ hai: toàn bộ thiết bị, dịch vụ, vật tư và chi phí được "
             "chép tự động từ phiếu khảo sát.")
    w.bullet("Theo dõi được chi phí công ty bỏ ra cho bảo hành, tách khỏi doanh thu dịch vụ.")

    w.h("3. NHỮNG AI THAM GIA", 1)
    w.table([
        ["Vai trò", "Làm gì trong luồng này"],
        ["Cán bộ kỹ thuật lập phiếu cung cấp thông tin",
         "Không thao tác trực tiếp trên màn này. Việc họ làm ở bước trước — đánh dấu thiết bị là "
         "Bảo hành rồi gửi phiếu đi — chính là thứ sinh ra phiếu bảo hành."],
        ["Người tra cứu (kỹ thuật, chăm sóc khách hàng, quản lý)",
         "Mở danh sách và xem chi tiết phiếu để biết phải làm những gì cho khách."],
        ["Khách hàng", "Không thao tác. Đây là bên được hưởng phần việc miễn phí ghi trên phiếu."],
    ])

    w.h("4. VÒNG ĐỜI CỦA PHIẾU", 1)
    w.table([
        ["Trạng thái", "Ý nghĩa", "Ai làm nó chuyển sang trạng thái này"],
        ["Đang tạo", "Bản nháp", "Không phát sinh trong thực tế, vì phiếu được sinh tự động và bỏ "
                                "qua bước nháp"],
        ["Đã duyệt", "Phiếu có hiệu lực, chờ giao việc cho kỹ thuật",
         "Hệ thống đặt ngay khi sinh phiếu, không cần ai duyệt"],
        ["Đang thực hiện", "Kỹ thuật đã nhận việc", "Chứng từ Phiếu giao việc cập nhật ngược về"],
        ["Đã hoàn thành", "Đã làm xong và nghiệm thu", "Các chứng từ phía sau cập nhật ngược về"],
    ])
    w.para("Trên toàn bộ dữ liệu đang có, mọi phiếu bảo hành đều ở trạng thái Đã duyệt. Hai trạng "
           "thái sau chỉ xuất hiện khi các bước giao việc và nghiệm thu được đưa vào sử dụng.",
           bold_prefix="Thực tế hiện nay: ")

    w.h("5. LUỒNG HOẠT ĐỘNG CHI TIẾT", 1)
    w.h("5.1. Bước 1 — Phiếu được sinh ra tự động", 2)
    w.para("Không ai bấm nút để tạo phiếu bảo hành. Khi người lập Phiếu cung cấp thông tin làm báo "
           "giá bấm Lưu và gửi, nếu phiếu đó có ít nhất một thiết bị nằm ở khối bảo hành thì hệ "
           "thống lập tức tạo một phiếu bảo hành và chép sang:")
    w.bullet("Thông tin khách hàng đầy đủ tại thời điểm lập: tên, mã số thuế, người đại diện, tài "
             "khoản ngân hàng, địa chỉ sửa chữa, người liên hệ.")
    w.bullet("Toàn bộ dòng thiết bị thuộc khối bảo hành, kèm dịch vụ và vật tư của từng thiết bị.")
    w.bullet("Năm khoản chi phí khác, giữ nguyên cả những khoản để 0.")
    w.bullet("Khối tổng hợp tiền của phần bảo hành, chốt tại thời điểm sinh phiếu.")
    w.para("Ngay sau đó, phiếu cung cấp thông tin được đánh dấu là đã tạo phiếu bảo hành, để không "
           "sinh lần thứ hai.")
    w.h("5.2. Bước 2 — Tra cứu và sử dụng", 2)
    w.table([
        ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
        ["Xem danh sách", "Vào menu Kiểm tra bảo hành sửa chữa, chọn Phiếu bảo hành",
         "Thấy các phiếu trong phạm vi quyền, lọc theo khách hàng, trạng thái, người tạo, khoảng "
         "ngày, tên hoặc số sê-ri thiết bị"],
        ["Xem chi tiết", "Bấm vào số phiếu ở cột đầu",
         "Thấy phần đầu phiếu, danh sách thiết bị kèm dịch vụ và vật tư, các khoản chi phí và khối "
         "tổng hợp"],
        ["Tra ngược nguồn gốc", "Bấm vào số phiếu cung cấp thông tin trên màn chi tiết",
         "Mở đúng phiếu khảo sát đã sinh ra phiếu bảo hành này"],
    ])
    w.h("5.3. Vì sao màn này không có nút Thêm, Sửa, Xoá", 2)
    w.para("Phiếu bảo hành sinh tự động và ra đời đã ở trạng thái Đã duyệt. Quy tắc của hệ thống "
           "cũ là chỉ sửa hoặc xoá được phiếu đang ở trạng thái nháp, nên trên thực tế không phiếu "
           "nào đủ điều kiện. Theo nguyên tắc chung, nút không dùng được thì ẩn hẳn thay vì hiện "
           "dạng chữ mờ, vì vậy màn này chỉ còn tra cứu. Muốn sửa nội dung bảo hành thì phải sửa "
           "từ gốc: nhờ người làm báo giá từ chối tiếp nhận phiếu cung cấp thông tin, sửa lại rồi "
           "gửi lại.")

    w.h("6. THÔNG BÁO — AI NHẬN, KHI NÀO, NỘI DUNG GÌ", 1)
    w.table([
        ["Sự kiện", "Ai nhận", "Nội dung", "Bấm vào thì đi đâu"],
        ["Phiếu bảo hành được sinh ra",
         "Đúng MỘT người: người vừa gửi phiếu cung cấp thông tin làm báo giá",
         "Thông báo cho biết đã có phiếu bảo hành kèm số phiếu",
         "Màn chi tiết đúng phiếu bảo hành vừa sinh"],
    ])
    w.h("Quy ước chung về thông báo", 2)
    w.bullet("Đây là lời nhắc cho chính người vừa thao tác, để họ biết hệ thống đã tách phần bảo "
             "hành ra chứng từ riêng — không phải yêu cầu người khác xử lý.")
    w.bullet("Gửi thông báo thất bại không làm hỏng nghiệp vụ: phiếu bảo hành vẫn được tạo.")
    w.para("Các sự kiện KHÔNG phát sinh thông báo: xem phiếu; lọc danh sách; phiếu cung cấp thông "
           "tin được gửi lại lần hai (vì không sinh thêm phiếu bảo hành nào).",
           bold_prefix="Lưu ý: ")

    w.h("7. PHÂN QUYỀN", 1)
    w.h("7.1. Nhìn thấy dữ liệu nào", 2)
    w.table([
        ["Quyền được cấp", "Nhìn thấy gì"],
        ["Xem phiếu bảo hành theo tổng công ty", "Toàn bộ phiếu của mọi công ty"],
        ["Xem phiếu bảo hành theo công ty", "Phiếu của công ty mình, cộng phiếu do chính mình tạo"],
        ["Xem phiếu bảo hành theo phòng ban",
         "Phiếu của các phòng ban mình quản lý, cộng phiếu do chính mình tạo"],
        ["Không có quyền nào ở trên", "Chỉ thấy phiếu do chính mình tạo"],
    ])
    w.para("Mở thẳng đường dẫn tới một phiếu ngoài phạm vi quyền vẫn bị chặn ở máy chủ, không chỉ "
           "ẩn trên giao diện.")
    w.h("7.2. Được phép làm gì", 2)
    w.table([
        ["Thao tác", "Điều kiện được phép"],
        ["Xem danh sách, xem chi tiết", "Có một trong các quyền ở mục 7.1, hoặc là người tạo phiếu"],
        ["Sửa, Xoá", "Không áp dụng — xem mục 5.3"],
        ["Tạo phiếu giao việc", "Chưa áp dụng vì bước giao việc chưa được đưa vào sử dụng"],
    ])

    w.h("8. QUY TẮC NGHIỆP VỤ BẮT BUỘC", 1)
    w.bullet("Mỗi phiếu cung cấp thông tin chỉ sinh ĐÚNG MỘT phiếu bảo hành. Phiếu bị từ chối rồi "
             "sửa và gửi lại cũng không sinh thêm phiếu thứ hai.")
    w.bullet("Phiếu bảo hành không có thuế: toàn bộ phần việc là miễn phí cho khách.")
    w.bullet("Số phiếu do hệ thống tự sinh theo khuôn: mã công ty, chữ viết tắt của loại phiếu, "
             "năm lập và số thứ tự sáu chữ số.")
    w.bullet("Thông tin khách hàng trên phiếu là bản chụp tại thời điểm lập. Sau này danh mục "
             "khách hàng đổi địa chỉ hay đổi người đại diện thì phiếu cũ vẫn giữ nguyên thông tin "
             "đã dùng, để đối chiếu về sau không bị sai lệch.")
    w.bullet("Số tiền trên phiếu được chốt tại thời điểm sinh phiếu và không tính lại khi xem, để "
             "phiếu cũ không bị đổi số khi danh mục thay đổi giá.")

    w.h("10. TRA CỨU, IN VÀ XUẤT DỮ LIỆU", 1)
    w.bullet("Ô tìm nhanh tìm theo: số phiếu bảo hành, tên khách hàng, người liên hệ.")
    w.bullet("Bộ lọc gồm: trạng thái, khách hàng, số phiếu cung cấp thông tin, người tạo, tên hoặc "
             "số sê-ri thiết bị, khoảng ngày tạo, và khối công ty – phòng ban theo phạm vi quyền.")
    w.bullet("Người dùng tự chọn cột hiển thị trên danh sách.")
    w.bullet("Chưa có chức năng in phiếu và xuất danh sách — xem mục 11.")

    w.h("11. LIÊN THÔNG VỚI HỆ THỐNG KHÁC", 1)
    w.para("Chức năng dùng chung kho dữ liệu với hệ thống cũ: phiếu bảo hành lập ở hệ thống mới "
           "tra cứu được ở hệ thống cũ và ngược lại. Cùng một kho dữ liệu đó còn phục vụ Hợp đồng "
           "dịch vụ và các loại phụ lục hợp đồng, nên khi xem phải phân biệt đúng loại chứng từ.")
    w.table([
        ["Điểm khác", "Hệ thống cũ", "Hệ thống mới", "Lý do"],
        ["Sinh phiếu trùng",
         "Phiếu bị từ chối rồi gửi lại sẽ sinh thêm một phiếu bảo hành nữa cho cùng lô thiết bị",
         "Chỉ sinh đúng một phiếu, gửi lại thì dùng lại phiếu cũ",
         "Phiếu trùng làm khối lượng bảo hành bị đếm hai lần"],
        ["Nguồn của số tiền", "Lấy số do trình duyệt của người dùng gửi lên",
         "Máy chủ tự tính lại từ dữ liệu đã lưu",
         "Không phụ thuộc vào trình duyệt, tránh sai số khi máy người dùng gặp sự cố"],
        ["Lập phiếu bằng tay từ màn Báo giá",
         "Có nút nhưng đã bị tắt từ lâu, không phiếu nào đi đường này",
         "Không làm nút đó", "Toàn bộ phiếu đang dùng đều sinh từ phiếu cung cấp thông tin"],
    ])

    w.h("12. GIỚI HẠN HIỆN TẠI", 1)
    w.bullet("Chưa có chức năng in phiếu và xuất danh sách ra Excel; hệ thống cũ cũng không có hai "
             "chức năng này cho màn Phiếu bảo hành.")
    w.bullet("Chưa có mục Lịch sử thay đổi, vì phiếu không sửa được nên chưa phát sinh thay đổi "
             "nào để ghi lại.")
    w.bullet("Thao tác Tạo phiếu giao việc chưa có, do bước giao việc chưa được đưa vào sử dụng. "
             "Khi bước đó hoàn thành thì bổ sung nút vào màn này.")
    w.bullet("Ba quyền xem của màn mới được gán cho vai trò quản trị cao nhất; các vai trò nghiệp "
             "vụ cần được cấp bổ sung thì mới thấy đủ dữ liệu.")

    out = os.path.join(HERE, "Mô tả nghiệp vụ - Phiếu bảo hành.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    files = [doc_cctt(), doc_bao_hanh()]

    blob = "\n".join(_texts)
    found = {p: len(re.findall(p, blob)) for p in BANNED if re.findall(p, blob)}
    print("!!! CON THUAT NGU KY THUAT:", found) if found else print("OK - sach thuat ngu ky thuat")
    for f in files:
        print("Da tao:", f)
