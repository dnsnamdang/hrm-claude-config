# -*- coding: utf-8 -*-
"""Sinh SRS man 'Danh muc ngan hang' (/human/banks) theo FORM CHUAN MOI (user chot 2026-08-17).

Form moi = 4 chuong (ban mau: .claude/skills/srs-documenter/assets/SRS_MAU.docx):
    Phan 1. Gioi thieu · Phan 2. Phan quyen · Phan 3. Dac ta chi tiet · Phan 4. Quy tac nghiep vu
Da bo so voi form cu: bang thong tin trang bia, dong "Phan he: ...", muc "Pham vi", chuong
"Tong quan", muc "Quy tac truy cap bat buoc", chuong "Danh muc chuc nang (Function list)",
muc "Tieu chi nghiem thu" cua tung chuc nang, dong "Chuc nang lien quan: FR-xx", 2 dong
"Menu:"/"Route (FE):" o muc Layout.

Nguon doi chieu (doc truc tiep tu code tren nhanh gop_db, 2026-08-17):
  BE  Modules/Human/Routes/api.php (prefix /human/banks — group /v1 chi co middleware auth:api)
      Modules/Human/Http/Controllers/Api/V1/BankController.php
      Modules/Human/Services/BankService.php
      Modules/Human/Entities/Bank.php, BankBranch.php
      Modules/Human/Http/Requests/CreateBankRequest.php, CreateBankBranchesRequest.php
      Modules/Human/Transformers/BankResource/*.php
      database/migrations/2025_06_30_135932_drop_table_banks.php (+ add_logo/status/created_by)
      app/Services/CatalogHistoryService.php + app/Services/Concerns/LogsCatalogHistory.php
  FE  hrm-client/pages/human/banks/index.vue + components/{BankModel,BankBranchesModel,
      BankBranchesAddModel,BankSearch}.vue · components/subsystem-menu/master-data.js
  Anh chup that: banks_shots/ (cong dev hrm-crm.eteksofts.com, 1440x900) — CHI DE LOCAL.

Chay:  python .plans/gop-db/banks-cut-mysql2/gen_srs.py
"""
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

sys.path.insert(0, os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "..", ".claude", "skills", "srs-documenter", "assets"))
from srs_docx_lib import SrsDoc  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, "banks_shots")
OUT = os.path.join(HERE, "SRS - Danh mục ngân hàng.docx")


def shot(name):
    return os.path.join(SHOTS, name)


ACTOR = 'Người dùng đã đăng nhập'
MENU = 'Phân hệ Danh mục chung → Ngân hàng'
FULL_URL = 'https://<host-hrm>/human/banks'


def sub(text):
    """Nhan muc con 2.x.y — ban mau in DAM (van la doan thuong, khong phai Heading)."""
    par = d.p()
    par.add_run(text).bold = True
    return par


def lay(shot=None, shot_caption=None, modal=None):
    """Muc "Layout man hinh" bam dung ban mau:
    dong dan nhap + 2 gach dau dong (Menu / URL day du) + anh chup that.

    Ban mau customer-docs/'SRS - Danh muc khach hang.docx' VAN GIU dong 'Menu: ...',
    chi bo dong 'Route (FE): ...' — nen khong dung lay() cua lib (ham do chi in URL).
    """
    d.p('Đường dẫn màn hình:')
    d.bullets(['Menu: %s' % MENU, 'URL đầy đủ: %s' % FULL_URL])
    if modal:
        d.p('Modal %s được mở ngay trên màn hình danh sách theo đường dẫn ở trên.' % modal)
    if shot:
        d.figure(shot, shot_caption or 'Màn hình thực tế', width_in=6.2)

d = SrsDoc(
    out=OUT,
    menu='Phân hệ Danh mục chung → Ngân hàng',
    route='/human/banks',
    full_url='https://<host-hrm>/human/banks',
    img_prefix='bank_')

# ================================================================ TRANG ĐẦU
d.title_block('Danh mục ngân hàng')
# Ban mau: 2 dong tieu de can giua 24pt nhung KHONG in dam -> bo bold cua lib.
for _p in d.doc.paragraphs[:2]:
    for _r in _p.runs:
        _r.bold = None

d.h2('Mục lục')
# Ban mau de truong TOC RONG (khong co dong huong dan 'Nhan chuot phai...'),
# nguoi dung bam Update Field trong Word de sinh muc luc.
d.toc(note='')

# ============================================================ PHẦN 1. GIỚI THIỆU
d.h1('Phần 1. Giới thiệu')

d.h2('1 Mục đích')
d.p('Tài liệu này đặc tả yêu cầu phần mềm cho màn hình Danh mục ngân hàng, nhằm:')
d.bullets([
    'Là căn cứ nghiệm thu chức năng và ràng buộc dữ liệu của màn hình.',
    'Làm rõ quan hệ hai cấp Ngân hàng → Chi nhánh và ảnh hưởng của danh mục này tới các màn tiêu '
    'thụ dữ liệu: hồ sơ nhân sự và danh mục tài khoản ngân hàng của công ty.',
    'Làm rõ điều kiện chặn xoá và khác biệt có chủ đích giữa thao tác Khóa và thao tác Xóa.',
    'Ghi nhận hiện trạng phân quyền: màn hình chưa khai báo quyền riêng nào, xem Phần 2.',
])

d.h2('2 Thuật ngữ và viết tắt')
d.table(['Thuật ngữ', 'Mô tả'], [
    ('Ngân hàng', 'Bản ghi cấp 1 của danh mục, định danh bằng Mã ngân hàng và Tên ngân hàng.'),
    ('Chi nhánh',
     'Bản ghi cấp 2, luôn thuộc đúng một ngân hàng; gồm Tên chi nhánh và Tỉnh/Thành phố.'),
    ('Tên giao dịch quốc tế', 'Tên tiếng Anh dùng trong giao dịch quốc tế của ngân hàng.'),
    ('Trạng thái Hoạt động', 'Ngân hàng đang dùng bình thường, sửa và xoá được.'),
    ('Trạng thái Khóa',
     'Ngân hàng ngừng sử dụng nhưng KHÔNG bị xoá; không sửa và không xoá được khi đang khóa.'),
    ('Tra cứu',
     'Chức năng lấy danh sách ngân hàng chuẩn từ dịch vụ ngoài để điền sẵn thông tin vào form.'),
    ('Tuỳ chỉnh cột', 'Cửa sổ cho mỗi người dùng tự chọn cột hiển thị và thứ tự cột của bảng.'),
    ('Lịch sử thay đổi',
     'Nhật ký thao tác của một bản ghi: ai tạo, sửa, khóa, mở khóa, xoá và giá trị trước - sau.'),
    ('SRS', 'Software Requirements Specification.'),
], widths=[1.8, 4.2])

# ============================================================ PHẦN 2. PHÂN QUYỀN
d.h1('Phần 2. Phân quyền')

d.h2('1 Danh sách quyền')

d.p('Nhóm quyền thao tác:')
d.table(['Ký hiệu', 'Tên quyền', 'Tác dụng trên màn hình'], [
    ('—', '(Màn hình không khai báo quyền riêng)',
     'Mọi tài khoản đã đăng nhập đều xem, tạo mới, sửa, khóa, mở khóa, xoá, quản lý chi nhánh và '
     'xem lịch sử được.'),
], widths=[0.8, 2.-1, 3.2])

d.p('Hiện trạng đã rà từ ba nguồn, cả ba đều thống nhất:')
d.bullets([
    'Khai báo quyền của hệ thống không có quyền nào cho danh mục ngân hàng; quyền “Quản lý danh '
    'mục tài khoản ngân hàng” thuộc màn Danh mục tài khoản ngân hàng của phân hệ Tài chính, không '
    'phải màn này.',
    'Toàn bộ 10 đường dẫn xử lý của màn chỉ yêu cầu đã đăng nhập, không gắn kiểm tra quyền.',
    'Khai báo menu để mục Ngân hàng hiển thị cho mọi người, không lọc theo quyền.',
    'Màn hình cũng không phân quyền theo cấp công ty / phòng ban / bộ phận: mọi người thấy chung '
    'một danh sách.',
])
d.p('⚠️ Rủi ro cần nghiệp vụ quyết định: mọi tài khoản đã đăng nhập đều sửa và xoá được danh mục '
    'dùng chung của toàn hệ thống. Muốn siết lại phải bổ sung quyền mới, gắn kiểm tra quyền cho các '
    'thao tác ghi dữ liệu và gắn khoá quyền vào khai báo menu; khi đó Phần 2 của tài liệu này phải '
    'cập nhật lại.')

d.h2('2 Ma trận phân quyền')
d.table(['Chức năng', 'Người dùng đã đăng nhập', 'Chưa đăng nhập / hết phiên'], [
    ('FR-01 Truy cập & xem danh sách', '✅', '❌'),
    ('FR-02 Tìm kiếm & lọc', '✅', '❌'),
    ('FR-03 Tuỳ chỉnh cột', '✅', '❌'),
    ('FR-04 Tạo mới ngân hàng', '✅', '❌'),
    ('FR-05 Tra cứu ngân hàng chuẩn', '✅', '❌'),
    ('FR-06 Sửa ngân hàng', '✅ (trừ ngân hàng đang Khóa)', '❌'),
    ('FR-07 Xem chi tiết ngân hàng', '✅', '❌'),
    ('FR-08 Khóa / Mở khóa', '✅', '❌'),
    ('FR-09 Xóa ngân hàng', '✅ (chỉ khi chưa được sử dụng)', '❌'),
    ('FR-10 Quản lý chi nhánh', '✅', '❌'),
    ('FR-11 Lịch sử thay đổi', '✅', '❌'),
], widths=[2.5, 1.9, 1.5])

# ================================================= PHẦN 3. ĐẶC TẢ CHI TIẾT
d.h1('Phần 3. Đặc tả chi tiết theo từng chức năng')

d.h2('1 Sơ đồ UML tổng quan')
d.overview_figure(
    'HỆ THỐNG HRM — DANH MỤC NGÂN HÀNG',
    [(ACTOR, list(range(11)))],
    [('FR-01', 'Truy cập & xem danh sách', 'view', None),
     ('FR-02', 'Tìm kiếm & Lọc', 'view', None),
     ('FR-03', 'Tuỳ chỉnh cột', 'view', None),
     ('FR-04', 'Tạo mới ngân hàng', 'crud', '«include» Kiểm tra trùng mã / tên'),
     ('FR-05', 'Tra cứu ngân hàng chuẩn', 'io', '«extend» Điền sẵn form tạo mới'),
     ('FR-06', 'Sửa ngân hàng', 'crud', None),
     ('FR-07', 'Xem chi tiết', 'view', None),
     ('FR-08', 'Khóa / Mở khóa', 'action', None),
     ('FR-09', 'Xóa ngân hàng', 'action', '«include» Kiểm tra ràng buộc sử dụng'),
     ('FR-10', 'Quản lý chi nhánh', 'crud', None),
     ('FR-11', 'Xem lịch sử thay đổi', 'view', None)],
    'Sơ đồ Use Case tổng quan màn Danh mục ngân hàng')

d.h2('2 Đặc tả chi tiết từng chức năng')

# ---------------------------------------------------------------- 2.1
d.h3('2.1 Xem danh sách ngân hàng')

sub('2.1.1 Giới thiệu')
d.intro_table(
    ten='Truy cập và xem danh sách ngân hàng',
    mota='Hiển thị danh sách ngân hàng kèm trạng thái, thông tin người tạo / người sửa và bộ nút '
         'thao tác của từng dòng.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đã đăng nhập vào hệ thống.',
    chinh='1. Người dùng vào menu Danh mục chung → Ngân hàng, hoặc gõ thẳng đường dẫn màn '
          'hình.\n'
          '2. Hệ thống nạp cấu hình cột hiển thị của tài khoản và khôi phục bộ lọc đã lưu nếu '
          'còn trong 10 phút.\n'
          '3. Hệ thống lấy danh sách ngân hàng theo bộ lọc và trang hiện tại.\n'
          '4. Hệ thống tính bộ nút Hành động cho từng dòng theo trạng thái và khả năng xóa.\n'
          '5. Người dùng chuyển trang hoặc đổi số dòng/trang nếu cần.',
    phu='• Không có dữ liệu → hiển thị “Không có dữ liệu phù hợp bộ lọc.”.\n'
        '• Ô dữ liệu trống (logo, tên giao dịch quốc tế, địa chỉ, người tạo…) → hiển thị dấu gạch '
        'ngang.\n'
        '• Ngân hàng đang Khóa → ẩn nút Sửa và nút Xóa.\n'
        '• Ngân hàng đang được sử dụng → ẩn nút Xóa.\n'
        '• Phiên đăng nhập hết hạn → điều hướng về màn đăng nhập.\n'
        '• Lỗi khi tải danh sách → hiển thị thông báo “Lỗi khi tải dữ liệu”, bảng để trống, '
        'không treo trang.',
    dacbiet=None)

sub('2.1.2 Layout màn hình')
lay(shot=shot('01-danh-sach.png'),
    shot_caption='Màn Danh mục ngân hàng lúc mới truy cập')
d.figure(shot('02-cot-hanh-dong.png'),
         'Bảng danh sách với đầy đủ cột Hành động khi thu gọn menu bên trái', width_in=6.2)

sub('2.1.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Bảng danh sách ngân hàng', 'Table/Grid', 'Read-only', '–', '–',
     'Hiển thị danh sách theo phân trang, cột Hành động chốt ở cuối bảng.'),
    ('Cột STT', 'Label', 'Read-only', '–', 'Số thứ tự theo trang',
     'Luôn hiển thị, không tắt được ở cửa sổ Tuỳ chỉnh cột.'),
    ('Cột Mã ngân hàng', 'Button', 'Enable', '0–255 ký tự', 'Theo dữ liệu',
     'Cột định danh, bấm vào mở cửa sổ Xem chi tiết. Luôn hiển thị, không tắt được.'),
    ('Cột Tên ngân hàng', 'Text', 'Read-only', '0–255 ký tự', 'Theo dữ liệu',
     'Tự xuống dòng khi nội dung dài.'),
    ('Cột Logo', 'Icon Button', 'Read-only', '–', 'Ẩn',
     'Ảnh logo cao 32 px, giữ tỉ lệ; mặc định tắt, bật ở cửa sổ Tuỳ chỉnh cột.'),
    ('Cột Tên viết tắt', 'Text', 'Read-only', '0–255 ký tự', 'Ẩn', 'Mặc định tắt.'),
    ('Cột Tên giao dịch quốc tế', 'Text', 'Read-only', '0–255 ký tự', 'Ẩn', 'Mặc định tắt.'),
    ('Cột Địa chỉ giao dịch', 'Text', 'Read-only', '0–255 ký tự', 'Ẩn', 'Mặc định tắt.'),
    ('Cột Chi nhánh', 'Button', 'Enable', '≥ 0', 'Ẩn',
     'Số chi nhánh của ngân hàng, bấm vào mở cửa sổ Chi nhánh; mặc định tắt.'),
    ('Cột Người tạo', 'Text', 'Read-only', '–', 'Theo dữ liệu', 'Họ tên người tạo bản ghi.'),
    ('Cột Ngày tạo', 'Text', 'Read-only', 'dd/mm/yyyy HH:mm', 'Theo dữ liệu',
     'Ngày giờ tạo bản ghi.'),
    ('Cột Người sửa', 'Text', 'Read-only', '–', 'Ẩn',
     'Họ tên người sửa gần nhất; mặc định tắt. Bản ghi chưa từng sửa thì để trống.'),
    ('Cột Ngày cập nhật', 'Text', 'Read-only', 'dd/mm/yyyy HH:mm', 'Ẩn', 'Mặc định tắt.'),
    ('Cột Trạng thái', 'Badge', 'Read-only', 'Hoạt động / Khóa', 'Theo dữ liệu',
     'Nhãn xanh cho Hoạt động, nhãn đỏ cho Khóa; nhãn do hệ thống trả về, giao diện không tự quy '
     'đổi số sang chữ.'),
    ('Cột Hành động', 'Icon Button', 'Enable / Ẩn', '–', 'Hiển thị',
     'Tối đa 2 nút chính và menu ba chấm; nút không dùng được thì ẩn hẳn chứ không làm mờ.'),
    ('Phân trang', 'Pagination', 'Enable', '–', 'Trang 1',
     'Có nút về đầu / lùi / số trang / tiến / về cuối.'),
    ('Ô Số dòng/trang', 'Dropdown', 'Enable', '5 / 10 / 20 / 50 / 100', '10',
     'Đổi số dòng thì danh sách quay về trang 1.'),
    ('Ô “Hiển thị a–b / N”', 'Label', 'Read-only', '–', 'Theo kết quả',
     'N là tổng số ngân hàng khớp bộ lọc, không phải tổng toàn danh mục.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', 'Ẩn',
     'Nội dung “Không có dữ liệu phù hợp bộ lọc.”.'),
    ('Vòng quay chờ', 'Loading', 'Hiển thị', '–', 'Ẩn', 'Hiện trong lúc nạp danh sách.'),
], required=False)

sub('2.1.4 Danh sách event và xử lý event')
d.event_table([
    ('Click menu Ngân hàng', 'Click', 'Điều hướng tới màn hình Danh mục ngân hàng.'),
    ('Mở màn hình', 'System',
     'Before:\n– Kiểm tra người dùng đã đăng nhập; hết phiên thì đưa về màn đăng nhập.\n'
     'During:\n– Nạp cấu hình cột của tài khoản, khôi phục bộ lọc đã lưu nếu chưa quá 10 phút.\n'
     'After:\n– Tải danh sách trang 1 với 10 dòng/trang và hiển thị bảng.'),
    ('Load danh sách', 'System',
     'Lấy danh sách theo bộ lọc, trang và số dòng/trang hiện tại; hiển thị vòng quay chờ trong lúc '
     'tải.'),
    ('Đổi trang', 'Click', 'Tải lại danh sách theo trang mới, giữ nguyên bộ lọc.'),
    ('Đổi số dòng/trang', 'Change', 'Đặt lại về trang 1 và tải lại danh sách.'),
    ('Bấm tiêu đề cột có mũi tên', 'Click',
     'Đổi cột và chiều sắp xếp rồi tải lại danh sách. ⚠️ Hiện tại danh sách KHÔNG đổi thứ tự do '
     'tên tham số sắp xếp giữa giao diện và máy chủ không khớp — lỗi đang chờ xử lý, xem BR-09.'),
    ('Bấm mã ngân hàng', 'Click', 'Mở cửa sổ Xem ngân hàng ở chế độ chỉ đọc.'),
    ('Bấm số ở cột Chi nhánh', 'Click', 'Mở cửa sổ Chi nhánh ngân hàng của đúng dòng đó.'),
])

# ---------------------------------------------------------------- 2.2
d.h3('2.2 Tìm kiếm và lọc danh sách')

sub('2.2.1 Giới thiệu')
d.intro_table(
    ten='Tìm kiếm và lọc danh sách ngân hàng',
    mota='Thu hẹp danh sách theo từ khoá mã / tên, theo tên giao dịch quốc tế và theo trạng thái.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang ở màn hình danh sách ngân hàng.',
    chinh='1. Người dùng nhập từ khoá vào ô tìm nhanh và/hoặc nhập, chọn các ô lọc.\n'
          '2. Người dùng bấm Tìm kiếm hoặc nhấn Enter trong ô tìm nhanh.\n'
          '3. Hệ thống đặt lại về trang 1 và tải danh sách khớp toàn bộ điều kiện.\n'
          '4. Hệ thống ghi nhớ bộ lọc trong 10 phút để dùng lại khi quay về màn hình.',
    phu='• Thay đổi ô Tên giao dịch quốc tế hoặc Trạng thái → hệ thống tự tải lại ngay, không cần '
        'bấm Tìm kiếm.\n'
        '• Chỉ gõ từ khoá mà chưa bấm Tìm kiếm / Enter → danh sách giữ nguyên.\n'
        '• Bấm Làm mới → xoá toàn bộ điều kiện và tải lại danh sách đầy đủ.\n'
        '• Không có kết quả → hiển thị trạng thái rỗng.',
    dacbiet='Tổng số ô lọc kể cả ô tìm nhanh là 3 nên bộ lọc bày hết trên một hàng ngang; màn hình '
            'không có nút “Tìm kiếm nâng cao” và không có cửa sổ “Cài đặt bộ lọc”.')

sub('2.2.2 Layout màn hình')
lay(shot=shot('03-tim-kiem.png'),
         shot_caption='Kết quả sau khi tìm nhanh theo từ khoá “Vietcom”')

sub('2.2.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Khối Bộ lọc danh sách', 'Modal', 'Enable', '–', '–', 'Mở sẵn',
     'Hàng ngang gồm 3 ô lọc và 2 nút.'),
    ('Ô tìm nhanh', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Gợi ý “Tìm theo mã, tên ngân hàng...”; khớp gần đúng đồng thời Mã và Tên; có nút x để xoá '
     'nhanh từ khoá.'),
    ('Ô Tên giao dịch quốc tế', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Khớp gần đúng theo tên giao dịch quốc tế.'),
    ('Ô Trạng thái', 'Dropdown', 'Enable', 'Hoạt động / Khóa', 'Không', 'Trống',
     'Để trống là lấy tất cả; khớp chính xác.'),
    ('Nút Tìm kiếm', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Áp dụng bộ lọc và quay về trang 1.'),
    ('Nút Làm mới', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Xoá toàn bộ điều kiện lọc, đặt lại trang 1 và tải lại danh sách.'),
])

sub('2.2.4 Danh sách event và xử lý event')
d.event_table([
    ('Nhập ô tìm nhanh', 'Keypress', 'Ghi nhận từ khoá, chưa tải lại danh sách.'),
    ('Nhấn Enter trong ô tìm nhanh', 'Keypress', 'Đặt lại trang 1 và tải lại danh sách.'),
    ('Bấm Tìm kiếm', 'Click', 'Đặt lại trang 1 và tải lại danh sách theo toàn bộ điều kiện.'),
    ('Đổi ô Tên giao dịch quốc tế hoặc Trạng thái', 'Change',
     'Đặt lại trang 1 và tải lại danh sách ngay.'),
    ('Bấm Làm mới', 'Click',
     'After:\n– Đặt toàn bộ ô lọc về giá trị khởi tạo và đặt lại mốc so sánh.\n'
     '– Tự tải lại danh sách đầy đủ, không dựa vào cơ chế theo dõi thay đổi (tránh trường hợp chỉ '
     'xoá ô lọc mà kết quả cũ vẫn nằm trên bảng).'),
    ('Rời màn hình rồi quay lại', 'System',
     'Khôi phục bộ lọc đã lưu nếu chưa quá 10 phút, ngược lại dùng bộ lọc rỗng.'),
])

# ---------------------------------------------------------------- 2.3
d.h3('2.3 Tuỳ chỉnh cột hiển thị')

sub('2.3.1 Biểu đồ Usecase')
d.uc_figure('FR-03', 'Tuỳ chỉnh cột hiển thị', 'view',
            [('include', 'Ghi nhớ cấu hình theo người dùng')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-03 Tuỳ chỉnh cột hiển thị')

sub('2.3.2 Giới thiệu')
d.intro_table(
    ten='Tuỳ chỉnh cột hiển thị',
    mota='Cho phép mỗi tài khoản tự chọn cột nào hiển thị trên bảng và thứ tự các cột.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang ở màn hình danh sách ngân hàng.',
    chinh='1. Người dùng bấm nút Tuỳ chỉnh cột hiển thị bên cạnh nút Tạo mới.\n'
          '2. Hệ thống mở cửa sổ “Tuỳ chỉnh cột” liệt kê 13 cột với trạng thái bật/tắt hiện tại.\n'
          '3. Người dùng tích hoặc bỏ tích cột, hoặc kéo thả để đổi vị trí.\n'
          '4. Người dùng bấm Lưu; hệ thống lưu cấu hình theo tài khoản và vẽ lại bảng.',
    phu='• Bấm Đóng → đóng cửa sổ, bỏ mọi thay đổi chưa lưu.\n'
        '• Cột STT, Mã ngân hàng và Hành động bị khoá → không bỏ tích được.',
    dacbiet='Cấu hình lưu riêng theo từng tài khoản và từng màn hình nên không ảnh hưởng tới người '
            'dùng khác.')

sub('2.3.3 Layout màn hình')
lay(modal='Tuỳ chỉnh cột', shot=shot('04-cau-hinh-cot.png'),
         shot_caption='Cửa sổ Tuỳ chỉnh cột — cột STT và Mã ngân hàng bị khoá')

sub('2.3.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Tuỳ chỉnh cột hiển thị', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Nằm cạnh nút Tạo mới, chỉ có biểu tượng cột.'),
    ('Danh sách cột', 'Table/Grid', 'Enable', '13 dòng', '–', 'Theo cấu hình đã lưu',
     'Mỗi dòng gồm ô tích chọn, tên cột và tay nắm kéo thả.'),
    ('Ô tích chọn cột', 'Icon Button', 'Enable / Disable', '–', 'Không', 'Theo cấu hình đã lưu',
     'Cột khoá (STT, Mã ngân hàng, Hành động) hiển thị ổ khoá và không đổi được.'),
    ('Tay nắm kéo thả', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Kéo để đổi thứ tự cột trên bảng.'),
    ('Nút Lưu', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Lưu cấu hình theo tài khoản và đóng cửa sổ.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Đóng cửa sổ, bỏ thay đổi chưa lưu.'),
])

sub('2.3.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Tuỳ chỉnh cột hiển thị', 'Click', 'Mở cửa sổ Tuỳ chỉnh cột với cấu hình hiện tại.'),
    ('Tích hoặc bỏ tích một cột', 'Change',
     'Cập nhật trạng thái tạm trong cửa sổ, chưa áp dụng lên bảng.'),
    ('Kéo thả đổi vị trí', 'Change', 'Cập nhật thứ tự tạm trong cửa sổ.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Bỏ qua thay đổi trên cột bị khoá.\n'
     'After:\n– Lưu cấu hình theo tài khoản.\n'
     '– Vẽ lại bảng theo cấu hình mới và đóng cửa sổ.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ và bỏ toàn bộ thay đổi chưa lưu.'),
])

# ---------------------------------------------------------------- 2.4
d.h3('2.4 Tạo mới ngân hàng')

sub('2.4.1 Biểu đồ Usecase')
d.uc_figure('FR-04', 'Tạo mới ngân hàng', 'crud',
            [('include', 'Kiểm tra bắt buộc và trùng mã / tên'),
             ('extend', 'Tra cứu ngân hàng chuẩn để điền sẵn'),
             ('extend', 'Tải logo lên')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-04 Tạo mới ngân hàng')

sub('2.4.2 Giới thiệu')
d.intro_table(
    ten='Tạo mới ngân hàng',
    mota='Thêm một ngân hàng mới vào danh mục dùng chung của hệ thống.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang ở màn hình danh sách ngân hàng.',
    chinh='1. Người dùng bấm nút Tạo mới.\n'
          '2. Hệ thống mở cửa sổ “Tạo ngân hàng” với toàn bộ ô nhập để trống.\n'
          '3. Người dùng nhập Mã ngân hàng, Tên ngân hàng, Tên viết tắt và các thông tin tuỳ chọn; '
          'có thể tải logo lên.\n'
          '4. Người dùng bấm Lưu.\n'
          '5. Hệ thống kiểm tra dữ liệu, ghi bản ghi mới với trạng thái Hoạt động và ghi một dòng '
          'lịch sử “Tạo mới”.\n'
          '6. Hệ thống hiển thị “Đã lưu thành công!”, đóng cửa sổ và tải lại danh sách.',
    phu='• Bỏ trống Tên ngân hàng → giao diện chặn ngay, hiện lỗi đỏ dưới ô và không gửi dữ liệu '
        'đi.\n'
        '• Bỏ trống Mã ngân hàng hoặc Tên viết tắt → hệ thống trả lỗi, hiện lỗi đỏ dưới đúng ô.\n'
        '• Trùng mã hoặc trùng tên → hiện lỗi đỏ tương ứng dưới ô, cửa sổ không đóng.\n'
        '• Bấm “Lưu và tiếp tục” → lưu xong giữ cửa sổ mở và xoá trắng các ô để nhập bản ghi kế '
        'tiếp.\n'
        '• Bấm Đóng → đóng cửa sổ, không lưu.',
    dacbiet='Cửa sổ Tạo mới KHÔNG có ô Trạng thái: bản ghi mới luôn ở trạng thái Hoạt động. Người '
            'tạo và Ngày tạo do hệ thống tự ghi.')

sub('2.4.3 Layout màn hình')
lay(modal='Tạo ngân hàng', shot=shot('05-tao-moi.png'),
         shot_caption='Cửa sổ Tạo ngân hàng khi mới mở (mọi ô còn trống)')
d.figure(shot('07-loi-trung-ma.png'),
         'Cửa sổ Tạo ngân hàng khi dữ liệu không hợp lệ: lỗi bắt buộc và lỗi trùng tên hiện đỏ ngay '
         'dưới từng ô', width_in=6.2)

sub('2.4.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '–', '“Tạo ngân hàng”',
     'Đổi thành “Sửa ngân hàng” / “Xem ngân hàng” tuỳ chế độ mở.'),
    ('Ô Gợi ý', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Từ khoá dùng cho chức năng Tra cứu; chỉ hiện ở chế độ Tạo mới và Sửa.'),
    ('Nút Tra cứu', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Mở cửa sổ danh sách ngân hàng chuẩn, xem mục 2.5.'),
    ('Mã ngân hàng', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Duy nhất toàn hệ thống; trống thì báo “Bắt buộc phải nhập”, trùng thì báo “Mã ngân hàng này '
     'đã tồn tại”.'),
    ('Tên ngân hàng', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Duy nhất toàn hệ thống; giao diện chặn trước khi gửi nếu để trống; trùng thì báo “Tên ngân '
     'hàng này đã tồn tại”.'),
    ('Tên viết tắt', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Không yêu cầu duy nhất; trống thì báo “Bắt buộc phải nhập”.'),
    ('Tên giao dịch quốc tế', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Thông tin tuỳ chọn.'),
    ('Địa chỉ giao dịch', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Thông tin tuỳ chọn.'),
    ('Khung Logo', 'Icon Button', 'Enable', '.jpg / .jpeg / .png, ≤ 5 MB', 'Không',
     '“Chưa có logo”', 'Khung xem trước 311×80; có chú thích định dạng ở biểu tượng cạnh nhãn.'),
    ('Nút Tải ảnh lên', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Sai định dạng báo “File không hợp lệ”; quá dung lượng báo “Dung lượng tối đa: 5MB”.'),
    ('Nút Xóa ảnh', 'Button', 'Enable / Ẩn', '–', '–', 'Ẩn',
     'Chỉ hiện khi đã có ảnh trong khung xem trước.'),
    ('Nút Lưu', 'Button', 'Enable / Disable', '–', '–', 'Enable',
     'Bị khóa trong lúc đang lưu để không tạo trùng bản ghi.'),
    ('Nút Lưu và tiếp tục', 'Button', 'Enable / Ẩn', '–', '–', 'Enable',
     'Chỉ hiện ở chế độ Tạo mới.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ, không lưu.'),
    ('Thông báo lỗi inline', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     'Chữ đỏ kèm biểu tượng cảnh báo, hiện ngay dưới ô tương ứng.'),
    ('Thông báo chung', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     '“Đã lưu thành công!” hoặc “Vui lòng kiểm tra lại thông tin”.'),
])

sub('2.4.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Tạo mới', 'Click', 'Đặt lại toàn bộ ô nhập về trống rồi mở cửa sổ Tạo ngân hàng.'),
    ('Chọn tệp logo', 'Change',
     'During:\n– Sai định dạng → hiện “File không hợp lệ”, bỏ ảnh.\n'
     '– Quá 5 MB → hiện “Dung lượng tối đa: 5MB”, bỏ ảnh.\n'
     'After:\n– Hợp lệ → hiện ảnh trong khung xem trước và bật nút Xóa ảnh.'),
    ('Bấm Xóa ảnh', 'Click', 'Xoá ảnh khỏi khung xem trước và bỏ tệp đã chọn.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Kiểm tra ô Tên ngân hàng; trống → hiện lỗi đỏ dưới ô, hiện thông báo “Vui lòng '
     'kiểm tra lại thông tin” và dừng xử lý.\n'
     '– Khóa nút Lưu và bật hiệu ứng chờ.\n'
     'During:\n– Nếu có tệp logo mới → tải ảnh lên trước rồi gán đường dẫn ảnh vào dữ liệu gửi '
     'đi.\n'
     '– Mã ngân hàng trống → hiển thị “Bắt buộc phải nhập” dưới ô Mã ngân hàng.\n'
     '– Tên viết tắt trống → hiển thị “Bắt buộc phải nhập” dưới ô Tên viết tắt.\n'
     '– Mã ngân hàng trùng → hiển thị “Mã ngân hàng này đã tồn tại”.\n'
     '– Tên ngân hàng trùng → hiển thị “Tên ngân hàng này đã tồn tại”.\n'
     '– Nếu có lỗi kiểm tra → không thực hiện bước After, cửa sổ vẫn mở, dữ liệu đã nhập giữ '
     'nguyên.\n'
     'After:\n– Ghi bản ghi mới với trạng thái Hoạt động.\n'
     '– Ghi một dòng lịch sử “Tạo mới”.\n'
     '– Hiển thị “Đã lưu thành công!”, đóng cửa sổ và tải lại danh sách.'),
    ('Bấm Lưu và tiếp tục', 'Click',
     'Xử lý như nút Lưu; sau khi lưu thành công thì giữ cửa sổ mở, xoá trắng các ô và tải lại danh '
     'sách nền.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ và đặt lại dữ liệu tạm của cửa sổ.'),
])

# ---------------------------------------------------------------- 2.5
d.h3('2.5 Tra cứu ngân hàng chuẩn để điền sẵn')

sub('2.5.1 Biểu đồ Usecase')
d.uc_figure('FR-05', 'Tra cứu ngân hàng chuẩn', 'io',
            [('include', 'Lọc theo từ khoá không dấu'),
             ('extend', 'Điền sẵn mã, tên, tên viết tắt, logo')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-05 Tra cứu ngân hàng chuẩn')

sub('2.5.2 Giới thiệu')
d.intro_table(
    ten='Tra cứu ngân hàng chuẩn',
    mota='Lấy thông tin ngân hàng từ danh mục ngân hàng chuẩn bên ngoài để điền sẵn vào form, '
         'tránh gõ tay sai mã và sai tên.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang mở cửa sổ Tạo ngân hàng hoặc Sửa ngân hàng.',
    chinh='1. Người dùng nhập từ khoá vào ô Gợi ý (không bắt buộc).\n'
          '2. Người dùng bấm nút Tra cứu.\n'
          '3. Hệ thống mở cửa sổ “Thông tin ngân hàng” hiển thị danh sách ngân hàng chuẩn đã lọc '
          'theo từ khoá.\n'
          '4. Người dùng bấm vào một dòng.\n'
          '5. Hệ thống điền sẵn Mã ngân hàng, Tên ngân hàng, Tên viết tắt và Logo vào form rồi '
          'đóng cửa sổ tra cứu.',
    phu='• Không lấy được dữ liệu từ dịch vụ ngoài → bảng hiện “Chưa có dữ liệu”, form không bị '
        'ảnh hưởng.\n'
        '• Bấm Đóng → chỉ đóng cửa sổ tra cứu, dữ liệu đang nhập dở trong form giữ nguyên.',
    dacbiet='Danh sách tra cứu là dữ liệu NGOÀI hệ thống, không phải danh mục hiện có; chọn một '
            'dòng chỉ điền sẵn form chứ chưa ghi dữ liệu.')

sub('2.5.3 Layout màn hình')
lay(modal='Thông tin ngân hàng', shot=shot('08-tra-cuu.png'),
         shot_caption='Cửa sổ Thông tin ngân hàng — bấm một dòng để điền nhanh vào form')

sub('2.5.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '–', '“Thông tin ngân hàng”',
     'Kèm dòng phụ “Click vào một dòng để điền nhanh thông tin”.'),
    ('Bảng ngân hàng chuẩn', 'Table/Grid', 'Enable', 'Danh sách', '–', 'Lấy từ dịch vụ ngoài',
     '5 cột: STT, Logo, Tên ngân hàng, Mã ngân hàng, Tên viết tắt; bấm cả dòng để chọn.'),
    ('Vòng quay chờ', 'Loading', 'Hiển thị', '–', '–', 'Ẩn', 'Chữ “Đang tải dữ liệu...”.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', '–', 'Ẩn', 'Chữ “Chưa có dữ liệu”.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ tra cứu.'),
])

sub('2.5.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm Tra cứu', 'Click', 'Mở cửa sổ tra cứu và lọc danh sách theo từ khoá đang có ở ô Gợi ý.'),
    ('Đổi từ khoá ở ô Gợi ý', 'Change',
     'Lọc lại danh sách trong cửa sổ tra cứu theo từ khoá mới, bỏ dấu tiếng Việt và không phân '
     'biệt hoa thường.'),
    ('Bấm một dòng trong bảng', 'Click',
     'After:\n– Điền Mã ngân hàng, Tên ngân hàng, Tên viết tắt và Logo vào form.\n'
     '– Đóng cửa sổ tra cứu; bản ghi chỉ được tạo khi người dùng bấm Lưu.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ tra cứu, giữ nguyên dữ liệu form.'),
])

# ---------------------------------------------------------------- 2.6
d.h3('2.6 Sửa ngân hàng')

sub('2.6.1 Biểu đồ Usecase')
d.uc_figure('FR-06', 'Sửa ngân hàng', 'crud',
            [('include', 'Kiểm tra trùng mã / tên, loại trừ chính bản ghi'),
             ('include', 'Ghi lịch sử thay đổi thông tin')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-06 Sửa ngân hàng')

sub('2.6.2 Giới thiệu')
d.intro_table(
    ten='Sửa ngân hàng',
    mota='Cập nhật thông tin của một ngân hàng đang ở trạng thái Hoạt động.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Ngân hàng cần sửa đang ở trạng thái Hoạt động; ngân hàng đang Khóa không hiển thị '
             'nút Sửa.',
    chinh='1. Người dùng bấm nút Sửa ở dòng tương ứng.\n'
          '2. Hệ thống nạp dữ liệu hiện tại của ngân hàng và mở cửa sổ “Sửa ngân hàng”.\n'
          '3. Người dùng chỉnh sửa thông tin cần thiết.\n'
          '4. Người dùng bấm Lưu.\n'
          '5. Hệ thống kiểm tra dữ liệu, cập nhật bản ghi và ghi một dòng lịch sử “Thay đổi thông '
          'tin” kèm giá trị cũ → giá trị mới.\n'
          '6. Hệ thống hiển thị “Đã lưu thành công!”, đóng cửa sổ và tải lại danh sách.',
    phu='• Vi phạm bắt buộc hoặc trùng mã / tên → hiện lỗi đỏ dưới ô tương ứng, cửa sổ không '
        'đóng.\n'
        '• Giữ nguyên mã / tên của chính bản ghi → không bị báo trùng.\n'
        '• Không thay đổi gì rồi bấm Lưu → vẫn báo lưu thành công nhưng KHÔNG phát sinh mốc lịch '
        'sử.\n'
        '• Bản ghi đã bị người khác xoá → hệ thống báo lỗi, không tạo lại bản ghi.',
    dacbiet='Cửa sổ Sửa không có nút “Lưu và tiếp tục”. Thay đổi Logo KHÔNG được ghi vào lịch sử '
            'vì logo không nằm trong danh sách trường theo dõi.')

sub('2.6.3 Layout màn hình')
lay(modal='Sửa ngân hàng', shot=shot('10-sua.png'),
         shot_caption='Cửa sổ Sửa ngân hàng với dữ liệu đã nạp sẵn')

sub('2.6.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '–', '“Sửa ngân hàng”',
     'Khác chế độ Tạo mới ở tiêu đề.'),
    ('Các ô thông tin', 'Textbox', 'Enable', '0–255 ký tự', 'Theo mục 2.4', 'Theo dữ liệu',
     'Nạp sẵn dữ liệu hiện tại của bản ghi; ràng buộc bắt buộc và không trùng giống Tạo mới.'),
    ('Khung Logo', 'Icon Button', 'Enable', '.jpg / .jpeg / .png, ≤ 5 MB', 'Không',
     'Logo hiện tại', 'Có nút Tải ảnh lên và nút Xóa ảnh.'),
    ('Nút Lưu', 'Button', 'Enable / Disable', '–', '–', 'Enable',
     'Bị khóa trong lúc đang lưu.'),
    ('Nút Lưu và tiếp tục', 'Button', 'Ẩn', '–', '–', 'Ẩn', 'Không hiển thị ở chế độ Sửa.'),
    ('Nút Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Đóng cửa sổ, không lưu.'),
])

sub('2.6.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Sửa', 'Click',
     'Gán bản ghi đang chọn rồi mở cửa sổ ở chế độ sửa; cửa sổ tự gọi lấy chi tiết và nạp vào '
     'form.'),
    ('Bấm Lưu', 'Click',
     'Before:\n– Kiểm tra ô Tên ngân hàng; trống → hiện lỗi đỏ và dừng xử lý.\n'
     'During:\n– Kiểm tra bắt buộc và trùng mã / tên, có loại trừ chính bản ghi đang sửa.\n'
     '– Nếu có lỗi → hiển thị lỗi đỏ dưới từng ô, không thực hiện bước After.\n'
     'After:\n– Chụp giá trị trước khi ghi rồi cập nhật bản ghi.\n'
     '– Ghi một dòng lịch sử “Thay đổi thông tin” nếu có ít nhất một trường theo dõi thay đổi.\n'
     '– Hiển thị “Đã lưu thành công!”, đóng cửa sổ và tải lại danh sách.'),
    ('Bấm Đóng', 'Click', 'Đóng cửa sổ, không ghi dữ liệu.'),
])

# ---------------------------------------------------------------- 2.7
d.h3('2.7 Xem chi tiết ngân hàng')

sub('2.7.1 Giới thiệu')
d.intro_table(
    ten='Xem chi tiết ngân hàng',
    mota='Xem toàn bộ thông tin của một ngân hàng ở chế độ chỉ đọc, kèm khối lịch sử thay đổi.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang ở màn hình danh sách ngân hàng.',
    chinh='1. Người dùng bấm vào mã ngân hàng ở cột định danh.\n'
          '2. Hệ thống mở cửa sổ “Xem ngân hàng” và nạp dữ liệu của bản ghi.\n'
          '3. Người dùng xem thông tin; có thể bấm “Xem lịch sử” để mở khối lịch sử ngay trong cửa '
          'sổ.\n'
          '4. Người dùng bấm Đóng để thoát.',
    phu='• Bản ghi đang ở trạng thái Khóa → vẫn xem được, ô Trạng thái hiển thị “Khoá”.\n'
        '• Bản ghi chưa có lịch sử → khối lịch sử hiển thị “Chưa có lịch sử thao tác nào.”.',
    dacbiet=None)

sub('2.7.2 Layout màn hình')
lay(modal='Xem ngân hàng', shot=shot('12-xem-lich-su.png'),
         shot_caption='Cửa sổ Xem ngân hàng với khối Lịch sử đã mở')

sub('2.7.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '“Xem ngân hàng”', 'Tiêu đề của cửa sổ chỉ đọc.'),
    ('Mã ngân hàng / Tên ngân hàng / Tên viết tắt', 'Textbox', 'Disable', '–', 'Theo dữ liệu',
     'Chỉ đọc, không sửa được.'),
    ('Tên giao dịch quốc tế / Địa chỉ giao dịch', 'Textbox', 'Disable', '–', 'Theo dữ liệu',
     'Chỉ đọc; trống thì hiện chữ gợi ý nhập màu mờ.'),
    ('Trạng thái', 'Textbox', 'Disable', 'Hoạt động / Khoá', 'Theo dữ liệu',
     'Ô này chỉ xuất hiện ở chế độ Xem.'),
    ('Khung Logo', 'Icon Button', 'Read-only', '–', 'Theo dữ liệu',
     'Chỉ hiện khung xem trước, không có nút tải hoặc xoá ảnh.'),
    ('Khối Lịch sử', 'Table/Grid', 'Enable', '–', 'Thu gọn',
     'Bấm “Xem lịch sử” để mở; có nút Làm mới và nút Thu gọn.'),
    ('Nút Đóng', 'Button', 'Enable', '–', 'Hiển thị', 'Nút duy nhất ở chân cửa sổ.'),
], required=False)

sub('2.7.4 Danh sách event và xử lý event')
d.event_table([
    ('Bấm mã ngân hàng', 'Click', 'Bật cờ chế độ xem, gán bản ghi rồi mở cửa sổ.'),
    ('Bấm Xem lịch sử', 'Click', 'Gọi lấy lịch sử của bản ghi và mở khối lịch sử.'),
    ('Bấm Đóng', 'Click',
     'Đóng cửa sổ và đặt lại cờ chế độ xem để lần mở sau không bị kẹt ở chế độ chỉ đọc.'),
])

# ---------------------------------------------------------------- 2.8
d.h3('2.8 Khóa và Mở khóa ngân hàng')

sub('2.8.1 Biểu đồ Usecase')
d.uc_figure('FR-08', 'Khóa / Mở khóa ngân hàng', 'action',
            [('include', 'Xác nhận thao tác'),
             ('include', 'Ghi lịch sử thay đổi trạng thái')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-08 Khóa / Mở khóa ngân hàng')

sub('2.8.2 Giới thiệu')
d.intro_table(
    ten='Khóa / Mở khóa ngân hàng',
    mota='Ngừng sử dụng một ngân hàng mà vẫn giữ nguyên dữ liệu, hoặc đưa ngân hàng đã khóa trở '
         'lại sử dụng.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Ngân hàng cần thao tác đang hiển thị trên danh sách.',
    chinh='1. Người dùng bấm nút Khóa hoặc Mở khóa ở dòng tương ứng.\n'
          '2. Hệ thống mở hộp thoại xác nhận có nêu tên ngân hàng.\n'
          '3. Người dùng bấm nút xác nhận.\n'
          '4. Hệ thống đổi trạng thái và ghi một dòng lịch sử nhóm “Thay đổi trạng thái”.\n'
          '5. Hệ thống hiển thị “Khoá thành công” hoặc “Mở khoá thành công” và tải lại danh sách.',
    phu='• Bấm Hủy → đóng hộp thoại, không thay đổi gì.\n'
        '• Lỗi khi xử lý → hiển thị thông báo lỗi trả về, trạng thái không đổi.',
    dacbiet='⚠️ Khóa KHÔNG kiểm tra ràng buộc sử dụng: ngân hàng đang được dùng ở hồ sơ nhân sự vẫn '
            'khóa được. Đây là khác biệt có chủ đích so với thao tác Xóa.')

sub('2.8.3 Layout màn hình')
lay(shot=shot('18-xac-nhan-khoa.png'),
         shot_caption='Hộp thoại Xác nhận khóa có nêu rõ tên ngân hàng')

sub('2.8.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Khóa', 'Icon Button', 'Enable / Ẩn', 'Hiển thị khi Hoạt động',
     'Biểu tượng ổ khoá, chỉ hiện với ngân hàng đang Hoạt động.'),
    ('Nút Mở khóa', 'Icon Button', 'Enable / Ẩn', 'Hiển thị khi Khóa',
     'Biểu tượng ổ khoá mở, chỉ hiện với ngân hàng đang Khóa.'),
    ('Hộp thoại xác nhận', 'Modal', 'Hiển thị', 'Ẩn',
     'Tiêu đề “Xác nhận khóa” hoặc “Xác nhận mở khóa”, nội dung nêu tên ngân hàng.'),
    ('Nút xác nhận', 'Button', 'Enable', 'Hiển thị', 'Nhãn “Khóa” hoặc “Mở khóa”.'),
    ('Nút Hủy', 'Button', 'Enable', 'Hiển thị', 'Đóng hộp thoại, không thực hiện gì.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', 'Ẩn',
     '“Khoá thành công” / “Mở khoá thành công” / “Thao tác thất bại”.'),
], required=False, scope=False)

sub('2.8.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm Khóa hoặc Mở khóa ở dòng', 'Click',
     'Ghi nhận dòng đang thao tác và mở hộp thoại xác nhận.'),
    ('Bấm nút xác nhận', 'Click',
     'Before:\n– Kiểm tra đã có dòng được chọn; chưa có thì dừng.\n'
     'During:\n– Gọi chức năng khóa hoặc mở khóa theo trạng thái hiện tại của dòng.\n'
     '– Lỗi → hiển thị thông báo lỗi trả về, giữ nguyên trạng thái.\n'
     'After:\n– Cập nhật trạng thái và ghi một dòng lịch sử nhóm Thay đổi trạng thái.\n'
     '– Hiển thị thông báo thành công và tải lại danh sách.'),
    ('Bấm Hủy', 'Click', 'Đóng hộp thoại, không gọi xử lý.'),
])

# ---------------------------------------------------------------- 2.9
d.h3('2.9 Xóa ngân hàng')

sub('2.9.1 Biểu đồ Usecase')
d.uc_figure('FR-09', 'Xóa ngân hàng', 'action',
            [('include', 'Kiểm tra ràng buộc sử dụng'),
             ('include', 'Xóa toàn bộ chi nhánh trực thuộc')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-09 Xóa ngân hàng')

sub('2.9.2 Giới thiệu')
d.intro_table(
    ten='Xóa ngân hàng',
    mota='Xoá hẳn một ngân hàng khỏi danh mục khi ngân hàng đó chưa được sử dụng ở bất kỳ đâu.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Ngân hàng đang ở trạng thái Hoạt động và chưa được tham chiếu ở hồ sơ nhân sự hay ở '
             'danh mục tài khoản ngân hàng của công ty.',
    chinh='1. Người dùng bấm nút Xóa ở dòng tương ứng.\n'
          '2. Hệ thống mở hộp thoại “Xác nhận xóa” nêu tên ngân hàng.\n'
          '3. Người dùng bấm Xóa.\n'
          '4. Hệ thống kiểm tra lại ràng buộc sử dụng ở phía máy chủ.\n'
          '5. Hệ thống ghi một dòng lịch sử “Xóa”, xoá toàn bộ chi nhánh rồi xoá ngân hàng.\n'
          '6. Hệ thống hiển thị “Xoá ngân hàng thành công” và tải lại danh sách.',
    phu='• Bấm Hủy → đóng hộp thoại, không xoá gì.\n'
        '• Ngân hàng đang được sử dụng mà gọi thẳng chức năng xoá, bỏ qua giao diện → hệ thống từ '
        'chối và báo “Không thể xóa bản ghi, ngân hàng đang được sử dụng trên hệ thống”.\n'
        '• Bản ghi đã bị người khác xoá trước đó → hệ thống báo dữ liệu đã thay đổi, không treo '
        'trang.',
    dacbiet='Điều kiện chặn xoá: đang được chọn ở tài khoản ngân hàng hoặc tài khoản uỷ quyền của '
            'nhân viên còn hiệu lực, hoặc đang được dùng ở danh mục tài khoản ngân hàng của công '
            'ty.')

sub('2.9.3 Layout màn hình')
lay(shot=shot('19-xac-nhan-xoa.png'),
         shot_caption='Hộp thoại Xác nhận xóa nêu rõ tên ngân hàng sẽ bị xoá')

sub('2.9.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Nút Xóa', 'Icon Button', 'Enable / Ẩn', 'Ẩn khi không đủ điều kiện',
     'Chỉ hiện khi trạng thái Hoạt động VÀ chưa được sử dụng; ẩn hẳn chứ không làm mờ.'),
    ('Hộp thoại Xác nhận xóa', 'Modal', 'Hiển thị', 'Ẩn',
     'Nội dung: Bạn có chắc muốn xóa ngân hàng “<tên>”?'),
    ('Nút Xóa trong hộp thoại', 'Button', 'Enable', 'Hiển thị', 'Thực hiện xoá.'),
    ('Nút Hủy', 'Button', 'Enable', 'Hiển thị', 'Đóng hộp thoại.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', 'Ẩn',
     '“Xoá ngân hàng thành công” hoặc thông báo lỗi trả về từ máy chủ.'),
], required=False, scope=False)

sub('2.9.5 Danh sách event và xử lý event')
d.event_table([
    ('Bấm nút Xóa ở dòng', 'Click', 'Ghi nhận dòng đang thao tác và mở hộp thoại xác nhận.'),
    ('Bấm Xóa trong hộp thoại', 'Click',
     'Before:\n– Kiểm tra đã có dòng được chọn.\n'
     'During:\n– Máy chủ kiểm tra lại ràng buộc sử dụng; vi phạm → trả lỗi “Không thể xóa bản ghi, '
     'ngân hàng đang được sử dụng trên hệ thống” và không xoá gì.\n'
     'After:\n– Ghi một dòng lịch sử “Xóa”, xoá toàn bộ chi nhánh rồi xoá ngân hàng.\n'
     '– Hiển thị “Xoá ngân hàng thành công” và tải lại danh sách.'),
    ('Bấm Hủy', 'Click', 'Đóng hộp thoại, không gọi xử lý.'),
])

# ---------------------------------------------------------------- 2.10
d.h3('2.10 Quản lý chi nhánh ngân hàng')

sub('2.10.1 Biểu đồ Usecase')
d.uc_figure('FR-10', 'Quản lý chi nhánh ngân hàng', 'crud',
            [('include', 'Kiểm tra trùng tên trong cùng ngân hàng'),
             ('extend', 'Lọc theo Tỉnh/Thành phố và tên chi nhánh')],
            actor=ACTOR,
            caption='Biểu đồ Use Case — FR-10 Quản lý chi nhánh ngân hàng')

sub('2.10.2 Giới thiệu')
d.intro_table(
    ten='Quản lý chi nhánh ngân hàng',
    mota='Xem, lọc, thêm, sửa và xoá chi nhánh của một ngân hàng ngay trên cửa sổ mở từ danh sách.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang ở màn hình danh sách và đã chọn một ngân hàng.',
    chinh='1. Người dùng mở cửa sổ Chi nhánh từ menu ba chấm hoặc bấm số ở cột Chi nhánh.\n'
          '2. Hệ thống tải danh sách chi nhánh của ngân hàng đó.\n'
          '3. Người dùng lọc theo Tỉnh/Thành phố và/hoặc Tên chi nhánh nếu cần.\n'
          '4. Người dùng bấm Tạo mới để thêm, hoặc bấm Sửa / Xóa trên từng dòng.\n'
          '5. Hệ thống lưu thay đổi và tải lại danh sách chi nhánh.',
    phu='• Ngân hàng chưa có chi nhánh → bảng hiển thị “Chưa có dữ liệu”.\n'
        '• Trùng tên chi nhánh trong cùng ngân hàng → hiện lỗi “Tên chi nhánh ngân hàng này đã tồn '
        'tại”.\n'
        '• Chi nhánh đang được sử dụng ở hồ sơ nhân sự → nút Xóa bị làm mờ kèm chú thích lý do.\n'
        '• Bấm Hủy ở hộp thoại xoá → không xoá gì.',
    dacbiet='⚠️ Danh sách trong cửa sổ chỉ hiển thị chi nhánh ĐÃ có Tỉnh/Thành phố, trong khi con '
            'số ở cột Chi nhánh ngoài danh sách đếm cả chi nhánh chưa gán Tỉnh/Thành phố — hai con '
            'số có thể lệch nhau.')

sub('2.10.3 Layout màn hình')
lay(modal='Chi nhánh ngân hàng', shot=shot('15-chi-nhanh.png'),
         shot_caption='Cửa sổ Chi nhánh ngân hàng của BIDV')
d.figure(shot('16-them-chi-nhanh.png'),
         'Cửa sổ Thêm chi nhánh ngân hàng với 2 trường bắt buộc', width_in=6.2)
d.figure(shot('17-xac-nhan-xoa-chi-nhanh.png'),
         'Hộp thoại xác nhận xoá chi nhánh; dòng có chi nhánh đang được sử dụng thì nút xoá bị làm '
         'mờ', width_in=6.2)

sub('2.10.4 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '–', '“Chi nhánh ngân hàng”',
     'Tiêu đề cửa sổ quản lý chi nhánh.'),
    ('Ô lọc Tỉnh/Thành phố', 'Dropdown', 'Enable', 'Danh sách tỉnh/thành', 'Không', 'Trống',
     'Cho phép xoá lựa chọn.'),
    ('Ô lọc Tên chi nhánh', 'Textbox', 'Enable', '0–255 ký tự', 'Không', 'Trống',
     'Khớp gần đúng; nhấn Enter để tìm.'),
    ('Nút Tìm kiếm', 'Button', 'Enable', '–', '–', 'Hiển thị', 'Áp dụng bộ lọc chi nhánh.'),
    ('Nút Làm mới', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Xoá điều kiện lọc và tải lại danh sách chi nhánh.'),
    ('Nút Tạo mới', 'Button', 'Enable', '–', '–', 'Hiển thị',
     'Mở cửa sổ Thêm chi nhánh ngân hàng.'),
    ('Bảng chi nhánh', 'Table/Grid', 'Read-only', 'Danh sách', '–', 'Theo dữ liệu',
     '4 cột: STT, Tên chi nhánh, Tỉnh/TP, Hành động; cuộn trong khung, tiêu đề dính trên.'),
    ('Nút Sửa dòng', 'Icon Button', 'Enable', '–', '–', 'Hiển thị',
     'Mở cửa sổ Sửa chi nhánh ngân hàng.'),
    ('Nút Xóa dòng', 'Icon Button', 'Enable / Disable', '–', '–', 'Hiển thị',
     'Làm mờ khi chi nhánh đang được sử dụng, chú thích “Không thể xóa bản ghi, chi nhánh đang '
     'được sử dụng trên hệ thống”.'),
    ('Ô Tên chi nhánh ngân hàng', 'Textbox', 'Enable', '0–255 ký tự', 'Có', 'Trống',
     'Duy nhất trong phạm vi một ngân hàng.'),
    ('Ô Tỉnh/Thành phố', 'Dropdown', 'Enable', 'Danh sách tỉnh/thành', 'Có', 'Trống',
     'Bắt buộc chọn.'),
    ('Nút Lưu / Lưu và tiếp tục / Đóng', 'Button', 'Enable', '–', '–', 'Hiển thị',
     '“Lưu và tiếp tục” chỉ hiện ở chế độ thêm mới.'),
    ('Thông báo kết quả', 'Toast / Alert', 'Hiển thị', '–', '–', 'Ẩn',
     '“Đã lưu thành công!” / “Xoá chi nhánh ngân hàng thành công” / “Vui lòng kiểm tra lại thông '
     'tin”.'),
])

sub('2.10.5 Danh sách event và xử lý event')
d.event_table([
    ('Mở cửa sổ Chi nhánh', 'Click',
     'Tải danh sách chi nhánh theo ngân hàng đang chọn và cập nhật lại số chi nhánh ở dòng tương '
     'ứng ngoài danh sách.'),
    ('Bấm Tìm kiếm trong cửa sổ', 'Click',
     'Tải lại danh sách chi nhánh theo Tỉnh/Thành phố và Tên chi nhánh.'),
    ('Bấm Làm mới trong cửa sổ', 'Click', 'Xoá hai ô lọc rồi tải lại danh sách chi nhánh.'),
    ('Bấm Tạo mới hoặc Sửa', 'Click',
     'Mở cửa sổ thêm hoặc sửa; chế độ sửa nạp sẵn dữ liệu chi nhánh đang chọn.'),
    ('Bấm Lưu ở cửa sổ chi nhánh', 'Click',
     'Before:\n– Tên chi nhánh trống → hiển thị “Bắt buộc phải nhập” dưới ô Tên chi nhánh.\n'
     '– Tỉnh/Thành phố chưa chọn → hiển thị “Bắt buộc phải nhập” dưới ô Tỉnh/Thành phố.\n'
     '– Nếu có lỗi → hiển thị “Vui lòng kiểm tra lại thông tin” và dừng xử lý.\n'
     'During:\n– Trùng tên trong cùng ngân hàng → hiển thị “Tên chi nhánh ngân hàng này đã tồn '
     'tại”.\n'
     'After:\n– Ghi chi nhánh, hiển thị “Đã lưu thành công!” và tải lại danh sách chi nhánh.'),
    ('Bấm Xóa dòng chi nhánh', 'Click',
     'Before:\n– Chi nhánh đang được sử dụng → nút bị làm mờ, không gọi xử lý.\n'
     'After:\n– Mở hộp xác nhận; xác nhận xong xoá chi nhánh, hiển thị “Xoá chi nhánh ngân hàng '
     'thành công” và tải lại danh sách.'),
])

# ---------------------------------------------------------------- 2.11
d.h3('2.11 Xem lịch sử thay đổi')

sub('2.11.1 Giới thiệu')
d.intro_table(
    ten='Xem lịch sử thay đổi của ngân hàng',
    mota='Truy vết các thao tác đã thực hiện trên một ngân hàng: ai làm, lúc nào, thay đổi gì.',
    tacnhan='Người quản lý danh mục; Người dùng đã đăng nhập',
    dieukien='Người dùng đang ở màn hình danh sách hoặc đang mở cửa sổ Xem của một ngân hàng.',
    chinh='1. Người dùng mở cửa sổ Lịch sử từ menu ba chấm, hoặc mở khối Lịch sử trong cửa sổ '
          'Xem.\n'
          '2. Hệ thống tải các mốc lịch sử của bản ghi, sắp xếp mới nhất trước.\n'
          '3. Người dùng lọc theo loại hoạt động nếu cần.\n'
          '4. Người dùng bấm Đóng để thoát.',
    phu='• Bản ghi chưa có lịch sử → hiển thị “Chưa có lịch sử thao tác nào.”.\n'
        '• Lọc theo một nhóm hoạt động → chỉ hiển thị các mốc thuộc nhóm đó.',
    dacbiet=None)

sub('2.11.2 Layout màn hình')
lay(modal='Lịch sử thay đổi', shot=shot('14-lich-su.png'),
         shot_caption='Cửa sổ Lịch sử thay đổi mở từ menu ba chấm của một dòng')

sub('2.11.3 Mô tả chi tiết giao diện')
d.ui_table([
    ('Tiêu đề cửa sổ', 'Label', 'Hiển thị', '–', '“Lịch sử thay đổi”',
     'Dòng phụ ghi “Ngân hàng: <mã> - <tên>”.'),
    ('Nút Bộ lọc', 'Button', 'Enable', '–', 'Hiển thị',
     'Mở khối chọn loại hoạt động và người thực hiện.'),
    ('Ô Loại hoạt động', 'Dropdown', 'Enable',
     'Tạo mới / Thay đổi thông tin / Thay đổi trạng thái', 'Trống',
     'Ba nhóm cố định dùng chung cho mọi màn danh mục.'),
    ('Dòng mốc lịch sử', 'Table/Grid', 'Read-only', '–', 'Theo dữ liệu',
     'Gồm ngày giờ, nhãn hành động có màu, người thực hiện kèm phòng ban và chi tiết thay đổi.'),
    ('Chi tiết thay đổi', 'Label', 'Read-only', '–', 'Theo dữ liệu',
     'Định dạng “<Tên trường>: <giá trị cũ> → <giá trị mới>” với nhãn trường bằng tiếng Việt.'),
    ('Trạng thái rỗng', 'Label', 'Hiển thị', '–', 'Ẩn', '“Chưa có lịch sử thao tác nào.”.'),
    ('Nút Đóng', 'Button', 'Enable', '–', 'Hiển thị', 'Đóng cửa sổ lịch sử.'),
], required=False)

sub('2.11.4 Danh sách event và xử lý event')
d.event_table([
    ('Bấm Lịch sử ở menu ba chấm', 'Click', 'Mở cửa sổ và tải lịch sử của đúng bản ghi đang chọn.'),
    ('Bấm Xem lịch sử trong cửa sổ Xem', 'Click',
     'Mở khối lịch sử ngay trong cửa sổ, cùng nội dung với cửa sổ Lịch sử.'),
    ('Chọn loại hoạt động', 'Change', 'Lọc danh sách mốc theo nhóm hoạt động đã chọn.'),
    ('Đóng cửa sổ', 'Click', 'Xoá bản ghi đang xem để lần mở sau luôn tải lại dữ liệu mới.'),
])

# ======================================================= PHẦN 4. QUY TẮC NGHIỆP VỤ
d.h1('Phần 4. Quy tắc nghiệp vụ')

sub('BR-01 — Định danh ngân hàng')
d.bullets([
    'Mã ngân hàng và Tên ngân hàng phải duy nhất trên toàn hệ thống.',
    'Khi sửa, phép kiểm tra trùng loại trừ chính bản ghi đang sửa.',
    'Ba trường bắt buộc khi ghi dữ liệu: Mã ngân hàng, Tên ngân hàng, Tên viết tắt.',
    'Tên giao dịch quốc tế, Địa chỉ giao dịch và Logo là tuỳ chọn.',
])

sub('BR-02 — Trạng thái ngân hàng')
d.bullets([
    'Ngân hàng có 2 trạng thái: Hoạt động (mặc định khi tạo) và Khóa.',
    'Ngân hàng đang Khóa không cho Sửa và không cho Xóa từ giao diện: hai nút bị ẩn.',
    'Khóa không kiểm tra ràng buộc sử dụng: ngân hàng đang được tham chiếu vẫn khóa được.',
    'Nhãn trạng thái do máy chủ trả về để mọi màn hiển thị thống nhất.',
])

sub('BR-03 — Điều kiện xoá ngân hàng')
d.bullets([
    'Không cho xoá khi ngân hàng đang được chọn ở tài khoản ngân hàng hoặc tài khoản uỷ quyền của '
    'nhân viên còn hiệu lực.',
    'Không cho xoá khi ngân hàng đang được dùng ở danh mục tài khoản ngân hàng của công ty.',
    'Điều kiện được kiểm tra hai lần: một lần để quyết định hiển thị nút, một lần tại máy chủ khi '
    'thực hiện xoá — giao diện ẩn nút không phải chốt chặn duy nhất.',
    'Xoá ngân hàng kéo theo xoá toàn bộ chi nhánh trực thuộc.',
])

sub('BR-04 — Quan hệ ngân hàng và chi nhánh')
d.bullets([
    'Mỗi chi nhánh thuộc đúng một ngân hàng; chi nhánh không có trạng thái riêng.',
    'Tên chi nhánh duy nhất TRONG PHẠM VI một ngân hàng; hai ngân hàng khác nhau được phép trùng '
    'tên chi nhánh.',
    'Tỉnh/Thành phố là bắt buộc khi thêm hoặc sửa chi nhánh.',
    'Chi nhánh đang được dùng ở hồ sơ nhân sự thì nút xoá bị làm mờ, khác quy ước ẩn nút của danh '
    'sách ngân hàng.',
    '⚠️ Danh sách chi nhánh nối bắt buộc với danh mục tỉnh/thành nên chi nhánh chưa gán Tỉnh/Thành '
    'phố không hiển thị trong cửa sổ, trong khi cột Chi nhánh vẫn đếm — cần thống nhất lại một '
    'trong hai.',
])

sub('BR-05 — Logo ngân hàng')
d.bullets([
    'Chỉ nhận tệp .jpg, .jpeg, .png và tối đa 5 MB.',
    'Tệp được tải lên trước, sau đó đường dẫn ảnh mới được lưu cùng bản ghi.',
    'Thay đổi logo không được ghi vào lịch sử thay đổi.',
])

sub('BR-06 — Lịch sử thay đổi')
d.bullets([
    'Các trường được theo dõi: Mã ngân hàng, Tên ngân hàng, Tên viết tắt, Tên giao dịch quốc tế, '
    'Địa chỉ giao dịch, Trạng thái.',
    'Ghi mốc cho các thao tác: Tạo mới, Thay đổi thông tin, Khóa, Mở khóa, Xóa.',
    'Không ghi mốc nếu nội dung không thay đổi.',
    'Trạng thái được lưu dưới dạng chữ (Hoạt động / Khóa) để bản ghi lịch sử tự chứa nghĩa.',
    'Bộ lọc loại hoạt động dùng chung 3 nhóm cố định cho mọi màn danh mục.',
    'Thao tác trên chi nhánh không được ghi lịch sử.',
])

sub('BR-07 — Phân quyền')
d.bullets([
    'Màn hình hiện không gắn quyền riêng: mọi tài khoản đã đăng nhập đều thao tác được đầy đủ.',
    'Nếu bổ sung quyền, phải làm đồng thời ở ba nơi: khai báo quyền của hệ thống, kiểm tra quyền '
    'trên các thao tác ghi dữ liệu và khoá quyền trong khai báo menu.',
])

sub('BR-08 — Nguồn dữ liệu trên nhánh gộp cơ sở dữ liệu')
d.bullets([
    'Danh mục ngân hàng và chi nhánh nằm trên kết nối mặc định của cơ sở dữ liệu gộp; toàn bộ khối '
    'đồng bộ sang cơ sở dữ liệu thứ hai đã được cắt bỏ để tránh ghi trùng hai lần vào cùng một '
    'bảng.',
    'Chức năng đồng bộ sang hệ thống CRM vẫn giữ và chỉ chạy khi cấu hình hệ thống bật; khi bật, '
    'lỗi kết nối CRM sẽ làm thao tác lưu thất bại kèm thông báo.',
    'Danh sách trong cửa sổ Tra cứu lấy từ dịch vụ ngoài, không phải dữ liệu của hệ thống.',
])

sub('BR-09 — Điểm cần xử lý đã ghi nhận khi lập tài liệu')
d.bullets([
    'Sắp xếp theo cột chưa có hiệu lực do tên tham số sắp xếp giữa giao diện và máy chủ không '
    'khớp; danh sách luôn giữ thứ tự bản ghi mới nhất trước.',
    'Cảnh báo “Thông tin chưa lưu” khi đóng cửa sổ đang nhập dở chưa hoạt động trên bản đang triển '
    'khai.',
    'Cột Chi nhánh và danh sách chi nhánh có thể lệch nhau khi chi nhánh chưa gán Tỉnh/Thành phố.',
])

d.save()

# ----------------------------------------------------------------- MỤC LỤC
# `d.toc()` chi CHEN TRUONG TOC — mo file len thay TRONG cho toi khi bam Update Field.
# Cho Word cap nhat that (giong hdsd_engine._update_fields_by_word) de muc luc hien san
# kem so trang dung.
def _update_toc_by_word(path, title):
    import subprocess
    ps = """
$p = "{path}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($p, $false, $false)
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
$doc.Repaginate()
$doc.BuiltInDocumentProperties("Title") = "{title}"
$doc.Save()
Write-Output ("Pages=" + $doc.ComputeStatistics(2))
$doc.Close(0)
$word.Quit()
""".format(path=path, title=title)
    res = subprocess.run(["powershell", "-NonInteractive", "-Command", ps],
                         capture_output=True, text=True)
    print('Cap nhat muc luc bang Word:', res.stdout.strip() or res.stderr.strip())


_update_toc_by_word(OUT, 'SRS - Danh mục ngân hàng')

# --------------------------------------- Bước 4 của skill: tự kiểm tra form mới
from docx import Document  # noqa: E402

_chk = Document(OUT)
# Luu y: KHONG kiem 'Menu:' — ban mau customer-docs VAN GIU dong "Menu: ..." o muc Layout,
# chi bo dong "Route (FE): ...". Skill ghi bo ca 2 dong la khong khop ban mau.
for _s in ['Tổng quan', 'Mini-Spec', 'Tiêu chí nghiệm thu', 'Ngoài phạm vi',
           'Chức năng liên quan', 'Route (FE)', 'Phân hệ:']:
    assert not any(_s in _p.text for _p in _chk.paragraphs), 'Con muc cua form cu: %s' % _s
print('OK - khong con muc nao cua form cu')
