from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# === Page setup ===
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# === Styles ===
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
        # Header background
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '4472C4',
            qn('w:val'): 'clear',
        })
        shading.append(bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ===================================================
# TITLE PAGE
# ===================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM\n(SRS)')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('THÀNH PHẦN TÍNH THƯỞNG\n(Bonus Component)')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('Module', 'Payroll'),
    ('Phiên bản', '1.0'),
    ('Ngày tạo', '2026-04-09'),
    ('Trạng thái', 'Draft'),
    ('Hệ thống', 'HRM - Quản lý nhân sự'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ===================================================
# 1. GIỚI THIỆU
# ===================================================
doc.add_heading('1. Giới thiệu', level=1)

doc.add_heading('1.1. Mục đích', level=2)
doc.add_paragraph(
    'Quản lý danh sách các thành phần dùng để tính thưởng cuối năm cho nhân viên. '
    'Mỗi thành phần có thể tính theo công thức tự động (dựa trên biến hệ thống từ HRM + ERP) '
    'hoặc nhập tay thủ công. Các thành phần này được sử dụng trong "Bảng chia thưởng cuối năm" (Bonus Distribution).'
)

doc.add_heading('1.2. Phạm vi', level=2)
doc.add_paragraph('Trong scope:', style='List Bullet')
doc.add_paragraph('CRUD thành phần tính thưởng', style='List Bullet 2')
doc.add_paragraph('Soạn công thức với biến hệ thống (HRM + ERP)', style='List Bullet 2')
doc.add_paragraph('Validate công thức trước khi lưu', style='List Bullet 2')
doc.add_paragraph('Tìm kiếm, phân trang danh sách', style='List Bullet 2')

doc.add_paragraph('Ngoài scope:', style='List Bullet')
doc.add_paragraph('Tính toán thực tế thưởng cho nhân viên (thuộc module Bảng chia thưởng)', style='List Bullet 2')
doc.add_paragraph('Quản lý biến hệ thống (hardcode trong frontend)', style='List Bullet 2')

doc.add_heading('1.3. Thuật ngữ', level=2)
add_table(doc,
    ['Thuật ngữ', 'Giải thích'],
    [
        ['Thành phần tính thưởng', 'Một cấu phần trong công thức chia thưởng (VD: "Theo doanh số", "Theo thâm niên")'],
        ['Mode formula', 'Thành phần tính theo công thức tự động'],
        ['Mode manual', 'Thành phần nhập tay (% chia + số tiền)'],
        ['Biến hệ thống', 'Giá trị lấy từ HRM (lương, thâm niên) hoặc ERP (doanh số, công nợ)'],
        ['Công thức', 'Biểu thức toán học kết hợp biến hệ thống và hàm (IF, MAX, MIN...)'],
    ],
    col_widths=[4, 12]
)

# ===================================================
# 2. ACTORS & PERMISSIONS
# ===================================================
doc.add_heading('2. Actors & Permissions', level=1)
add_table(doc,
    ['Actor', 'Mô tả', 'Permissions'],
    [
        ['Admin / Kế toán trưởng', 'Người quản lý thành phần tính thưởng', 'Xem, tạo, sửa, xoá'],
        ['Nhân viên kế toán', 'Người dùng thông thường', 'Xem danh sách (nếu có quyền Payroll)'],
    ],
    col_widths=[4, 6, 6]
)

# ===================================================
# 3. USE CASES
# ===================================================
doc.add_heading('3. Use Cases', level=1)

use_cases = [
    {
        'id': 'UC-01', 'name': 'Xem danh sách thành phần tính thưởng',
        'actor': 'Admin / Kế toán',
        'pre': 'Đã đăng nhập, có quyền vào module Payroll',
        'steps': [
            'Vào menu Tính lương > Thành phần tính thưởng',
            'Hệ thống hiển thị danh sách với cột: STT, Mã TP, Tên, Loại, Công thức, Mô tả, Thao tác',
            'Có thể tìm kiếm theo mã hoặc tên',
            'Phân trang với tuỳ chọn 10/25/50/100 bản ghi',
        ],
        'post': 'Hiển thị danh sách sắp xếp theo thời gian cập nhật mới nhất',
        'alt': '', 'exc': '',
    },
    {
        'id': 'UC-02', 'name': 'Tạo mới thành phần tính thưởng',
        'actor': 'Admin / Kế toán',
        'pre': 'Đã đăng nhập, có quyền tạo',
        'steps': [
            'Nhấn nút "Tạo thành phần"',
            'Hệ thống mở modal, tự sinh mã TP (format: TPxxx)',
            'Nhập tên thành phần (bắt buộc)',
            'Chọn cách tính: "Theo công thức" hoặc "Nhập thủ công"',
            'Nếu chọn công thức: soạn công thức bằng biến hệ thống, hàm, toán tử',
            'Có thể validate công thức trước khi lưu',
            'Nhấn "Lưu"',
        ],
        'post': 'Thành phần được tạo, hiển thị trong danh sách',
        'alt': 'Mã TP trùng → báo lỗi "Mã thành phần đã tồn tại"',
        'exc': 'Validation lỗi → hiển thị message lỗi tương ứng',
    },
    {
        'id': 'UC-03', 'name': 'Sửa thành phần tính thưởng',
        'actor': 'Admin / Kế toán',
        'pre': 'Thành phần đã tồn tại',
        'steps': [
            'Nhấn nút sửa (icon bút) trên dòng thành phần',
            'Hệ thống mở modal với dữ liệu hiện tại',
            'Sửa thông tin cần thay đổi',
            'Nhấn "Lưu"',
        ],
        'post': 'Thành phần được cập nhật, updated_by ghi nhận user hiện tại',
        'alt': '', 'exc': 'Mã TP trùng với thành phần khác → báo lỗi',
    },
    {
        'id': 'UC-04', 'name': 'Xoá thành phần tính thưởng',
        'actor': 'Admin / Kế toán',
        'pre': 'Thành phần đã tồn tại',
        'steps': [
            'Nhấn nút xoá (icon thùng rác)',
            'Hệ thống hiển thị modal xác nhận',
            'Nhấn "Xoá"',
        ],
        'post': 'Thành phần bị xoá khỏi DB',
        'alt': '', 'exc': '',
    },
    {
        'id': 'UC-05', 'name': 'Soạn công thức',
        'actor': 'Admin / Kế toán',
        'pre': 'Đang ở modal tạo/sửa, mode = "Theo công thức"',
        'steps': [
            'Gõ tên biến → hệ thống gợi ý autocomplete',
            'Hoặc nhấn vào biến trong danh sách "Biến hệ thống" để chèn',
            'Hoặc nhấn vào hàm trong danh sách "Hàm tính toán" để chèn',
            'Có thể tham chiếu thành phần khác bằng mã TP',
            'Nhấn "Kiểm tra công thức" để validate',
        ],
        'post': 'Công thức được soạn và validate thành công',
        'alt': 'Công thức lỗi → hiển thị danh sách lỗi',
        'exc': '',
    },
]

for uc in use_cases:
    doc.add_heading(f'{uc["id"]}: {uc["name"]}', level=2)
    rows = [
        ['Actor', uc['actor']],
        ['Precondition', uc['pre']],
    ]
    for i, step in enumerate(uc['steps'], 1):
        rows.append([f'Bước {i}', step])
    rows.append(['Postcondition', uc['post']])
    if uc['alt']:
        rows.append(['Alternative Flow', uc['alt']])
    if uc['exc']:
        rows.append(['Exception', uc['exc']])
    add_table(doc, ['', ''], rows, col_widths=[4, 12])

# ===================================================
# 4. BUSINESS RULES
# ===================================================
doc.add_heading('4. Business Rules', level=1)
add_table(doc,
    ['ID', 'Rule', 'Mô tả', 'Áp dụng tại'],
    [
        ['BR-01', 'Mã TP unique', 'Mã thành phần không được trùng trong hệ thống', 'UC-02, UC-03'],
        ['BR-02', 'Mã TP max 50 ký tự', '', 'UC-02, UC-03'],
        ['BR-03', 'Tên max 255 ký tự', '', 'UC-02, UC-03'],
        ['BR-04', 'Mode 2 giá trị', 'Chỉ có formula hoặc manual', 'UC-02, UC-03'],
        ['BR-05', 'Formula nullable', 'Chỉ bắt buộc khi mode = formula (validate FE)', 'UC-02, UC-03'],
        ['BR-06', 'Auto-gen mã TP', 'Frontend tự sinh TPxxx khi mở modal tạo mới', 'UC-02'],
        ['BR-07', 'Audit trail', 'created_by khi tạo, updated_by khi sửa', 'UC-02, UC-03'],
        ['BR-08', 'Sắp xếp', 'Theo updated_at DESC', 'UC-01'],
        ['BR-09', 'Tìm kiếm', 'LIKE %keyword% theo code hoặc name', 'UC-01'],
        ['BR-10', 'Phân trang', 'Mặc định 100/trang, tuỳ chọn 10/25/50/100', 'UC-01'],
    ],
    col_widths=[2, 3.5, 6, 3]
)

# ===================================================
# 5. DATA MODEL
# ===================================================
doc.add_heading('5. Data Model', level=1)

doc.add_heading('5.1. Entity Relationship', level=2)
doc.add_paragraph('[BonusComponent] 1──N [BonusDistributionComponent]')
doc.add_paragraph('[BonusComponent] N──1 [Employee] (created_by, updated_by)')

doc.add_heading('5.2. Bảng: bonus_components', level=2)
add_table(doc,
    ['Cột', 'Type', 'Nullable', 'Default', 'Mô tả'],
    [
        ['id', 'bigint PK', 'No', 'auto', ''],
        ['code', 'varchar(50)', 'No', '', 'Mã thành phần (unique)'],
        ['name', 'varchar(255)', 'No', '', 'Tên thành phần'],
        ['mode', 'varchar(20)', 'No', 'formula', 'formula / manual'],
        ['formula', 'text', 'Yes', 'null', 'Công thức tính'],
        ['note', 'text', 'Yes', 'null', 'Ghi chú'],
        ['created_by', 'bigint', 'No', '', 'ID người tạo'],
        ['updated_by', 'bigint', 'Yes', 'null', 'ID người cập nhật'],
        ['created_at', 'timestamp', 'Yes', '', ''],
        ['updated_at', 'timestamp', 'Yes', '', ''],
    ],
    col_widths=[3, 2.5, 2, 2, 5]
)

doc.add_heading('5.3. Enum Values', level=2)
add_table(doc,
    ['Field', 'Value', 'Meaning'],
    [
        ['mode', 'formula', 'Tính theo công thức'],
        ['mode', 'manual', 'Nhập tay thủ công'],
    ],
    col_widths=[4, 4, 6]
)

# ===================================================
# 6. API SPECIFICATION
# ===================================================
doc.add_heading('6. API Specification', level=1)

apis = [
    {
        'name': '6.1. Danh sách thành phần',
        'method': 'GET /api/v1/payroll/bonus-components',
        'params': [
            ['page', 'integer', 'No', 'Trang hiện tại'],
            ['limit', 'integer', 'No', 'Số bản ghi/trang (default: 10)'],
            ['keyword', 'string', 'No', 'Tìm theo mã hoặc tên'],
        ],
        'response': 'Trả về mảng data + total, lastPage, currentPage, perPage',
    },
    {
        'name': '6.2. Chi tiết thành phần',
        'method': 'GET /api/v1/payroll/bonus-components/{id}',
        'params': [],
        'response': 'Trả về object: id, code, name, mode, formula, note, created_by, updated_by',
    },
    {
        'name': '6.3. Tạo mới',
        'method': 'POST /api/v1/payroll/bonus-components',
        'params': [
            ['code', 'string', 'Yes', 'max:50, unique'],
            ['name', 'string', 'Yes', 'max:255'],
            ['mode', 'string', 'Yes', 'in:formula,manual'],
            ['formula', 'string', 'No', 'Công thức'],
            ['note', 'string', 'No', 'Ghi chú'],
        ],
        'response': '200: success | 422: validation errors',
    },
    {
        'name': '6.4. Cập nhật',
        'method': 'PUT /api/v1/payroll/bonus-components/{id}',
        'params': [
            ['code', 'string', 'Yes', 'max:50, unique (trừ chính nó)'],
            ['name', 'string', 'Yes', 'max:255'],
            ['mode', 'string', 'Yes', 'in:formula,manual'],
            ['formula', 'string', 'No', 'Công thức'],
            ['note', 'string', 'No', 'Ghi chú'],
        ],
        'response': '200: success | 404: not found | 422: validation errors',
    },
    {
        'name': '6.5. Xoá',
        'method': 'DELETE /api/v1/payroll/bonus-components/{id}',
        'params': [],
        'response': '200: delete success | 404: not found',
    },
]

for api in apis:
    doc.add_heading(api['name'], level=2)
    p = doc.add_paragraph()
    run = p.add_run(api['method'])
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.add_run('Auth: Bearer Token (JWT)').font.size = Pt(10)

    if api['params']:
        add_table(doc,
            ['Field', 'Type', 'Required', 'Mô tả'],
            api['params'],
            col_widths=[3, 2.5, 2, 7]
        )

    doc.add_paragraph(f'Response: {api["response"]}')

# ===================================================
# 7. BIẾN HỆ THỐNG
# ===================================================
doc.add_heading('7. Biến hệ thống', level=1)

doc.add_heading('7.1. Biến HRM', level=2)
add_table(doc,
    ['Code', 'Tên', 'Nguồn'],
    [
        ['luong_p1', 'Lương P1', 'employee_salary_histories.base_salary'],
        ['luong_p2', 'Lương P2', 'employee_salary_histories.p2_salary'],
        ['luong_p3', 'Lương P3', 'employee_salary_histories.p3_salary'],
        ['tham_nien', 'Thâm niên (tháng)', 'Tính từ employee_infos.enter_date'],
        ['so_thang_lam_viec', 'Số tháng làm việc', 'timesheet_summaries'],
    ],
    col_widths=[4, 4, 7]
)

doc.add_heading('7.2. Biến ERP — Nhân viên', level=2)
erp_nv = [
    ['dstc_dk_dau_nam_nv', 'DSTC đăng ký đầu năm của NV'],
    ['dstc_quyet_toan_nv', 'DSTC đã quyết toán của NV'],
    ['dstc_hd_hieu_luc_nv', 'DSTC hợp đồng hiệu lực của NV'],
    ['gia_tri_hd_hieu_luc_nv', 'Giá trị HĐ hiệu lực của NV'],
    ['gia_tri_hd_truoc_thue_nv', 'Giá trị HĐ trước thuế của NV'],
    ['ds_net_nv', 'Doanh số net của NV'],
    ['ds_thuc_xuat_nv', 'Doanh số thực xuất của NV'],
    ['so_tien_thuc_thu_nv', 'Tiền thực thu của NV'],
    ['cong_no_kh_nv', 'Công nợ KH của NV'],
    ['ty_le_ht_kh_ds_hd_nv', 'Tỷ lệ HT DS theo HĐ của NV'],
    ['ty_le_ht_kh_dstc_hd_nv', 'Tỷ lệ HT DSTC theo HĐ của NV'],
    ['ty_le_ht_kh_dstc_qt_nv', 'Tỷ lệ HT DSTC quyết toán của NV'],
    ['ty_le_tong_th_ds_hd_nv', 'Tỷ lệ tổng TH DS theo HĐ của NV'],
    ['ty_le_tong_th_dstc_hd_nv', 'Tỷ lệ tổng TH DSTC theo HĐ của NV'],
    ['ty_le_tong_th_dstc_qt_nv', 'Tỷ lệ tổng TH DSTC QT của NV'],
]
add_table(doc, ['Code', 'Tên'], erp_nv, col_widths=[5.5, 9])

doc.add_heading('7.3. Biến ERP — Phòng ban', level=2)
erp_phong = [
    ['dstc_dk_dau_nam_phong', 'DSTC đăng ký đầu năm của phòng'],
    ['dstc_quyet_toan_phong', 'DSTC đã quyết toán của phòng'],
    ['dstc_hd_hieu_luc_phong', 'DSTC hợp đồng hiệu lực của phòng'],
    ['gia_tri_hd_hieu_luc_phong', 'Giá trị HĐ hiệu lực của phòng'],
    ['gia_tri_hd_truoc_thue_phong', 'Giá trị HĐ trước thuế của phòng'],
    ['ds_net_phong', 'Doanh số net của phòng'],
    ['ds_thuc_xuat_phong', 'Doanh số thực xuất của phòng'],
    ['so_tien_thuc_thu_phong', 'Tiền thực thu của phòng'],
    ['cong_no_kh_phong', 'Công nợ KH của phòng'],
]
add_table(doc, ['Code', 'Tên'], erp_phong, col_widths=[5.5, 9])

doc.add_heading('7.4. Hàm tính toán', level=2)
add_table(doc,
    ['Hàm', 'Cú pháp', 'Mô tả'],
    [
        ['IF', 'IF(điều_kiện, giá_trị_đúng, giá_trị_sai)', 'Hàm điều kiện'],
        ['MAX', 'MAX(a, b)', 'Giá trị lớn nhất'],
        ['MIN', 'MIN(a, b)', 'Giá trị nhỏ nhất'],
        ['ROUND', 'ROUND(number, digits)', 'Làm tròn'],
        ['SUM', 'SUM(a, b, c)', 'Tổng'],
        ['AVG', 'AVG(a, b, c)', 'Trung bình'],
        ['ABS', 'ABS(number)', 'Giá trị tuyệt đối'],
        ['CEILING', 'CEILING(number)', 'Làm tròn lên'],
        ['FLOOR', 'FLOOR(number)', 'Làm tròn xuống'],
    ],
    col_widths=[2.5, 7, 4]
)

# ===================================================
# 8. UI SPECIFICATION
# ===================================================
doc.add_heading('8. UI Specification', level=1)

doc.add_heading('8.1. Màn hình Danh sách', level=2)
doc.add_paragraph('Route: /payroll/bonus-component')
doc.add_paragraph('Layout: Header (tiêu đề + nút tạo) → Ô tìm kiếm → Bảng danh sách → Phân trang')

doc.add_heading('User interactions:', level=3)
add_table(doc,
    ['Action', 'Behavior', 'API call'],
    [
        ['Nhập keyword', 'Tìm kiếm debounce 300ms', 'GET /bonus-components?keyword='],
        ['Click "Tạo thành phần"', 'Mở modal tạo mới', '-'],
        ['Click icon sửa', 'Mở modal sửa với data', 'GET /bonus-components/{id}'],
        ['Click icon xoá', 'Mở modal xác nhận', '-'],
        ['Xác nhận xoá', 'Xoá và reload danh sách', 'DELETE /bonus-components/{id}'],
    ],
    col_widths=[4, 5, 5.5]
)

doc.add_heading('8.2. Modal tạo/sửa', level=2)
doc.add_paragraph('Layout: Modal XL với form fields + formula editor')
doc.add_paragraph('Fields: Mã TP (auto-gen), Tên thành phần, Cách tính (dropdown), Ghi chú')
doc.add_paragraph('Formula editor (khi mode = Công thức):')
doc.add_paragraph('Preview công thức với syntax highlighting', style='List Bullet')
doc.add_paragraph('Textarea soạn công thức với autocomplete', style='List Bullet')
doc.add_paragraph('Danh sách biến hệ thống (click để chèn)', style='List Bullet')
doc.add_paragraph('Danh sách thành phần khác (click để tham chiếu)', style='List Bullet')
doc.add_paragraph('Danh sách hàm tính toán (click để chèn)', style='List Bullet')
doc.add_paragraph('Nút "Kiểm tra công thức" để validate', style='List Bullet')

# ===================================================
# 9. FILE REFERENCES
# ===================================================
doc.add_heading('9. Phụ lục — File references', level=1)
add_table(doc,
    ['Layer', 'File path'],
    [
        ['Migration', 'Modules/Payroll/Database/Migrations/2026_04_01_000000_create_bonus_components_table.php'],
        ['Entity', 'Modules/Payroll/Entities/BonusComponent.php'],
        ['Controller', 'Modules/Payroll/Http/Controllers/Api/V1/BonusComponentController.php'],
        ['Service', 'Modules/Payroll/Services/BonusComponentService.php'],
        ['Create Request', 'Modules/Payroll/Http/Requests/CreateBonusComponentRequest.php'],
        ['Update Request', 'Modules/Payroll/Http/Requests/UpdateBonusComponentRequest.php'],
        ['List Resource', 'Modules/Payroll/Transformers/BonusComponentResource/BonusComponentListResource.php'],
        ['Detail Resource', 'Modules/Payroll/Transformers/BonusComponentResource/BonusComponentDetailResource.php'],
        ['Routes', 'Modules/Payroll/Routes/api.php (lines 334-340)'],
        ['FE Page', 'pages/payroll/bonus-component/index.vue'],
    ],
    col_widths=[3, 12]
)

# === SAVE ===
output_path = '/Users/nguyentrancu/DEV/code/HRM/hrm-claude-config/docs/srs/SRS_Bonus_Component.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
