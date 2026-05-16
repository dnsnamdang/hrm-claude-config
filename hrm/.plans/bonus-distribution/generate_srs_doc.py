from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): '4472C4', qn('w:val'): 'clear'})
        shading.append(bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_uc(doc, uc):
    doc.add_heading(f'{uc["id"]}: {uc["name"]}', level=2)
    rows = [['Actor', uc['actor']], ['Precondition', uc['pre']]]
    for i, s in enumerate(uc['steps'], 1):
        rows.append([f'Bước {i}', s])
    rows.append(['Postcondition', uc['post']])
    if uc.get('alt'):
        rows.append(['Alternative Flow', uc['alt']])
    if uc.get('exc'):
        rows.append(['Exception', uc['exc']])
    add_table(doc, ['', ''], rows, col_widths=[4, 12])

# === TITLE PAGE ===
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM\n(SRS)')
run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79); run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('BẢNG CHIA THƯỞNG CUỐI NĂM\n(Bonus Distribution)')
run.bold = True; run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4); run.font.name = 'Times New Roman'

doc.add_paragraph(); doc.add_paragraph()
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
for i, (k, v) in enumerate([('Module', 'Payroll'), ('Phiên bản', '1.0'), ('Ngày tạo', '2026-04-09'), ('Trạng thái', 'Draft'), ('Hệ thống', 'HRM - Quản lý nhân sự')]):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11); run.font.name = 'Times New Roman'
    info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# === 1. GIỚI THIỆU ===
doc.add_heading('1. Giới thiệu', level=1)
doc.add_heading('1.1. Mục đích', level=2)
doc.add_paragraph('Quản lý bảng chia thưởng cuối năm cho nhân viên theo từng phòng ban. Hệ thống cho phép thiết lập các thành phần tính thưởng (công thức hoặc thủ công), tự động tính toán phân bổ thưởng cho từng nhân viên dựa trên dữ liệu HRM (lương, thâm niên) và ERP (doanh số, công nợ), với workflow duyệt phiếu.')

doc.add_heading('1.2. Phạm vi', level=2)
doc.add_paragraph('Trong scope:', style='List Bullet')
for s in ['CRUD bảng chia thưởng (tạo, xem, sửa, xoá)', 'Thiết lập thành phần tính thưởng cho mỗi bảng', 'Tính toán tự động phân bổ thưởng theo công thức', 'Điều chỉnh thủ công số tiền thưởng', 'Workflow duyệt: Nháp → Chờ duyệt → Đã duyệt / Không duyệt', 'Lấy quỹ thưởng phòng ban từ ERP', 'Lấy biến tính thưởng từ ERP (doanh số, công nợ...)']:
    doc.add_paragraph(s, style='List Bullet 2')
doc.add_paragraph('Ngoài scope:', style='List Bullet')
for s in ['Quản lý thành phần tính thưởng (module riêng BonusComponent)', 'Quyết toán thưởng cuối năm bên ERP', 'Xuất báo cáo thưởng']:
    doc.add_paragraph(s, style='List Bullet 2')

doc.add_heading('1.3. Thuật ngữ', level=2)
add_table(doc, ['Thuật ngữ', 'Giải thích'], [
    ['Bảng chia thưởng', 'Phiếu phân bổ thưởng cuối năm cho 1 phòng ban trong 1 năm'],
    ['Thành phần tính thưởng', 'Cấu phần trong công thức (VD: Theo lương, Theo doanh số)'],
    ['Quỹ thưởng phòng', 'Tổng tiền thưởng phòng ban được duyệt từ ERP'],
    ['Căn cứ tính toán', 'Giá trị base dùng để phân bổ thưởng (kết quả evaluate công thức)'],
    ['Preview calculate', 'Tính toán xem trước, không lưu DB'],
    ['ERP (TanPhatDev)', 'Hệ thống quản lý doanh nghiệp, cung cấp dữ liệu doanh số, công nợ'],
], col_widths=[4, 12])

# === 2. ACTORS & PERMISSIONS ===
doc.add_heading('2. Actors & Permissions', level=1)
add_table(doc, ['Actor', 'Mô tả', 'Permissions'], [
    ['Kế toán / Nhân sự', 'Người tạo và quản lý bảng chia thưởng', 'Tạo, sửa, xoá, gửi duyệt (id 1050)'],
    ['BGĐ / Người duyệt', 'Người có quyền duyệt bảng chia thưởng', 'Duyệt / Không duyệt (id 1054)'],
    ['Quản lý', 'Xem theo phạm vi', 'Xem theo tổng CT (1051), CT (1052), PB (1053)'],
], col_widths=[3.5, 5, 6])

doc.add_paragraph()
doc.add_paragraph('Danh sách permissions (PermissionsTableSeeder):')
add_table(doc, ['ID', 'Name', 'Group'], [
    ['1050', 'Tạo bảng chia thưởng', 'Bảng chia thưởng'],
    ['1051', 'Xem bảng chia thưởng theo tổng công ty', 'Bảng chia thưởng'],
    ['1052', 'Xem bảng chia thưởng theo công ty', 'Bảng chia thưởng'],
    ['1053', 'Xem bảng chia thưởng theo phòng ban', 'Bảng chia thưởng'],
    ['1054', 'BGĐ duyệt bảng chia thưởng', 'Bảng chia thưởng'],
], col_widths=[2, 7, 5])

# === 3. USE CASES ===
doc.add_heading('3. Use Cases', level=1)

use_cases = [
    {'id': 'UC-01', 'name': 'Xem danh sách bảng chia thưởng', 'actor': 'Kế toán / Quản lý',
     'pre': 'Đã đăng nhập, có quyền xem module Payroll',
     'steps': ['Vào menu Tính lương > Bảng chia thưởng cuối năm', 'Hệ thống hiển thị danh sách với bộ lọc: Công ty, Phòng ban, Năm, Trạng thái, Người tạo, Người duyệt, Khoảng thời gian', 'Bảng hiển thị: STT, Mã phiếu, Công ty, Phòng ban, Năm, Tổng quỹ, Người lập, Ngày lập, Người duyệt, Ngày duyệt, Trạng thái, Thao tác', 'Phân trang 10/25/50/100 bản ghi'],
     'post': 'Hiển thị danh sách sắp xếp theo updated_at DESC'},
    {'id': 'UC-02', 'name': 'Tạo bảng chia thưởng', 'actor': 'Kế toán',
     'pre': 'Có quyền tạo (id 1050)',
     'steps': ['Nhấn "Tạo bảng chia thưởng"', 'Chọn phòng ban và năm thưởng', 'Hệ thống gọi ERP lấy quỹ thưởng phòng (từ bảng quyết toán đã duyệt)', 'Thêm thành phần tính thưởng, thiết lập % quỹ (tổng = 100%)', 'Nhấn "Tính bảng thưởng" → hệ thống preview tính toán (không lưu DB)', 'Xem kết quả phân bổ, điều chỉnh nếu cần', 'Nhấn "Lưu nháp" hoặc "Lưu & Gửi duyệt"'],
     'post': 'Bảng chia thưởng được tạo với mã tự sinh (format: {CtyCode}.BCT{Năm}.{Số})',
     'alt': 'Phòng ban + năm đã có bảng → báo lỗi trùng',
     'exc': 'ERP không có quyết toán đã duyệt → poolError hiển thị'},
    {'id': 'UC-03', 'name': 'Sửa bảng chia thưởng', 'actor': 'Kế toán',
     'pre': 'Bảng ở trạng thái Nháp hoặc Không duyệt',
     'steps': ['Nhấn icon sửa trên danh sách', 'Hệ thống load dữ liệu hiện tại', 'Sửa thành phần, % quỹ, tính lại bảng thưởng', 'Nhấn "Lưu nháp" hoặc "Lưu & Gửi duyệt"'],
     'post': 'Bảng được cập nhật',
     'exc': 'Trạng thái Chờ duyệt hoặc Đã duyệt → không cho sửa'},
    {'id': 'UC-04', 'name': 'Xoá bảng chia thưởng', 'actor': 'Kế toán',
     'pre': 'Bảng ở trạng thái Nháp hoặc Không duyệt',
     'steps': ['Nhấn icon xoá', 'Xác nhận trong modal', 'Nhấn "Xoá"'],
     'post': 'Bảng bị xoá cùng components và employees',
     'exc': 'Trạng thái khác → không cho xoá'},
    {'id': 'UC-05', 'name': 'Tính toán bảng thưởng (Preview)', 'actor': 'Kế toán',
     'pre': 'Đã thiết lập thành phần, tổng % = 100%, có quỹ thưởng',
     'steps': ['Nhấn "Tính bảng thưởng"', 'Hệ thống gọi API preview-calculate với thông tin components', 'Backend lấy danh sách NV trong phòng ban', 'Lấy biến HRM (lương, thâm niên, ngày công)', 'Lấy biến ERP (doanh số, công nợ...)', 'Evaluate công thức cho mỗi NV → base_value', 'Tính percent_share = base_value / totalBase * 100', 'Tính amount = pool_amount * percent_share', 'Trả kết quả về FE (không lưu DB)'],
     'post': 'Bảng phân bổ hiển thị với dữ liệu tính toán'},
    {'id': 'UC-06', 'name': 'Gửi duyệt', 'actor': 'Kế toán',
     'pre': 'Bảng ở trạng thái Nháp, tổng % = 100%',
     'steps': ['Nhấn "Lưu & Gửi duyệt"', 'Hệ thống lưu dữ liệu', 'Chuyển trạng thái từ Nháp → Chờ duyệt'],
     'post': 'Bảng chuyển sang trạng thái Chờ duyệt, không cho sửa nữa',
     'exc': 'Tổng % ≠ 100 → báo lỗi'},
    {'id': 'UC-07', 'name': 'Duyệt bảng chia thưởng', 'actor': 'BGĐ / Người duyệt',
     'pre': 'Bảng ở trạng thái Chờ duyệt, user có quyền "BGĐ duyệt bảng chia thưởng" (id 1054)',
     'steps': ['Mở chi tiết bảng chia thưởng', 'Xem bảng phân bổ', 'Nhấn "DUYỆT"'],
     'post': 'Trạng thái → Đã duyệt, ghi approved_by và approved_at'},
    {'id': 'UC-08', 'name': 'Từ chối duyệt', 'actor': 'BGĐ / Người duyệt',
     'pre': 'Bảng ở trạng thái Chờ duyệt, user có quyền duyệt',
     'steps': ['Mở chi tiết bảng chia thưởng', 'Nhấn "KHÔNG DUYỆT"'],
     'post': 'Trạng thái → Không duyệt, cho phép người tạo sửa lại'},
]

for uc in use_cases:
    add_uc(doc, uc)

# === 4. BUSINESS RULES ===
doc.add_heading('4. Business Rules', level=1)
add_table(doc, ['ID', 'Rule', 'Mô tả', 'Áp dụng'], [
    ['BR-01', 'Unique phòng ban + năm', 'Mỗi phòng ban chỉ có 1 bảng chia thưởng / năm', 'UC-02'],
    ['BR-02', 'Mã phiếu tự sinh', '{CtyCode}.BCT{Năm}.{Số 5 chữ số}', 'UC-02'],
    ['BR-03', 'Tổng % = 100', 'Tổng % quỹ các thành phần phải bằng 100% (±0.01)', 'UC-05, UC-06'],
    ['BR-04', 'Chỉ sửa/xoá Nháp/Không duyệt', 'Status 0, 3 mới cho sửa và xoá', 'UC-03, UC-04'],
    ['BR-05', 'Chỉ duyệt Chờ duyệt', 'Status 1 mới cho duyệt/từ chối', 'UC-07, UC-08'],
    ['BR-06', 'Quyền duyệt', 'Cần permission "BGĐ duyệt bảng chia thưởng" (id 1054)', 'UC-07, UC-08'],
    ['BR-07', 'Preview không lưu DB', 'Tính bảng thưởng chỉ preview, lưu khi bấm Lưu', 'UC-05'],
    ['BR-08', 'Quỹ thưởng từ ERP', 'Lấy từ bảng quyết toán đã duyệt bên TanPhatDev', 'UC-02'],
    ['BR-09', 'Biến HRM + ERP', 'Merge biến từ 2 nguồn để evaluate công thức', 'UC-05'],
    ['BR-10', 'TP thủ công', 'Mode manual: NV có base_value=0, nhập tay % chia và số tiền', 'UC-05'],
    ['BR-11', 'TP công thức', 'Mode formula: evaluate công thức → base_value → phân bổ theo tỷ lệ', 'UC-05'],
    ['BR-12', 'Không trùng thành phần', 'Trong 1 bảng, mỗi bonus_component chỉ chọn 1 lần', 'UC-02, UC-03'],
    ['BR-13', '% quỹ 0-100', 'Mỗi thành phần % quỹ từ 0 đến 100, tổng không vượt 100', 'UC-02, UC-03'],
], col_widths=[1.5, 4, 6.5, 2.5])

# === 5. DATA MODEL ===
doc.add_heading('5. Data Model', level=1)
doc.add_heading('5.1. Entity Relationship', level=2)
doc.add_paragraph('[BonusDistribution] 1──N [BonusDistributionComponent] 1──N [BonusDistributionEmployee]')
doc.add_paragraph('[BonusDistributionComponent] N──1 [BonusComponent]')
doc.add_paragraph('[BonusDistributionEmployee] N──1 [EmployeeInfo]')
doc.add_paragraph('[BonusDistribution] N──1 [Employee] (created_by, approved_by)')
doc.add_paragraph('[BonusDistribution] N──1 [Department], [Company]')

doc.add_heading('5.2. Bảng: bonus_distributions', level=2)
add_table(doc, ['Cột', 'Type', 'Nullable', 'Default', 'Mô tả'], [
    ['id', 'bigint PK', 'No', 'auto', ''], ['code', 'varchar(30)', 'No', '', 'Mã phiếu (unique)'],
    ['company_id', 'bigint', 'No', '', 'ID công ty'], ['department_id', 'bigint', 'No', '', 'ID phòng ban'],
    ['year', 'smallint', 'No', '', 'Năm thưởng'], ['pool_amount', 'decimal(20,2)', 'No', '0', 'Tổng quỹ thưởng'],
    ['status', 'tinyint', 'No', '0', '0:Nháp, 1:Chờ duyệt, 2:Đã duyệt, 3:Không duyệt'],
    ['note', 'text', 'Yes', '', 'Ghi chú'], ['created_by', 'bigint', 'No', '', 'Người tạo'],
    ['approved_by', 'bigint', 'Yes', '', 'Người duyệt'], ['approved_at', 'timestamp', 'Yes', '', 'Ngày duyệt'],
    ['created_at', 'timestamp', 'Yes', '', ''], ['updated_at', 'timestamp', 'Yes', '', ''],
], col_widths=[2.8, 2.5, 1.5, 1.5, 6])
doc.add_paragraph('Unique constraint: (department_id, year)')

doc.add_heading('5.3. Bảng: bonus_distribution_components', level=2)
add_table(doc, ['Cột', 'Type', 'Nullable', 'Default', 'Mô tả'], [
    ['id', 'bigint PK', 'No', 'auto', ''], ['bonus_distribution_id', 'bigint', 'No', '', 'FK → bonus_distributions'],
    ['bonus_component_id', 'bigint', 'No', '', 'FK → bonus_components'],
    ['percent', 'decimal(5,2)', 'No', '0', '% quỹ'], ['pool_amount', 'decimal(20,2)', 'No', '0', 'Quỹ thành phần'],
    ['allow_adjustment', 'boolean', 'No', 'false', 'Cho phép điều chỉnh số tiền'],
    ['sort_order', 'smallint', 'No', '0', 'Thứ tự hiển thị'],
], col_widths=[3.5, 2.5, 1.5, 1.5, 5])

doc.add_heading('5.4. Bảng: bonus_distribution_employees', level=2)
add_table(doc, ['Cột', 'Type', 'Nullable', 'Default', 'Mô tả'], [
    ['id', 'bigint PK', 'No', 'auto', ''],
    ['bonus_distribution_id', 'bigint', 'No', '', 'FK → bonus_distributions'],
    ['bonus_distribution_component_id', 'bigint', 'No', '', 'FK → bonus_distribution_components'],
    ['employee_info_id', 'bigint', 'No', '', 'FK → employee_infos'],
    ['base_value', 'decimal(20,4)', 'No', '0', 'Giá trị căn cứ'], ['percent_share', 'decimal(10,4)', 'No', '0', '% phân bổ'],
    ['amount', 'decimal(20,2)', 'No', '0', 'Số tiền thưởng'],
], col_widths=[4, 2.5, 1.5, 1, 5])

doc.add_heading('5.5. Enum Values', level=2)
add_table(doc, ['Field', 'Value', 'Meaning', 'Cho phép'], [
    ['status', '0', 'Nháp', 'Sửa, xoá, gửi duyệt'],
    ['status', '1', 'Chờ duyệt', 'Duyệt, từ chối (cần quyền)'],
    ['status', '2', 'Đã duyệt', 'Chỉ xem'],
    ['status', '3', 'Không duyệt', 'Sửa, xoá, gửi duyệt lại'],
], col_widths=[3, 2, 3, 6])

# === 6. API SPECIFICATION ===
doc.add_heading('6. API Specification', level=1)

apis = [
    ('6.1. Danh sách', 'GET /api/v1/payroll/bonus-distributions',
     [['page', 'int', 'No', 'Trang'], ['limit', 'int', 'No', 'Số/trang (default 10)'],
      ['company_id', 'int', 'No', 'Lọc công ty'], ['department_id', 'int', 'No', 'Lọc phòng ban'],
      ['year', 'int', 'No', 'Lọc năm'], ['status', 'int', 'No', 'Lọc trạng thái'],
      ['created_by', 'int', 'No', 'Lọc người tạo'], ['approved_by', 'int', 'No', 'Lọc người duyệt'],
      ['date_from', 'date', 'No', 'Từ ngày'], ['date_to', 'date', 'No', 'Đến ngày']],
     'Mảng data + total, lastPage, currentPage, perPage'),
    ('6.2. Chi tiết', 'GET /api/v1/payroll/bonus-distributions/{id}', [], 'Object chi tiết + components + employees + can_approve'),
    ('6.3. Tạo mới', 'POST /api/v1/payroll/bonus-distributions',
     [['company_id', 'int', 'Yes', 'ID công ty'], ['department_id', 'int', 'Yes', 'ID phòng ban'],
      ['year', 'int', 'Yes', '2000-2100'], ['pool_amount', 'numeric', 'Yes', '≥ 0'],
      ['note', 'string', 'No', 'Ghi chú'], ['components', 'array', 'Yes', 'Min 1 thành phần'],
      ['components.*.bonus_component_id', 'int', 'Yes', 'ID thành phần'],
      ['components.*.percent', 'numeric', 'Yes', '0-100'], ['components.*.allow_adjustment', 'bool', 'No', '']],
     '200: {id} | 422: validation | 400: trùng PB+năm'),
    ('6.4. Cập nhật', 'PUT /api/v1/payroll/bonus-distributions/{id}',
     [['pool_amount', 'numeric', 'Yes', '≥ 0'], ['note', 'string', 'No', ''],
      ['components', 'array', 'Yes', 'Min 1'], ['employees', 'array', 'No', 'Phân bổ NV'],
      ['employees.*.bonus_component_id', 'int', 'Yes*', ''], ['employees.*.employee_info_id', 'int', 'Yes*', ''],
      ['employees.*.base_value', 'numeric', 'No', ''], ['employees.*.percent_share', 'numeric', 'No', ''],
      ['employees.*.amount', 'numeric', 'No', '']],
     '200: success | 404 | 422 | 400: sai trạng thái'),
    ('6.5. Xoá', 'DELETE /api/v1/payroll/bonus-distributions/{id}', [], '200 | 404 | 400: sai trạng thái'),
    ('6.6. Gửi duyệt', 'POST /api/v1/payroll/bonus-distributions/{id}/submit', [], '200 | 400: tổng % ≠ 100'),
    ('6.7. Duyệt', 'POST /api/v1/payroll/bonus-distributions/{id}/approve', [], '200 | 403: không có quyền | 400: sai trạng thái'),
    ('6.8. Từ chối', 'POST /api/v1/payroll/bonus-distributions/{id}/reject', [], '200 | 403 | 400'),
    ('6.9. Preview tính toán', 'POST /api/v1/payroll/bonus-distributions/preview-calculate',
     [['department_id', 'int', 'Yes', ''], ['year', 'int', 'Yes', ''], ['pool_amount', 'numeric', 'Yes', ''],
      ['components', 'array', 'Yes', 'Danh sách thành phần + % + allow_adjustment']],
     '200: {components: [{..., employees: [...]}]} — không lưu DB'),
    ('6.10. Lấy quỹ thưởng', 'GET /api/v1/payroll/bonus-distributions/department-pool',
     [['department_id', 'int', 'Yes', ''], ['year', 'int', 'Yes', '']],
     '200: {pool_amount, department_id, year} | 404: chưa có quyết toán'),
]

for name, method, params, resp in apis:
    doc.add_heading(name, level=2)
    p = doc.add_paragraph(); run = p.add_run(method); run.bold = True; run.font.size = Pt(11)
    doc.add_paragraph('Auth: Bearer Token (JWT)').runs[0].font.size = Pt(10)
    if params:
        add_table(doc, ['Field', 'Type', 'Required', 'Mô tả'], params, col_widths=[4.5, 2, 1.5, 6])
    doc.add_paragraph(f'Response: {resp}')

# === 7. UI SPECIFICATION ===
doc.add_heading('7. UI Specification', level=1)

doc.add_heading('7.1. Trang danh sách', level=2)
doc.add_paragraph('Route: /payroll/bonus-distribution')
doc.add_paragraph('Bộ lọc: Công ty, Phòng ban, Năm, Trạng thái, Người tạo, Người duyệt, Từ ngày, Đến ngày (tất cả Select2/DatePicker)')
doc.add_paragraph('Bảng: STT, Mã phiếu (link), Công ty, Phòng ban, Năm, Tổng quỹ (format dấu phẩy), Người lập, Ngày lập, Người duyệt, Ngày duyệt, Trạng thái (badge), Thao tác (xem/sửa/xoá)')
doc.add_paragraph('Phân trang: 10/25/50/100, mặc định 100')

doc.add_heading('7.2. Trang tạo/sửa/chi tiết', level=2)
doc.add_paragraph('Route: /payroll/bonus-distribution/{id} (id=0 tạo mới)')

doc.add_paragraph('Header buttons:', style='List Bullet')
doc.add_paragraph('Nháp/Không duyệt: HỦY, LƯU NHÁP, LƯU & GỬI DUYỆT', style='List Bullet 2')
doc.add_paragraph('Chờ duyệt (có quyền): HỦY, KHÔNG DUYỆT, DUYỆT', style='List Bullet 2')
doc.add_paragraph('Đã duyệt: HỦY (chỉ xem)', style='List Bullet 2')

doc.add_paragraph('Bảng thiết lập thành phần:', style='List Bullet')
doc.add_paragraph('Cột: STT, Thành phần (Select2 lọc trùng), Loại (badge), Công thức, % quỹ (input 0-100), Quỹ TP (auto), Điều chỉnh (checkbox), Xoá', style='List Bullet 2')
doc.add_paragraph('Footer: Tổng % (xanh/đỏ) + Tổng quỹ', style='List Bullet 2')

doc.add_paragraph('Bảng phân bổ thưởng cuối:', style='List Bullet')
doc.add_paragraph('Header cố định khi scroll dọc (max-height 70vh)', style='List Bullet 2')
doc.add_paragraph('Dòng tổng quỹ phòng (sticky): % phân bổ, quỹ TP, đã chia (xanh/đỏ)', style='List Bullet 2')
doc.add_paragraph('Dòng NV: tên + mã, chức vụ, căn cứ/% chia, số tiền, tổng thưởng (% quỹ)', style='List Bullet 2')
doc.add_paragraph('Loading spinner khi tính toán', style='List Bullet 2')
doc.add_paragraph('Empty state: "Chưa có dữ liệu. Vui lòng nhấn Tính bảng thưởng"', style='List Bullet 2')

# === 8. WORKFLOW ===
doc.add_heading('8. Workflow trạng thái', level=1)
doc.add_paragraph('Nháp (0) → [Gửi duyệt] → Chờ duyệt (1) → [Duyệt] → Đã duyệt (2)')
doc.add_paragraph('Chờ duyệt (1) → [Không duyệt] → Không duyệt (3) → [Sửa + Gửi lại] → Chờ duyệt (1)')

# === 9. FILE REFERENCES ===
doc.add_heading('9. Phụ lục — File references', level=1)
add_table(doc, ['Layer', 'File path'], [
    ['Migration', 'Modules/Payroll/Database/Migrations/2026_04_06_000001_create_bonus_distributions_table.php'],
    ['Migration', 'Modules/Payroll/Database/Migrations/2026_04_06_000002_create_bonus_distribution_components_table.php'],
    ['Migration', 'Modules/Payroll/Database/Migrations/2026_04_06_000003_create_bonus_distribution_employees_table.php'],
    ['Entity', 'Modules/Payroll/Entities/BonusDistribution.php'],
    ['Entity', 'Modules/Payroll/Entities/BonusDistributionComponent.php'],
    ['Entity', 'Modules/Payroll/Entities/BonusDistributionEmployee.php'],
    ['Controller', 'Modules/Payroll/Http/Controllers/Api/V1/BonusDistributionController.php'],
    ['Service', 'Modules/Payroll/Services/BonusDistributionService.php'],
    ['Service', 'Modules/Payroll/Services/BonusVariableService.php'],
    ['Service', 'app/Services/ErpApiService.php'],
    ['Request', 'Modules/Payroll/Http/Requests/CreateBonusDistributionRequest.php'],
    ['Request', 'Modules/Payroll/Http/Requests/UpdateBonusDistributionRequest.php'],
    ['Resource', 'Modules/Payroll/Transformers/BonusDistributionResource/BonusDistributionDetailResource.php'],
    ['Resource', 'Modules/Payroll/Transformers/BonusDistributionResource/BonusDistributionListResource.php'],
    ['Routes', 'Modules/Payroll/Routes/api.php (lines 342-354)'],
    ['FE List', 'pages/payroll/bonus-distribution/index.vue'],
    ['FE Detail', 'pages/payroll/bonus-distribution/_id/index.vue'],
    ['Permission', 'Modules/Timesheet/Database/Seeders/PermissionsTableSeeder.php (id 1050-1054)'],
], col_widths=[2.5, 13])

output = '/Users/nguyentrancu/DEV/code/HRM/hrm-claude-config/docs/srs/SRS_Bonus_Distribution.docx'
doc.save(output)
print(f'Saved: {output}')
