from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# === Page setup ===
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# === Styles ===
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): '4472C4', qn('w:val'): 'clear'})
        shading.append(bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    return table


# === TITLE ===
title = doc.add_heading('SRS: Phân công quản lý giải pháp', level=1)

# Meta info
add_table(doc, ['Thông tin', 'Chi tiết'], [
    ['Module', 'Giao việc (Assign)'],
    ['Phiên bản', '1.0'],
    ['Ngày tạo', '2026-05-17'],
    ['Người tạo', '@dnsnamdang'],
    ['Trạng thái', 'Implemented'],
])

# === 1. GIỚI THIỆU ===
doc.add_heading('1. Giới thiệu', level=2)

doc.add_heading('1.1. Mục đích', level=3)
doc.add_paragraph('Cho phép Trưởng phòng GP / PM phân công lại PM hoặc Leader hạng mục cho giải pháp khi có thay đổi nhân sự. Lưu lịch sử phân công đầy đủ và gửi thông báo push cho các bên liên quan.')

doc.add_heading('1.2. Phạm vi', level=3)
doc.add_paragraph('Trong scope:')
for item in ['Phân công PM giải pháp', 'Phân công Leader hạng mục', 'Lịch sử phân công (popup)', 'Gửi thông báo cho PM/Leader cũ và mới', 'Xử lý vai trò mới cho người cũ']:
    doc.add_paragraph(f'• {item}')

doc.add_paragraph('Ngoài scope:')
for item in ['Phân công hàng loạt (batch)', 'Phân công tự động theo rule', 'Chuyển task từ người cũ sang mới']:
    doc.add_paragraph(f'• {item}')

doc.add_heading('1.3. Thuật ngữ', level=3)
add_table(doc, ['Thuật ngữ', 'Giải thích'], [
    ['Trưởng phòng GP', 'Người tạo giải pháp (solutions.created_by)'],
    ['PM', 'Project Manager — quản lý dự án giải pháp (solutions.pm_id)'],
    ['Leader', 'Người phụ trách hạng mục (solution_modules.leader_id)'],
    ['GP', 'Giải pháp'],
    ['Hạng mục', 'Module con trong giải pháp (solution_modules)'],
])

# === 2. ACTORS ===
doc.add_heading('2. Actors & Permissions', level=2)
add_table(doc, ['Actor', 'Mô tả', 'Permissions'], [
    ['Trưởng phòng GP', 'Người tạo giải pháp (created_by)', 'Phân công PM + Leader'],
    ['PM hiện tại', 'PM đang phụ trách (pm_id)', 'Phân công Leader'],
    ['PM/Leader cũ', 'Người bị thay thế', 'Nhận thông báo'],
    ['PM/Leader mới', 'Người được phân công', 'Nhận thông báo'],
    ['Thành viên', 'Nhân sự trong giải pháp', 'Xem lịch sử (readonly)'],
])

# === 3. USE CASES ===
doc.add_heading('3. Use Cases', level=2)

doc.add_heading('UC-01: Phân công PM giải pháp', level=3)
add_table(doc, ['', ''], [
    ['Actor', 'Trưởng phòng GP'],
    ['Precondition', 'GP status ≠ Đóng (2); User là created_by'],
    ['Main Flow', '1. Click "Phân công" → Popup\n2. Chọn loại PM\n3. Xem PM hiện tại\n4. (Tuỳ chọn) Chọn vai trò mới cho PM cũ\n5. (Nếu GP có module + chọn vai trò) Chọn module cho PM cũ\n6. Chọn PM mới (cùng phòng)\n7. Nhập ghi chú ≥50 ký tự\n8. "Xác nhận" → Confirm → Submit'],
    ['Postcondition', 'solutions.pm_id = PM mới; PM cũ thêm vào nhân sự (nếu có vai trò); Log; Notification'],
    ['Alternative', 'Không chọn vai trò → PM cũ rời dự án'],
    ['Exception', 'Validation fail; 403 không có quyền'],
])

doc.add_heading('UC-02: Phân công Leader hạng mục', level=3)
add_table(doc, ['', ''], [
    ['Actor', 'Trưởng phòng GP HOẶC PM hiện tại'],
    ['Precondition', 'GP status ≠ Đóng; GP có hạng mục; User là created_by hoặc pm_id'],
    ['Main Flow', '1. Click "Phân công" → Popup\n2. Chọn loại Leader\n3. Chọn hạng mục → Hiện Leader hiện tại\n4. (Tuỳ chọn) Vai trò mới cho Leader cũ\n5. Chọn Leader mới (cùng phòng)\n6. Ghi chú ≥50 ký tự\n7. "Xác nhận" → Confirm → Submit'],
    ['Postcondition', 'solution_modules.leader_id = Leader mới; Leader cũ thêm vào member (nếu có vai trò); Log; Notification'],
    ['Alternative', 'Không chọn vai trò → Leader cũ rời hạng mục'],
    ['Exception', 'Module không thuộc solution (404); Validation fail'],
])

doc.add_heading('UC-03: Xem lịch sử phân công', level=3)
add_table(doc, ['', ''], [
    ['Actor', 'Tất cả user có quyền xem tab Nhân sự'],
    ['Precondition', 'User có quyền truy cập trang manager'],
    ['Main Flow', '1. Click "Xem lịch sử" → Popup bảng 8 cột\n2. Sort mới nhất lên trên\n3. Click "Đóng" để thoát'],
    ['Postcondition', 'Không thay đổi dữ liệu'],
])

# === 4. BUSINESS RULES ===
doc.add_heading('4. Business Rules', level=2)
add_table(doc, ['ID', 'Rule', 'Mô tả', 'Áp dụng'], [
    ['BR-01', 'Quyền PC PM', 'Chỉ created_by', 'UC-01'],
    ['BR-02', 'Quyền PC Leader', 'created_by HOẶC pm_id', 'UC-02'],
    ['BR-03', 'GP Đóng', 'Ẩn button khi status=2', 'UC-01,02'],
    ['BR-04', 'PM cũ + không module', 'Thêm vào solution_members', 'UC-01'],
    ['BR-05', 'PM cũ + có module', 'Bắt buộc chọn module → solution_module_members', 'UC-01'],
    ['BR-06', 'Không chọn vai trò', 'Người cũ rời dự án', 'UC-01,02'],
    ['BR-07', 'Leader cũ → member', 'Thêm vào module_members cùng hạng mục', 'UC-02'],
    ['BR-08', 'Leader mới là member', 'Giữ nguyên record member', 'UC-02'],
    ['BR-09', 'DS chọn', 'Cùng department_id, loại trừ người hiện tại', 'UC-01,02'],
    ['BR-10', 'Ghi chú', 'Min 50 ký tự, required', 'UC-01,02'],
    ['BR-11', 'Notification', 'Gửi cho cũ + mới', 'UC-01,02'],
    ['BR-12', 'Confirm', 'Hiện confirm popup trước submit', 'UC-01,02'],
])

# === 5. DATA MODEL ===
doc.add_heading('5. Data Model', level=2)
doc.add_heading('Bảng: solution_manager_logs', level=3)
add_table(doc, ['Cột', 'Type', 'Nullable', 'Mô tả'], [
    ['id', 'bigint PK', 'No', 'Auto increment'],
    ['solution_id', 'unsignedBigInteger', 'No', 'FK → solutions.id'],
    ['solution_module_id', 'unsignedBigInteger', 'Yes', 'FK → solution_modules (null=PM)'],
    ['type', "enum('pm','leader')", 'No', 'Loại phân công'],
    ['old_manager_id', 'unsignedBigInteger', 'No', 'FK → employees.id'],
    ['new_manager_id', 'unsignedBigInteger', 'No', 'FK → employees.id'],
    ['old_manager_new_role_id', 'unsignedBigInteger', 'Yes', 'FK → project_roles.id'],
    ['old_manager_module_id', 'unsignedBigInteger', 'Yes', 'Module cho PM cũ'],
    ['note', 'text', 'No', 'Min 50 chars'],
    ['company_id', 'unsignedBigInteger', 'Yes', ''],
    ['department_id', 'unsignedBigInteger', 'Yes', ''],
    ['part_id', 'unsignedBigInteger', 'Yes', ''],
    ['created_by', 'unsignedBigInteger', 'Yes', ''],
    ['updated_by', 'unsignedBigInteger', 'Yes', ''],
    ['created_at', 'timestamp', 'Yes', ''],
    ['updated_at', 'timestamp', 'Yes', ''],
])

# === 6. API ===
doc.add_heading('6. API Specification', level=2)

doc.add_heading('6.1. POST /assign/solutions/{id}/manager/assign-manager', level=3)
add_table(doc, ['Field', 'Type', 'Required', 'Validate', 'Mô tả'], [
    ['type', 'string', 'Yes', 'in:pm,leader', 'Loại phân công'],
    ['solution_module_id', 'integer', 'Khi type=leader', 'exists', 'ID hạng mục'],
    ['new_manager_id', 'integer', 'Yes', 'exists:employees,id', 'ID người mới'],
    ['old_manager_new_role_id', 'integer', 'No', 'exists:project_roles,id', 'Vai trò mới cho người cũ'],
    ['old_manager_module_id', 'integer', 'Conditional', '', 'Module cho PM cũ'],
    ['note', 'string', 'Yes', 'min:50', 'Ghi chú'],
])

doc.add_heading('6.2. GET /assign/solutions/{id}/manager/assign-manager-logs', level=3)
doc.add_paragraph('Response: Array of logs sorted by created_at DESC. Fields: id, type, type_name, module_name, old_manager_name, new_manager_name, old_manager_new_role_name, note, created_by_name, created_at.')

# === 7. UI ===
doc.add_heading('7. UI Specification', level=2)
doc.add_paragraph('Vị trí: /assign/solutions/{id}/manager → Tab "Nhân sự"')
doc.add_paragraph('Button "Phân công" (secondary) bên trái "Thêm nhân sự"')
doc.add_paragraph('Button "Xem lịch sử" (light) luôn hiện')
doc.add_paragraph('Popup Phân công: size lg, form 7 fields, confirm trước submit')
doc.add_paragraph('Popup Lịch sử: size xl, bảng 8 cột readonly')

# === 8. NFR ===
doc.add_heading('8. Non-functional Requirements', level=2)
for item in [
    'Atomic: Transaction DB — rollback nếu fail',
    'Notification: Async, lỗi không ảnh hưởng response',
    'Confirm: Popup xác nhận trước submit',
    'Security: JWT auth + kiểm tra quyền',
    'Compatibility: PHP 7.4, Laravel 8, Nuxt 2, Vue 2',
]:
    doc.add_paragraph(f'• {item}')

# === 9. FILE REFERENCES ===
doc.add_heading('9. File References', level=2)
add_table(doc, ['Layer', 'File Path'], [
    ['Migration', 'database/migrations/2026_05_17_000001_create_solution_manager_logs_table.php'],
    ['Entity', 'Modules/Assign/Entities/SolutionManagerLog.php'],
    ['Controller', 'Modules/Assign/Http/Controllers/Api/V1/SolutionController.php'],
    ['Service', 'Modules/Assign/Services/SolutionService.php'],
    ['Request', 'Modules/Assign/Http/Requests/Solution/AssignManagerRequest.php'],
    ['Resource', 'Modules/Assign/Transformers/SolutionManagerLogResource.php'],
    ['Routes', 'Modules/Assign/Routes/api.php'],
    ['FE Modal', 'pages/assign/solutions/components/manager/AssignManagerModal.vue'],
    ['FE History', 'pages/assign/solutions/components/manager/AssignManagerHistoryModal.vue'],
    ['FE Tab', 'pages/assign/solutions/components/manager/HumanResourceTab.vue'],
])

# Save
doc.save('/Users/dnsnamdang/Documents/DNSMEDIA/websites/hrm/.plans/solution-manager-assignment/srs.docx')
print('SRS docx saved successfully!')
