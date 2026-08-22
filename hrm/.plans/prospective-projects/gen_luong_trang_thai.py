# -*- coding: utf-8 -*-
"""Sinh tài liệu Mô tả nghiệp vụ - Luồng trạng thái Dự án tiền khả thi."""
import sys
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
CENTER = WD_ALIGN_PARAGRAPH.CENTER

doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Pt(50)


def h(text, level=1):
    doc.add_paragraph(text, style="Heading %d" % level)


def para(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = 1
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
    doc.add_paragraph()
    return t


# ---------------- Bìa ----------------
p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("MÔ TẢ NGHIỆP VỤ"); r.bold = True; r.font.size = Pt(20)
p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("Luồng trạng thái Dự án tiền khả thi"); r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("Phân hệ Giao việc · Màn Dự án tiền khả thi · Cập nhật ngày 20/08/2026")
r.font.size = Pt(11)
doc.add_page_break()

# ---------------- 1 ----------------
h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
para("Tài liệu mô tả toàn bộ vòng đời trạng thái của một Dự án tiền khả thi: có những trạng thái "
     "nào, trạng thái nào chuyển sang trạng thái nào, ai/thao tác nào làm nó chuyển, và ở mỗi "
     "trạng thái người dùng còn được làm gì. Dành cho:")
bullet("Cán bộ kinh doanh, trưởng phòng kinh doanh — nắm dự án của mình đang ở bước nào.")
bullet("Phòng giải pháp, phòng xây dựng giá — biết thao tác của mình đẩy dự án sang bước nào.")
bullet("QA và người nghiệm thu — có bảng đối chiếu để kiểm thử đúng/sai.")
para("Tài liệu KHÔNG hướng dẫn thao tác từng nút bấm (đó là tài liệu Hướng dẫn sử dụng).")

# ---------------- 2 ----------------
h("2. HAI BỘ TRẠNG THÁI KHÁC NHAU", 1)
para("Hệ thống dùng hai bộ trạng thái riêng biệt, tuỳ loại dự án:")
bullet("Dự án thường (dự án độc lập và dự án con): 12 trạng thái, đi qua đủ các bước làm giải "
       "pháp kỹ thuật, dự toán, báo giá, hợp đồng.")
bullet("Dự án cha (gói thầu tổng): 8 trạng thái, lược bỏ hoàn toàn phần giải pháp kỹ thuật vì "
       "dự án cha không tự làm giải pháp — phần đó do các dự án con thực hiện.")
para("Lưu ý quan trọng: hai bộ này có những trạng thái trùng số thứ tự nhưng KHÁC TÊN và KHÁC "
     "Ý NGHĨA (ví dụ bước thứ 7: dự án thường là “Thương thảo giá”, dự án cha là “Trình duyệt "
     "hợp đồng”). Khi đọc báo cáo hoặc lọc danh sách phải xem dự án đó là loại nào rồi mới đọc "
     "tên trạng thái.")

h("2.1. Bộ trạng thái của dự án thường", 2)
table([
    ["Bước", "Tên trạng thái", "Ý nghĩa nghiệp vụ"],
    ["1", "Đang tạo", "Bản nháp. Được lưu thiếu thông tin, chưa sinh mã dự án."],
    ["2", "Thu thập thông tin dự án", "Dự án đã lưu chính thức, đang khảo sát/họp với khách hàng, ghi phiếu thu thập thông tin."],
    ["3", "Chờ tiếp nhận làm giải pháp", "Đã gửi Yêu cầu làm giải pháp sang phòng giải pháp, đang chờ tiếp nhận."],
    ["4", "Đang làm giải pháp", "Phòng giải pháp (hoặc chính đơn vị kinh doanh nếu tự triển khai) đang xây dựng giải pháp."],
    ["5", "Đã duyệt giải pháp", "Hồ sơ giải pháp đã được duyệt."],
    ["6", "Dự toán", "Đã phát sinh Yêu cầu xây dựng giá hoặc đã có báo giá — đang làm giá."],
    ["7", "Thương thảo giá", "Báo giá đã được duyệt, đang thương thảo giá với khách hàng."],
    ["8", "Thương thảo dự án hợp đồng", "Đã chốt giải pháp, chuyển sang thương thảo dự án/hợp đồng."],
    ["9", "Thực hiện hợp đồng", "Hợp đồng đang triển khai."],
    ["10", "Nghiệm thu và thanh lý hợp đồng", "Đang nghiệm thu, thanh lý hợp đồng."],
    ["11", "Đóng/Không thực hiện dự án", "Dự án dừng, không theo tiếp. Có lưu nguyên nhân thất bại, ghi chú, người đóng và thời điểm đóng."],
    ["12", "Kết thúc và lưu trữ", "Dự án hoàn tất, chỉ còn để tra cứu."],
])

h("2.2. Bộ trạng thái của dự án cha", 2)
table([
    ["Bước", "Tên trạng thái", "Ý nghĩa nghiệp vụ"],
    ["1", "Đang tạo", "Bản nháp của gói thầu tổng."],
    ["2", "Đang thực hiện", "Đã có ít nhất một dự án con trực thuộc; các con đang chạy luồng riêng."],
    ["7", "Trình duyệt hợp đồng", "Báo giá tổng của gói thầu đã được tạo hợp đồng bên hệ thống ERP."],
    ["8", "Thương thảo DA/Hợp đồng", "Đang thương thảo dự án/hợp đồng cho cả gói."],
    ["9", "HĐ đủ điều kiện thực hiện", "Hợp đồng đã đủ điều kiện triển khai."],
    ["10", "Nghiệm thu & Thanh lý", "Đang nghiệm thu, thanh lý cả gói."],
    ["11", "Đóng/Không thực hiện dự án", "Đóng cả gói thầu — kéo theo đóng toàn bộ dự án con chưa đóng."],
    ["12", "Kết thúc & lưu trữ", "Gói thầu hoàn tất, chỉ còn để tra cứu."],
])

# ---------------- 3 ----------------
h("3. MA TRẬN CHUYỂN TRẠNG THÁI — DỰ ÁN THƯỜNG", 1)
para("Bảng dưới liệt kê đầy đủ các đường chuyển trạng thái hiện hệ thống đang tự thực hiện. Một "
     "dòng = một sự kiện nghiệp vụ. Trạng thái nào không xuất hiện ở cột “Trạng thái đến” nghĩa "
     "là hiện chưa có thao tác nào trong hệ thống đưa dự án tới đó (xem chương 7).")
table([
    ["Trạng thái trước", "Trạng thái sau", "Sự kiện làm nó chuyển", "Ai thực hiện", "Thao tác ở đâu"],
    ["(chưa có)", "Đang tạo", "Bấm Lưu nháp khi tạo mới dự án", "Nhân viên kinh doanh", "Màn Thêm mới dự án tiền khả thi"],
    ["(chưa có)", "Thu thập thông tin dự án", "Bấm Lưu (lưu chính thức) khi tạo mới — hệ thống sinh luôn mã dự án", "Nhân viên kinh doanh", "Màn Thêm mới dự án tiền khả thi"],
    ["Đang tạo", "Thu thập thông tin dự án", "Bấm Lưu (lưu chính thức) ở màn Sửa dự án", "Nhân viên kinh doanh", "Màn Sửa dự án"],
    ["Đang tạo", "Thu thập thông tin dự án", "Lưu cuộc họp hoặc lưu phiếu thu thập thông tin gắn với dự án", "Người chủ trì cuộc họp / nhân viên kinh doanh", "Tab Cuộc họp trong màn chi tiết dự án"],
    ["Thu thập thông tin dự án", "Chờ tiếp nhận làm giải pháp", "Gửi Yêu cầu làm giải pháp (phiếu ở trạng thái Chờ tiếp nhận)", "Nhân viên kinh doanh", "Màn Yêu cầu làm giải pháp"],
    ["Chờ tiếp nhận làm giải pháp", "Thu thập thông tin dự án", "Phòng giải pháp TỪ CHỐI yêu cầu làm giải pháp", "Trưởng phòng giải pháp", "Màn Yêu cầu làm giải pháp"],
    ["Chờ tiếp nhận làm giải pháp", "Thu thập thông tin dự án", "Người lập HUỶ yêu cầu làm giải pháp", "Nhân viên kinh doanh", "Màn Yêu cầu làm giải pháp"],
    ["Thu thập thông tin dự án / Chờ tiếp nhận làm giải pháp", "Đang làm giải pháp", "Giải pháp được tạo và gửi đi, hoặc đang ở các bước chờ duyệt nội bộ của giải pháp", "Phòng giải pháp (dự án liên phòng/theo phòng) hoặc chính nhân viên kinh doanh (dự án Tự triển khai)", "Màn Giải pháp"],
    ["Đã duyệt giải pháp", "Đang làm giải pháp", "Hồ sơ trình duyệt giải pháp bị TỪ CHỐI, hoặc tạo phiên bản giải pháp mới", "Người duyệt hồ sơ / phòng giải pháp", "Màn Giải pháp"],
    ["Đang làm giải pháp", "Đã duyệt giải pháp", "Hồ sơ trình duyệt giải pháp được DUYỆT. Dự án Tự triển khai thì duyệt tự động ngay khi trình hồ sơ", "Người duyệt hồ sơ giải pháp", "Màn Giải pháp — Hồ sơ trình duyệt"],
    ["Đã duyệt giải pháp (hoặc bất kỳ bước nào trước Dự toán)", "Dự toán", "Gửi Yêu cầu xây dựng giá", "Nhân viên kinh doanh", "Màn Yêu cầu xây dựng giá"],
    ["Bất kỳ bước nào trước Dự toán", "Dự toán", "Tạo báo giá cho dự án (kể cả tạo từ BOM)", "Nhân viên xây dựng giá / kinh doanh", "Màn Báo giá"],
    ["Dự toán", "Thương thảo giá", "Báo giá của dự án được DUYỆT (duyệt cuối)", "Trưởng phòng / Ban giám đốc theo phân cấp duyệt giá", "Màn Báo giá"],
    ["Đã duyệt giải pháp / Thương thảo giá", "Thương thảo dự án hợp đồng", "Bấm Chốt giải pháp trên hồ sơ đã duyệt (hoặc hết hiệu lực)", "Chỉ nhân viên kinh doanh phụ trách dự án", "Màn chi tiết dự án — nút Chốt giải pháp"],
    ["Mọi trạng thái trừ Đóng", "Đóng/Không thực hiện dự án", "Bấm Đóng dự án, chọn nguyên nhân thất bại", "Chỉ nhân viên kinh doanh phụ trách dự án", "Màn chi tiết dự án — nút Đóng dự án"],
    ["Mọi trạng thái trừ Đóng", "Đóng/Không thực hiện dự án", "Dự án CHA bị đóng — toàn bộ dự án con chưa đóng bị đóng theo", "Nhân viên kinh doanh phụ trách dự án cha", "Màn chi tiết dự án cha — nút Đóng dự án"],
])

para("Ghi chú về nhánh “Tự triển khai”: dự án chọn cách triển khai Tự triển khai thì không đi qua "
     "Yêu cầu làm giải pháp, nên không bao giờ ở trạng thái Chờ tiếp nhận làm giải pháp — từ "
     "Thu thập thông tin dự án nhảy thẳng sang Đang làm giải pháp khi tạo giải pháp, và khi trình "
     "hồ sơ thì được duyệt luôn nên sang Đã duyệt giải pháp ngay.")

# ---------------- 4 ----------------
h("4. MA TRẬN CHUYỂN TRẠNG THÁI — DỰ ÁN CHA", 1)
table([
    ["Trạng thái trước", "Trạng thái sau", "Sự kiện làm nó chuyển", "Ai thực hiện"],
    ["(chưa có)", "Đang tạo", "Bấm Lưu nháp khi tạo dự án cha", "Nhân viên kinh doanh"],
    ["(chưa có)", "Đang thực hiện", "Bấm Lưu chính thức khi tạo dự án cha", "Nhân viên kinh doanh"],
    ["Đang tạo", "Đang thực hiện", "Lưu một dự án con bất kỳ trỏ về dự án cha này (dự án cha có con đầu tiên)", "Nhân viên kinh doanh của dự án con"],
    ["Đang thực hiện", "Trình duyệt hợp đồng", "Báo giá tổng của gói thầu được tạo hợp đồng bên hệ thống ERP", "Người tạo hợp đồng bên ERP"],
    ["Mọi trạng thái trừ Đóng", "Đóng/Không thực hiện dự án", "Bấm Đóng dự án ở dự án cha (chỉ hiện khi cha đang ở Đang thực hiện)", "Chỉ nhân viên kinh doanh phụ trách dự án cha"],
])
para("Khi đóng dự án cha, hệ thống làm liên hoàn: đóng dự án cha, chuyển các báo giá tổng của gói "
     "sang Đóng, rồi đóng lần lượt từng dự án con chưa đóng (kèm theo giải pháp, hạng mục giải "
     "pháp, yêu cầu làm giải pháp, yêu cầu xây dựng giá và báo giá của con). Một dự án con lỗi "
     "không chặn các con còn lại.")

# ---------------- 5 ----------------
h("5. SƠ ĐỒ LUỒNG RÚT GỌN", 1)
para("Dự án thường (đường đi thuận lợi):", bold=True)
para("Đang tạo → Thu thập thông tin dự án → Chờ tiếp nhận làm giải pháp → Đang làm giải pháp → "
     "Đã duyệt giải pháp → Dự toán → Thương thảo giá → Thương thảo dự án hợp đồng → "
     "Thực hiện hợp đồng → Nghiệm thu và thanh lý hợp đồng → Kết thúc và lưu trữ")
para("Các đường quay lui:", bold=True)
bullet("Chờ tiếp nhận làm giải pháp → Thu thập thông tin dự án (từ chối hoặc huỷ yêu cầu làm giải pháp).")
bullet("Đã duyệt giải pháp → Đang làm giải pháp (hồ sơ bị từ chối hoặc mở phiên bản giải pháp mới).")
para("Đường thoát:", bold=True)
bullet("Bất kỳ trạng thái nào (trừ khi đã đóng) → Đóng/Không thực hiện dự án. Đây là trạng thái "
       "cuối, không có đường quay lại.")
para("Dự án cha:", bold=True)
para("Đang tạo → Đang thực hiện → Trình duyệt hợp đồng → Thương thảo DA/Hợp đồng → "
     "HĐ đủ điều kiện thực hiện → Nghiệm thu & Thanh lý → Kết thúc & lưu trữ, "
     "và nhánh thoát Đóng/Không thực hiện dự án.")

# ---------------- 6 ----------------
h("6. Ở MỖI TRẠNG THÁI CÒN LÀM ĐƯỢC GÌ", 1)
h("6.1. Ma trận trạng thái × thao tác (dự án thường)", 2)
table([
    ["Thao tác", "Được phép khi", "Điều kiện kèm theo"],
    ["Sửa dự án", "Mọi trạng thái trừ Đóng/Không thực hiện dự án", "Khi dự án đã đóng thì nút Sửa và Xoá bị ẩn ở màn chi tiết"],
    ["Sửa ô Cách triển khai dự án và Có cần làm giải pháp", "Đang tạo, Thu thập thông tin dự án", "Ngoài ra còn bị khoá hẳn khi dự án đã có giải pháp hoặc đã có yêu cầu làm giải pháp"],
    ["Đổi Loại dự án (độc lập / cha / con)", "Chỉ Đang tạo", "Và chưa có dự án con nào trực thuộc"],
    ["Xoá dự án", "Chỉ Đang tạo", "Chỉ người tạo dự án, và dự án chưa có con trực thuộc"],
    ["Tạo cuộc họp", "Mọi trạng thái trừ Đang tạo và Đóng/Không thực hiện dự án", "—"],
    ["Tạo giải pháp trực tiếp", "Chỉ Thu thập thông tin dự án", "Dự án Tự triển khai, có làm giải pháp, chưa có giải pháp, phiếu thu thập thông tin đã nhập đủ trường bắt buộc"],
    ["Tạo yêu cầu làm giải pháp", "Thu thập thông tin dự án", "Dự án chưa có yêu cầu làm giải pháp, do chính người đang đăng nhập tạo, không phải dự án cha và không phải dự án Tự triển khai"],
    ["Chốt giải pháp", "Khi đã có hồ sơ trình duyệt ở trạng thái Đã duyệt hoặc Hết hiệu lực", "Chỉ nhân viên kinh doanh phụ trách; dự án chưa đóng"],
    ["Đóng dự án", "Mọi trạng thái trừ Đóng/Không thực hiện dự án", "Chỉ nhân viên kinh doanh phụ trách dự án"],
])

h("6.2. Ma trận trạng thái × thao tác (dự án cha)", 2)
table([
    ["Thao tác", "Được phép khi", "Điều kiện kèm theo"],
    ["Sửa Tổng ngân sách dự kiến", "Trước bước Trình duyệt hợp đồng", "Từ Trình duyệt hợp đồng trở đi hệ thống chặn, báo dự án đã sang bước hợp đồng nên không đổi được tổng ngân sách"],
    ["Đổi Khách hàng của dự án cha", "Khi chưa có dự án con", "Đã có con thì khoá vì các con đang kế thừa khách hàng từ cha"],
    ["Được chọn làm dự án cha cho dự án con mới", "Mọi trạng thái trừ Đóng/Không thực hiện dự án và Kết thúc & lưu trữ", "—"],
    ["Đóng dự án cha", "Chỉ khi đang ở Đang thực hiện", "Chỉ nhân viên kinh doanh phụ trách dự án cha"],
])

# ---------------- 7 ----------------
h("7. QUY TẮC ĐI KÈM TRẠNG THÁI", 1)
bullet("Sinh mã dự án: dự án chỉ được cấp mã khi rời khỏi trạng thái Đang tạo. Bản nháp không có mã.")
bullet("Lưu nháp bỏ qua kiểm tra bắt buộc nhập, trừ Tên dự án. Lưu chính thức mới kiểm đủ các "
       "trường bắt buộc.")
bullet("Mọi lần đổi trạng thái đều được ghi lại nhật ký (trạng thái cũ, trạng thái mới, thời "
       "điểm) để phục vụ báo cáo tiến độ theo mốc thời gian.")
bullet("Dự án ở trạng thái Đang tạo bị loại khỏi các báo cáo và một số danh sách chọn — bản nháp "
       "không được coi là dự án thật.")
bullet("Đóng dự án bắt buộc chọn nguyên nhân thất bại; hệ thống lưu nguyên nhân, ghi chú, người "
       "đóng, thời điểm đóng và hiển thị dải cảnh báo đỏ ở màn chi tiết.")
bullet("Đóng dự án gửi thông báo tới những người liên quan của các giải pháp đang mở, kèm liên "
       "kết mở thẳng dự án vừa đóng.")
bullet("Không đóng lại dự án đã đóng — hệ thống báo dự án đã đóng.")

# ---------------- 8 ----------------
h("8. GIỚI HẠN HIỆN TẠI CẦN LƯU Ý", 1)
para("Đây là các điểm nên biết trước khi nghiệm thu, không phải lỗi cần sửa ngay:")
bullet("Ba trạng thái cuối của dự án thường — Thực hiện hợp đồng, Nghiệm thu và thanh lý hợp "
       "đồng, Kết thúc và lưu trữ — hiện chưa có thao tác nào trong hệ thống tự đưa dự án tới. "
       "Chúng mới được dùng để hiển thị và để lọc/báo cáo; phần hợp đồng nằm bên hệ thống ERP.")
bullet("Tương tự ở dự án cha: các bước Thương thảo DA/Hợp đồng, HĐ đủ điều kiện thực hiện, "
       "Nghiệm thu & Thanh lý, Kết thúc & lưu trữ hiện chưa có thao tác nào tự chuyển tới.")
bullet("Màn Sửa dự án luôn gửi lên trạng thái Thu thập thông tin dự án khi bấm Lưu. Với dự án đã "
       "đi xa hơn (ví dụ đang ở Dự toán hay Thương thảo giá), thao tác sửa và lưu lại sẽ kéo "
       "trạng thái dự án về Thu thập thông tin dự án. Cần thống nhất mong muốn nghiệp vụ ở điểm "
       "này: hoặc khoá nút Sửa từ một bước nào đó trở đi, hoặc giữ nguyên trạng thái cũ khi lưu.")
bullet("Hai bộ trạng thái dùng chung số thứ tự nhưng khác tên. Khi làm báo cáo hoặc lọc danh "
       "sách theo trạng thái, phải tách riêng dự án cha và dự án thường, nếu không số liệu sẽ bị "
       "gộp nhầm giữa hai nghĩa khác nhau.")

out = "/Users/manhcuong/Desktop/dns/HRM/.plans/prospective-projects/Mô tả nghiệp vụ - Luồng trạng thái Dự án tiền khả thi.docx"
doc.save(out)
print("Đã lưu:", out)
