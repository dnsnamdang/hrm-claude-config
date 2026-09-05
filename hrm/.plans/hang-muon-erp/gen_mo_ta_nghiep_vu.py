# -*- coding: utf-8 -*-
"""Sinh tai lieu mo ta nghiep vu "Phieu yeu cau xuat hang muon" (ERP)."""
import os
import re
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Mô tả nghiệp vụ - Phiếu yêu cầu xuất hàng mượn.docx")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)


def h(text, level):
    doc.add_paragraph(text, style="Heading %d" % level)


def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = 1
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.cell(i, j).text = val
    return t


# ───────────────────────────── BÌA ─────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("MÔ TẢ NGHIỆP VỤ")
r.bold = True
r.font.size = Pt(20)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Phiếu yêu cầu xuất hàng mượn")
r.bold = True
r.font.size = Pt(16)

ip = doc.add_paragraph()
ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
ip.add_run("Hệ thống ERP · Phân hệ Kho → Quản lý hàng mượn · Cập nhật ngày 04/09/2026")

doc.add_page_break()

# ───────────────────────────── 1 ─────────────────────────────
h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
para("Tài liệu mô tả đầy đủ nghiệp vụ của chức năng “Phiếu yêu cầu xuất hàng mượn”: chức năng này "
     "dùng để làm gì, ai tham gia, phiếu đi qua những bước nào, mỗi bước hệ thống gửi thông báo cho "
     "ai, và những quy tắc bắt buộc phải tuân thủ. Tài liệu không hướng dẫn thao tác từng nút bấm — "
     "phần đó thuộc về tài liệu hướng dẫn sử dụng.")
bullet("Nhân viên kinh doanh — người lập phiếu, cần biết khi nào được lập và lập xong thì chờ ai.")
bullet("Kế toán kho — người duyệt, cần biết duyệt phiếu này thì hàng mượn được xử lý ra sao.")
bullet("Kế toán công nợ và quản lý kho — cần biết vì sao hàng mượn biến mất khỏi danh sách theo dõi.")
bullet("Trưởng phòng kinh doanh, ban giám đốc — cần biết phạm vi dữ liệu mình xem được.")
bullet("Nhân sự nghiệm thu và kiểm thử — cần biết đủ nhánh xử lý và đủ thông báo để kiểm tra.")

# ───────────────────────────── 2 ─────────────────────────────
h("2. CHỨC NĂNG NÀY DÙNG ĐỂ LÀM GÌ", 1)

h("2.1. Mục đích", 2)
para("Khi công ty cho khách hàng mượn hàng, hệ thống ghi nhận số hàng đó là “hàng đang mượn” và theo "
     "dõi cho tới khi khách trả về kho. Trên thực tế có nhiều trường hợp khách không trả hàng về kho "
     "nữa: hàng đã được giao chính thức cho khách qua một chứng từ xuất khác, hoặc số hàng mượn đó đã "
     "được xử lý bằng một đường khác ngoài việc nhập lại kho. Nếu không có chức năng này, số hàng đó "
     "sẽ treo mãi ở trạng thái đang mượn, kho và kế toán phải theo dõi một khoản không bao giờ tất "
     "toán được.")
para("“Phiếu yêu cầu xuất hàng mượn” là đề nghị của nhân viên kinh doanh gửi kế toán kho, xin tất "
     "toán số hàng khách đang mượn mà không phải nhập lại kho. Kế toán kho duyệt bằng cách lập “Phiếu "
     "xuất hàng mượn”. Sau khi duyệt, số hàng được ghi nhận là đã xử lý xong; khi toàn bộ hàng trên "
     "phiếu xuất mượn gốc đã xử lý hết, phiếu xuất mượn đó chuyển sang trạng thái Đã trả và không còn "
     "nằm trong danh sách hàng mượn phải theo dõi.")

h("2.2. Vị trí trong luồng lớn", 2)
para("Chức năng nằm ở đoạn cuối của vòng đời hàng mượn. Trình tự đầy đủ:")
bullet("Nhân viên kinh doanh lập phiếu yêu cầu xuất hàng, chọn hình thức Xuất mượn; hàng ra khỏi kho "
       "và được đánh dấu Đã mượn.")
bullet("Trong thời gian khách giữ hàng, có bốn hướng xử lý: khách trả hàng về kho, khách xin gia hạn "
       "mượn, khách mua luôn số hàng đang mượn, hoặc tất toán số hàng đó bằng phiếu yêu cầu xuất hàng "
       "mượn — chính là chức năng tài liệu này mô tả.")
bullet("Sau khi tất toán hết, phiếu xuất mượn gốc đóng lại ở trạng thái Đã trả.")

h("2.3. Giá trị mang lại", 2)
bullet("Kho và kế toán không phải theo dõi những khoản hàng mượn thực chất đã xử lý xong.")
bullet("Mỗi lần tất toán đều có chứng từ, có người đề nghị, người duyệt, thời điểm duyệt và tài liệu "
       "đính kèm chứng minh — tra ngược lại được khi đối chiếu công nợ.")
bullet("Kế toán kho là người chốt cuối, nên số lượng tất toán luôn được rà soát lại chứ không do một "
       "mình nhân viên kinh doanh quyết định.")
bullet("Hệ thống tự khóa phần số lượng đang chờ xử lý, tránh việc cùng một lô hàng mượn bị vừa xin "
       "trả, vừa xin bán, vừa xin tất toán.")

# ───────────────────────────── 3 ─────────────────────────────
h("3. NHỮNG AI THAM GIA", 1)
table([
    ["Vai trò", "Làm gì trong luồng này"],
    ["Nhân viên kinh doanh (người lập phiếu xuất mượn)",
     "Lập phiếu yêu cầu xuất hàng mượn cho chính những phiếu xuất mượn do mình tạo; chọn phiếu mượn, "
     "nhập số lượng đề nghị, ghi chú và đính kèm tài liệu; theo dõi kết quả duyệt."],
    ["Kế toán kho",
     "Nhận thông báo, mở phiếu kiểm tra tài liệu đính kèm và số lượng; từ chối kèm lý do, hoặc duyệt "
     "bằng cách lập phiếu xuất hàng mượn với số lượng được duyệt cho từng dòng hàng."],
    ["Trưởng phòng / lãnh đạo có quyền xem theo cấp",
     "Không thao tác trên phiếu. Chỉ theo dõi phiếu của phòng ban, công ty mình để nắm tình hình hàng "
     "mượn của bộ phận."],
    ["Khách hàng",
     "Không thao tác trên hệ thống. Là bên đang giữ số hàng mượn được tất toán; tài liệu thỏa thuận "
     "hoặc chứng từ giao nhận với khách chính là tệp đính kèm bắt buộc của phiếu."],
])
para("Luồng này không có bước duyệt của trưởng phòng hay ban giám đốc. Người duyệt duy nhất là kế "
     "toán kho.")

# ───────────────────────────── 4 ─────────────────────────────
h("4. VÒNG ĐỜI CỦA PHIẾU", 1)

h("4.1. Trạng thái của phiếu yêu cầu xuất hàng mượn", 2)
table([
    ["Trạng thái", "Ý nghĩa", "Ai làm phiếu chuyển sang trạng thái này"],
    ["Đang tạo",
     "Trạng thái nháp. Hiện tại luồng không sinh ra trạng thái này: bấm lưu là phiếu đi thẳng sang Chờ "
     "duyệt, không có bước lưu nháp.",
     "Không ai — trạng thái tồn tại trên bộ lọc nhưng không phát sinh trong thực tế."],
    ["Chờ duyệt",
     "Phiếu đã gửi đi, đang chờ kế toán kho xử lý. Phiếu không sửa, không xóa được nữa.",
     "Nhân viên kinh doanh, ngay khi bấm lưu phiếu."],
    ["Đã duyệt",
     "Kế toán kho đã đồng ý và đã lập phiếu xuất hàng mượn kèm theo. Số hàng tương ứng được tất toán.",
     "Kế toán kho, tại thời điểm lưu phiếu xuất hàng mượn thành công."],
    ["Không duyệt",
     "Kế toán kho từ chối. Lý do từ chối được lưu lại và hiển thị ở mục Ghi chú duyệt trên phiếu. Đây "
     "là điểm kết thúc, phiếu không quay lại được.",
     "Kế toán kho, khi bấm Từ chối và nhập lý do."],
])

h("4.2. Trạng thái hàng mượn trên phiếu xuất mượn gốc", 2)
para("Phiếu yêu cầu xuất hàng mượn không chỉ đổi trạng thái của chính nó, mà còn cập nhật ngược về "
     "phiếu xuất mượn gốc — đây là phần người dùng hay bỏ sót:")
table([
    ["Trạng thái hàng mượn", "Ý nghĩa", "Ai làm nó chuyển sang trạng thái đó"],
    ["Chờ mượn", "Đã lập phiếu xuất mượn nhưng hàng chưa ra khỏi kho xong.",
     "Do luồng xuất kho của phiếu xuất mượn quyết định, không liên quan chức năng này."],
    ["Đã mượn", "Khách đang giữ hàng. Đây là điều kiện bắt buộc để lập phiếu yêu cầu xuất hàng mượn.",
     "Do luồng xuất kho của phiếu xuất mượn quyết định."],
    ["Đã trả", "Toàn bộ hàng trên phiếu xuất mượn đã được xử lý xong, không còn phải theo dõi.",
     "Cập nhật ngược tự động khi kế toán kho lưu phiếu xuất hàng mượn và dòng hàng cuối cùng của phiếu "
     "mượn được tất toán hết. Ngoài chức năng này, luồng trả hàng về kho và luồng xuất bán hàng mượn "
     "cũng có thể đưa phiếu về trạng thái Đã trả."],
])

h("4.3. Chứng từ sinh ra kèm theo", 2)
para("Khi kế toán kho duyệt, hệ thống sinh thêm một chứng từ riêng là “Phiếu xuất hàng mượn”. Chứng "
     "từ này được lập ở trạng thái Đã duyệt ngay từ đầu, không có bước duyệt riêng, và là nơi ghi số "
     "lượng thực tế được duyệt cho từng dòng hàng.")

# ───────────────────────────── 5 ─────────────────────────────
h("5. LUỒNG HOẠT ĐỘNG CHI TIẾT", 1)

h("5.1. Điều kiện để bắt đầu", 2)
para("Chỉ lập được phiếu yêu cầu xuất hàng mượn khi có phiếu xuất mượn thỏa đồng thời các điều kiện "
     "sau. Danh sách trong ô chọn phiếu đã lọc sẵn theo đúng các điều kiện này, nên người dùng không "
     "chọn nhầm được:")
bullet("Phiếu yêu cầu xuất hàng phải là hình thức Xuất mượn.")
bullet("Phiếu đã ở trạng thái Đã hạch toán — tức hàng đã xuất kho và hoàn tất thủ tục.")
bullet("Trạng thái hàng mượn của phiếu là Đã mượn — khách vẫn đang giữ hàng.")
bullet("Người đang đăng nhập chính là người đã lập phiếu xuất mượn đó. Người khác, kể cả trưởng "
       "phòng, không lập thay được.")

h("5.2. Nhân viên kinh doanh lập phiếu", 2)
para("Người lập vào menu Kho → Quản lý hàng mượn → Phiếu Yêu cầu xuất hàng mượn, bấm thêm mới và "
     "thực hiện tuần tự:")
bullet("Chọn phiếu xuất mượn: mở ô Phiếu xuất mượn và chọn một hoặc nhiều phiếu. Một phiếu yêu cầu "
       "có thể gom nhiều phiếu mượn cùng lúc, miễn là đều do chính mình lập.")
bullet("Hệ thống nạp danh sách hàng hóa đang mượn của các phiếu vừa chọn, kèm số lượng còn được "
       "phép đề nghị của từng dòng.")
bullet("Nhập số lượng xuất cho từng dòng hàng, theo từng phiếu mượn. Cùng một mặt hàng mượn từ nhiều "
       "phiếu thì nhập riêng cho từng phiếu, hệ thống cộng lại thành số lượng của mặt hàng đó.")
bullet("Chọn đơn vị tính. Hệ thống tự quy đổi khi so sánh với số lượng đang mượn, nên chọn thùng hay "
       "chọn chiếc đều kiểm soát đúng.")
bullet("Nhập ghi chú nêu lý do đề nghị tất toán — bắt buộc.")
bullet("Đính kèm tối thiểu một tệp PDF làm căn cứ — bắt buộc.")
bullet("Bấm lưu. Hệ thống sinh mã phiếu, chuyển trạng thái Chờ duyệt và gửi thông báo cho kế toán kho.")
para("Cách tính số lượng còn được đề nghị của một dòng hàng: lấy số lượng đã xuất mượn, trừ đi phần "
     "đã tất toán trước đó, trừ tiếp phần đang bị khóa bởi các chứng từ chưa xử lý xong — gồm phiếu "
     "trả hàng về kho chưa hoàn tất, yêu cầu xuất bán hàng mượn đang chờ duyệt và các yêu cầu xuất "
     "hàng mượn khác đang chờ duyệt. Nhập vượt con số này thì hệ thống chặn không cho lưu.", "Lưu ý: ")

h("5.3. Kế toán kho tiếp nhận", 2)
para("Kế toán kho nhận thông báo qua chuông thông báo, hoặc chủ động vào một trong hai lối: menu Kho "
     "→ Yêu cầu xuất hàng mượn, hoặc menu Duyệt → Phiếu yêu cầu xuất hàng mượn chờ duyệt. Mở phiếu ra "
     "sẽ thấy thông tin chung, danh sách phiếu xuất mượn liên quan, phòng ban yêu cầu, tệp đính kèm "
     "và bảng chi tiết hàng hóa kèm số lượng đề nghị của từng phiếu mượn.")
para("Từ đây có hai hướng xử lý:")
table([
    ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
    ["Từ chối",
     "Bấm nút Từ chối, hệ thống mở ô nhập lý do. Lý do bắt buộc nhập, tối đa 255 ký tự.",
     "Phiếu chuyển sang Không duyệt, lưu lại người duyệt, thời điểm duyệt và lý do vào mục Ghi chú "
     "duyệt. Người lập nhận thông báo. Hàng mượn giữ nguyên, phần số lượng đang bị khóa được nhả ra "
     "để dùng cho chứng từ khác. Luồng kết thúc."],
    ["Duyệt",
     "Bấm nút Tạo phiếu xuất hàng mượn. Hệ thống chuyển sang màn lập phiếu xuất hàng mượn với dữ liệu "
     "lấy sẵn từ phiếu yêu cầu. Kế toán kho nhập số lượng được duyệt cho từng dòng — có thể duyệt ít "
     "hơn số đề nghị — rồi lưu.",
     "Sinh phiếu xuất hàng mượn, phiếu yêu cầu chuyển sang Đã duyệt, hàng mượn được tất toán theo số "
     "lượng duyệt. Người lập nhận thông báo. Luồng kết thúc."],
])
para("Bấm “Tạo phiếu xuất hàng mượn” chính là thao tác duyệt — trên màn hình không có nút tên là "
     "Duyệt. Phiếu yêu cầu chỉ thực sự chuyển sang Đã duyệt sau khi phiếu xuất hàng mượn được lưu "
     "thành công; nếu kế toán kho mở màn lập phiếu rồi bỏ ngang, phiếu yêu cầu vẫn ở Chờ duyệt.",
     "Điểm dễ hiểu nhầm: ")

h("5.4. Hệ thống xử lý sau khi duyệt", 2)
para("Khi phiếu xuất hàng mượn được lưu, hệ thống thực hiện liền một mạch, hoặc thành công trọn vẹn "
     "hoặc không ghi nhận gì:")
bullet("Sinh phiếu xuất hàng mượn với mã riêng, ở trạng thái Đã duyệt.")
bullet("Ghi số lượng được duyệt của từng dòng ngược về phiếu yêu cầu, để người lập đối chiếu được "
       "phần đề nghị và phần được duyệt.")
bullet("Cộng số lượng vừa duyệt vào phần đã tất toán của từng dòng hàng trên phiếu xuất mượn gốc, "
       "đồng thời đánh dấu đây là phần được xử lý bằng đường khác chứ không phải nhập trả về kho.")
bullet("Kiểm tra lại phiếu xuất mượn gốc: nếu mọi dòng hàng đã tất toán hết, chuyển phiếu đó sang "
       "trạng thái Đã trả và gỡ khỏi danh sách hàng mượn phải theo dõi.")
bullet("Chuyển phiếu yêu cầu sang Đã duyệt, ghi nhận người duyệt và thời điểm duyệt, rồi gửi thông "
       "báo cho người lập.")
para("Phiếu xuất hàng mượn không làm thay đổi tồn kho. Hàng đã ra khỏi kho từ phiếu xuất mượn "
     "trước đó; chứng từ này chỉ tất toán phần hàng đang cho mượn.", "Rất quan trọng: ")

h("5.5. Sau khi kết thúc", 2)
bullet("Người lập và kế toán kho tra cứu lại phiếu ở danh sách, xem được người duyệt, ngày duyệt và "
       "lý do từ chối nếu có.")
bullet("Kế toán kho đối chiếu phần hàng mượn còn lại ở màn Danh sách hàng mượn.")
bullet("Phiếu xuất hàng mượn đã lập được tra ở menu Kho → Phiếu xuất hàng mượn, có cột dẫn ngược về "
       "phiếu yêu cầu tương ứng.")
bullet("Không có bước hủy hay thu hồi sau khi đã duyệt. Duyệt nhầm phải xử lý bằng nghiệp vụ khác.")

# ───────────────────────────── 6 ─────────────────────────────
h("6. THÔNG BÁO — AI NHẬN, KHI NÀO, NỘI DUNG GÌ", 1)
table([
    ["Sự kiện", "Ai nhận", "Nội dung thông báo", "Bấm vào thì đi đâu"],
    ["Người lập lưu phiếu yêu cầu xuất hàng mượn",
     "Toàn bộ nhân sự có quyền “Kế toán kho” và làm việc cùng công ty với người lập phiếu. Không giới "
     "hạn theo kho hay theo phòng ban, nên mọi kế toán kho của công ty đó đều nhận được.",
     "“<Tên người lập> vừa tạo yêu cầu xuất hàng mượn: <mã phiếu>”",
     "Mở đúng màn chi tiết của phiếu yêu cầu vừa lập."],
    ["Kế toán kho từ chối phiếu",
     "Chỉ một người: người lập phiếu.",
     "“<Tên người duyệt> vừa từ chối yêu cầu xuất hàng mượn: <mã phiếu>”",
     "Mở đúng màn chi tiết của phiếu bị từ chối, xem được lý do ở mục Ghi chú duyệt."],
    ["Kế toán kho duyệt phiếu (lưu phiếu xuất hàng mượn)",
     "Chỉ một người: người lập phiếu.",
     "“<Tên người duyệt> vừa duyệt yêu cầu xuất hàng mượn: <mã phiếu>”",
     "Mở đúng màn chi tiết của phiếu vừa được duyệt."],
])

h("6.1. Quy ước chung về thông báo", 2)
bullet("Thông báo hiện ở chuông thông báo trên thanh công cụ và hiện ngay không cần tải lại trang. "
       "Mã phiếu trong nội dung được in đậm.")
bullet("Người lập phiếu nếu đồng thời có quyền kế toán kho thì vẫn nhận thông báo do chính mình tạo "
       "ra ở sự kiện lập phiếu, vì danh sách người nhận lấy theo quyền chứ không loại trừ người thao "
       "tác.")
bullet("Kế toán kho của công ty khác không nhận được thông báo, kể cả khi phiếu mượn liên quan tới "
       "kho của họ.")
bullet("Các sự kiện sau KHÔNG phát sinh thông báo, không phải lỗi: mở xem phiếu, in phiếu, xuất danh "
       "sách ra tệp, và việc phiếu xuất mượn gốc chuyển sang trạng thái Đã trả. Người theo dõi hàng "
       "mượn phải tự vào màn Danh sách hàng mượn để biết.")
bullet("Thông báo chỉ là kênh nhắc việc. Nếu việc gửi thông báo gặp trục trặc thì phiếu vẫn nằm "
       "đúng trạng thái của nó và vẫn xử lý được qua danh sách chờ duyệt.")

# ───────────────────────────── 7 ─────────────────────────────
h("7. PHÂN QUYỀN", 1)

h("7.1. Quyền xem — nhìn thấy dữ liệu nào", 2)
table([
    ["Quyền", "Nhìn thấy gì ở danh sách phiếu yêu cầu xuất hàng mượn"],
    ["Không có quyền xem nào",
     "Chỉ thấy phiếu do chính mình lập. Đây là trường hợp của phần lớn nhân viên kinh doanh và họ vẫn "
     "lập phiếu bình thường."],
    ["Xem phiếu hàng mượn theo phòng ban",
     "Thấy phiếu của các phòng ban mình được phân công quản lý, cộng thêm phiếu do chính mình lập."],
    ["Xem phiếu hàng mượn theo công ty",
     "Thấy phiếu của công ty mình, cộng thêm phiếu do chính mình lập."],
    ["Xem phiếu hàng mượn theo tổng công ty",
     "Theo thiết kế là thấy phiếu của mọi công ty. Trên bản đang chạy quyền này chưa phát huy tác dụng "
     "— xem chương 12."],
    ["Kế toán kho",
     "Thấy mọi phiếu đã gửi đi của công ty mình, không phụ thuộc người lập là ai; đồng thời là người "
     "duy nhất vào được màn Yêu cầu xuất hàng mượn và màn Phiếu xuất hàng mượn."],
])

h("7.2. Quyền thao tác — điều kiện được phép", 2)
table([
    ["Thao tác", "Ai được làm", "Điều kiện bắt buộc"],
    ["Lập phiếu yêu cầu xuất hàng mượn", "Nhân viên kinh doanh, không cần quyền riêng",
     "Phải là người đã lập phiếu xuất mượn, và phiếu mượn đó đang ở trạng thái Đã mượn."],
    ["Xem chi tiết phiếu", "Người lập phiếu và nhân sự có quyền Kế toán kho",
     "Người ngoài hai nhóm này mở đường dẫn trực tiếp sẽ nhận màn báo không tìm thấy."],
    ["Từ chối phiếu", "Kế toán kho", "Phiếu đang ở trạng thái Chờ duyệt và phải nhập lý do."],
    ["Duyệt phiếu (lập phiếu xuất hàng mượn)", "Kế toán kho",
     "Phiếu đang ở trạng thái Chờ duyệt. Điều kiện được kiểm tra lại lần nữa lúc lưu, nên hai người "
     "cùng duyệt một phiếu thì người sau bị chặn."],
    ["Sửa, xóa phiếu", "Không ai", "Chức năng đã bị gỡ khỏi hệ thống — xem chương 12."],
    ["Xuất danh sách ra tệp", "Người xem được danh sách", "Không có điều kiện thêm."],
])

# ───────────────────────────── 8 ─────────────────────────────
h("8. QUY TẮC NGHIỆP VỤ BẮT BUỘC", 1)

h("8.1. Bắt buộc nhập khi lưu phiếu yêu cầu", 2)
bullet("Ít nhất một phiếu xuất mượn được chọn.")
bullet("Ít nhất một dòng hàng có số lượng lớn hơn 0; nếu tất cả để trống hoặc bằng 0, hệ thống báo "
       "không có thay đổi và không lưu.")
bullet("Số lượng của từng dòng: bắt buộc nhập, là số, không âm, tối đa 6 chữ số.")
bullet("Ghi chú: bắt buộc nhập, tối đa 255 ký tự.")
bullet("Tệp đính kèm: bắt buộc, tối thiểu một tệp và chỉ nhận định dạng PDF.")

h("8.2. Bắt buộc nhập khi kế toán kho từ chối", 2)
bullet("Lý do từ chối: bắt buộc nhập, tối đa 255 ký tự. Lý do này hiển thị lại cho người lập.")

h("8.3. Bắt buộc nhập khi kế toán kho duyệt", 2)
bullet("Số lượng duyệt của từng dòng: bắt buộc nhập, là số, không âm, tối đa 6 chữ số.")
bullet("Ít nhất một dòng có số lượng duyệt lớn hơn 0.")
bullet("Ghi chú của phiếu xuất hàng mượn: không bắt buộc, tối đa 255 ký tự.")

h("8.4. Ràng buộc số lượng", 2)
bullet("Số lượng đề nghị của một dòng không vượt quá phần hàng còn đang mượn của dòng đó, sau khi đã "
       "trừ phần đã tất toán và phần đang bị khóa bởi chứng từ khác chưa xử lý xong.")
bullet("Số lượng duyệt cũng bị kiểm tra lại theo đúng công thức trên tại thời điểm duyệt, nên nếu "
       "trong lúc chờ có chứng từ khác xử lý mất phần hàng đó thì kế toán kho sẽ bị chặn.")
bullet("Mọi so sánh số lượng đều quy đổi về đơn vị tính cơ bản của mặt hàng, phần lẻ bị làm tròn "
       "xuống, nên không đề nghị được số lượng lẻ hơn một đơn vị đóng gói.")
bullet("Đơn giá lấy theo hàng hóa, người lập không sửa. Thành tiền là số lượng nhân đơn giá, dùng để "
       "tham chiếu giá trị lô hàng tất toán.")

h("8.5. Khóa chỉnh sửa", 2)
bullet("Phiếu đã lưu là không sửa và không xóa được, ở bất kỳ trạng thái nào.")
bullet("Tệp đính kèm chỉ đưa lên được lúc lập phiếu, sau đó không thay và không gỡ được.")
bullet("Phiếu đã bị từ chối không mở lại được; muốn tiếp tục phải lập phiếu mới.")
bullet("Phiếu đã duyệt không thu hồi được; số hàng đã tất toán không tự hoàn lại.")

h("8.6. Quy tắc sinh mã", 2)
bullet("Mã phiếu yêu cầu xuất hàng mượn có dạng PYCXHM kèm số thứ tự 5 chữ số, ví dụ PYCXHM-00123.")
bullet("Mã phiếu xuất hàng mượn có dạng PXHM kèm số thứ tự 5 chữ số, ví dụ PXHM-00087.")
bullet("Mã sinh tự động sau khi lưu, không sửa được, không trùng nhau và không cấp lại khi phiếu bị "
       "từ chối.")

h("8.7. Ghi nhận đơn vị tổ chức", 2)
bullet("Phiếu ghi nhận công ty, phòng ban và bộ phận theo hồ sơ của người lập tại thời điểm lập. "
       "Người lập chuyển phòng sau đó thì phiếu cũ vẫn thuộc phòng ban cũ.")

# ───────────────────────────── 9 ─────────────────────────────
h("9. CÁC LỐI VÀO MÀN HÌNH", 1)
para("Màn danh sách phiếu yêu cầu xuất hàng mượn có nhiều mục menu cùng trỏ vào, khác nhau ở phần "
     "đuôi đường dẫn, và mỗi lối vào cho ra một danh sách khác hẳn nhau:")
table([
    ["Vào bằng", "Danh sách hiện ra", "Dùng khi nào"],
    ["Menu Kho → Quản lý hàng mượn → Phiếu Yêu cầu xuất hàng mượn (kèm ?type=all)",
     "Toàn bộ phiếu trong phạm vi quyền xem của tôi, kèm phiếu do tôi lập",
     "Lối vào chính của nhân viên kinh doanh và lãnh đạo theo dõi"],
    ["Menu Duyệt → Phiếu yêu cầu xuất hàng mượn chờ duyệt (kèm ?type=for-approve)",
     "Chỉ phiếu đang Chờ duyệt của công ty tôi",
     "Danh sách việc kế toán kho phải xử lý"],
    ["Menu Kho → Yêu cầu xuất hàng mượn (màn dành cho kế toán)",
     "Mọi phiếu đã gửi đi, đủ mọi trạng thái trừ trạng thái nháp",
     "Kế toán kho tra cứu lại lịch sử xử lý, có bộ lọc theo ngày và theo hàng hóa"],
    ["Đường dẫn danh sách không kèm gì",
     "Chỉ phiếu do chính tôi lập",
     "Xem lại việc của mình"],
    ["Đường dẫn kèm ?type=return",
     "Phiếu đã duyệt do chính tôi lập",
     "Không có mục menu nào trỏ vào; hệ thống dùng nội bộ cho luồng trả hàng"],
])
bullet("Đường dẫn quyết định phạm vi xem, không phải quyền — người không có quyền xem theo cấp mà mở "
       "link “xem tất cả” thì vẫn chỉ thấy phiếu của chính mình.")
bullet("Nút Làm mới chỉ xóa điều kiện lọc, không đưa người dùng sang phạm vi khác.")
bullet("Link kèm giá trị lạ thì hệ thống bỏ qua, giữ phạm vi mặc định là chỉ phiếu của chính mình.")
para("Ngoài ra, phiếu xuất hàng mượn sinh ra khi duyệt được tra ở menu Kho → Phiếu xuất hàng mượn; "
     "màn này chỉ nhân sự có quyền Kế toán kho mở được.")

# ───────────────────────────── 10 ─────────────────────────────
h("10. TRA CỨU, IN VÀ XUẤT DỮ LIỆU", 1)

h("10.1. Danh sách và bộ lọc", 2)
para("Danh sách hiển thị các cột: số thứ tự, mã phiếu, người lập kèm mã phòng ban, ngày lập, trạng "
     "thái, người duyệt, ngày duyệt và cột hành động. Mã phiếu bấm vào mở thẳng chi tiết.")
bullet("Lọc theo mã phiếu, gõ một phần mã cũng ra.")
bullet("Lọc theo trạng thái: Đã duyệt, Chờ duyệt, Đang tạo, Không duyệt.")
bullet("Lọc theo người lập và theo người duyệt.")
bullet("Lọc theo tên hàng hóa hoặc mã hàng hóa — tìm ra những phiếu có chứa mặt hàng đó.")
bullet("Lọc theo khoảng thời gian lập phiếu.")
bullet("Lọc theo công ty, phòng ban — chỉ hiện với người có quyền xem theo cấp tương ứng.")
para("Ở cột hành động, chỉ phiếu đang Chờ duyệt và người xem là kế toán kho mới thấy mục “Tạo phiếu "
     "xuất hàng mượn”. Các trường hợp còn lại menu hành động để trống.")

h("10.2. Xuất dữ liệu", 2)
bullet("Xuất danh sách phiếu ra tệp Excel gồm: số thứ tự, mã phiếu, người lập, ngày lập, trạng thái, "
       "người duyệt, ngày duyệt. Tệp lấy đúng điều kiện lọc đang áp dụng trên màn hình và ghi lại "
       "khoảng thời gian đã lọc ở đầu tệp.")
bullet("Danh sách phiếu xuất hàng mượn cũng xuất được Excel tương tự, gồm mã phiếu, phiếu yêu cầu "
       "tương ứng, người yêu cầu, người lập và ngày lập.")

h("10.3. In phiếu", 2)
para("Hệ thống có sẵn mẫu in “Yêu cầu xuất hàng mượn” gồm thông tin công ty, số phiếu, danh sách "
     "phiếu mượn liên quan, người lập, ghi chú và bảng chi tiết hàng hóa kèm số lượng theo từng phiếu "
     "mượn. Tuy nhiên nút gọi bản in hiện chưa được bật trên màn danh sách và màn chi tiết — xem "
     "chương 12.")

# ───────────────────────────── 11 ─────────────────────────────
h("11. LIÊN THÔNG VỚI PHẦN HÀNH KHÁC", 1)

h("11.1. Dùng chung dữ liệu với phân hệ nào", 2)
bullet("Phân hệ Kho — phiếu yêu cầu xuất hàng loại Xuất mượn là nguồn dữ liệu đầu vào; màn Danh sách "
       "hàng mượn là nơi đối chiếu kết quả sau khi tất toán.")
bullet("Danh mục hàng hóa — tên hàng, mã hàng, model, thương hiệu, đơn vị tính và hệ số quy đổi lấy "
       "từ danh mục chung; phiếu lưu lại bản sao tại thời điểm lập nên đổi tên hàng về sau không làm "
       "sai phiếu cũ.")
bullet("Hồ sơ nhân sự — công ty, phòng ban, bộ phận của người lập và danh sách phòng ban do một người "
       "quản lý, dùng để quyết định phạm vi xem.")
bullet("Chuông thông báo dùng chung toàn hệ thống.")
bullet("Kho lưu trữ tệp trên máy chủ đám mây — nơi giữ các tệp PDF đính kèm.")

h("11.2. Quan hệ với ba luồng hàng mượn còn lại", 2)
table([
    ["Luồng", "Dùng khi nào", "Tác động lên hàng mượn"],
    ["Yêu cầu xuất hàng mượn (tài liệu này)",
     "Chốt số hàng khách đang mượn mà không nhập lại kho",
     "Tất toán phần hàng mượn tương ứng; hết hàng thì phiếu mượn chuyển Đã trả"],
    ["Trả hàng mượn",
     "Khách trả hàng thật về kho",
     "Nhập hàng về kho và tất toán phần hàng mượn tương ứng"],
    ["Yêu cầu xuất bán hàng mượn",
     "Khách quyết định mua luôn số hàng đang mượn",
     "Chuyển hàng mượn thành hàng bán, ghi nhận doanh thu và công nợ"],
    ["Yêu cầu gia hạn hàng mượn",
     "Khách cần mượn thêm thời gian",
     "Chỉ kéo dài hạn mượn, không thay đổi số lượng"],
])
para("Bốn luồng cùng trừ vào một nguồn là số hàng đang mượn của phiếu xuất mượn gốc. Chứng từ nào "
     "đang chờ xử lý thì phần số lượng của nó bị khóa lại, luồng khác không dùng trùng được. Riêng "
     "luồng gia hạn còn chặn thêm: đang có yêu cầu gia hạn chưa duyệt thì không lập thêm yêu cầu gia "
     "hạn mới cho cùng phiếu mượn.")

h("11.3. Các điểm cố ý làm khác", 2)
para("Không áp dụng vì đây là mô tả hệ thống ERP hiện hành, không phải màn được chuyển đổi từ hệ "
     "thống khác nên không có điểm khác biệt có chủ đích cần đối chiếu.")

# ───────────────────────────── 12 ─────────────────────────────
h("12. GIỚI HẠN HIỆN TẠI", 1)
table([
    ["Giới hạn", "Ảnh hưởng tới người dùng", "Cách xoay xở hiện nay"],
    ["Không có chức năng sửa và xóa phiếu",
     "Nhập nhầm số lượng hay đính kèm nhầm tệp là không sửa được.",
     "Nhờ kế toán kho từ chối, sau đó lập phiếu mới. Trạng thái Đang tạo trên bộ lọc vì vậy luôn không "
     "có dữ liệu."],
    ["Chỉ người lập phiếu xuất mượn mới lập được phiếu yêu cầu",
     "Nhân viên nghỉ việc hoặc bàn giao khách thì không ai lập được phiếu tất toán cho phiếu mượn cũ.",
     "Phải nhờ quản trị chuyển người lập trên phiếu mượn gốc; chưa có màn bàn giao chính thức."],
    ["Quyền Xem phiếu hàng mượn theo tổng công ty chưa phát huy tác dụng",
     "Người được cấp quyền này vẫn chỉ thấy phiếu của công ty mình do màn danh sách còn chặn cứng theo "
     "công ty của người đăng nhập.",
     "Xem theo từng công ty, hoặc chờ chỉnh sửa. Cần lưu ý khi nghiệm thu để không kết luận nhầm là "
     "mất dữ liệu."],
    ["Chưa bật nút in phiếu",
     "Mẫu in đã có nhưng người dùng không gọi được từ màn hình.",
     "Dùng chức năng xuất Excel danh sách, hoặc chụp màn hình chi tiết."],
    ["Nhãn trên màn lập phiếu xuất hàng mượn còn ghi nhầm",
     "Màn lập phiếu xuất hàng mượn đang hiển thị nhãn “Chọn phiếu yêu cầu xuất bán hàng mượn” và "
     "“Người lập phiếu YCXBHM”, dễ khiến kế toán kho tưởng mình đang mở nhầm chức năng xuất bán.",
     "Căn cứ vào mã phiếu để phân biệt: phiếu yêu cầu xuất hàng mượn bắt đầu bằng PYCXHM."],
    ["Không có thông báo khi phiếu mượn được tất toán hết",
     "Người theo dõi công nợ hàng mượn không được báo khi một phiếu mượn đóng lại.",
     "Chủ động kiểm tra ở màn Danh sách hàng mượn."],
    ["Thông báo lập phiếu gửi cho mọi kế toán kho trong công ty",
     "Kế toán kho không phụ trách kho liên quan vẫn nhận thông báo, dễ nhiễu.",
     "Lọc lại theo danh sách phiếu chờ duyệt thay vì dựa vào chuông thông báo."],
    ["Duyệt rồi không thu hồi được",
     "Duyệt nhầm số lượng thì hàng mượn đã bị tất toán sai.",
     "Xử lý bằng nghiệp vụ điều chỉnh khác, cần phối hợp kế toán."],
])

doc.save(OUT)
print("Đã lưu:", OUT)

# ───────────────────── Kiểm tra thuật ngữ kỹ thuật ─────────────────────
BAD = [
    r"\bstatus\b", r"\bborrow_", r"\bproduct_export", r"\bqty\b", r"\bid\b(?!\w)",
    r"Controller", r"\.php", r"\.blade", r"Eloquent", r"middleware",
    r"\bAPI\b", r"endpoint", r"\bbảng\s+[a-z_]{4,}\b", r"\bcột\s+[a-z_]{4,}\b",
    r"\btrạng thái\s*=\s*\d", r"\bnull\b", r"\broute\b", r"Redis", r"\bS3\b",
]
text = "\n".join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            text += "\n" + c.text

found = []
for pat in BAD:
    for m in re.finditer(pat, text, flags=re.IGNORECASE):
        seg = text[max(0, m.start() - 40):m.end() + 40].replace("\n", " ")
        found.append("%-24s | ...%s..." % (pat, seg))

if found:
    print("CẢNH BÁO — còn thuật ngữ kỹ thuật:")
    for f in found:
        print("  ", f)
else:
    print("Kiểm tra thuật ngữ: sạch")
