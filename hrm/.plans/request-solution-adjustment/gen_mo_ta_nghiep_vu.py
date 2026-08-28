# -*- coding: utf-8 -*-
"""Sinh tài liệu Mô tả nghiệp vụ — Yêu cầu điều chỉnh giải pháp."""
import sys
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
CENTER = WD_ALIGN_PARAGRAPH.CENTER

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)


def h(text, level=1):
    doc.add_paragraph(text, style="Heading %d" % level)


def para(text=""):
    p = doc.add_paragraph()
    p.alignment = JUSTIFY
    p.add_run(text)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = 1
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.cell(i, j).text = val
    doc.add_paragraph()
    return t


# ------------------------------------------------------------------ BÌA
def cover():
    p = doc.add_paragraph(); p.alignment = CENTER
    r = p.add_run("MÔ TẢ NGHIỆP VỤ"); r.bold = True; r.font.size = Pt(20)
    p = doc.add_paragraph(); p.alignment = CENTER
    r = p.add_run("Yêu cầu điều chỉnh giải pháp"); r.bold = True; r.font.size = Pt(16)
    p = doc.add_paragraph(); p.alignment = CENTER
    r = p.add_run("Phân hệ Giao việc — Quản lý dự án · Cập nhật ngày 26/08/2026")
    r.font.size = Pt(12)
    p = doc.add_paragraph(); p.alignment = CENTER
    r = p.add_run("Đường dẫn: Dự án tiềm năng → Chi tiết dự án → tab Giải pháp → tab con Yêu cầu điều chỉnh GP")
    r.font.size = Pt(12); r.italic = True
    doc.add_page_break()


cover()

# ------------------------------------------------------- 1
h("1. TÀI LIỆU NÀY DÀNH CHO AI", 1)
para("Tài liệu mô tả toàn bộ nghiệp vụ của chức năng Yêu cầu điều chỉnh giải pháp: vì sao có chức "
     "năng này, ai tham gia, phiếu chạy qua những bước nào, ai nhận thông báo ở từng bước và những "
     "quy tắc bắt buộc phải tuân theo. Tài liệu không hướng dẫn thao tác từng nút bấm (phần đó nằm "
     "ở tài liệu Hướng dẫn sử dụng).")
bullet("Nhân viên kinh doanh phụ trách dự án — người khởi tạo yêu cầu điều chỉnh.")
bullet("Trưởng phòng giải pháp và Quản lý dự án (PM) — người tiếp nhận hoặc từ chối yêu cầu.")
bullet("Nhân viên lập báo giá và nhân viên xây dựng giá — bên bị ảnh hưởng dây chuyền khi yêu cầu được tiếp nhận.")
bullet("Ban lãnh đạo, người nghiệm thu và bộ phận kiểm thử — cần hiểu luồng để đối chiếu khi nghiệm thu.")

# ------------------------------------------------------- 2
h("2. CHỨC NĂNG NÀY DÙNG ĐỂ LÀM GÌ", 1)
h("2.1. Mục đích", 2)
para("Giải pháp sau khi được duyệt là bản cam kết kỹ thuật để phòng kinh doanh mang đi làm giá và "
     "báo giá cho khách hàng. Trong thực tế, sau khi giải pháp đã duyệt vẫn thường phát sinh nhu cầu "
     "sửa: khách hàng đổi yêu cầu, khảo sát lại thấy thiếu thiết bị, hoặc phương án cũ vượt ngân sách. "
     "Trước đây các đề nghị sửa này được trao đổi miệng hoặc qua tin nhắn, không để lại dấu vết, nên "
     "không ai biết ai đã đề nghị sửa gì, phòng giải pháp có đồng ý hay không, và bộ phận làm giá vẫn "
     "tiếp tục dựng giá trên phương án cũ đã lỗi thời.")
para("Chức năng Yêu cầu điều chỉnh giải pháp biến đề nghị sửa đó thành một phiếu chính thức có mã, "
     "có người gửi, có nội dung, có file đính kèm và có kết quả xử lý được ghi lại vĩnh viễn.")

h("2.2. Vị trí trong luồng lớn", 2)
para("Chức năng nằm ở đoạn giữa của luồng dự án tiềm năng, sau khi giải pháp đã được duyệt và trước "
     "hoặc trong lúc phòng kinh doanh đang làm giá — báo giá cho khách hàng:")
bullet("Bước trước: giải pháp được xây dựng, trình duyệt và đạt trạng thái Đã duyệt giải pháp trở lên.")
bullet("Bước này: nhân viên kinh doanh phát hiện cần sửa giải pháp, gửi phiếu yêu cầu điều chỉnh.")
bullet("Bước sau: nếu phòng giải pháp tiếp nhận, mọi hoạt động làm giá và báo giá đang chạy trên "
       "phương án cũ bị dừng lại, phòng giải pháp dựng phiên bản giải pháp mới rồi làm giá lại từ đầu.")

h("2.3. Giá trị mang lại", 2)
bullet("Đề nghị sửa giải pháp có bằng chứng: mọi phiếu đều lưu người gửi, ngày gửi, nội dung và "
       "phiên bản giải pháp tại thời điểm gửi — về sau tra lại được vì sao phương án đổi.")
bullet("Ngăn lãng phí công sức làm giá: ngay khi phòng giải pháp tiếp nhận, hệ thống tự dừng các "
       "yêu cầu xây dựng giá và báo giá đang dở dang, không để ai tiếp tục dựng giá trên phương án sắp bị thay.")
bullet("Người có thẩm quyền được báo ngay: phiếu vừa gửi là Quản lý dự án và Trưởng phòng tiếp nhận "
       "nhận được thông báo, không phải chờ ai nhắc.")
bullet("Từ chối cũng phải nêu lý do, nên phòng kinh doanh biết vì sao đề nghị không được chấp thuận.")

# ------------------------------------------------------- 3
h("3. NHỮNG AI THAM GIA", 1)
table([
    ["Vai trò", "Làm gì trong luồng này"],
    ["Nhân viên kinh doanh phụ trách dự án",
     "Là người duy nhất được tạo phiếu yêu cầu điều chỉnh. Nhập nội dung cần điều chỉnh, đính kèm "
     "tài liệu minh chứng và gửi đi. Sau khi gửi chỉ được xem lại, không sửa và không thu hồi."],
    ["Quản lý dự án (PM) của giải pháp",
     "Nhận thông báo khi có phiếu mới, xem xét nội dung và quyết định Tiếp nhận hoặc Từ chối. Nếu "
     "tiếp nhận thì tự dựng phiên bản giải pháp mới."],
    ["Trưởng phòng của phòng tiếp nhận yêu cầu làm giải pháp",
     "Có thẩm quyền ngang PM: cũng nhận thông báo, cũng được Tiếp nhận hoặc Từ chối phiếu. Bao gồm "
     "cả người được cấu hình quản lý tất cả phòng ban."],
    ["Nhân viên lập báo giá của dự án",
     "Không thao tác trên phiếu, nhưng chịu ảnh hưởng: báo giá đang làm dở bị chuyển sang trạng thái "
     "Dừng và người này nhận thông báo báo giá tạm dừng."],
    ["Bộ phận xây dựng giá",
     "Không thao tác trên phiếu. Yêu cầu xây dựng giá đang chờ hoặc đang làm của dự án bị chuyển "
     "sang trạng thái Dừng."],
    ["Khách hàng",
     "Không tham gia hệ thống. Là bên phát sinh nhu cầu đổi phương án, thường là nguồn gốc của "
     "nội dung ghi trong phiếu."],
])

# ------------------------------------------------------- 4
h("4. VÒNG ĐỜI CỦA PHIẾU YÊU CẦU ĐIỀU CHỈNH", 1)
para("Phiếu chỉ có ba trạng thái và đi một chiều, không quay lui. Phiếu không có trạng thái nháp: "
     "bấm Gửi là phiếu tồn tại chính thức ngay.")
table([
    ["Trạng thái", "Ý nghĩa", "Ai làm phiếu chuyển sang trạng thái này"],
    ["Đã gửi",
     "Phiếu vừa được gửi, đang chờ phòng giải pháp xem xét. Đây là trạng thái khởi tạo duy nhất.",
     "Nhân viên kinh doanh phụ trách dự án, khi bấm Gửi trong màn tạo yêu cầu."],
    ["Tiếp nhận",
     "Phòng giải pháp đồng ý xem xét điều chỉnh. Đây là trạng thái kết thúc — kéo theo việc dừng "
     "các yêu cầu xây dựng giá và báo giá đang dở dang của dự án.",
     "Quản lý dự án của giải pháp hoặc Trưởng phòng tiếp nhận, khi bấm Tiếp nhận."],
    ["Từ chối",
     "Phòng giải pháp không chấp thuận đề nghị, kèm lý do bắt buộc nhập. Đây là trạng thái kết thúc "
     "và không kéo theo thay đổi nào khác trong hệ thống.",
     "Quản lý dự án của giải pháp hoặc Trưởng phòng tiếp nhận, khi bấm Từ chối và nhập lý do."],
])
para("Lưu ý về các trạng thái do chứng từ khác quyết định: bản thân phiếu yêu cầu điều chỉnh không "
     "bao giờ bị chứng từ khác cập nhật ngược về. Ngược lại, chính phiếu này khi được tiếp nhận sẽ "
     "cập nhật trạng thái của yêu cầu xây dựng giá và của báo giá thuộc dự án — chi tiết ở mục 5.4.")

# ------------------------------------------------------- 5
h("5. LUỒNG HOẠT ĐỘNG CHI TIẾT", 1)

h("5.1. Bước 1 — Phát sinh nhu cầu điều chỉnh và kiểm tra điều kiện gửi", 2)
para("Nhân viên kinh doanh phụ trách dự án mở chi tiết dự án tiềm năng, vào tab Giải pháp, chọn tab "
     "con Yêu cầu điều chỉnh GP. Hệ thống chỉ cho phép gửi yêu cầu khi thoả đồng thời các điều kiện sau:")
table([
    ["Điều kiện", "Lý do nghiệp vụ", "Không thoả thì sao"],
    ["Dự án đã có giải pháp tương ứng",
     "Chưa có giải pháp thì không có gì để điều chỉnh.",
     "Tab con hiển thị dòng Dự án chưa có giải pháp tương ứng, không có nút nào."],
    ["Giải pháp đang ở một trong các trạng thái: Đã duyệt giải pháp, Đã duyệt giá, Chờ làm giá, Chốt giải pháp",
     "Giải pháp chưa duyệt thì phòng giải pháp vẫn đang tự sửa được, không cần cơ chế phiếu.",
     "Nút Tạo yêu cầu không hiển thị. Nếu cố gửi, hệ thống chặn và báo giải pháp chưa được duyệt "
     "nên không thể tạo yêu cầu điều chỉnh."],
    ["Người thao tác đúng là nhân viên kinh doanh phụ trách dự án",
     "Chỉ người chịu trách nhiệm thương mại của dự án mới được thay mặt khách hàng đề nghị đổi phương án.",
     "Nút Tạo yêu cầu không hiển thị. Nếu cố gửi, hệ thống chặn và báo chỉ nhân viên kinh doanh "
     "phụ trách dự án mới được tạo yêu cầu điều chỉnh."],
    ["Giải pháp được chọn đúng là giải pháp của dự án đang mở",
     "Tránh gửi nhầm yêu cầu sang giải pháp của dự án khác.",
     "Hệ thống chặn và báo giải pháp không thuộc dự án này."],
])

h("5.2. Bước 2 — Lập và gửi phiếu", 2)
para("Nhân viên kinh doanh khai báo ba nội dung: giải pháp cần điều chỉnh (hệ thống điền sẵn, không "
     "sửa được, gồm mã giải pháp, tên giải pháp và phiên bản hiện hành), nội dung điều chỉnh (bắt "
     "buộc nhập, mô tả rõ cần sửa gì) và file đính kèm (không bắt buộc, dùng cho bản vẽ, email khách "
     "hàng, biên bản khảo sát lại).")
para("Khi gửi, hệ thống thực hiện đồng thời các việc sau và chỉ ghi nhận nếu tất cả thành công:")
bullet("Sinh mã phiếu tự động theo dạng YCDCGP kèm số thứ tự năm chữ số, ví dụ YCDCGP.00001.")
bullet("Gán trạng thái Đã gửi.")
bullet("Chụp lại phiên bản giải pháp tại thời điểm gửi, để sau này biết phiếu được viết trên phương án nào.")
bullet("Lưu các file đính kèm và gắn vào phiếu.")
bullet("Gửi thông báo cho Quản lý dự án và Trưởng phòng tiếp nhận — chi tiết ở chương 6.")
para("Phiếu vừa gửi hiện ngay trên bảng danh sách của tab con, ở dòng đầu tiên vì danh sách xếp "
     "phiếu mới nhất lên trên.")

h("5.3. Bước 3 — Phòng giải pháp xem xét", 2)
para("Quản lý dự án hoặc Trưởng phòng tiếp nhận mở phiếu từ thông báo hoặc vào màn Quản lý giải pháp, "
     "tab YC Điều chỉnh. Người xem đọc nội dung điều chỉnh, tải file đính kèm rồi chọn một trong hai hướng:")
table([
    ["Hướng xử lý", "Diễn ra thế nào", "Kết quả"],
    ["Tiếp nhận",
     "Người xử lý xác nhận lại một lần trong hộp xác nhận trước khi hệ thống ghi nhận.",
     "Phiếu chuyển sang Tiếp nhận, lưu người tiếp nhận và thời điểm tiếp nhận. Hệ thống chạy tiếp "
     "các thay đổi dây chuyền ở mục 5.4."],
    ["Từ chối",
     "Người xử lý bắt buộc nhập lý do từ chối rồi mới gửi được.",
     "Phiếu chuyển sang Từ chối, lưu lý do, người từ chối và thời điểm từ chối. Không có thay đổi "
     "nào khác trong hệ thống. Nhân viên kinh doanh muốn theo đuổi tiếp thì lập phiếu mới, không "
     "sửa được phiếu cũ."],
])
para("Trong lúc phiếu còn ở trạng thái Đã gửi, hai nút Tiếp nhận và Từ chối chỉ hiển thị với đúng "
     "người có thẩm quyền. Phiếu đã sang Tiếp nhận hoặc Từ chối thì hai nút này biến mất với mọi "
     "người, ai mở phiếu cũng chỉ còn xem được.")

h("5.4. Bước 4 — Hệ quả dây chuyền khi phiếu được tiếp nhận", 2)
para("Đây là phần quan trọng nhất của nghiệp vụ và cũng là phần dễ gây bất ngờ nhất cho người dùng. "
     "Vì giải pháp sắp được sửa nên mọi công việc làm giá dựa trên phương án cũ đều trở thành vô "
     "nghĩa. Ngay khi phiếu được tiếp nhận, hệ thống tự động rà toàn bộ dự án:")
table([
    ["Chứng từ liên quan", "Đang ở trạng thái", "Bị chuyển thành"],
    ["Yêu cầu xây dựng giá của dự án", "Chờ xây dựng giá hoặc Đang xây dựng giá", "Dừng"],
    ["Báo giá gắn với yêu cầu xây dựng giá đang được làm dở",
     "Đang tạo, Chờ TP duyệt hoặc Chờ BGĐ duyệt", "Dừng"],
])
bullet("Yêu cầu xây dựng giá đã có báo giá hoàn chỉnh, đã đóng hoặc đã dừng trước đó thì giữ nguyên, không bị đụng tới.")
bullet("Báo giá đã duyệt xong hoặc đã trúng thầu thì giữ nguyên, không bị dừng.")
bullet("Trạng thái Dừng là trạng thái kết thúc: chứng từ đã dừng không mở lại được. Muốn tiếp tục thì "
       "lập yêu cầu xây dựng giá và báo giá mới trên phiên bản giải pháp mới.")
bullet("Người lập từng báo giá bị dừng nhận được thông báo riêng — chi tiết ở chương 6.")
para("Nếu quá trình dừng dây chuyền gặp sự cố kỹ thuật, hệ thống vẫn giữ kết quả tiếp nhận phiếu; "
     "sự cố được ghi lại để bộ phận kỹ thuật xử lý chứ không làm hỏng việc tiếp nhận.")

h("5.5. Bước 5 — Dựng phiên bản giải pháp mới", 2)
para("Việc tiếp nhận phiếu chỉ mang ý nghĩa xác nhận đồng ý điều chỉnh; hệ thống không tự sinh phiên "
     "bản giải pháp mới. Quản lý dự án phải tự vào giải pháp, tạo phiên bản mới và cho chạy lại quy "
     "trình duyệt như bình thường. Sau khi phiên bản mới được duyệt, phòng kinh doanh lập yêu cầu "
     "xây dựng giá mới để làm giá lại.")

# ------------------------------------------------------- 6
h("6. THÔNG BÁO — AI NHẬN, KHI NÀO, NỘI DUNG GÌ", 1)
table([
    ["Sự kiện", "Ai nhận", "Nội dung người nhận thấy", "Bấm vào thì đi đâu"],
    ["Nhân viên kinh doanh gửi phiếu yêu cầu điều chỉnh",
     "Hai nhóm cộng lại, trùng nhau thì chỉ nhận một lần: (1) đúng một người là Quản lý dự án của "
     "giải pháp; (2) toàn bộ những người được cấu hình quản lý phòng tiếp nhận yêu cầu làm giải pháp "
     "— gồm cả người được tích quản lý tất cả phòng ban. Không gửi cho toàn bộ nhân viên của phòng.",
     "Tiêu đề: Yêu cầu điều chỉnh giải pháp mới. Nội dung: Có yêu cầu điều chỉnh giải pháp: "
     "<mã phiếu> - GP: <mã giải pháp> - <tên giải pháp>.",
     "Mở màn Quản lý giải pháp của đúng giải pháp đó và nhảy thẳng vào tab YC Điều chỉnh."],
    ["Phiếu được tiếp nhận, kéo theo một báo giá bị chuyển sang Dừng",
     "Đúng một người là người lập chính của báo giá bị dừng. Mỗi báo giá bị dừng phát sinh một thông "
     "báo riêng. Không gửi cho người duyệt báo giá, không gửi cho phòng kinh doanh.",
     "Tiêu đề: Báo giá tạm dừng. Nội dung: Báo giá <mã báo giá> tạm dừng do có yêu cầu điều chỉnh "
     "giải pháp mới (<mã phiếu>).",
     "Mở chi tiết đúng báo giá bị dừng."],
])

h("6.1. Quy ước chung về thông báo", 2)
bullet("Thông báo đến qua chuông thông báo trên thanh công cụ, cùng cơ chế với các thông báo khác của hệ thống.")
bullet("Người gửi phiếu đồng thời là người có thẩm quyền xử lý (trường hợp hiếm, ví dụ được cấu hình "
       "quản lý phòng tiếp nhận) thì vẫn nhận thông báo của chính mình; danh sách người nhận đã lọc "
       "trùng nên không ai nhận hai lần cho cùng một phiếu.")
bullet("Sự cố khi gửi thông báo không làm hỏng nghiệp vụ: phiếu vẫn được tạo, việc tiếp nhận vẫn "
       "được ghi nhận, chỉ riêng thông báo là không đến.")
bullet("Không xác định được người nhận nào (giải pháp chưa có Quản lý dự án và phòng tiếp nhận chưa "
       "cấu hình người quản lý) thì hệ thống bỏ qua bước gửi, phiếu vẫn được tạo bình thường.")

h("6.2. Những sự kiện KHÔNG phát sinh thông báo", 2)
bullet("Phiếu được tiếp nhận: người gửi phiếu KHÔNG nhận được thông báo. Muốn biết kết quả phải tự "
       "vào tab con xem trạng thái phiếu.")
bullet("Phiếu bị từ chối: người gửi phiếu cũng KHÔNG nhận được thông báo, kể cả lý do từ chối.")
bullet("Yêu cầu xây dựng giá bị chuyển sang Dừng: không ai nhận thông báo về việc này; chỉ báo giá "
       "bị dừng mới sinh thông báo.")
bullet("Xem chi tiết phiếu, tải file đính kèm: không sinh thông báo.")

# ------------------------------------------------------- 7
h("7. PHÂN QUYỀN", 1)
h("7.1. Nhìn thấy dữ liệu nào", 2)
para("Chức năng này KHÔNG dùng quyền cấu hình trong danh mục phân quyền của hệ thống. Việc ai được "
     "làm gì hoàn toàn dựa trên vai trò thực tế của người đó đối với dự án và giải pháp đang mở.")
table([
    ["Người xem", "Nhìn thấy gì"],
    ["Nhân viên kinh doanh phụ trách dự án",
     "Vào được tab con, thấy toàn bộ phiếu của dự án, xem được chi tiết và file đính kèm. Không thấy "
     "cột Hành động vì không có quyền xử lý phiếu."],
    ["Quản lý dự án của giải pháp và Trưởng phòng tiếp nhận",
     "Thấy toàn bộ phiếu của dự án kèm cột Hành động, xử lý được phiếu đang ở trạng thái Đã gửi."],
    ["Người dùng khác mở được màn chi tiết dự án hoặc màn quản lý giải pháp",
     "Thấy danh sách phiếu và xem được chi tiết, nhưng nút xử lý không hiện nên không thao tác được "
     "gì; nút Tạo yêu cầu cũng không hiện."],
    ["Người không có quyền vào màn chi tiết dự án hoặc màn quản lý giải pháp",
     "Không tiếp cận được chức năng này, vì tab con nằm bên trong hai màn đó."],
])

h("7.2. Được phép làm thao tác nào", 2)
table([
    ["Thao tác", "Ai được làm", "Điều kiện kèm theo"],
    ["Tạo yêu cầu",
     "Chỉ nhân viên kinh doanh phụ trách dự án.",
     "Dự án đã có giải pháp và giải pháp đang ở Đã duyệt giải pháp, Đã duyệt giá, Chờ làm giá hoặc "
     "Chốt giải pháp. Chỉ làm được ở màn chi tiết dự án; ở màn quản lý giải pháp nút này luôn ẩn."],
    ["Xem chi tiết phiếu", "Mọi người mở được tab con.", "Không có điều kiện."],
    ["Tiếp nhận",
     "Quản lý dự án của giải pháp, hoặc người quản lý phòng tiếp nhận yêu cầu làm giải pháp.",
     "Phiếu đang ở trạng thái Đã gửi. Phiếu đã tiếp nhận hoặc đã từ chối thì nút biến mất."],
    ["Từ chối", "Cùng nhóm người như thao tác Tiếp nhận.",
     "Phiếu đang ở trạng thái Đã gửi và phải nhập lý do từ chối."],
    ["Sửa phiếu, xoá phiếu, thu hồi phiếu", "Không ai.",
     "Hệ thống không có các thao tác này — xem chương 11."],
])

# ------------------------------------------------------- 8
h("8. QUY TẮC NGHIỆP VỤ BẮT BUỘC", 1)
h("8.1. Bắt buộc nhập theo từng nút bấm", 2)
table([
    ["Nút bấm", "Bắt buộc nhập", "Không bắt buộc"],
    ["Gửi (trong màn Tạo yêu cầu điều chỉnh giải pháp)",
     "Nội dung điều chỉnh. Bỏ trống hoặc chỉ gõ khoảng trắng thì hệ thống báo nội dung điều chỉnh "
     "không được để trống và không gửi đi.",
     "File đính kèm. Giải pháp là ô chỉ để đọc, hệ thống tự điền."],
    ["Tiếp nhận", "Không có trường nào phải nhập; chỉ cần xác nhận trong hộp hỏi lại.", "—"],
    ["Gửi (trong màn Từ chối yêu cầu điều chỉnh)",
     "Lý do từ chối. Bỏ trống thì hệ thống báo lý do từ chối không được để trống.", "—"],
])
para("Chức năng không có nút lưu nháp: phiếu chỉ tồn tại khi đã gửi chính thức, nên mọi trường bắt "
     "buộc đều phải nhập ngay từ lần lưu đầu tiên.")

h("8.2. Ràng buộc dữ liệu và khoá chỉnh sửa", 2)
bullet("Mã phiếu do hệ thống sinh theo dạng YCDCGP kèm năm chữ số, duy nhất trên toàn hệ thống, "
       "người dùng không nhập và không sửa được.")
bullet("Phiên bản giải pháp ghi trên phiếu là bản chụp tại thời điểm gửi. Giải pháp về sau lên phiên "
       "bản mới thì phiếu cũ vẫn giữ nguyên phiên bản cũ, không bị đổi theo.")
bullet("Phiếu đã gửi bị khoá hoàn toàn về nội dung: không sửa nội dung điều chỉnh, không thêm bớt "
       "file đính kèm, không xoá phiếu. Chỉ trạng thái mới thay đổi được, và chỉ đúng một lần, do "
       "người có thẩm quyền thực hiện.")
bullet("Phiếu đã ở trạng thái Tiếp nhận hoặc Từ chối là kết thúc, không quay lại Đã gửi được và "
       "không xử lý lại được.")
bullet("Đề nghị bị từ chối mà vẫn muốn theo đuổi thì lập phiếu mới; hệ thống không giới hạn số phiếu "
       "trên một dự án và không chặn việc gửi nhiều phiếu cùng lúc.")
bullet("Việc tiếp nhận phiếu không tự đưa giải pháp về trạng thái sửa được và không tự sinh phiên bản "
       "mới; đó là thao tác thủ công của Quản lý dự án.")
bullet("Không có danh mục dùng chung nào bị chọn trong phiếu, nên không phát sinh tình huống danh "
       "mục bị khoá làm mất giá trị đã chọn.")

# ------------------------------------------------------- 9
h("9. TRA CỨU, IN VÀ XUẤT DỮ LIỆU", 1)
para("Danh sách phiếu được xem theo từng dự án, nằm trong tab con Yêu cầu điều chỉnh GP, gồm các cột: "
     "số thứ tự, mã yêu cầu, phiên bản giải pháp, người yêu cầu, ngày gửi, trạng thái và cột hành động "
     "(chỉ hiện với người có thẩm quyền xử lý). Phiếu mới nhất xếp trên cùng, có phân trang.")
bullet("Tìm kiếm và bộ lọc: giao diện hiện chưa có ô tìm nhanh và chưa có bộ lọc theo trạng thái. "
       "Người dùng lật trang để tìm phiếu. Bù lại số lượng phiếu trên một dự án thường rất ít.")
bullet("Xem chi tiết: bấm vào mã yêu cầu hoặc biểu tượng con mắt để mở màn chi tiết, trong đó có đầy "
       "đủ nội dung điều chỉnh, file đính kèm, người tiếp nhận hoặc người từ chối kèm lý do.")
bullet("In: chức năng không có bản in. Phiếu là chứng từ nội bộ, không dùng để ký hay gửi khách hàng.")
bullet("Xuất dữ liệu: chức năng không có xuất Excel. Cần thống kê thì tra trên từng dự án.")

# ------------------------------------------------------- 10
h("10. LIÊN THÔNG VỚI HỆ THỐNG KHÁC", 1)
para("Chức năng này là nghiệp vụ mới của phân hệ Giao việc, không có màn tương ứng bên hệ thống ERP "
     "cũ và không đồng bộ dữ liệu sang ERP. Toàn bộ liên thông đều nằm trong nội bộ các phân hệ của "
     "hệ thống hiện tại:")
table([
    ["Liên thông với", "Chiều tác động", "Nội dung"],
    ["Giải pháp", "Đọc",
     "Lấy mã, tên, phiên bản hiện hành, Quản lý dự án và phòng tiếp nhận của giải pháp để xác định "
     "điều kiện gửi và danh sách người nhận thông báo."],
    ["Dự án tiềm năng", "Đọc",
     "Xác định ai là nhân viên kinh doanh phụ trách để quyết định ai được tạo phiếu."],
    ["Yêu cầu xây dựng giá", "Ghi",
     "Khi phiếu được tiếp nhận, các yêu cầu đang chờ hoặc đang xây dựng giá của dự án bị chuyển sang Dừng."],
    ["Báo giá", "Ghi",
     "Khi phiếu được tiếp nhận, báo giá đang tạo hoặc đang chờ duyệt bị chuyển sang Dừng và người "
     "lập báo giá nhận thông báo."],
    ["Quản lý tài liệu dùng chung", "Ghi",
     "File đính kèm của phiếu lưu chung với kho tài liệu của hệ thống, không tạo kho riêng."],
    ["Cấu hình phòng ban và người quản lý phòng", "Đọc",
     "Xác định ai là Trưởng phòng của phòng tiếp nhận, bao gồm cả người được tích quản lý tất cả phòng ban."],
])
para("Điểm cố ý làm khác so với cách làm cũ:")
table([
    ["Điểm khác", "Cách làm cũ", "Cách làm hiện tại", "Lý do"],
    ["Kênh đề nghị sửa giải pháp",
     "Trao đổi miệng hoặc qua tin nhắn, không lưu vết.",
     "Phiếu chính thức có mã, người gửi, nội dung và kết quả xử lý.",
     "Cần bằng chứng để quy trách nhiệm và tra cứu về sau."],
    ["Dừng việc làm giá trên phương án cũ",
     "Người làm giá tự biết mà dừng, thường biết muộn nên làm thừa.",
     "Hệ thống tự dừng ngay khi phiếu được tiếp nhận.",
     "Tránh dựng giá và báo giá trên phương án sắp bị thay."],
    ["Nơi ghi nhận việc đổi phương án",
     "Ghi lẫn vào ghi chú của giải pháp.",
     "Tách thành nghiệp vụ riêng, không sửa vào hồ sơ yêu cầu làm giải pháp.",
     "Giữ luồng gọn, không làm rối nghiệp vụ yêu cầu làm giải pháp vốn đã phức tạp."],
])

# ------------------------------------------------------- 11
h("11. GIỚI HẠN HIỆN TẠI", 1)
para("Những điểm dưới đây là hiện trạng đã biết, nêu ra để người dùng và bộ phận kiểm thử không hiểu "
     "nhầm thành lỗi:")
bullet("Người gửi phiếu không nhận được thông báo khi phiếu được tiếp nhận hoặc bị từ chối. Tạm thời "
       "nhân viên kinh doanh phải chủ động vào tab con kiểm tra trạng thái phiếu, hoặc phòng giải "
       "pháp báo lại bằng kênh khác.")
bullet("Không có thao tác thu hồi phiếu. Gửi nhầm thì phải nhờ người có thẩm quyền từ chối phiếu, "
       "ghi lý do là gửi nhầm, rồi gửi phiếu mới.")
bullet("Không sửa và không xoá được phiếu đã gửi, kể cả khi mới gửi xong và chưa ai xử lý.")
bullet("Việc tiếp nhận không tự tạo phiên bản giải pháp mới. Quản lý dự án phải nhớ tự tạo, nếu quên "
       "thì phiếu đã tiếp nhận nhưng giải pháp vẫn giữ nguyên phương án cũ.")
bullet("Yêu cầu xây dựng giá và báo giá đã bị dừng thì không mở lại được. Nếu phòng giải pháp tiếp "
       "nhận nhầm phiếu, các chứng từ đã dừng phải lập lại từ đầu.")
bullet("Hệ thống không cảnh báo khi dự án đã có phiếu đang chờ xử lý, nên có thể tồn tại nhiều phiếu "
       "Đã gửi cùng lúc trên một dự án; mỗi phiếu được xử lý độc lập.")
bullet("Danh sách phiếu chưa có ô tìm kiếm và bộ lọc trạng thái trên giao diện.")
bullet("Chức năng chưa có màn lịch sử thay đổi riêng; dấu vết chỉ gồm người gửi, ngày gửi, người xử "
       "lý và ngày xử lý hiển thị trong chi tiết phiếu.")
bullet("Chức năng chưa có bản in và chưa có xuất Excel.")

# --------------------------------------------------- KIỂM TRA THUẬT NGỮ
BANNED = [
    "status =", "status=", "controller", "service", "entity", "migration",
    "resource", "endpoint", " api ", "request(", "http", "null", "true", "false",
    "database", "query", "cascade", "id =", "table ", "column", "field",
    "solution_id", "pm_id", "main_sale", "employee_info", "423", "422", "vue",
    "laravel", "middleware", "checkpermission", "notifier", "queue", "json",
]
texts = [p.text for p in doc.paragraphs]
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            texts.append(c.text)
hits = []
for txt in texts:
    low = txt.lower()
    for b in BANNED:
        if b in low:
            hits.append((b, txt[:90]))
if hits:
    print("!!! LỌT THUẬT NGỮ KỸ THUẬT:")
    for b, txt in hits:
        print("   -", b, "|", txt)
else:
    print("Kiểm tra thuật ngữ: sạch")

OUT = "Mô tả nghiệp vụ - Yêu cầu điều chỉnh giải pháp.docx"
doc.save(OUT)
print("Đã lưu:", OUT)
